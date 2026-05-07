from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "TDLWebshop_szakdolgozat_javitott_alap.docx"


def set_document_defaults(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.space_after = Pt(6)

    for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
        style = styles[style_name]
        style.font.name = "Times New Roman"
        style.font.color.rgb = RGBColor(17, 24, 39)
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(6)

    styles["Heading 1"].font.size = Pt(16)
    styles["Heading 2"].font.size = Pt(14)
    styles["Heading 3"].font.size = Pt(12)


def add_title_page(document: Document) -> None:
    for _ in range(4):
        document.add_paragraph()

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("TDLWebshop\n")
    run.bold = True
    run.font.size = Pt(22)
    run = title.add_run("Épületgépészeti webshop és adminisztrációs rendszer")
    run.bold = True
    run.font.size = Pt(18)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Szakdolgozat")
    run.font.size = Pt(16)

    for _ in range(7):
        document.add_paragraph()

    meta = document.add_table(rows=4, cols=2)
    meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta.columns[0].width = Cm(5)
    meta.columns[1].width = Cm(8)
    rows = [
        ("Készítette:", "Tóth Dávid László"),
        ("Témavezető:", "Dr. Bilicki Vilmos"),
        ("Képzés / szak:", "[ide írd be a pontos szakot]"),
        ("Év:", "2026"),
    ]
    for row, (label, value) in zip(meta.rows, rows):
        row.cells[0].text = label
        row.cells[1].text = value
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(12)

    for _ in range(6):
        document.add_paragraph()

    note = document.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note.add_run("[Intézmény neve]\n[Kar / tanszék neve]").font.size = Pt(12)
    document.add_page_break()


def add_paragraph(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(text)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def add_placeholder(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    run.italic = True
    run.font.color.rgb = RGBColor(37, 99, 235)


def add_bullets(document: Document, items: list[str]) -> None:
    for item in items:
        document.add_paragraph(item, style="List Bullet")


def add_numbered(document: Document, items: list[str]) -> None:
    for item in items:
        document.add_paragraph(item, style="List Number")


def add_table(document: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for idx, header in enumerate(headers):
        table.rows[0].cells[idx].text = header
    for row_data in rows:
        row = table.add_row()
        for idx, value in enumerate(row_data):
            row.cells[idx].text = value
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(10)


def add_summary(document: Document) -> None:
    document.add_heading("Tartalmi összefoglaló", level=1)
    add_paragraph(
        document,
        "A szakdolgozat célja egy épületgépészeti termékekre specializált webshop és "
        "adminisztrációs rendszer megtervezése és megvalósítása. A TDLWebshop olyan "
        "webalkalmazásként készült, amely egyszerre támogatja a vásárlói folyamatokat, "
        "a termékkatalógus kezelését, a rendeléskezelést, a készletfigyelést, valamint "
        "az adminisztratív és dolgozói feladatokat. A rendszer Angular alapú kliensből, "
        "Firebase szolgáltatásokból, Firestore adatbázisból, biztonsági szabályokból és "
        "egy külön OpenRouter proxyval működő AI asszisztensből áll.",
    )
    add_paragraph(
        document,
        "A dolgozat bemutatja a probléma hátterét, az MVP határát, a követelményeket, "
        "a felhasználói eseteket, az adatmodellt, a jogosultsági rendszert, a biztonsági "
        "megfontolásokat, valamint a tesztelési és validációs folyamatot. A fejlesztés "
        "során fontos szempont volt, hogy a rendszer ne csak funkciók halmaza legyen, "
        "hanem bírálható, reprodukálható és termékszerű szakdolgozati munka.",
    )
    document.add_heading("Abstract", level=1)
    add_paragraph(
        document,
        "The thesis presents the design and implementation of TDLWebshop, a web-based "
        "e-commerce and administration system focused on building engineering products. "
        "The application supports product browsing, cart management, checkout, user "
        "profiles, order tracking, administrative product and order management, stock "
        "monitoring, CSV import, invoice generation and an AI assistant connected through "
        "a server-side OpenRouter proxy.",
    )
    add_paragraph(
        document,
        "The main emphasis of the thesis is not only the implementation itself, but also "
        "the engineering justification of the system: requirements, use cases, architecture, "
        "data model, access control, security, testing and reproducible deployment are "
        "documented as part of the final deliverable.",
    )
    document.add_page_break()


def add_toc(document: Document) -> None:
    document.add_heading("Tartalomjegyzék", level=1)
    add_paragraph(
        document,
        "A tartalomjegyzéket Wordben a véglegesítés előtt frissíteni kell: "
        "Hivatkozások -> Tartalomjegyzék -> Frissítés. Ez a dokumentum szöveges alap, "
        "a végleges oldalszámokat a képernyőképek és ábrák beillesztése után érdemes "
        "rögzíteni.",
    )
    document.add_page_break()


def add_intro(document: Document) -> None:
    document.add_heading("1. Bevezetés, problémafelvetés és célkitűzés", level=1)
    add_paragraph(
        document,
        "Az épületgépészeti termékek értékesítése sajátos terület, mert a vásárlók "
        "jelentős része nem egyszerűen terméket keres, hanem megoldást egy konkrét "
        "szerelési, fűtési, vízvezetékes vagy klímatechnikai feladatra. Egy általános "
        "webshopban gyakran nehéz áttekinteni, hogy egy termék milyen kategóriába "
        "tartozik, van-e készleten, hogyan rendelhető, és milyen adminisztratív folyamat "
        "kapcsolódik hozzá. A TDLWebshop ezt a problémát egy célzott, épületgépészeti "
        "fókuszú rendszerrel kezeli.",
    )
    add_paragraph(
        document,
        "A fejlesztés célja egy olyan MVP megvalósítása volt, amely bemutatja a vásárlói "
        "és az adminisztratív oldal fő folyamatait. A vásárlói oldalon a termékböngészés, "
        "keresés, kosár, checkout, profil, kívánságlista és rendeléskövetés kapott "
        "hangsúlyt. Az admin oldalon a termékfeltöltés, CSV import, rendeléskezelés, "
        "helyszíni vásárlás, készletfigyelés, jogosultságkezelés és bizonylatgenerálás "
        "jelenti a rendszer szakdolgozati magját.",
    )
    add_placeholder(document, "[KÉPERNYŐKÉP HELYE: Kezdőlap dark mode, kategória lenyílóval]")
    add_placeholder(document, "[KÉPERNYŐKÉP HELYE: Kezdőlap AI asszisztenssel, nyitott AI ablakkal]")
    document.add_heading("1.1. MVP-határ", level=2)
    add_paragraph(
        document,
        "Az MVP határa tudatosan úgy lett meghúzva, hogy a rendszer működő webshopként "
        "és adminisztrációs felületként is bemutatható legyen, de ne vállaljon túl nagy "
        "külső integrációs kockázatot. A bankkártyás fizetés ebben az állapotban nem "
        "valódi fizetési gateway-en keresztül történik, hanem fizetési módként jelenik "
        "meg. A számla/bizonylat PDF formában generálódik, de nem helyettesít jogszabályi "
        "értelemben teljes körű számlázóprogramot.",
    )
    add_bullets(
        document,
        [
            "Vásárlói út: termékek böngészése, kosár, checkout, rendelésleadás, profil és státuszkövetés.",
            "Admin út: termékek kezelése, CSV import, készletfigyelés, rendeléskezelés, helyszíni vásárlás.",
            "Biztonság: Firebase Auth, Firestore szabályok, szerepkörök, tiltott felhasználók kezelése.",
            "Dokumentált korlát: a fizetés és az AI asszisztens csak MVP-szintű, kontrollált működésként szerepel.",
        ],
    )


def add_market_requirements(document: Document) -> None:
    document.add_heading("2. Piaci és területi összehasonlítás", level=1)
    add_paragraph(
        document,
        "A piaci vizsgálat célja az volt, hogy a TDLWebshop ne elszigetelt gyakorlófeladatként, "
        "hanem valós webshop-folyamatokhoz viszonyítva legyen értelmezhető. A vizsgálat "
        "épületgépészeti és általános webshopok fő funkcióira koncentrált: terméklista, "
        "szűrés, kosár, checkout, profil, adminisztráció, készletkezelés, dokumentáció és "
        "ügyfélkommunikáció.",
    )
    add_placeholder(document, "[ÁBRA HELYE: Piackutatási táblázat összefoglalója az Excel fájlok alapján]")
    add_paragraph(
        document,
        "A saját rendszer értéke elsősorban abban jelenik meg, hogy a vásárlói és admin "
        "folyamatokat egy szakdolgozati MVP-ben együtt kezeli. A helyszíni vásárlás, "
        "mentett vásárlók, dolgozói jogosultságok, CSV import és AI asszisztens olyan "
        "kiegészítések, amelyek a rendszer termékszerűségét erősítik.",
    )

    document.add_heading("3. Követelmények és use case-ek", level=1)
    add_paragraph(
        document,
        "A követelmények meghatározásánál három fő szerepkör jelent meg: vendég vagy "
        "vásárló, dolgozó és admin. A szerepkörök közötti különbség nem csak a felületben, "
        "hanem a Firestore biztonsági szabályokban és az admin oldali jogosultságkezelésben "
        "is megjelenik.",
    )
    add_table(
        document,
        ["Azonosító", "Követelmény", "Ellenőrzés"],
        [
            ["K-01", "A vásárló tudjon terméket keresni, kosárba tenni és rendelést leadni.", "Kézi checkout teszt"],
            ["K-02", "A rendszer validálja a hibás emailt és telefonszámot.", "Checkout validációs teszt"],
            ["K-03", "Az admin tudjon terméket feltölteni és CSV-ből importálni.", "Admin CSV import teszt"],
            ["K-04", "A dolgozó csak korlátozott admin funkciókat érjen el.", "Jogosultsági teszt"],
            ["K-05", "A rendelés státuszváltása auditálható és készletváltozással járjon.", "Order service teszt és kódrészlet"],
            ["K-06", "Az AI asszisztens ne találjon ki nem létező terméket.", "AI asszisztens domain teszt"],
        ],
    )
    add_placeholder(document, "[ÁBRA HELYE: Use case diagram - vendég/vásárló, dolgozó, admin]")
    add_placeholder(document, "[ÁBRA HELYE: Checkout és rendelésleadás szekvencia diagram]")


def add_ux_tech(document: Document) -> None:
    document.add_heading("4. Felhasználói felület és UX", level=1)
    add_paragraph(
        document,
        "A felület kialakításánál cél volt, hogy a webshop modern, sötét tónusú, technikai "
        "hangulatú arculatot kapjon, miközben a fő vásárlói folyamatok áttekinthetők "
        "maradjanak. A dark/light megjelenés CSS változókkal kezelhető, így a layout "
        "nem változik meg témaváltáskor.",
    )
    for item in [
        "Terméklista szűréssel/kereséssel: termékek oldal.",
        "Termékadatlap: egy konkrét termék oldala.",
        "Kosár: több termékkel.",
        "Checkout validáció: hibás email/telefonszám példával.",
        "Checkout sikeres rendelés: sikeres leadás után.",
        "Profil/rendeléskövetés: vásárlói profil oldal.",
        "Kívánságlista: kívánságlista oldal.",
    ]:
        add_placeholder(document, f"[KÉPERNYŐKÉP HELYE: {item}]")

    document.add_heading("5. Technológiai háttér", level=1)
    add_paragraph(
        document,
        "A kliensoldal Angular alapokra épül, mert komponensalapú szerkezete jól illeszkedik "
        "a webshop különálló nézeteihez: kezdőlap, terméklista, termékadatlap, kosár, "
        "checkout, profil és admin felület. A Firebase a hitelesítést, Firestore adatbázist, "
        "biztonsági szabályokat és hostingot adja. Az AI asszisztenshez külön Cloudflare "
        "Worker proxy készült, hogy az OpenRouter API kulcs ne kerüljön kliensoldali kódba.",
    )
    add_table(
        document,
        ["Technológia", "Szerepe", "Döntés indoka"],
        [
            ["Angular", "Frontend keretrendszer", "Komponensalapú, jól strukturálható kliensoldal"],
            ["Firebase Auth", "Hitelesítés", "Gyors MVP, Firebase ökoszisztéma"],
            ["Firestore", "Adattárolás", "Valós idejű, dokumentumalapú webshop-adatok"],
            ["Firestore Rules", "Adatvédelmi és jogosultsági szabályok", "Kliensoldali appnál kritikus védelmi réteg"],
            ["Cloudflare Worker", "OpenRouter proxy", "API kulcs szerveroldali kezelése Spark Firebase mellett"],
            ["GitHub Actions", "CI ellenőrzés", "Build és teszt reprodukálható bizonyítása"],
        ],
    )


def add_architecture_implementation(document: Document) -> None:
    document.add_heading("6. Architektúra és adatmodell", level=1)
    add_paragraph(
        document,
        "A rendszer architektúrája kliensoldali Angular alkalmazásból, Firebase backend "
        "szolgáltatásokból és külön AI proxyból áll. Az üzleti logika egy része Angular "
        "service-ekben jelenik meg, míg az adatelérési és jogosultsági korlátokat a "
        "Firestore szabályok biztosítják.",
    )
    add_placeholder(document, "[ÁBRA HELYE: Komponens-architektúra ábra - Angular, Firebase, Worker, OpenRouter]")
    add_placeholder(document, "[ÁBRA HELYE: Adatmodell diagram - Product, Cart, Order, OrderItem, UserProfile, Coupon, Invoice, SavedCustomer]")
    add_table(
        document,
        ["Entitás", "Szerep", "Fontos mezők"],
        [
            ["Product", "Termékadat és készlet", "name, sku, category, price, stock, images"],
            ["Order", "Rendelés fejadatai", "customer, items, totals, status, paymentMethod"],
            ["OrderItem", "Rendelési tétel", "productId, name, quantity, unitPrice"],
            ["UserProfile", "Felhasználói profil és szerepkör", "uid, email, role, disabled"],
            ["SavedCustomer", "Mentett helyszíni vásárló", "name, email, phone, company, taxNumber, disabled"],
            ["Coupon", "Kupon és kedvezmény", "code, value, usageLimit, active"],
            ["Invoice", "PDF bizonylat adatai", "invoiceNumber, orderId, totals, issueDate"],
        ],
    )

    document.add_heading("7. Megvalósítás", level=1)
    add_paragraph(
        document,
        "A megvalósítás fejezet a legfontosabb működési egységeket mutatja be: checkout, "
        "rendeléskezelés, készletváltozás, számlagenerálás, admin felület, CSV import, "
        "jogosultságkezelés és AI asszisztens. A kódrészleteknél nem teljes fájlokat, "
        "hanem rövid, magyarázható részleteket érdemes bemutatni.",
    )
    code_refs = [
        "src/pages/checkout/checkout.ts 367-555: rendelés véglegesítése és validáció.",
        "src/app/services/order.service.ts 41-127: státusz, audit és készlet tranzakció.",
        "src/app/services/order.service.ts 222-267: helyszíni rendelés tranzakció.",
        "src/app/services/order.service.ts 269-302: számlaszám generálás.",
        "src/app/services/invoice.service.ts 9-154: PDF számla felépítése.",
        "src/pages/admin/admin.ts 607-748: admin/dolgozói jogosultságok.",
        "src/pages/admin/admin.ts 1181-1257: CSV import validáció és mentés.",
        "src/app/services/chatbot-llm.service.ts 26-90 és 211-236: AI asszisztens domain- és katalóguslogika.",
        "workers/openrouter-proxy/src/index.js 82-197: OpenRouter proxy és szerveroldali kulcskezelés.",
    ]
    for ref in code_refs:
        add_placeholder(document, f"[KÓDRÉSZLET HELYE: {ref}]")
    for item in [
        "Admin áttekintés: admin főnézet.",
        "Admin termékkezelés/CSV import: admin termékek fül.",
        "Admin készletfigyelés: készlet fül.",
        "Helyszíni vásárlás: mentett vásárló kiválasztással.",
        "PDF számla/bizonylat: generált PDF.",
        "Admin felhasználó/jogosultság kezelés: admin felhasználók fül.",
    ]:
        add_placeholder(document, f"[KÉPERNYŐKÉP HELYE: {item}]")


def add_security_testing_ai(document: Document) -> None:
    document.add_heading("8. Biztonság és jogosultságkezelés", level=1)
    add_paragraph(
        document,
        "A biztonsági minimum a rendszer egyik legfontosabb része, mert a webshop "
        "személyes adatokat, rendeléseket, szerepköröket és adminisztratív műveleteket "
        "kezel. A Firebase kliensoldali konfigurációja webalkalmazásnál önmagában nem "
        "titok, viszont a Firestore szabályoknak és a szerepköröknek szigorúan kell "
        "védeniük az adatokat.",
    )
    add_placeholder(document, "[KÓDRÉSZLET HELYE: firestore.rules 25-76 - aktív felhasználó, admin és dolgozó jogosultság]")
    add_placeholder(document, "[KÓDRÉSZLET HELYE: firestore.rules 294-361 - products, orders, users, savedCustomers és audit szabályok]")
    add_bullets(
        document,
        [
            "Admin és dolgozói szerepkörök külön kezelése.",
            "Tiltott felhasználó írási jogosultságának megvonása.",
            "Audit napló törlésének és módosításának tiltása.",
            "OpenRouter API kulcs szerveroldali Worker secretként kezelve.",
            ".env és valódi tokenek kizárása a repóból.",
            "Kuponok és rendelési adatok jogosultsági védelme.",
        ],
    )

    document.add_heading("9. Tesztelés és validáció", level=1)
    add_paragraph(
        document,
        "A tesztelés célja annak bizonyítása, hogy a kritikus vásárlói és adminisztratív "
        "folyamatok működnek. Az automata build és teszt mellett kézi tesztjegyzőkönyv is "
        "készüljön, mert a GUI/UX és jogosultsági folyamatok egy része így ellenőrizhető "
        "a legjobban.",
    )
    add_table(
        document,
        ["Teszt", "Elvárt eredmény", "Bizonyíték"],
        [
            ["Regisztráció és bejelentkezés", "A felhasználó belép, tiltott fiók elutasítva", "Kézi teszt"],
            ["Termékkeresés és kategória", "A terméklista szűrhető és kereshető", "Képernyőkép"],
            ["Kosár", "Mennyiség módosítható, tétel törölhető", "Kézi teszt"],
            ["Checkout validáció", "Hibás email/telefon nem enged tovább", "Képernyőkép"],
            ["Rendelés létrehozás", "Rendelés létrejön, összegzés helyes", "Kézi teszt"],
            ["Admin státuszváltás", "Státusz és készlet konzisztensen változik", "Kézi/kódteszt"],
            ["CSV import", "Valid sorok menthetők, hibás sorok jelöltek", "Admin képernyőkép"],
            ["AI asszisztens", "Csak domain témában válaszol, nem talál ki terméket", "Kézi teszt"],
        ],
    )
    add_placeholder(document, "[KÉPERNYŐKÉP HELYE: GitHub Actions zöld CI futás]")

    document.add_heading("10. Reprodukálhatóság, CI és üzemeltetés", level=1)
    add_paragraph(
        document,
        "A repó akkor bírálható nyugodtan, ha tiszta környezetben is elindítható. Ehhez "
        "README, .env.example, telepítési lépések, demo szerepkörök, CI futás és deploy "
        "leírás szükséges. A node_modules, build mappák, lokális cache-ek és valódi "
        "secret értékek nem részei a beadandó forráskódnak.",
    )
    add_numbered(
        document,
        [
            "Függőségek telepítése: npm install.",
            "Lokális indítás: npm run start.",
            "Build ellenőrzés: npm run build.",
            "Automata tesztek: npm test -- --watch=false.",
            "Firebase hosting deploy: firebase deploy --only hosting.",
            "OpenRouter Worker deploy: workers/openrouter-proxy mappából npx wrangler deploy.",
        ],
    )

    document.add_heading("11. Mesterséges intelligencia használata a fejlesztés során", level=1)
    add_paragraph(
        document,
        "A szakdolgozat készítése során mesterséges intelligenciát támogató eszközöket is "
        "használtam. Ezeket nem önálló fejlesztőként kezeltem, hanem asszisztensként: "
        "ötletelésre, hibakeresési irányok keresésére, kódrészletek ellenőrzésére, "
        "dokumentációs vázlatok rendezésére és tesztelési szempontok összegyűjtésére. "
        "A végső döntések, a kód átnézése, a futtatás és a validáció az én feladatom "
        "maradt.",
    )
    add_paragraph(
        document,
        "Az AI használata leginkább akkor volt hasznos, amikor több lehetséges megoldás "
        "közül kellett választani, például a Firebase Spark csomag korlátai miatt az "
        "OpenRouter API kulcsot nem Firebase Function secretként, hanem Cloudflare Worker "
        "oldalon kellett kezelni. Ilyenkor az AI segített alternatívákat megfogalmazni, "
        "de a választott megoldást a projekt korlátai, a biztonsági szempontok és a "
        "futtathatóság alapján ellenőriztem.",
    )
    add_paragraph(
        document,
        "A fejlesztés során előfordult, hogy egy AI által javasolt megoldás túl általános "
        "vagy a projekt szempontjából nem megfelelő volt. Ilyen példa volt az AI asszisztens "
        "termékajánlási logikája: a cél nem az volt, hogy a modell tetszőleges terméket "
        "kitaláljon, hanem hogy a saját katalógus alapján, ellenőrzött módon ajánljon. "
        "Ezért a kliensoldali logikában külön szűrés készült, amely csak releváns "
        "katalógustalálat esetén jelenít meg terméket, más esetben pedig szakmai irányt "
        "és kapcsolatfelvételi javaslatot ad.",
    )
    add_paragraph(
        document,
        "Tudatosan nem használtam AI-t a dolgozat személyes következtetéseinek végleges "
        "megfogalmazására és a projektértékelés felelősségének átvállalására. Ezeknél a "
        "részeknél fontos, hogy saját tapasztalat alapján jelenjen meg, mit sikerült "
        "megvalósítani, milyen korlátok maradtak, és mit fejlesztenék tovább. Az AI által "
        "adott szövegeket ezért át kell írni saját nyelvezetre, és csak ellenőrzött, "
        "megértett tartalom kerülhet a végleges dolgozatba.",
    )
    add_paragraph(
        document,
        "A rendszer maga is tartalmaz AI funkciót: a vásárlói AI asszisztens OpenRouter "
        "modellen keresztül működik, de a hívás nem közvetlenül a böngészőből történik. "
        "A kliens egy Cloudflare Worker proxyt hív, ahol az API kulcs secretként van "
        "kezelve. A megoldás korlátja, hogy a modell válaszai nem tekinthetők hivatalos "
        "műszaki tanácsadásnak, ezért a felületnek óvatosan kell kommunikálnia: pontos "
        "termékajánlat vagy beszerezhetőség esetén emailes vagy személyes egyeztetés "
        "javasolt.",
    )


def add_conclusion(document: Document) -> None:
    document.add_heading("12. Összefoglalás és továbbfejlesztési lehetőségek", level=1)
    add_paragraph(
        document,
        "A TDLWebshop szakdolgozati MVP-ként bemutatja, hogyan lehet egy épületgépészeti "
        "webshopot vásárlói és adminisztratív oldalról is felépíteni. A rendszer erőssége, "
        "hogy a termékböngészés, kosár, checkout, profil, admin rendeléskezelés, CSV import, "
        "készletfigyelés, jogosultságkezelés, PDF bizonylat és AI asszisztens egy közös "
        "folyamatba illeszkedik.",
    )
    add_paragraph(
        document,
        "Továbbfejlesztési irányként valódi online fizetési integráció, hivatalos számlázó "
        "rendszerhez kapcsolódás, részletesebb készlet-előrejelzés, fejlettebb B2B árképzés, "
        "valamint szerveroldali üzleti logika bővítése jelölhető meg. Ezek nem az MVP "
        "hiányosságai, hanem olyan irányok, amelyek egy éles kereskedelmi rendszer felé "
        "vinnék tovább a projektet.",
    )
    add_paragraph(
        document,
        "[SAJÁT ZÁRÓGONDOLAT HELYE: ide írj 8-12 mondatot arról, mit tanultál a projektből, "
        "melyik része volt a legnehezebb, és mit csinálnál másként egy következő verzióban.]",
    )


def main() -> None:
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    set_document_defaults(document)
    add_title_page(document)
    add_summary(document)
    add_toc(document)
    add_intro(document)
    add_market_requirements(document)
    add_ux_tech(document)
    add_architecture_implementation(document)
    add_security_testing_ai(document)
    add_conclusion(document)
    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
