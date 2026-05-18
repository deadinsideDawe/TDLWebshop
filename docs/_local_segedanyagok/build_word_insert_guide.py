from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parents[2]
OUT_DIR = ROOT / "docs" / "_local_segedanyagok" / "word_beillesztesi_utmutato"
OUT_DIR.mkdir(parents=True, exist_ok=True)
OUT = OUT_DIR / "TDLWebshop_word_beillesztesi_utmutato_kepek_kodok.docx"


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(9)


def add_note(doc, text):
    p = doc.add_paragraph()
    p.style = "Intense Quote"
    p.add_run(text)


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, header in enumerate(headers):
        set_cell_text(hdr[i], header, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
    doc.add_paragraph()


doc = Document()
section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.top_margin = Cm(2.2)
section.bottom_margin = Cm(2.2)
section.left_margin = Cm(2.5)
section.right_margin = Cm(2.5)

styles = doc.styles
styles["Normal"].font.name = "Times New Roman"
styles["Normal"].font.size = Pt(11)
for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
    styles[style_name].font.name = "Times New Roman"

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("TDLWebshop - Word beillesztési útmutató")
run.bold = True
run.font.size = Pt(18)
subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.add_run("Képernyőképek, ábrák, kódrészletek és ábrafeliratok a szakdolgozathoz").italic = True

add_note(
    doc,
    "Ez a dokumentum csak helyi segédanyag. Nem kell beadni és nem kell GitHubra feltölteni. "
    "A célja az, hogy a TDLWebshop_szakdolgozat_javitott_v2.docx fájlban gyorsan tudd cserélni a placeholdereket valódi tartalomra."
)

doc.add_heading("1. Általános szabály", level=1)
doc.add_paragraph(
    "A dolgozatban nem az a cél, hogy minden funkcióról legyen külön kép és hosszú kódrészlet. "
    "A konzulensi visszajelzés alapján inkább azt kell bizonyítani, hogy a rendszer fő folyamatai működnek, "
    "a repo reprodukálható, a jogosultság és adatkezelés átgondolt, valamint a tesztelés dokumentált."
)
doc.add_paragraph(
    "Javasolt mennyiség: 12-15 képernyőkép, 4-6 ábra/táblázat és 5-7 rövid kódrészlet. "
    "Egy kódrészlet lehetőleg 15-35 sor legyen; 40 sor fölé csak akkor menj, ha tényleg szükséges."
)

doc.add_heading("2. Ábrák és táblázatok beillesztése", level=1)
diagram_rows = [
    (
        "Piaci összehasonlítás",
        "2.2 Piaci kitekintés / hasonló rendszerek",
        r"C:\Users\Dell\webshop\docs\_local_segedanyagok\word_abrak\07_piaci_osszehasonlito_abra.svg vagy 07_piaci_osszehasonlito_tablazat.md",
        "2.1. ábra: Épületgépészeti webshopok funkcióinak összehasonlítása saját piackutatás alapján.",
    ),
    (
        "Követelmény-traceability táblázat",
        "3. Követelmények fejezet végére, a use case-ek elé",
        r"C:\Users\Dell\webshop\docs\_local_segedanyagok\word_abrak\09_kovetelmeny_traceability.md",
        "3. táblázat: Követelmény-traceability összefoglaló.",
    ),
    (
        "Use case diagram",
        "3.1 Fő use case-ek elejére",
        r"C:\Users\Dell\webshop\docs\_local_segedanyagok\word_abrak\08_use_case_diagram.svg",
        "3.1. ábra: A TDLWebshop fő use case-ei és szereplői.",
    ),
    (
        "Komponens-architektúra",
        "5. Architektúra / rendszerfelépítés",
        r"C:\Users\Dell\webshop\docs\_local_segedanyagok\word_abrak\01_komponens_architektura.svg",
        "5.1. ábra: A TDLWebshop komponens-architektúrája.",
    ),
    (
        "Adatmodell diagram",
        "5. Adatmodell alfejezet",
        r"C:\Users\Dell\webshop\docs\_local_segedanyagok\word_abrak\03_adatmodell_diagram.svg",
        "5.2. ábra: A TDLWebshop fő Firestore entitásai és kapcsolatai.",
    ),
    (
        "Checkout szekvencia",
        "6. Megvalósítás / rendelési folyamat",
        r"C:\Users\Dell\webshop\docs\_local_segedanyagok\word_abrak\02_checkout_szekvencia.svg",
        "6.1. ábra: A checkout folyamat fő lépései.",
    ),
    (
        "AI asszisztens adatfolyam",
        "6. Megvalósítás / AI asszisztens",
        r"C:\Users\Dell\webshop\docs\_local_segedanyagok\word_abrak\06_ai_asszisztens_adatfolyam.svg",
        "6.2. ábra: Az AI asszisztens adatfolyama és a szerveroldali proxy szerepe.",
    ),
]
add_table(doc, ["Tartalom", "Hova kerüljön", "Forrásfájl", "Javasolt felirat"], diagram_rows)

doc.add_heading("3. Képernyőképek pontos listája", level=1)
shot_rows = [
    ("Kezdőlap dark módban", "GUI/UX fejezet", "Legyen nyitva a kategória lenyíló; látszódjon a kereső, a navbar és a hero.", "A TDLWebshop kezdőlapja dark módban, kategória menüvel."),
    ("Terméklista", "GUI/UX vagy termékböngészés use case", "Használj keresést vagy kategóriaszűrést; látszódjanak termékkártyák és készlet/ár.", "Terméklista kereséssel és kategória szerinti böngészéssel."),
    ("Termékadatlap", "GUI/UX fejezet", "Nyiss meg egy konkrét terméket; legyen kép, ár, készlet, kosárba gomb.", "Termékadatlap részletes termékinformációkkal."),
    ("Kosár oldal", "Kosárkezelés / checkout előtti állapot", "Legalább két termék legyen benne, különböző mennyiséggel.", "Kosár oldal több termékkel és mennyiségkezeléssel."),
    ("Checkout validáció", "Checkout megvalósítás és tesztelés", "Adj meg hibás emailt vagy telefonszámot; látszódjon a hibaüzenet.", "Checkout űrlap mezőszintű validációval."),
    ("Sikeres rendelés", "Rendelési folyamat", "Rendelés leadása után a sikeres visszajelzés vagy rendelési összegzés látszódjon.", "Sikeres rendelésleadás visszajelzése."),
    ("Profil / rendeléstörténet", "Felhasználói funkciók", "Bejelentkezett vásárló rendelései és állapota látszódjon.", "Vásárlói profil és rendeléskövetés."),
    ("Kívánságlista", "Felhasználói funkciók", "Tegyél bele 2-3 terméket, majd fotózd az oldalt.", "Kívánságlista bejelentkezett felhasználóhoz kötve."),
    ("Admin áttekintés", "Admin funkciók", "Látszódjanak stat kártyák vagy fő admin navigáció.", "Admin áttekintő felület."),
    ("Admin termékkezelés / CSV import", "CSV import megvalósítás", "Látszódjon a CSV feltöltő rész vagy import összegzés.", "Admin termékkezelés és CSV import felülete."),
    ("Admin rendeléskezelés", "Rendelés státuszváltás", "Legyen látható státusz módosítása vagy rendelés részletei.", "Admin rendeléskezelés státuszmódosítással."),
    ("Helyszíni vásárlás", "B2B/helyszíni értékesítés", "Mentett vásárló kiválasztva, termékek hozzáadva.", "Helyszíni vásárlás rögzítése mentett vásárlóval."),
    ("PDF bizonylat", "PDF generálás", "Nyisd meg a generált PDF-et, látszódjon fejléc, vevő, tételek, végösszeg.", "Generált PDF bizonylat / számla nézete."),
    ("AI asszisztens", "AI asszisztens fejezet", "Tegyél fel katalógushoz kötött kérdést; látszódjon, hogy nem random terméket talál ki.", "AI asszisztens katalógushoz kötött válasszal."),
    ("GitHub Actions zöld CI", "Tesztelés és reprodukálhatóság", "A legutóbbi zöld workflow run látszódjon.", "GitHub Actions sikeres CI futás."),
]
add_table(doc, ["Képernyőkép", "Hova kerüljön", "Hogyan készítsd", "Javasolt ábrafelirat"], shot_rows)

doc.add_heading("4. Kódrészletek, amiket érdemes betenni", level=1)
code_rows = [
    ("Checkout véglegesítés", r"C:\Users\Dell\webshop\src\pages\checkout\checkout.ts", "376-570", "Rendelésleadás: validáció, tiltott felhasználó kezelése, rendelés mentése, profiladatok frissítése. Ebből csak kb. 25-35 sort válassz, ne az egész blokkot."),
    ("Checkout validáció", r"C:\Users\Dell\webshop\src\pages\checkout\checkout.ts", "590-651", "Email, telefonszám, kötelező mezők és kuponvalidáció. Ez jó kódrészlet a GUI/UX és tesztelés fejezethez."),
    ("Státusz, audit, készlet tranzakció", r"C:\Users\Dell\webshop\src\app\services\order.service.ts", "41-124", "Admin státuszváltáskor audit napló és készletváltozás egy tranzakcióban. Ez erős szakmai pont."),
    ("Helyszíni rendelés tranzakció", r"C:\Users\Dell\webshop\src\app\services\order.service.ts", "244-275", "Helyszíni vásárlásnál készletellenőrzés és rendelés létrehozása tranzakcióban."),
    ("Számlaszám generálás", r"C:\Users\Dell\webshop\src\app\services\order.service.ts", "313-345", "Éves számlaszámláló és rendeléshez kapcsolt számlaszám tranzakcióval."),
    ("PDF bizonylat felépítése", r"C:\Users\Dell\webshop\src\app\services\invoice.service.ts", "81-154", "Kliensoldali PDF bizonylat felépítése, demo/MVP korlát megjegyzésével."),
    ("Firestore jogosultsági alapok", r"C:\Users\Dell\webshop\firestore.rules", "30-67", "Aktív felhasználó, admin és dolgozói jogosultságok alapfüggvényei."),
    ("Firestore fő collection szabályok", r"C:\Users\Dell\webshop\firestore.rules", "288-356", "Products, orders, users, savedCustomers és audit szabályok."),
    ("CSV import validáció", r"C:\Users\Dell\webshop\src\pages\admin\admin.ts", "1183-1258", "CSV előnézet, validáció, import mentése. Elég belőle egy rövidebb részlet."),
    ("Dolgozói jogosultságok", r"C:\Users\Dell\webshop\src\pages\admin\admin.ts", "654-734", "Admin/dolgozó jogosultságok kezelése és alapértelmezések."),
    ("AI asszisztens klienslogika", r"C:\Users\Dell\webshop\src\app\services\chatbot-llm.service.ts", "31-84", "A modell nem állítható kliensoldalról, domain/katalógus logika indítása."),
    ("OpenRouter proxy", r"C:\Users\Dell\webshop\workers\openrouter-proxy\src\index.js", "153-211", "Szerveroldali OpenRouter hívás, secret kezelés, CORS ellenőrzés. Ez a legfontosabb AI-proxy kódrészlet."),
]
add_table(doc, ["Téma", "Fájl", "Sorok", "Mit magyarázz mellé"], code_rows)

doc.add_heading("5. Mit érdemes kihagyni", level=1)
for item in [
    "Ne tegyél be teljes fájlokat kódrészletként. A dolgozatban a kód csak bizonyíték, nem forráskód-lista.",
    "Ne maradjon benne [KÉP / ÁBRA HELYE], [KÓDRÉSZLET HELYE] vagy saját munkautasítás.",
    "Ne kerüljön be segédanyag, AI prompt, belső jegyzet vagy olyan fájl, ami csak neked készült.",
    "Ne legyen túl sok hasonló képernyőkép. Egy funkcióról elég egy jól megválasztott állapot.",
    "Ne írj olyat, hogy a rendszer éles számlázó vagy teljes NAV-kompatibilis megoldás. Ezt MVP/demo korlátként kezeld.",
]:
    doc.add_paragraph(item, style="List Bullet")

doc.add_heading("6. Javasolt mai munkasorrend", level=1)
steps = [
    "Nyisd meg a TDLWebshop_szakdolgozat_javitott_v2.docx fájlt.",
    "Keresd meg az összes [KÉP / ÁBRA HELYE] és [KÓDRÉSZLET HELYE] jelölést.",
    "Először az ábrákat cseréld: piaci összehasonlítás, traceability, use case, architektúra, adatmodell, checkout szekvencia.",
    "Utána készítsd el a 12-15 képernyőképet, és minden kép alá tegyél ábrafeliratot.",
    "A kódrészletekből csak rövid, szakmailag magyarázható blokkokat illessz be.",
    "Frissítsd a tartalomjegyzéket és az ábra-/táblázatfeliratokat Wordben.",
    "A végén olvasd át, hogy ne maradjon munkautasítás vagy placeholder.",
]
for i, step in enumerate(steps, 1):
    doc.add_paragraph(f"{i}. {step}")

doc.add_heading("7. Rövid válasz a kérdésedre", level=1)
doc.add_paragraph(
    "Nem muszáj óriási kódrészleteket betenni, sőt nem is ajánlott. A dolgozat akkor lesz jobb, ha kevesebb, de jól kiválasztott "
    "kódrészlet szerepel benne, és mindegyikhez 1-2 bekezdés saját magyarázat társul. Képernyőképből sem kell minden kattintásról kép: "
    "a fenti 12-15 darab már jól lefedi a vásárlói, admin, PDF, AI és CI bizonyítékokat."
)

doc.add_paragraph()
footer = doc.add_paragraph()
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer.add_run("Helyi segédanyag - nem beadandó dokumentum").italic = True

doc.save(OUT)
print(OUT)
