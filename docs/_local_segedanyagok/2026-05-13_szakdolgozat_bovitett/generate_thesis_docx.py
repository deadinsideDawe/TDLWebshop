from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


OUT_DIR = Path(__file__).resolve().parent
DOCX_PATH = OUT_DIR / "TDLWebshop_szakdolgozat_40_50_oldalas_munkapeldany.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Times New Roman"
    run.font.size = Pt(10)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_borders(table):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for border_name in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = OxmlElement(f"w:{border_name}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "BFC7D5")
        borders.append(border)
    tbl_pr.append(borders)


def add_toc(paragraph):
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = 'TOC \\o "1-3" \\h \\z \\u'
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    paragraph._p.append(fld_begin)
    paragraph._p.append(instr_text)
    paragraph._p.append(fld_sep)
    run = paragraph.add_run("A tartalomjegyzek frissitesehez Wordben jobb klikk -> Mezo frissitese.")
    run.italic = True
    paragraph._p.append(fld_end)


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = "Times New Roman"
        run.font.color.rgb = RGBColor(17, 24, 39)
    return p


def add_para(doc, text, italic=False, bold=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    run.italic = italic
    run.bold = bold
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(item)
        run.font.name = "Times New Roman"
        run.font.size = Pt(11)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(item)
        run.font.name = "Times New Roman"
        run.font.size = Pt(11)


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_borders(table)
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_text(hdr[i], h, bold=True)
        set_cell_shading(hdr[i], "D9EAFD")
        if widths:
            hdr[i].width = Cm(widths[i])
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], str(value))
            if widths:
                cells[i].width = Cm(widths[i])
    doc.add_paragraph()
    return table


def add_figure(doc, caption, instruction, height_lines=12):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_borders(table)
    cell = table.cell(0, 0)
    set_cell_shading(cell, "EEF2F7")
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("IDE KERUL A TENYLEGES ABRA / KEPERNYOKEP\n")
    r.bold = True
    r.font.name = "Times New Roman"
    r.font.size = Pt(12)
    r2 = p.add_run(instruction)
    r2.font.name = "Times New Roman"
    r2.font.size = Pt(10)
    for _ in range(height_lines):
        cell.add_paragraph("")
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap.add_run(caption)
    run.italic = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(10)


def add_code_placeholder(doc, title, file_ref, why):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_borders(table)
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F8FAFC")
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(title + "\n")
    r.bold = True
    r.font.name = "Consolas"
    r.font.size = Pt(10)
    r2 = p.add_run(f"Kodreszlet helye: {file_ref}\n")
    r2.font.name = "Consolas"
    r2.font.size = Pt(9)
    r3 = p.add_run(why)
    r3.font.name = "Times New Roman"
    r3.font.size = Pt(10)
    for _ in range(6):
        cell.add_paragraph("")
    doc.add_paragraph()


def new_chapter(doc, title):
    doc.add_page_break()
    add_heading(doc, title, 1)


def add_expanded_evidence_section(doc):
    """Local thesis helper content: detailed proof blocks and full-page figure slots."""
    new_chapter(doc, "Reszletes bizonyitasi es abra-beillesztesi terv")
    add_para(
        doc,
        "A konzulensi visszajelzes alapjan a dolgozat erteket nem tovabbi uj funkciok, hanem a meglovo rendszer "
        "bizonyitasa noveli. Ezert ebben a reszben a fo kepernyok, kodreszletek, tesztek es dontesi pontok olyan "
        "sorrendben szerepelnek, ahogyan a vegleges dolgozatban is erdemes bemutatni oket. A cel az, hogy a biro "
        "ne csak azt lassa, hogy a TDLWebshop mukodik, hanem azt is, hogy a mukodes mogott tudatos kovetelmenyek, "
        "adatmodell, jogosultsagi logika, validacio es teszteles all.",
    )
    add_para(
        doc,
        "A kovetkezo alfejezetek nem puszta keplistak. Minden abrahelyhez rovid magyarazat tartozik: milyen "
        "felhasznaloi folyamatot bizonyit, melyik kovetelmenyhez kapcsolodik, es miert fontos a szakdolgozat "
        "szempontjabol. A vegleges leadas elott a szurke helyorzo dobozokat tenyleges kepernyokepekre kell "
        "cserelni, a magyarazo szoveget pedig sajat megfogalmazasra erdemes atirni.",
    )

    add_heading(doc, "Kiemelt GUI/UX kepernyok", 2)
    gui_figures = [
        ("Kezdolap dark modban, kategoria lenyiloval", "A nyitokepernyo a brandet, a keresest, a kategoriakat es a modern dark/light temarendszert mutatja be."),
        ("Kezdolap AI asszisztenssel", "Ez bizonyitja, hogy a webshopban nem kulonallo demo, hanem felhasznaloi oldalrol elerheto AI tamogatas mukodik."),
        ("Termeklista kategoria- es kereses-szurovel", "A termeklista mutatja a webshop alapvetu bongeszesi folyamatat es a katalogus kezelhetoseget."),
        ("Termekadatlap tobb keppel", "A termekadatlap a vasarloi donteshez szukseges informaciokat es a kosarba helyezest bizonyitja."),
        ("Kosar tobb termekkel", "A kosar a mennyisegmodositast, torlest es osszesitest mutatja be."),
        ("Checkout validacios hibaval", "A hibas email vagy telefonszam peldaja bizonyitja, hogy a rendszer nem engedi vakon tovabb a hibas bemenetet."),
        ("Checkout sikeres rendelessel", "A sikeres rendeles leadasa a vasarloi ut egyik legfontosabb bizonyiteka."),
        ("Profil es rendeleskovetes", "A regisztralt vasarlo itt lathatja korabbi rendeleseit es azok allapotat."),
        ("Kivansaglista", "A kedvenc termekek mentese a vasarloi elmenyt es visszaterest tamogatja."),
        ("Kapcsolat oldal", "A kapcsolat sablon mutatja, hogy a webshop nem csak termekkatalogus, hanem ugyfelszolgalati felulet is."),
    ]
    for index, (caption, explanation) in enumerate(gui_figures, start=1):
        add_para(doc, f"{index}. kepernyokep: {caption}. {explanation}")
        add_figure(
            doc,
            f"{index}. abra: {caption}",
            "Ide illeszd be a webshop megfelelo oldalarol keszult kepernyokepet. A kep alatt maradjon rovid, targyszeru abrafelirat.",
            18,
        )
        doc.add_page_break()

    add_heading(doc, "Adminisztracios es dolgozoi folyamatok kepernyoi", 2)
    admin_figures = [
        ("Admin attekintes", "A statisztikai kartyak es rendelesi osszesitok a vezeto/admin szerepkor fo kepernyojet adjak."),
        ("Admin termekkezeles es CSV import", "Ez mutatja a tomeges termekfeltoltes szakdolgozati erteket, mert nagyobb katalogus kezeleset bizonyitja."),
        ("Keszletfigyeles", "A keszletmodul a domainhez kapcsolodo tobbletfunkcio, amely elter az egyszeru webshop CRUD-tol."),
        ("Helyszini vasarlas mentett vasarloval", "A B2B/helyszini ertekesitesi folyamat a TDLWebshop egyik legerosebb sajatossaga."),
        ("Mentett vasarlok kezelese", "A szerkesztes, torles es tiltasi logika a valos adminisztracios igenyekhez kapcsolodik."),
        ("Felhasznalo- es jogosultsagkezeles", "Itt latszik az admin, dolgozo es vasarlo szerepkorok elvalasztasa."),
        ("Ertesitesek es fizetesi hatarido", "A jovahagyasi helyzetek kezeleset es a dolgozoi korlatokat szemlelteti."),
        ("PDF szamla vagy bizonylat", "A generalt dokumentum a rendelesi folyamat lezart, letoltheto kimenete."),
    ]
    for index, (caption, explanation) in enumerate(admin_figures, start=1):
        add_para(doc, f"{index}. admin kep: {caption}. {explanation}")
        add_figure(
            doc,
            f"Admin abra {index}: {caption}",
            "Ide keruljon a megfelelo admin feluletrol keszult kepernyokep. Figyelj arra, hogy tesztadat vagy demo adat szerepeljen rajta.",
            18,
        )
        doc.add_page_break()

    add_heading(doc, "Kodreszletek magyarazata", 2)
    add_para(
        doc,
        "A dolgozatba nem erdemes tul hosszu kodblokkokat beilleszteni. A cel az, hogy nehany szakmailag fontos "
        "reszlet latszodjon: tranzakcios rendeleskezeles, inputvalidacio, jogosultsag, PDF generalas, CSV import "
        "es AI proxy. A kodreszletek alatt mindig rovid magyarazat kell, amely leirja, hogy a reszlet milyen "
        "kovetelmenyt valosit meg es milyen hibakat eloz meg.",
    )
    code_evidence = [
        ("Checkout veglegesites", "src/pages/checkout/checkout.ts 367-555 es 580-637", "A webes rendeles leadasa, validacio, kupon es osszesites osszekapcsolasa."),
        ("Rendeles statusz, audit es keszlet", "src/app/services/order.service.ts 41-128", "A statuszvaltas es keszletkezeles egyutt mutatja a rendeles eletciklusat."),
        ("Helyszini rendeles tranzakcio", "src/app/services/order.service.ts 229-272", "Az admin/dolgozoi eladasnal fontos, hogy a termekek foglalasa es a rendeles egyutt tortenjen."),
        ("Szamlaszam generalas", "src/app/services/order.service.ts 276-310", "A PDF bizonylat egyedi azonositasa es kovethetosege miatt lenyeges."),
        ("PDF bizonylat felepitese", "src/app/services/invoice.service.ts 8-183", "A rendelesbol letoltheto dokumentum keszul, amely a helyszini vasarlasnal kulcsfontossagu."),
        ("Firestore jogosultsagok", "firestore.rules 25-76 es 288-356", "A szerveroldali szabalyok vedik a termekeket, rendelest, felhasznaloi adatokat es auditot."),
        ("Admin/dolgozo UI jogosultsag", "src/pages/admin/admin.ts 606-735", "A felulet a szerepkor szerint korlatozza az elerheto muveleteket."),
        ("CSV import validacio", "src/pages/admin/admin.ts 1183-1258", "A tomeges feltoltesnel hibas sorok szurese es elonezet szukseges."),
        ("AI kataloguslogika", "src/app/services/chatbot-llm.service.ts 31-88, 92-120, 214-250", "A valaszadas domainhez es sajat termekkatalogushoz kotott."),
        ("OpenRouter proxy", "workers/openrouter-proxy/src/index.js 1-59 es 153-211", "Az API kulcs nem kerul kliensoldalra, hanem kulso Worker kezeli a hivast."),
    ]
    for index, (title, ref, why) in enumerate(code_evidence, start=1):
        add_code_placeholder(doc, f"Kodreszlet {index}: {title}", ref, why)
        add_para(doc, f"Magyarazat: {why} A vegleges dolgozatban ide 6-12 soros kodkep vagy rovid kodblokk elegendo.")
        if index % 2 == 0:
            doc.add_page_break()

    add_heading(doc, "Tesztelesi bizonyitekok es kezi ellenorzes", 2)
    add_para(
        doc,
        "A teszteles fejezetben kulon kell valasztani az automata teszteket, a build/CI bizonyitekait es a kezi "
        "tesztjegyzokonyvet. A binalo szamara az a fontos, hogy a kritikus folyamatok ellenorzese reprodukalhato "
        "legyen: latszodjon, milyen bemenettel, milyen jogosultsaggal es milyen vart eredmennyel tortent a teszt.",
    )
    test_rows = [
        ("Regisztracio es bejelentkezes", "Vasarloi fiok letrejon, hibas adatnal visszajelzes jelenik meg.", "Kezi teszt + auth kepernyokep"),
        ("Tiltott felhasznalo", "Tiltott profil nem tud rendelni vagy belepni a vedett funkciokba.", "Kezi teszt + Firestore szabaly"),
        ("Termekkereses", "Nev, kategoria vagy cikkszam alapjan szukitheto a lista.", "Kezi teszt"),
        ("Kosar", "Mennyiseg modosithato, torles es vegosszeg frissul.", "Automata/kezi teszt"),
        ("Kupon", "Ervenyes kupon kedvezmenyt ad, ervenytelen hibat jelez.", "Automata teszt"),
        ("Checkout", "Hibas email/telefon megallitja a folyamatot, jo adatnal rendelest hoz letre.", "Automata/kezi teszt"),
        ("Admin statuszvaltas", "Statusz modosul, audit rekord keszul.", "Automata/kezi teszt"),
        ("Keszlet", "Helyszini rendelesnel a keszlet csokken, hiany eseten hiba kezelheto.", "Kezi teszt"),
        ("CSV import", "Hibas sorok nem mentodnek, ervenyes sorok feltolthetok.", "Kezi teszt"),
        ("PDF bizonylat", "Letoltheto, olvashato dokumentum keszul.", "Automata/kezi teszt"),
        ("AI asszisztens", "Relevans domain kerdesre valaszol, irrelevans kerdesnel korlatoz.", "Kezi teszt"),
    ]
    add_table(doc, ["Tesztelt funkcio", "Elvart eredmeny", "Bizonyitek"], test_rows, widths=[4, 8, 4])
    add_para(
        doc,
        "A vegleges dolgozatba ide keruljon egy kepernyokep a sikeres npm run build futasrol, az npm test "
        "eredmenyerol, valamint a GitHub Actions zold CI futasarol. Ha az npm audit tovabbra is jelez talalatot, "
        "azt nem szabad elhallgatni: MVP-korlatozaskent es tovabbfejlesztesi feladatkent kell leirni.",
    )
    add_figure(
        doc,
        "CI bizonyitek: GitHub Actions zold futas",
        "Ide illeszd be a GitHub Actions kepernyot, ahol a legfrissebb workflow zold pipaval latszik.",
        18,
    )
    doc.add_page_break()

    add_heading(doc, "Konzulensi kockazati pontok kezelese", 2)
    risk_rows = [
        ("npm audit talalatok", "A fuggosegek kozul nehany serulekenyseget jelezhet.", "Frissites vagy tudatos MVP-korlatozaskent dokumentalas."),
        ("Kliensoldali checkout osszesites", "A webes kosar es vegosszeg a kliensoldali allapotbol indul.", "Szerveroldali ellenorzes jovobeli fejleszteskent, Firestore szabalyokkal vedett iras."),
        ("Webes keszletkezeles", "A helyszini eladas szigorubb, mint a webes rendeles.", "Teljesiteskori ellenorzes es keszlethiba kezelese tovabbfejlesztes."),
        ("Dolgozoi jogosultsag", "A tul tag alapertelmezes kockazat lehet.", "Szerepkorok es lathatosag dokumentalasa, least privilege irany tovabbfejlesztesben."),
        ("Vendeg email azonositas", "Email alapjan torteno vendeg kapcsolat adatvedelmi kockazatot jelenthet.", "Minimalis adatkezeles es kesobbi erosebb azonositasi folyamat."),
        ("AI proxy visszaeles", "A CORS onmagaban nem teljes kvotavedelem.", "Rate limit vagy kvota bevezetese jovobeli eles uzemhez."),
    ]
    add_table(doc, ["Kockazat", "Mi a problema?", "Dolgozati kezeles"], risk_rows, widths=[4, 7, 5])
    add_para(
        doc,
        "Ezek a pontok nem feltetlenul rontjak a szakdolgozat erteket, ha oszinten es mernoki modon szerepelnek a "
        "korlatok kozott. Egy MVP eseteben nem az a cel, hogy minden ipari uzemi kockazat teljesen megoldott legyen, "
        "hanem hogy a fejleszto felismerje a kockazatokat, es meg tudja mondani, milyen iranyban lehet oket kezelni.",
    )

    add_heading(doc, "Onallo zaroertekeles vazlata", 2)
    add_para(
        doc,
        "A sajat zaro reflexioban erdemes konkret peldakat hozni: melyik hiba okozott sok munkat, hogyan valtozott "
        "a Firestore jogosultsag, miert lett kulon Worker az AI proxyhoz, es mit tanultal a tesztelesbol. Ezt a reszt "
        "nem szabad sablonosan hagyni, mert a vedesen is hitelesebb, ha a sajat tapasztalataid jelennek meg benne.",
    )
    add_bullets(
        doc,
        [
            "Mit sikerult a projekt eleji tervhez kepest megvalositani?",
            "Melyik funkcio lett erosebb, mint eredetileg gondoltad?",
            "Melyik pont maradt tudatos MVP-korlatozas?",
            "Miben segitett az AI, es hol kellett sajat ellenorzes?",
            "Mit fejlesztenel tovabb, ha a rendszer eles uzemi webshop lenne?",
        ],
    )


def build_document():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    styles = doc.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"].font.size = Pt(12)
    for name in ("Title", "Heading 1", "Heading 2", "Heading 3"):
        styles[name].font.name = "Times New Roman"

    # Cover
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Szegedi Tudomanyegyetem\n").bold = True
    p.add_run("[Kar / Intezet / Tanszek pontos neve]\n")
    p.add_run("[Szak pontos neve]\n\n\n")
    title = p.add_run("TDLWebshop\n")
    title.bold = True
    title.font.size = Pt(24)
    subtitle = p.add_run("Epuletegepeszeti webshop es adminisztracios rendszer fejlesztese\n")
    subtitle.font.size = Pt(16)
    p.add_run("\nSzakdolgozat\n\n\n\n")
    p.add_run("Keszitette: Toth David Laszlo\n")
    p.add_run("Temavezeto: Dr. Bilicki Vilmos, egyetemi docens\n\n\n")
    p.add_run("Szeged\n2026")
    doc.add_page_break()

    add_heading(doc, "Feladatkiiras helye", 1)
    add_para(
        doc,
        "Ezen az oldalon a hivatalos szakdolgozati feladatkiiras vagy annak roviditett, "
        "a temavezetovel egyeztetett valtozata helyezheto el. A vegleges leadas elott "
        "ide kell bemasolni vagy kepkent beilleszteni az intezmeny altal elfogadott feladatkiirast.",
    )
    add_figure(
        doc,
        "1. abra: A hivatalos feladatkiiras helye",
        "Ide illeszd be a hivatalos feladatkiirast vagy annak scannelt/Wordbol exportalt kepet.",
        15,
    )
    doc.add_page_break()

    add_heading(doc, "Tartalmi osszefoglalo", 1)
    add_para(
        doc,
        "A szakdolgozat celja egy epuletegepeszeti termekeket kezelo, vasarloi es adminisztracios "
        "folyamatokat egy rendszerben osszefogo webshop prototipusanak bemutatasa. A TDLWebshop "
        "nem pusztan termekek listazasara keszult, hanem olyan MVP-kent, amelyben a termekkereses, "
        "kosarkezeles, megrendeles, profil alapu rendeleskovetes, adminisztracios termek- es "
        "rendeleskezeles, keszletfigyeles, PDF bizonylat, CSV import es AI asszisztens egyutt jelenik meg.",
    )
    add_para(
        doc,
        "A dolgozat bemutatja a rendszer celkituzeset, a hasonlo megoldasokhoz viszonyitott helyet, "
        "a kovetelmenyeket es use case-eket, majd reszletesen ismerteti az Angular es Firebase alapu "
        "megvalositast. Kulon fejezet foglalkozik a biztonsagi minimumokkal, a tesztelessel, a "
        "reprodukalhatosaggal es a fejlesztes soran alkalmazott mesterseges intelligencia szerepevel.",
    )
    add_para(
        doc,
        "Kulcsszavak: webshop, Angular, Firebase, Firestore, adminisztracio, epuletegepeszet, "
        "rendeleskezeles, keszletkezeles, AI asszisztens, szakdolgozati MVP.",
        italic=True,
    )
    doc.add_page_break()

    add_heading(doc, "Tartalomjegyzek", 1)
    add_toc(doc.add_paragraph())

    # Chapter 1
    new_chapter(doc, "1. Bevezetes")
    add_para(
        doc,
        "A webaruhazak fejlesztese ma mar nem csak termekek megjeleniteset jelenti. Egy hasznalhato "
        "rendszerben a vasarloi folyamat, az adminisztracio, a keszletadatok, a jogosultsagkezeles, "
        "a hibas bemenetek kezelese es a megrendelesek utokovetese egyarant fontos szerepet kap. "
        "Az epuletegepeszeti termekek kereskedelme kulonosen olyan terulet, ahol a vasarlo gyakran "
        "nem egyetlen termeket keres, hanem megoldast egy konkret szerelesi vagy felujitasi feladatra.",
    )
    add_para(
        doc,
        "A TDLWebshop fejlesztese soran az volt a celom, hogy egy olyan webshop prototipust keszitsek, "
        "amely a hagyomanyos vasarloi mukodes mellett tartalmaz adminisztracios es helyszini ertekesitesi "
        "folyamatokat is. Ez azert fontos, mert egy kisebb epuletegepeszeti vallalkozasnal a webes "
        "ertekesites es a szemelyes, helyszini kiszolgalas sokszor ugyanahhoz a keszlethez es ugyanahhoz "
        "a rendelesi nyilvantartashoz kapcsolodik.",
    )
    add_para(
        doc,
        "A dolgozatban nem egy teljes ertekes, jogilag minden reszleteben eles uzemre felkeszitett "
        "webshopot mutatok be, hanem egy szakdolgozati MVP-t. Az MVP mar bizonyitja a rendszer "
        "legfontosabb mukodeseit, de nehany teruleten tudatos korlatokat tartalmaz. Ilyen korlat peldaul "
        "a webes checkout teljes szerveroldali ar- es keszletellenorzesenek kesobbi megerositese, valamint "
        "az AI asszisztens hasznalatakor a kvota- es rate limit kezeles tovabbfejlesztese.",
    )
    add_heading(doc, "1.1. Motivacio", 2)
    add_para(
        doc,
        "A projekt motivacioja az volt, hogy egy valos domainhez kapcsolodo, tobb szereplos rendszert "
        "keszitsek. Az epuletegepeszeti webshop tema lehetoseget ad arra, hogy egyszerre jelenjen meg "
        "a felhasznaloi felulet, az adminisztracios felulet, az adatmodell, a jogosultsagkezeles, a "
        "rendelesi folyamat es a keszletmodositas. Ez szakdolgozati szempontbol kedvezo, mert nem egy "
        "elszigetelt CRUD modulrol van szo, hanem osszekapcsolt folyamatokrol.",
    )
    add_heading(doc, "1.2. A dolgozat felepitese", 2)
    add_para(
        doc,
        "A dolgozat eloszor a piaci es domain hatteret mutatja be, majd meghatarozza a rendszer celjat es "
        "MVP-hatarat. Ezutan kovetkeznek a kovetelmenyek, use case-ek, GUI/UX szempontok, technologiai "
        "dontesek, architektura es adatmodell. A megvalositas fejezet a fontosabb mukodeseket reszletezi, "
        "majd a biztonsagi es tesztelesi fejezetek mutatjak be, hogyan ellenorizheto es vedheto a rendszer. "
        "Kulon fejezetben szerepel a mesterseges intelligencia hasznalata, vegul az eredmenyek es a "
        "tovabbfejlesztesi lehetosegek zarjak a dolgozatot.",
    )
    add_figure(
        doc,
        "2. abra: A TDLWebshop kezdolapja dark modban",
        "Kepernyokep: fooldal, dark mode, kategoria lenyilo megnyitva. Ez mutatja az elso benyomast es a navigaciot.",
        10,
    )

    # Chapter 2
    new_chapter(doc, "2. Piaci es domain kitekintes")
    add_para(
        doc,
        "Az epuletegepeszeti webaruhazak kozott gyakori, hogy a felulet nagy termekkatalogust kezel, "
        "de a szakmai dontestamogatas vagy a helyszini ertekesites tamogatasa kevesbe hangsulyos. "
        "A TDLWebshop fejlesztese soran ezert nem csak egy altalanos webshop mintat kovettem, hanem "
        "olyan funkciokat is figyelembe vettem, amelyek egy szereloi vagy kisvallalkozoi mukodesben hasznosak.",
    )
    add_para(
        doc,
        "A vizsgalt piaci kornyezetben a legtobb webshop eros a termeklistazasban es a kosar/checkout "
        "folyamatban, de a dolgozoi jogosultsag, a helyszini vasarlas rogzitese, a keszletfigyeles es az "
        "adminisztracios bizonyithatosag mar kevesbe lathato a kulso felhasznaloi feluleten. A sajat "
        "rendszer erteke abban van, hogy ezekbol tobb elemet egyetlen szakdolgozati prototipusba kapcsol ossze.",
    )
    add_heading(doc, "2.1. Hasonlo rendszerek osszehasonlitasa", 2)
    add_table(
        doc,
        ["Szempont", "Altalanos epuletegepeszeti webshop", "Barkacsaruhaz webshop", "TDLWebshop MVP"],
        [
            ["Termekkatalogus", "Nagy valasztek, sok szuro", "Nagyon szeles termekkinalat", "60 termekes, domainhez igazodo katalogus"],
            ["Admin funkciok", "Kivulrol nem lathato", "Belsos rendszerben lehet", "Sajat admin panel termekkel, rendelesevel, CSV-vel"],
            ["Helyszini eladas", "Nem jellemzo", "Bolti kasszarendszer", "Adminbol rogzitett helyszini vasarlas PDF bizonylattal"],
            ["Keszletfigyeles", "Altalaban van keszletjelzes", "Raktari integracio", "Alacsony keszlet es admin keszletnezet"],
            ["AI tamogatas", "Ritkan lathato", "Altalanos chatbot lehet", "Katalogushoz kotott, domainkorlatos AI asszisztens"],
        ],
        widths=[4, 4.2, 4.2, 4.2],
    )
    add_para(
        doc,
        "A tablazat alapjan a TDLWebshop nem abban probal versenyezni, hogy termekszamban vagy logisztikai "
        "hatterszolgaltatasban felulmuljon egy nagyvallalati webaruhazat. A szakdolgozati erteket inkabb "
        "az adja, hogy egy kisebb, szakmai celcsoporthoz kapcsolodo rendszerben egyutt jelenik meg a vasarloi "
        "ut, az adminisztracio, a jogosultsagkezeles es a bizonyithato mukodes.",
    )
    add_heading(doc, "2.2. Domain sajatossagok", 2)
    add_para(
        doc,
        "Az epuletegepeszeti termekeknel a vasarloi kerdes gyakran szakmai kontextushoz kotodik. Egy radiator, "
        "kazan, bojler, klima vagy szerelveny kivalasztasa nem mindig dontheto el pusztan ar alapjan. A "
        "felhasznalonak sokszor meretre, teljesitmenyre, kompatibilitasra, keszletre vagy elerhetosegre van "
        "szuksege. Emiatt a rendszerben a termekkatalogus mellett fontos a kereshetoseg, a kategoria szerinti "
        "rendezes es az olyan kommunikacio, amely nem iger tul sokat, hanem szakmai egyeztetesre is teret hagy.",
    )
    add_figure(
        doc,
        "3. abra: Termeklista szuresi es keresesi nezettel",
        "Kepernyokep: Termekek oldal, kereses vagy kategoria szerinti szures aktiv allapotban.",
        10,
    )

    # Chapter 3
    new_chapter(doc, "3. Celkituzes es MVP-hatar")
    add_para(
        doc,
        "A projekt celja egy olyan webes alkalmazas megvalositasa volt, amely a TDLWebshop nev alatt "
        "epuletegepeszeti termekeket jelenit meg, kezeli a vasarloi kosarat es rendelest, valamint adminisztracios "
        "feluletet biztosit a termekek, rendelesek, felhasznalok es keszletadatok kezelesere.",
    )
    add_para(
        doc,
        "Az MVP-hatar meghatarozasakor azt tekintettem kesznek, ha a rendszerben egy vasarlo el tud jutni "
        "a termek megtekintesetol a rendelese leadasaig, egy admin vagy dolgozo pedig kepes termeket kezelni, "
        "rendelest megtekinteni, statuszt modositani es helyszini vasarlast rogzitani. A rendszerben kulon "
        "figyelmet kapott a jogosultsagkezeles, mert a vasarlo, dolgozo es admin szerepkor nem ugyanazokat "
        "a muveleteket vegezheti el.",
    )
    add_heading(doc, "3.1. MVP-be tartozo funkciok", 2)
    add_bullets(
        doc,
        [
            "Kezdolap, termeklista, termekadatlap es kategoria szerinti bongeszes.",
            "Kosar, mennyisegmodositas, kuponlogika es checkout validacio.",
            "Rendeles letrehozasa, statuszkovetes es vasarloi profil oldal.",
            "Kivansaglista es akcios termekek megjelenitese.",
            "Admin termekkezeles, CSV import, keszletfigyeles es rendeleskezeles.",
            "Helyszini vasarlas mentett vasarloval es PDF bizonylat generalassal.",
            "Dolgozoi es admin szerepkorok elkulonitese.",
            "AI asszisztens, amely a sajat katalogushoz es epuletegepeszeti temakhoz kotott.",
        ],
    )
    add_heading(doc, "3.2. Tudatosan vallalt korlatok", 2)
    add_para(
        doc,
        "A rendszer szakdolgozati MVP, ezert bizonyos funkciok nem teljes vallalati szintu megoldaskent keszultek el. "
        "A webes rendeleseinel az arak es kosartetelek kliensoldali folyamatbol erkeznek, amit eles uzem elott "
        "szerveroldali ujraszamolassal es szigorubb keszletellenorzessel kellene megerositeni. Az AI asszisztens "
        "eseten az OpenRouter kulcs nem kerul kliensoldalra, de eles hasznalatnal tovabbi rate limit, naplozas es "
        "kvotakezeles lenne szukseges.",
    )
    add_figure(
        doc,
        "4. abra: MVP funkcionalis hatara",
        "Ide illessz be egy egyszeru abrat: Vasarloi folyamat, Admin folyamat, Dolgozoi folyamat, AI asszisztens.",
        12,
    )

    # Chapter 4
    new_chapter(doc, "4. Kovetelmenyek")
    add_para(
        doc,
        "A kovetelmenyeket ugy fogalmaztam meg, hogy azok ellenorizhetoek legyenek. A kovetelmeny nem csupan "
        "funkciootlet, hanem olyan allitas, amelyhez use case, modul es teszt vagy kepernyokepi bizonyitek "
        "kapcsolhato. Ez segit abban, hogy a dolgozat ne csak leirja a rendszert, hanem igazolja is a megvalositast.",
    )
    add_heading(doc, "4.1. Funkcionalis kovetelmenyek", 2)
    add_table(
        doc,
        ["Azonosito", "Kovetelmeny", "Modul", "Bizonyitek"],
        [
            ["K-01", "A vasarlo tudjon termeket keresni es kategoriak szerint bongeszni.", "Termeklista", "Termeklista kepernyokep"],
            ["K-02", "A vasarlo tudjon termeket kosarba tenni es mennyiseget modositani.", "Kosar", "Kosar teszt es kepernyokep"],
            ["K-03", "A checkout ellenorizze az email es telefonszam formajat.", "Checkout", "Validacios kepernyokep"],
            ["K-04", "A rendeles letrehozasa utan statuszkovetes legyen elerheto.", "Order/Profile", "Profil oldal"],
            ["K-05", "Az admin tudjon termeket CSV-bol importalni.", "Admin Products", "CSV import kepernyokep"],
            ["K-06", "Az admin es dolgozo jogosultsagai kulonuljenek el.", "Auth/Admin", "Firestore rules es UI"],
            ["K-07", "Helyszini vasarlas es PDF bizonylat keszuljon.", "Admin Orders/PDF", "PDF kepernyokep"],
            ["K-08", "Az AI asszisztens ne lepjen ki a webshop/domain temabol.", "AI service/Worker", "AI tesztkerdes"],
        ],
        widths=[2, 7, 3.5, 4],
    )
    add_heading(doc, "4.2. Nem funkcionalis kovetelmenyek", 2)
    add_para(
        doc,
        "A nem funkcionalis kovetelmenyek koze tartozik a reprodukalhato futtatas, a verziozott forraskod, a "
        "titkok elkulonitett kezelese, a reszponziv felulet, a hibas bemenetek kezelese es a kritikus folyamatok "
        "tesztelhetosege. A projektben ezek kozul tobb elem dokumentacioval es CI futassal is igazolt.",
    )
    add_heading(doc, "4.3. Traceability tablazat", 2)
    add_table(
        doc,
        ["Kovetelmeny", "Use case", "Forrasfajl / modul", "Teszt vagy kep"],
        [
            ["K-02", "Termek kosarba helyezese", "cart.service.ts, product-list", "Kosar teszt"],
            ["K-03", "Rendeles leadasa hibas adatokkal", "checkout.ts", "Checkout validacios kep"],
            ["K-05", "CSV import", "admin.ts", "Admin CSV kep"],
            ["K-06", "Admin/dolgozo jogosultsag", "firestore.rules, admin.ts", "Jogosultsagi teszt"],
            ["K-07", "Helyszini vasarlas", "order.service.ts, invoice.service.ts", "PDF kep"],
            ["K-08", "AI asszisztens kerdes", "chatbot-llm.service.ts, Worker", "AI chat kep"],
        ],
        widths=[3, 4, 5, 4],
    )

    # Chapter 5
    new_chapter(doc, "5. Use case-ek")
    add_para(
        doc,
        "A rendszer use case-ei harom fo szerepkor kore rendezhetok: vasarlo, dolgozo es admin. A vasarlo a "
        "termekek megtekintesetol a rendelese leadasaig es koveteseig jut el. A dolgozo elsosorban ertekesitesi "
        "es keszlethez kapcsolodo feladatokat lat el. Az admin teljesebb attekintest kap, beleertve a "
        "jogosultsagokat, felhasznalokat, rendeleseket es termekimportot.",
    )
    add_figure(
        doc,
        "5. abra: Use case diagram",
        "Ide illeszd be a keszitett use case diagramot: Vasarlo, Dolgozo, Admin szerepkorokkal.",
        16,
    )
    add_heading(doc, "5.1. Vasarloi folyamat", 2)
    add_numbered(
        doc,
        [
            "A vasarlo megnyitja a kezdolapot es kategoria vagy kereses alapjan termeket keres.",
            "A termekadatlap vagy termekkartya alapjan kosarba helyezi a kivalasztott termeket.",
            "A kosarban modositja a mennyiseget, majd a checkout oldalra lep.",
            "Megadja a szallitasi es kapcsolattartasi adatokat, a rendszer validalja azokat.",
            "Sikeres leadas utan a rendeles bekerul az adatbazisba, es a vasarlo a profil oldalon kovetheti.",
        ],
    )
    add_heading(doc, "5.2. Adminisztracios folyamat", 2)
    add_numbered(
        doc,
        [
            "Az admin bejelentkezik es megnyitja az admin feluletet.",
            "Attekinti a rendeleseket, termekeket, keszletadatokat es felhasznalokat.",
            "CSV fajlbol tomegesen importal termekeket, vagy egyedileg modosit termekadatokat.",
            "Rendeles statuszt modosit, amely auditbejegyzest es adott esetben keszletvaltozast eredmenyez.",
            "Helyszini vasarlast rogzit mentett vasarloval, majd PDF bizonylatot general.",
        ],
    )
    add_figure(
        doc,
        "6. abra: Helyszini vasarlas mentett vasarloval",
        "Kepernyokep: Admin rendelesek/helyszini vasarlas, mentett vasarlo kivalasztva, termekek listazva.",
        12,
    )

    # Chapter 6
    new_chapter(doc, "6. GUI es felhasznaloi elmeny")
    add_para(
        doc,
        "A felulet kialakitasanal a cel egy modern, de nem tulzottan jatekos hatasu webshop megjelenes volt. "
        "A dark mode ipari, technologiai hangulatot ad, amely illeszkedik a TDLWebshop logo szinvilagahoz. "
        "A light mode tisztabb, uzletiesebb alternativat biztosit. A ket mod kozott a layout nem valtozik, "
        "csak a szinek, hatterek es arnyekok igazodnak.",
    )
    add_para(
        doc,
        "A navigacio kozponti eleme a fejlec, ahol a logo, kereses, kategoria lenyilo, kosar, profil es tema "
        "valto kapcsolo egyutt jelenik meg. A liquid glass hatast csak ott alkalmaztam, ahol UI elemkent indokolt, "
        "peldaul dropdownban vagy keresesi panelben. A termekkartyaknal es fo tartalmi blokkoknal stabilabb, "
        "iparibb kartyaelrendezest hasznaltam, mert egy webshopban az olvashatosag es a termekadatok gyors "
        "ertekelese fontosabb a latvanyos dekoracional.",
    )
    add_figure(doc, "7. abra: Kezdolap AI asszisztenssel", "Kepernyokep: kezdolap, AI asszisztens ablak nyitva.", 12)
    add_figure(doc, "8. abra: Termekadatlap", "Kepernyokep: egy konkret termek oldala nagy termekkepes nezetben.", 12)
    add_figure(doc, "9. abra: Kosar tobb termekkel", "Kepernyokep: kosar oldal, legalabb ket termekkel es mennyisegallitassal.", 12)
    add_heading(doc, "6.1. Hibakezeles es ures allapotok", 2)
    add_para(
        doc,
        "A felhasznaloi elmeny fontos resze, hogy a rendszer ne csak idealis esetben legyen ertheto. Ezert "
        "a checkout oldalon hibas email vagy telefonszam eseten a felhasznalo visszajelzest kap, es nem tud "
        "tovabblepni addig, amig az adatot nem javitja. Az admin feluleten az import es mentesei muveletek is "
        "visszajelzest adnak, igy lathato, ha egy CSV sor hibas vagy egy muvelet nem sikerult.",
    )
    add_figure(doc, "10. abra: Checkout validacio hibas email/telefonszam eseten", "Kepernyokep: checkout, hibas mezokkel.", 12)

    # Chapter 7
    new_chapter(doc, "7. Technologiai hatter")
    add_heading(doc, "7.1. Angular", 2)
    add_para(
        doc,
        "A frontend fejlesztes Angular keretrendszerben keszult. Az Angular elonye, hogy komponensalapu "
        "feluletfejlesztest tesz lehetove, jol tamogatja a szolgaltatasokba szervezett uzleti logikat, es "
        "TypeScript alapon fejlesztheto. A projektben az oldalak es szolgaltatasok elkulonitese segitette, "
        "hogy a kosar, rendeles, autentikacio, PDF generalas es AI asszisztens mukodese attekintheto maradjon.",
    )
    add_heading(doc, "7.2. Firebase es Firestore", 2)
    add_para(
        doc,
        "A backend oldali adattarolast es hitelesitest Firebase szolgaltatasokra epitettem. A Firestore "
        "dokumentumorientalt adatmodellje illeszkedik a termekek, rendelesek, felhasznalok es mentett vasarlok "
        "tarolasahoz. A biztonsagi szabalyok kulonosen fontosak, mert kliensoldali alkalmazasnal nem eleg a "
        "feluleten elrejteni egy funkciot; az adatbazis hozzaferest is szerepkorok szerint kell korlatozni.",
    )
    add_heading(doc, "7.3. Cloudflare Worker es OpenRouter", 2)
    add_para(
        doc,
        "Az AI asszisztensnel a kliens nem kozvetlenul hivja az OpenRouter API-t. A hivas egy Cloudflare Worker "
        "proxyn keresztul tortenik, igy az API kulcs szerveroldalon marad. Ez szakdolgozati es biztonsagi "
        "szempontbol is fontos, mert a kulcs nem kerul a bongeszobe es nem commitolodik a repositoryba. "
        "Az MVP-ben a Worker az alap domainkorlatokat es CORS beallitast kezeli, eles uzem elott pedig "
        "rate limit es reszletesebb naplozas is indokolt lenne.",
    )
    add_heading(doc, "7.4. CI es reprodukalhatosag", 2)
    add_para(
        doc,
        "A GitHub Actions workflow celja, hogy a repositoryba kerulo kod ne csak lokalis gepen, hanem tiszta "
        "kornyezetben is ellenorizheto legyen. A zold CI futas bizonyitja, hogy a telepites, build es tesztek "
        "automatizaltan lefutnak. A dolgozatban erdemes kepkent is bemutatni a sikeres workflow futast.",
    )
    add_figure(doc, "11. abra: GitHub Actions zold CI futas", "Kepernyokep: GitHub Actions oldal, legfrissebb zold CI futas.", 10)

    # Chapter 8
    new_chapter(doc, "8. Architektura")
    add_para(
        doc,
        "A TDLWebshop architekturaja kliensoldali Angular alkalmazasra, Firebase szolgaltatasokra es kulso "
        "AI proxyra epul. A frontend komponensek jelenitik meg a vasarloi es admin feluletet, mig a szolgaltatasok "
        "kezelik az adatelerest, kosarat, rendelest, szamlageneralast es AI kommunikaciot. A Firestore az adatok "
        "perzisztens taroloja, az Auth a felhasznalok azonositasaert felel, a Hosting pedig a statikus alkalmazas "
        "kiszolgalasat vegzi.",
    )
    add_figure(
        doc,
        "12. abra: Komponens-architektura",
        "Ide illeszd be a komponens architektura diagramot: Angular app, services, Firebase Auth/Firestore/Hosting, Cloudflare Worker, OpenRouter.",
        16,
    )
    add_heading(doc, "8.1. Adatfolyamok", 2)
    add_para(
        doc,
        "A vasarloi rendeles adatfolyama a termeklista es kosar felol indul, majd a checkout oldalon valik "
        "rendeles dokumentumma. Admin oldalon a statuszmodositas es helyszini vasarlas mar szorosabban kapcsolodik "
        "a keszletvaltozashoz es auditnaplohoz. Az AI asszisztens adatfolyama elter ettol: ott a kliens katalogus- "
        "es domainkontextust allit ossze, majd a Worker a kulso modell fele tovabbitja a kerest.",
    )
    add_figure(
        doc,
        "13. abra: Rendelesi szekvencia",
        "Ide illeszd be a szekvencia diagramot: Vasarlo -> Checkout -> OrderService -> Firestore -> Profile/Admin.",
        14,
    )

    # Chapter 9
    new_chapter(doc, "9. Adatmodell")
    add_para(
        doc,
        "Az adatmodell fo entitasai a Product, CartItem, Order, OrderItem, UserProfile, Coupon, Invoice, "
        "SavedCustomer es Audit bejegyzesek. A dokumentumorientalt tarolas miatt az adatok egy resze "
        "denormalizaltan jelenik meg, peldaul a rendeles tetelei tartalmazzak az adott pillanatban ervenyes "
        "termeknevet es arat. Ez megkonnyiti a kesobbi rendelesmegtekintest akkor is, ha a termekadat idokozben valtozik.",
    )
    add_figure(
        doc,
        "14. abra: Adatmodell diagram",
        "Ide illeszd be az adatmodell diagramot: Product, Order, OrderItem, UserProfile, Coupon, SavedCustomer, Invoice.",
        16,
    )
    add_table(
        doc,
        ["Entitas", "Fobb mezok", "Szerep"],
        [
            ["Product", "name, sku, category, price, stock, images", "Termekkatalogus es keszlet alapja"],
            ["Order", "customer, items, total, status, paymentMethod", "Webes vagy helyszini rendeles"],
            ["OrderItem", "productId, name, quantity, unitPrice", "Rendeles tetel sor"],
            ["UserProfile", "uid, email, role, disabled", "Jogosultsag es profiladat"],
            ["SavedCustomer", "name, email, phone, company, taxNumber", "Helyszini vasarlok ujrafelhasznalasa"],
            ["Coupon", "code, discount, active, limits", "Kedvezmeny kezeles"],
            ["Audit", "orderId, oldStatus, newStatus, actor", "Statuszmodositas nyomkovetes"],
        ],
        widths=[3, 6, 7],
    )
    add_para(
        doc,
        "A modell tervezesenel fontos szempont volt, hogy a vasarloi es admin folyamatok ugyanahhoz az adatbazishoz "
        "kapcsolodjanak, de a hozzaferes ne legyen azonos. Ezert a Firestore szabalyokban szerepkor alapu "
        "ellenorzesek vannak, es a tiltott felhasznalokhoz kapcsolodo muveletek is korlatozottak.",
    )

    # Chapter 10
    new_chapter(doc, "10. Megvalositas")
    add_para(
        doc,
        "A megvalositas fejezetben a rendszer azon reszeit emelem ki, amelyek szakmailag a legjobban mutatjak "
        "a projekt osszetettseget. Nem minden komponens es fuggveny reszletes bemutatasa a cel, hanem azoknak "
        "a folyamatoknak a magyarazata, amelyek a szakdolgozat erteket adjak: checkout, rendelesmentes, keszlet, "
        "statusz/audit, PDF generalas, CSV import, jogosultsag es AI asszisztens.",
    )
    add_heading(doc, "10.1. Checkout es rendeles veglegesitese", 2)
    add_para(
        doc,
        "A checkout oldalon a felhasznalo altal megadott adatok validalasa utan tortenik a rendeles veglegesitese. "
        "A rendszer ellenorzi a kapcsolattartasi adatokat, kezeli a kuponkedvezmenyt, osszegzi a kosarat, majd "
        "letrehozza a rendeles dokumentumot. A felhasznalo szamara a folyamat sikeres leadaskent jelenik meg, "
        "admin oldalon pedig a rendeles bekerul a listaba.",
    )
    add_code_placeholder(
        doc,
        "Checkout rendeles veglegesitese",
        "src/pages/checkout/checkout.ts 367-555 es 580-637",
        "Ez a resz mutatja a rendeles leadasa elotti ellenorzeseket, osszegzest es a szolgaltatas hivasat.",
    )
    add_figure(doc, "15. abra: Sikeres checkout utan megjeleno allapot", "Kepernyokep: sikeres rendelés leadasa utan.", 10)
    add_heading(doc, "10.2. Rendeles statusz, audit es keszlet", 2)
    add_para(
        doc,
        "A rendeles statuszanak modositasa nem csupan egy szoveges mezot erint. A rendszer audit bejegyzest "
        "keszit, es bizonyos statuszvaltasokhoz keszletvaltozas is kapcsolodik. Ez azert fontos, mert egy "
        "webshopban a rendelesek allapota es a keszletadatok kozotti konzisztencia alapveto szakmai kovetelmeny.",
    )
    add_code_placeholder(
        doc,
        "Statuszvaltas, audit es keszletlogika",
        "src/app/services/order.service.ts 41-128",
        "A kodreszlet a rendeles statuszmodositasanak, auditjanak es keszlethez kotott mukodesenek bemutatasara alkalmas.",
    )
    add_heading(doc, "10.3. Helyszini vasarlas", 2)
    add_para(
        doc,
        "A helyszini vasarlas az admin oldali funkcionalitas egyik fontos resze. Itt az admin vagy dolgozo "
        "mentett vasarlot valaszthat, termekeket adhat a vasarlashoz, majd PDF bizonylatot keszithet. Ez a "
        "funkcio kulonbozik a klasszikus webes checkouttol, mert az ertekesito rogzit mindent, es a folyamat "
        "inkabb egy belso adminisztracios munkafolyamatra hasonlit.",
    )
    add_code_placeholder(
        doc,
        "Helyszini rendeles tranzakcio",
        "src/app/services/order.service.ts 229-272",
        "A kodreszlet azt mutatja, hogyan jon letre helyszini rendeles es hogyan kapcsolodik hozza a keszletmodositas.",
    )
    add_heading(doc, "10.4. Szamlaszam es PDF bizonylat", 2)
    add_para(
        doc,
        "A rendszer PDF bizonylatot general a rendelesekhez. A PDF nem teljes jogi erteku szamlazo rendszer, "
        "hanem szakdolgozati celra keszult bizonylat/prototipus, amely megmutatja a rendeles adatait, a vevot, "
        "a tetelek listajat es a fizetendo osszeget. A dolgozatban ezt erdemes igy megfogalmazni, hogy ne "
        "keltsen teljes NAV-kompatibilis szamlazo rendszerre vonatkozo elvarast.",
    )
    add_code_placeholder(
        doc,
        "PDF bizonylat felepitese",
        "src/app/services/invoice.service.ts 8-183",
        "A kodreszlet a PDF struktura, fejlec, vevoadatok, tetelek es osszesito blokk felepiteset mutatja.",
    )
    add_figure(doc, "16. abra: General PDF bizonylat", "Kepernyokep vagy PDF export: egy helyszini vagy webes rendeles bizonylata.", 12)
    add_heading(doc, "10.5. CSV import", 2)
    add_para(
        doc,
        "Az admin feluleten elerheto CSV import azert kerult a rendszerbe, mert egy webshopnal a termekek "
        "egyenkenti felvitele sok ido lenne. A validacio megmutatja, hogy mely sorok feldolgozhatok es melyek "
        "hibasak. Ez a funkcio szakdolgozati szempontbol azert hasznos, mert egyszerre erinti az adatformatumot, "
        "validaciot, admin feluletet es Firestore mentest.",
    )
    add_code_placeholder(
        doc,
        "CSV import validacio es mentes",
        "src/pages/admin/admin.ts 1183-1258",
        "A kodreszlet a tomeges termekfeltoltes ellenorzeset es menthetoseget tamasztja ala.",
    )
    add_figure(doc, "17. abra: Admin termekkezeles es CSV import", "Kepernyokep: admin termekek ful, CSV import panel.", 12)
    add_heading(doc, "10.6. Jogosultsagkezeles", 2)
    add_para(
        doc,
        "A szerepkorok kozul az admin rendelkezik a legszelesebb jogosultsaggal, a dolgozo korlatozottabb "
        "adminisztracios muveleteket vegezhet, a vasarlo pedig sajat adataihoz es rendeleseihez fer hozza. "
        "A jogosultsagok nemcsak a feluleten jelennek meg, hanem az adatbazis szabalyokban is. Ez kulonosen "
        "fontos, mert kliensoldali alkalmazasnal a felulet elrejtese onmagaban nem jelent biztonsagot.",
    )
    add_code_placeholder(
        doc,
        "Firestore szerepkorok es aktiv felhasznalo ellenorzes",
        "firestore.rules 25-76 es 288-356",
        "Ez a kodreszlet mutatja az admin, dolgozo, vasarlo es tiltott felhasznalo hozzaferesi logikajat.",
    )
    add_code_placeholder(
        doc,
        "Admin/dolgozoi jogosultsagok UI oldalon",
        "src/pages/admin/admin.ts 606-735",
        "A kodreszlet azt mutatja, hogyan igazodik az admin felulet a bejelentkezett szerepkorhoz.",
    )
    add_figure(doc, "18. abra: Admin felhasznalo- es jogosultsagkezeles", "Kepernyokep: admin felhasznalok ful, szerepkorokkal.", 12)
    add_heading(doc, "10.7. AI asszisztens", 2)
    add_para(
        doc,
        "Az AI asszisztens celja nem az, hogy minden kerdesre korlatlanul valaszoljon, hanem hogy a TDLWebshop "
        "katalogusa es az epuletegepeszeti temakor alapjan segitse a felhasznalot. Ha nincs pontos termektalalat, "
        "a rendszer inkabb szakmai iranyt ad, es jelzi, hogy pontos ajanlatert vagy beszerezhetosegert emailben "
        "vagy szemelyesen erdemes egyeztetni. Ez csokkenti annak kockazatat, hogy a chatbot veletlenszeruen "
        "ajanljon nem megfelelo termeket.",
    )
    add_code_placeholder(
        doc,
        "AI domain- es kataloguslogika",
        "src/app/services/chatbot-llm.service.ts 31-88, 92-120 es 214-250",
        "A kodreszlet az AI valaszok temakorhoz kotott feldolgozasat es termekkatalogushoz kapcsolodasat mutatja.",
    )
    add_code_placeholder(
        doc,
        "OpenRouter proxy szerveroldali kulcskezelessel",
        "workers/openrouter-proxy/src/index.js 1-59 es 153-211",
        "Ez a resz mutatja, hogy az API kulcs nem a kliensoldali alkalmazasban van, hanem Worker kornyezetben.",
    )

    # Chapter 11
    new_chapter(doc, "11. Biztonsag es adatvedelem")
    add_para(
        doc,
        "A webes alkalmazasoknal a biztonsag nem egyetlen funkcio, hanem tobb retegu megkozelites. A projektben "
        "a legfontosabb biztonsagi teruletek a hitelesites, szerepkor alapu hozzaferes, tiltott felhasznalok "
        "kezelese, inputvalidacio, titkok elkulonitese, PDF es kupon visszaelesek korlatozasa, valamint az AI "
        "proxy kulcskezelesenek biztonsaga.",
    )
    add_heading(doc, "11.1. Titkok es konfiguracio", 2)
    add_para(
        doc,
        "A repositoryban nem szerepelhet valodi jelszo, API kulcs vagy token. A konfiguraciohoz .env.example "
        "hasznalhato, amely megmutatja, milyen valtozokra van szukseg, de nem tartalmaz titkot. A Firebase webes "
        "config onmagaban nem ugyanaz, mint egy szerveroldali titok, de a Firestore szabalyoknak megfeleloen "
        "vedeniuk kell az adatbazist.",
    )
    add_heading(doc, "11.2. Jogosultsag es tiltott felhasznalo", 2)
    add_para(
        doc,
        "A szerepkorok es a disabled allapot kezelese azert fontos, mert egy mar regisztralt vagy korabban "
        "bejelentkezett felhasznalo nem kaphat tovabbra is hozzaferest akkor, ha az admin letiltotta. A rendszer "
        "ezt nem csak kliensoldali uzenettel, hanem szabaly- es logikai szinten is kezeli.",
    )
    add_heading(doc, "11.3. MVP-korlatok biztonsagi szempontbol", 2)
    add_para(
        doc,
        "A konzulensi visszajelzes alapjan fontos rogziteni, hogy a webes checkout kliensoldali adatokbol dolgozik. "
        "Szakdolgozati MVP-kent ez bemutatja a folyamatot, de eles uzemben szerveroldali ar- es keszletujraszamolasra "
        "lenne szukseg. Ugyanigy az AI proxy CORS beallitasa nem helyettesiti a rate limitet; ez egy tovabbi "
        "fejlesztesi irany.",
    )
    add_table(
        doc,
        ["Kockazat", "Jelenlegi kezeles", "Tovabbfejlesztes"],
        [
            ["Kliensoldali ar/kosar adat", "Checkout validacio es dokumentalt MVP-korlat", "Szerveroldali ujraszamolas"],
            ["Keszlethiany teljesiteskor", "Admin/OrderService logika", "Szigorubb hibaval megallas"],
            ["Kupon visszaeles", "Kuponvalidacio es tesztek", "Szerveroldali kuponellenorzes"],
            ["AI API visszaeles", "Worker proxy, kulcs nem kliensoldalon", "Rate limit, kvota, naplozas"],
            ["Vendeg email azonositasa", "MVP-ben email alapu kezeles", "Ertesitesi/token alapu igazolas"],
        ],
        widths=[4.2, 5.8, 5.8],
    )

    # Chapter 12
    new_chapter(doc, "12. Teszteles es validacio")
    add_para(
        doc,
        "A teszteles celja annak bizonyitasa, hogy a rendszer kritikus folyamatai mukodnek. A projektben automata "
        "tesztek es kezi teszteles egyarant szerepet kapnak. Az automata tesztek gyorsan ellenorzik a fontosabb "
        "logikai reszeket, a kezi teszteles pedig a teljes felhasznaloi utakat es feluleti allapotokat vizsgalja.",
    )
    add_heading(doc, "12.1. Automata ellenorzesek", 2)
    add_para(
        doc,
        "A build es tesztek futtatasa a fejlesztoi gepen es GitHub Actions kornyezetben is ellenorizheto. A dolgozat "
        "vegleges valtozataban ide kell beirni a legfrissebb futtatas pontos eredmenyet, peldaul: npm run build "
        "sikeres, npm test -- --watch=false sikeres, GitHub Actions CI zold.",
    )
    add_table(
        doc,
        ["Ellenorzes", "Parancs / hely", "Elvart eredmeny", "Bizonyitek"],
        [
            ["Build", "npm run build", "Sikeres Angular build", "Terminal vagy CI kep"],
            ["Unit tesztek", "npm test -- --watch=false", "Tesztek zolden lefutnak", "Terminal vagy CI kep"],
            ["CI", "GitHub Actions", "Legfrissebb workflow zold", "GitHub screenshot"],
            ["Firestore szabalyok", "firestore.rules attekintese", "Nincs nyitott read/write szabaly", "Kodreszlet"],
        ],
        widths=[3.5, 4.5, 4.5, 3.5],
    )
    add_heading(doc, "12.2. Kezi tesztforgatokonyvek", 2)
    add_table(
        doc,
        ["Teszt", "Lepesek", "Elvart eredmeny"],
        [
            ["Regisztracio/bejelentkezes", "Uj user letrehozasa, login, logout", "Profil elerheto, jogosultsag helyes"],
            ["Termekkereses", "Kategoria es kulcsszo szerinti kereses", "Szurt termeklista jelenik meg"],
            ["Kosar", "Termek hozzaadasa, mennyiseg modositasa, torles", "Vegosszeg valtozik"],
            ["Checkout validacio", "Hibas email/telefon megadasa", "Rendszer nem enged tovabb"],
            ["Sikeres rendeles", "Ervenyes adatokkal rendeles leadasa", "Rendeles megjelenik profilban/adminban"],
            ["Admin statusz", "Rendeles statusz modositasa", "Audit es statusz frissul"],
            ["CSV import", "Termek CSV feltoltese", "Ervenyes sorok mentodnek"],
            ["Helyszini vasarlas", "Mentett vasarlo + termekek + PDF", "Bizonylat letoltheto"],
            ["AI asszisztens", "Domain, termek es nem relevans kerdes", "Csak megfelelo temaban valaszol"],
        ],
        widths=[3.7, 7, 5],
    )
    add_figure(doc, "19. abra: Admin attekintes", "Kepernyokep: admin fo oldal statisztikai kartyakkal.", 10)
    add_figure(doc, "20. abra: Keszletfigyeles", "Kepernyokep: admin keszlet ful alacsony keszlet jelzessel.", 10)

    add_expanded_evidence_section(doc)

    # Chapter 13
    new_chapter(doc, "13. Mesterseges intelligencia hasznalata a fejlesztes soran")
    add_para(
        doc,
        "A szakdolgozat keszitese soran mesterseges intelligenciat ket kulonbozo szinten hasznaltam. Az elso "
        "szint a fejlesztest tamogato AI-hasznalat volt: otleteles, hibakereses, kodreview, dokumentacios "
        "vazlatok keszitese es ellenorzesi listak osszeallitasa. A masodik szint maga a rendszerbe beepitett "
        "AI asszisztens, amely az OpenRouteren keresztul valaszol epuletegepeszeti es katalogushoz kapcsolodo "
        "kerdesekre.",
    )
    add_para(
        doc,
        "A fejlesztesi folyamatban az AI kimeneteit nem vegleges igazsagkent kezeltem, hanem javaslatkent. "
        "A kodmodositasokat futtatassal, tesztekkel, kepernyon torteno ellenorzessel es sajat atnezessel "
        "validaltam. Peldaul a jogosultsagi es Firestore szabalyoknal kulon figyelni kellett arra, hogy a "
        "frontend oldali elrejtes ne helyettesitse az adatbazis szintu tiltast. Ugyanigy az AI asszisztens "
        "bekotesenel fontos volt, hogy az API kulcs ne keruljon kliensoldalra.",
    )
    add_para(
        doc,
        "Az AI segitsege elsosorban gyorsitotta a munkat, mert segitett strukturaltabban vegiggondolni a "
        "hibahelyzeteket, a tesztelesi forgatokonyveket es a dokumentacios kovetelmenyeket. Ugyanakkor lassithatta "
        "is a folyamatot, amikor egy javaslat tul altalanos volt, vagy nem illett pontosan a projekt szerkezetebe. "
        "Ilyenkor a megoldast at kellett alakitani a sajat kodbazis mintaihoz.",
    )
    add_para(
        doc,
        "A dolgozat szoveges reszenel az AI-t vazlatkeszitesre es nyelvi ellenorzesre hasznaltam, de a vegleges "
        "szoveget sajat megfogalmazasra kell atirni. Ez azert fontos, mert a szakdolgozatnak nem csak szakmailag, "
        "hanem szemelyes munkakent is vedheto formaban kell megjelennie. A dontesekert, a kodert, az ellenorzesert "
        "es a vegleges beadott tartalomert en felelek.",
    )
    add_para(
        doc,
        "A rendszerbe beepitett AI asszisztensnel a legfontosabb korlat, hogy nem helyettesit szakembert es nem "
        "adhat bizonytalan termekajanlast biztos allitaskent. Ha nincs pontos katalogustalalat, akkor altalanos "
        "szakmai iranyt ad, es felhivja a figyelmet arra, hogy pontos ajanlatert, keszletinformacioert vagy "
        "beszerezhetosegert emailben vagy szemelyesen erdemes egyeztetni. Ez a mukodes tudatosan ovatosabb, "
        "mint egy korlatlan altalanos chatbot.",
    )
    add_figure(doc, "21. abra: AI asszisztens domainkorlatos valasza", "Kepernyokep: AI chat relevans es nem relevans kerdesre adott valasszal.", 10)

    # Chapter 14
    new_chapter(doc, "14. Ertekeles, korlatok es tovabbfejlesztes")
    add_para(
        doc,
        "A TDLWebshop szakdolgozati MVP-kent elerte azt a celt, hogy egy osszetett, tobb szerepkoros webshop "
        "mukodeset bemutassa. A vasarloi folyamat, adminisztracios felulet, CSV import, keszletfigyeles, PDF "
        "bizonylat, jogosultsagkezeles es AI asszisztens egyutt mar olyan projektvolument ad, amely tulmutat "
        "egy egyszeru CRUD alkalmazason.",
    )
    add_para(
        doc,
        "A rendszer legfontosabb erossege az, hogy a vasarloi es admin folyamatok osszekapcsolodnak. Egy rendeles "
        "nem csak a vasarlo oldalan jelenik meg, hanem adminisztraciosan is kezelheto. A helyszini vasarlas kulon "
        "erteket ad, mert a webshopot nem csak online ertekesitesi feluletkent, hanem belso munkafolyamatokat "
        "tamogato rendszerkent is kezeli.",
    )
    add_para(
        doc,
        "A korlatok kozul a legfontosabb az, hogy eles uzem elott a webes rendelesei folyamatot szerveroldali "
        "ar- es keszletellenorzessel kellene megerositeni. Emellett az AI asszisztensnel rate limitre, naplozasra "
        "es tovabbi vedelmekre lenne szukseg. A fizetesi modok jelenleg nem teljes online bankkartyas fizetesi "
        "integraciot jelentenek, hanem rendelesei/fizetesi mod valasztasi logikat.",
    )
    add_heading(doc, "14.1. Tovabbfejlesztesi lehetosegek", 2)
    add_bullets(
        doc,
        [
            "Szerveroldali ar- es keszletujraszamolas webes checkoutnal.",
            "Teljes fizetesi szolgaltato integracio teszt es eles kornyezettel.",
            "AI asszisztens rate limit, kvota es admin naplozas.",
            "Reszletesebb termekajanlo szereloi csomagok adminbol osszeallithato logikaval.",
            "Kimutatasok es riportok vasarloi, rendelesi es keszletadatok alapjan.",
            "Teljesebb mobil admin felulet es akadalymentessegi ellenorzes.",
        ],
    )

    # Chapter 15
    new_chapter(doc, "15. Osszefoglalas")
    add_para(
        doc,
        "A szakdolgozatban bemutatott TDLWebshop egy epuletegepeszeti temaju webshop es adminisztracios rendszer "
        "prototipusa. A projekt celja nem pusztan egy termeklista letrehozasa volt, hanem egy olyan MVP "
        "megvalositasa, amelyben a vasarloi, adminisztracios es dolgozoi folyamatok egy rendszerben jelennek meg.",
    )
    add_para(
        doc,
        "A fejlesztes soran Angular, Firebase, Firestore, GitHub Actions, Cloudflare Worker es OpenRouter alapu "
        "megoldasok kapcsolodtak ossze. A dolgozat bemutatta a rendszer kovetelmenyeit, use case-eit, GUI/UX "
        "donteseit, architekturajat, adatmodelljet, biztonsagi megfontolasait es teszteleset. A kiegészito "
        "AI asszisztens azt mutatja meg, hogyan lehet egy webshopot szakmai kerdesekben is tamogatni, ugyanakkor "
        "a valaszok korlatainak jelzese fontos resze a felelos mukodesnek.",
    )
    add_para(
        doc,
        "Sajat zaroreflexio helye: ide 8-12 mondatban ird le sajat hangodon, hogy mit tanultal a projektbol, "
        "melyik resz volt a legnehezebb, milyen hibakat javitottal ki a fejlesztes soran, es mit csinalnal "
        "maskent, ha ujrakezdened. Ezt a reszt mindenkeppen sajat szovegre kell atirni, mert a vedesen is "
        "ez lesz az egyik legszemelyesebb pont.",
        italic=True,
    )

    # Bibliography
    new_chapter(doc, "Irodalomjegyzek")
    add_para(
        doc,
        "A vegleges dolgozatban ide keruljenek a felhasznalt dokumentaciok, hivatalos technologiai forrasok es "
        "piaci osszehasonlitas alapjai. A pontos hivatkozasi formatumot az intezmenyi elvaras szerint kell rendezni.",
    )
    add_bullets(
        doc,
        [
            "Angular hivatalos dokumentacio: https://angular.dev/",
            "Firebase hivatalos dokumentacio: https://firebase.google.com/docs",
            "Cloudflare Workers dokumentacio: https://developers.cloudflare.com/workers/",
            "OpenRouter dokumentacio: https://openrouter.ai/docs",
            "GitHub Actions dokumentacio: https://docs.github.com/actions",
            "A piaci osszehasonlitashoz vizsgalt epuletegepeszeti es barkacsaruhaz webshopok.",
        ],
    )

    # Appendix
    new_chapter(doc, "Mellekletek")
    add_heading(doc, "A. Kepernyokep-lista", 2)
    screenshots = [
        "Kezdolap dark mode, kategoria lenyiloval.",
        "Kezdolap AI asszisztens nyitott ablakkal.",
        "Termeklista kereses vagy kategoria szures kozben.",
        "Termekadatlap konkret termekkel.",
        "Kosar tobb termekkel.",
        "Checkout validacio hibas email/telefonszam peldaval.",
        "Checkout sikeres rendeles utan.",
        "Profil oldal rendeleskovetessel.",
        "Kivansaglista oldal.",
        "Admin attekintes.",
        "Admin termekkezeles es CSV import.",
        "Admin keszletfigyeles.",
        "Helyszini vasarlas mentett vasarloval.",
        "General PDF bizonylat.",
        "Admin felhasznalo- es jogosultsagkezeles.",
        "GitHub Actions zold CI.",
    ]
    add_table(doc, ["Sorszam", "Kepernyokep", "Allapot"], [[i + 1, s, "Beillesztendo"] for i, s in enumerate(screenshots)], widths=[2, 10, 4])

    add_heading(doc, "B. Kodreszlet-lista", 2)
    code_refs = [
        ("Checkout veglegesites", "src/pages/checkout/checkout.ts 367-555 es 580-637"),
        ("Statusz, audit, keszlet", "src/app/services/order.service.ts 41-128"),
        ("Helyszini rendeles", "src/app/services/order.service.ts 229-272"),
        ("Szamlaszam generalas", "src/app/services/order.service.ts 276-310"),
        ("PDF bizonylat", "src/app/services/invoice.service.ts 8-183"),
        ("Firestore jogosultsag", "firestore.rules 25-76 es 288-356"),
        ("Admin/dolgozo UI jogosultsag", "src/pages/admin/admin.ts 606-735"),
        ("CSV import", "src/pages/admin/admin.ts 1183-1258"),
        ("AI kataloguslogika", "src/app/services/chatbot-llm.service.ts 31-88, 92-120, 214-250"),
        ("OpenRouter proxy", "workers/openrouter-proxy/src/index.js 1-59 es 153-211"),
    ]
    add_table(doc, ["Tema", "Fajl es sorok", "Miért érdekes?"], [[t, r, "Szakmai megvalositas bizonyiteka"] for t, r in code_refs], widths=[4, 8, 4])

    add_heading(doc, "C. Kezi teszt checklist", 2)
    add_table(
        doc,
        ["Tesztpont", "Sikerult?", "Megjegyzes"],
        [[s, "[ ]", ""] for s in [
            "Regisztracio es bejelentkezes",
            "Tiltott felhasznalo ellenorzese",
            "Termekkereses es kategoriak",
            "Kosar mennyisegmodositas",
            "Kupon validacio",
            "Checkout hibas adatokkal",
            "Sikeres webes rendeles",
            "Profil rendeleskovetes",
            "Admin statuszvaltas",
            "CSV import",
            "Helyszini vasarlas",
            "PDF bizonylat letoltes",
            "Dolgozoi jogosultsag",
            "AI relevans kerdes",
            "AI nem relevans kerdes",
        ]],
        widths=[8, 3, 5],
    )

    # Add several figure-only appendix pages to make the document realistic for 40-50 pages after screenshots.
    add_heading(doc, "D. Beillesztendo nagy meretu kepernyokepek", 2)
    for idx, name in enumerate(screenshots[:12], start=1):
        add_figure(
            doc,
            f"D.{idx}. abra: {name}",
            "A vegleges dolgozatban ezt a blokkot csereld le a tenyleges, jo felbontasu kepernyokepre.",
            14,
        )
        if idx % 2 == 0:
            doc.add_page_break()

    doc.save(DOCX_PATH)
    return DOCX_PATH


if __name__ == "__main__":
    path = build_document()
    print(path)
