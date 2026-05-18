from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt


OUT = Path(__file__).with_name("TDLWebshop_szakdolgozat_osszedolgozott_40_50_munkapeldany.docx")


def set_run_font(run, size=12, bold=False):
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.bold = bold


def add_paragraph(doc, text="", style=None, align=None, bold=False):
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    set_run_font(r, bold=bold)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        set_run_font(run, 14 if level == 1 else 12, bold=True)
    p.paragraph_format.space_before = Pt(12 if level == 1 else 6)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_table(doc, headers, rows, caption=None):
    if caption:
        add_paragraph(doc, caption, bold=True)
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                set_run_font(r, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
            for p in cells[i].paragraphs:
                for r in p.runs:
                    set_run_font(r, size=10)
    doc.add_paragraph()
    return table


def add_figure_placeholder(doc, caption, instruction):
    p = add_paragraph(doc, "[KEP / ABRA HELYE] " + caption, bold=True)
    p.paragraph_format.left_indent = Cm(0.5)
    add_paragraph(doc, "Beillesztendo tartalom: " + instruction)


def add_code_placeholder(doc, file_path, line_range, reason):
    add_paragraph(doc, f"[KODRESZLET HELYE] {file_path} - {line_range}", bold=True)
    add_paragraph(doc, "Mi legyen lathato a kodreszletben: " + reason)


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_after = Pt(6)

    for name in ("Heading 1", "Heading 2", "Heading 3"):
        style = styles[name]
        style.font.name = "Times New Roman"
        style.font.bold = True
        style.font.size = Pt(14 if name == "Heading 1" else 12)


def add_front_matter(doc):
    add_paragraph(doc, "SZEGEDI TUDOMANYEGYETEM", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
    add_paragraph(doc, "Informatikai Intezet", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
    doc.add_paragraph()
    doc.add_paragraph()
    p = add_paragraph(doc, "SZAKDOLGOZAT", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
    for run in p.runs:
        run.font.size = Pt(20)
    doc.add_paragraph()
    doc.add_paragraph()
    add_paragraph(doc, "TDLWebshop - epulelgepeszeti webshop es adminisztracios rendszer", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
    doc.add_paragraph()
    doc.add_paragraph()
    add_paragraph(doc, "Toth David Laszlo", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
    doc.add_paragraph()
    doc.add_paragraph()
    add_paragraph(doc, "Szeged, 2026", align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_page_break()

    add_paragraph(doc, "TDLWebshop - epulelgepeszeti webshop es adminisztracios rendszer", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
    doc.add_paragraph()
    add_table(
        doc,
        ["Mezo", "Tartalom"],
        [
            ["Keszitette", "Toth David Laszlo"],
            ["Szak", "[pontos szak megadasa]"],
            ["Intezet / tanszek", "[intezmenyi kovetelmeny szerinti pontos megnevezes]"],
            ["Temavezeto", "Dr. Bilicki Vilmos, egyetemi docens"],
            ["Hely es ev", "Szeged, 2026"],
        ],
    )
    add_paragraph(doc, "Megjegyzes: a szogletes zarojelekben szereplo adatokat a vegleges formai kovetelmenyek alapjan kell kitolteni.")
    doc.add_page_break()

    add_heading(doc, "Feladatkiiras", 1)
    add_paragraph(
        doc,
        "A szakdolgozat feladata egy epulelgepeszeti termekeket kezelo webaruhaz es a hozza kapcsolodo adminisztracios rendszer megtervezese es megvalositasa. "
        "A rendszer celja, hogy egy kisebb vagy kozepes meretu szakmai kereskedes digitalis mukodeset tamogassa: a vasarloi oldalon termekbongeszes, kosar, rendelesleadas, "
        "profil es rendelestortenet erheto el, mig az admin oldalon termek-, keszlet-, rendeles-, felhasznalo- es helyszini vasarlaskezeles jelenik meg.",
    )
    add_paragraph(
        doc,
        "A munka soran kiemelt szempont volt, hogy a webshop ne csak latvanyos felulet legyen, hanem mernokileg is ertelmezheto MVP-kent mukodjon. Ennek resze a jogosultsagi modell, "
        "a Firestore biztonsagi szabalyok, a tranzakcios rendeleskezeles, a PDF bizonylat/szamla generalasa, a CSV alapu termekimport, valamint egy katalogushoz kotott AI asszisztens "
        "prototipusa is. A feladat hatara tudatosan nem egy teljes erteku, bankkartya-elfogadassal es vallalati ERP-integracioval rendelkezo eles webshop, hanem egy bizonyithatoan "
        "mukodo, szakdolgozati szinten bemutathato webes rendszer.",
    )
    doc.add_page_break()

    add_heading(doc, "Tartalmi osszefoglalo", 1)
    add_table(
        doc,
        ["Szempont", "Leiras"],
        [
            ["Tema megnevezese", "TDLWebshop - epulelgepeszeti webshop es adminisztracios rendszer"],
            ["Feladat", "Egy olyan webes alkalmazas letrehozasa, amely vasarloi es adminisztracios folyamatokat is kezel."],
            ["Megoldasi mod", "Angular alapu frontend, Firebase/Firestore adattarolas es jogosultsagkezeles, Cloudflare Worker alapu AI proxy."],
            ["Alkalmazott eszkozok", "Angular, TypeScript, Firebase Authentication, Firestore, Firebase Hosting, Cloudflare Workers, OpenRouter API, GitHub Actions."],
            ["Elert eredmenyek", "Mukodo vasarloi ut, admin felulet, rendeleskezeles, keszletfigyeles, CSV import, PDF bizonylat, AI asszisztens es dokumentalt teszteles."],
            ["Kulcsszavak", "webshop, Angular, Firebase, Firestore, admin rendszer, epulelgepeszet, AI asszisztens, CI, MVP"],
        ],
    )
    doc.add_page_break()

    add_heading(doc, "Tartalomjegyzek", 1)
    add_paragraph(doc, "A vegleges dokumentumban ezt Word automatikus tartalomjegyzekkel kell frissiteni.")
    add_paragraph(doc, "Javaslat: Hivatkozasok -> Tartalomjegyzek -> Automatikus tartalomjegyzek, majd teljes frissites.")
    doc.add_page_break()


def add_intro(doc):
    add_heading(doc, "1. Bevezetes, problemafelvetes es celkituzes", 1)
    paragraphs = [
        "Az online kereskedelemben a felhasznalok ma mar nem csak egy egyszeru termeklistat varnak el, hanem gyors keresest, egyertelmu keszletinformaciot, megbizhato rendelesi folyamatot es atlathato kommunikaciot. Ez kulonosen igaz az epulelgepeszeti teruletre, ahol a termekek gyakran szakmai dontest igenyelnek, a vasarlo pedig nem mindig tudja pontosan, melyik alkatresz vagy berendezes illik a sajat helyzetehez.",
        "A TDLWebshop alapotlete ebbol a problemabol indult ki. A cel egy olyan webaruhaz volt, amely nem altalanos fogyasztasi cikkekre, hanem futeshez, vizszereleshez, hutestechnikahoz, szellozeshez es szerelesi anyagokhoz kapcsolodo termekekre epul. A rendszer egyszerre celozza a lakossagi vasarlokat es azokat a szakmai vagy uzleti szereploket, akik gyakran helyszini vagy B2B jellegu vasarlast igenyelnek.",
        "A szakdolgozat nem csupan a frontend megjelenitesre osszpontosit, hanem a teljesebb folyamatra: a termekadatok kezelesere, a kosar es checkout logikajara, a rendelesek kovetesere, az adminisztracios munkafolyamatokra, a jogosultsagokra, a biztonsagi szabalyokra es a tesztelhetosegre. A dolgozat celja az is, hogy bemutassa, hogyan lehet egy hallgatoi projektet termekszeru MVP iranyaba vinni.",
        "A dolgozatban bemutatott rendszer fejlesztesenel fontos szempont volt a reprodukalhatosag. A repo tartalmazza a futtatasi leirast, a kornyezeti valtozok mintajat, a CI ellenorzest es a tesztelesi dokumentaciot. Ez azert lenyeges, mert a szakdolgozat ertekelese soran nem eleg az, hogy a rendszer egyszer mar mukodott: bizonyitani kell, hogy mas kornyezetben is ertelmezhetoen elindithato es ellenorizheto.",
    ]
    for p in paragraphs:
        add_paragraph(doc, p)
    add_heading(doc, "1.1. A megoldando problema", 2)
    for p in [
        "Az epulelgepeszeti termekek kinalata sokfele kategoriat erint, es a termekek kozotti kulonbseg gyakran nem csak arban vagy markaban jelenik meg. Egy radiator, csaptelep, szivattyu, klima vagy szellozteto eseteben fontos lehet a meret, teljesitmeny, alkalmazasi terulet, kompatibilitas es keszlet. Emiatt a felhasznalo szamara kulonosen ertekes, ha az oldal nem csak listaz, hanem segiti a valasztast.",
        "A masik problema az adminisztracio. Egy kisebb kereskedesnel a webes rendelesek mellett elofordulhat helyszini vasarlas, telefonos egyeztetes, mentett vasarlo, ceges kedvezmeny, keszletfigyeles es PDF bizonylat igeny is. Ha ezek kulon rendszerekben vagy kezzel vezetett tablazatokban jelennek meg, az hibat es adateltorest okozhat.",
        "A TDLWebshop ezeket a problemakat egy kozponti rendszerben kezeli. A vasarloi oldal a termekek megtalalasara es a rendeles leadasa felol kozelit, az admin felulet pedig a termekek, rendelesek, keszlet es vasarlok kezelesere. A ket oldal kozos adatmodellre epul, igy ugyanaz a termek- es rendelesadat jelenik meg a kulonbozo szerepkorok szamara.",
    ]:
        add_paragraph(doc, p)
    add_heading(doc, "1.2. Celkituzes es MVP-hatar", 2)
    add_paragraph(doc, "A projekt celja nem egy teljes ipari webshoprendszer lemasolasa volt, hanem egy olyan MVP letrehozasa, amely a legfontosabb vasarloi es adminisztracios folyamatokat bizonyithatoan lefedi. Az MVP-hatar meghatarozasa tudatos dontes volt, mert egy szakdolgozati projektben a tul sok felig kesz funkcio gyengitheti a rendszer megitelest.")
    add_table(
        doc,
        ["MVP resze", "Rovid indoklas", "Allapot"],
        [
            ["Termeklista es termekadatlap", "A webshop alapja, a vasarloi tajekozodas kiindulopontja.", "Megvalositva"],
            ["Kosar es checkout", "A vasarloi ut legfontosabb uzleti folyamata.", "Megvalositva"],
            ["Rendeleskovetes profilbol", "Bizalmat es attekinthetoseget ad a vasarlonak.", "Megvalositva"],
            ["Admin termek- es keszletkezeles", "A bolt mukodtetesehez szukseges belso folyamat.", "Megvalositva"],
            ["Helyszini vasarlas", "Domainhez illeszkedo B2B/uzleti kulonbseg.", "Megvalositva"],
            ["PDF bizonylat/szamla", "Adminisztracios bizonyitek es szakdolgozati ertek.", "Megvalositva"],
            ["AI asszisztens", "Katalogushoz kotott, korlatos szakmai seged.", "Prototipus"],
            ["Bankkartya-elfogadas", "Valos penzugyi integracio magasabb kockazattal.", "Tovabbfejlesztes"],
            ["ERP integracio", "Vallalati szintu illesztes, kulso rendszerekkel.", "Tovabbfejlesztes"],
        ],
        "1. tablazat - Az MVP-hatar osszefoglalasa",
    )
    add_figure_placeholder(doc, "A TDLWebshop kezdolapja dark modban", "Keszits kepernyokepet a fooldalrol ugy, hogy a kategoriak legordulo menuje nyitva van.")


def add_market(doc):
    add_heading(doc, "2. Piaci es teruleti osszehasonlitas", 1)
    for p in [
        "A TDLWebshop pozicionalasahoz erdemes megvizsgalni, hogy egy epulelgepeszeti vagy muszaki webshop milyen funkciokat kinal a gyakorlatban. A piaci osszehasonlitas celja nem az, hogy a projekt mindenben felulmulja a letezo rendszereket, hanem hogy megmutassa: a valasztott funkciok valos teruleti igenyekhez kapcsolodnak.",
        "Az altalanos webshopok tobbnyire eros termeklistaval, kategoriakkal, keresovel es kosarral rendelkeznek, de a belso adminisztracios folyamatok vagy nem lathatok, vagy kulon vallalati rendszerben mukodnek. A szakdolgozat szempontjabol a TDLWebshop erteke eppen abban jelenik meg, hogy a vasarloi oldal mellett az admin oldal es a helyszini vasarlas is bemutatott resze a rendszernek.",
        "A domain kulonlegessege, hogy az epulelgepeszeti termekeknel a vasarlo gyakran nem csak gyorsan rendelni akar, hanem szakmai dontesben is segitseget var. Emiatt a katalogushoz kotott AI asszisztens jo szakdolgozati kiegeszites, de csak akkor, ha a korlatai vilagosak: nem helyettesiti a szakembert, es nem talalhat ki nem letezo termekeket.",
    ]:
        add_paragraph(doc, p)
    add_table(
        doc,
        ["Szempont", "Altalanos webshop", "Epulelgepeszeti webshop", "TDLWebshop megoldasa"],
        [
            ["Termeklista", "Altalanos kategoriak es keresesi lehetoseg.", "Szakmai kategoriak, meretek, keszlet.", "Futes, viz, hutes, szellozes, szerelesi anyagok, lakossagi megoldasok."],
            ["Admin folyamat", "Tobbnyire nem lathato a felhasznalonak.", "Keszlet es rendeles gyors kezelese fontos.", "Admin felulet, CSV import, keszletfigyeles, helyszini vasarlas."],
            ["Rendeleskovetes", "Alap statuszok.", "Fontos a keszlet es teljesites kovetese.", "Profilbol lathato rendelesek es admin statuszvaltas."],
            ["B2B jelleg", "Gyakran kulon rendszer.", "Ceges kedvezmeny es mentett vasarlo fontos.", "Mentett vasarlok, ceges jeloles, helyszini rogzitese."],
            ["AI tamogatas", "Ritkan katalogushoz kotott.", "Szakmai kerdeseknel hasznos lehet.", "Katalogus-alapu, korlatozott AI asszisztens OpenRouter proxyn keresztul."],
        ],
        "2. tablazat - Piaci funkciok osszehasonlitasa",
    )
    add_heading(doc, "2.1. A sajat rendszer erteke", 2)
    for p in [
        "A TDLWebshop fo erteke nem egyetlen kulonleges funkcioban, hanem a folyamatok egyuttmukodeseben van. A vasarlo termeket keres, kosarba tesz, rendelest ad le, majd a profiljaban koveti az allapotot. Az admin ugyanebben az adatkorben dolgozik: modositja a rendelest, figyeli a keszletet, PDF-et general, illetve helyszini vasarlast rogzit.",
        "A rendszer oktatasi szempontbol azert is erdekes, mert tobb, egymassal osszefuggo technikai teruletet mutat be. Ilyen a kliensoldali Angular alkalmazas, a Firestore alapu adattarolas, a Firebase Authentication, a biztonsagi szabalyok, a CI ellenorzes, valamint a szerveroldali titokkezelessel kialakitott AI proxy.",
        "A dolgozatban ezt az erteket ugy kell bizonyitani, hogy nem csak a kesz kepernyok latszanak, hanem az is, hogyan jut el az adat a felhasznaloi muvelettol a Firestore dokumentumig, hogyan valtozik a rendelestortenet, es milyen ellenorzesek vedik az adatokat.",
    ]:
        add_paragraph(doc, p)
    add_figure_placeholder(doc, "Piaci osszehasonlito tablazat", "Ide beillesztheto a sajat piackutatasbol keszult rovid tablazat vagy diagram.")


def add_requirements(doc):
    add_heading(doc, "3. Kovetelmenyek es use case-ek", 1)
    for p in [
        "A kovetelmenyek megfogalmazasa azert fontos, mert a szakdolgozatban a rendszer mukodeset nem eleg felsorolni. Meg kell mutatni, hogy az egyes funkciok milyen felhasznaloi vagy adminisztracios igenyre adnak valaszt, es hogyan ellenorizheto, hogy ezek tenylegesen mukodnek.",
        "A TDLWebshop kovetelmenyei harom nagy csoportba sorolhatok. Az elso csoport a vasarloi folyamatokat tartalmazza: termekkereses, kosar, checkout, profil es rendeleskovetes. A masodik az adminisztracios folyamatokat fedi le: termekfeltoltes, rendeleskezeles, keszlet, vasarlo kezeles, kupon es PDF. A harmadik csoport a technikai es minosegi kovetelmenyeket rogziti, peldaul a jogosultsagokat, biztonsagot, tesztelhetoseget es reprodukalhatot.",
    ]:
        add_paragraph(doc, p)
    add_table(
        doc,
        ["Azonosito", "Kovetelmeny", "Use case", "Modul / bizonyitek"],
        [
            ["K1", "A felhasznalo tudjon termeket keresni es kategoriak szerint bongeszni.", "Termekek bongeszese", "Termeklista oldal, keresesi mezo, kategoriak"],
            ["K2", "A felhasznalo tudjon termeket kosarba tenni es mennyiseget modositani.", "Kosar kezeles", "CartService, kosar oldal"],
            ["K3", "A checkout ellenorizze az emailt, telefonszamot es kotelezo adatokat.", "Rendeles leadasa", "checkout.ts validacios logika"],
            ["K4", "A rendeles letrejotte utan az admin lassa es modositani tudja az allapotot.", "Admin rendeleskezeles", "admin.ts, order.service.ts"],
            ["K5", "A statuszvaltas audit es keszletvaltozas mellett tortenjen.", "Rendeles teljesitese", "OrderService tranzakcios logika"],
            ["K6", "A dolgozo csak korlatozott admin funkciokat erjen el.", "Dolgozoi felulet", "Firestore rules, admin jogosultsagi logika"],
            ["K7", "A PDF bizonylat tartalmazza a rendelest, vevot, tetelek es osszegeket.", "Szamla / bizonylat letoltese", "invoice.service.ts"],
            ["K8", "Az AI asszisztens ne talaljon ki termeket, csak katalogushoz kototten ajanljon.", "AI kerdes megvalaszolasa", "chatbot-llm.service.ts, Worker proxy"],
        ],
        "3. tablazat - Kovetelmeny-traceability osszefoglalo",
    )
    add_heading(doc, "3.1. Fo use case-ek", 2)
    use_cases = [
        ("UC1", "Vasarlo termeket keres es kosarba teszi", "A vasarlo kategoriat valaszt vagy keresoszot ad meg, megnyitja a termeket, majd kosarba helyezi."),
        ("UC2", "Vasarlo rendelest ad le", "A vasarlo kitolti a checkout urlapot, a rendszer validal, majd letrehozza a rendelest."),
        ("UC3", "Admin teljesiti a rendelest", "Az admin megnyitja a rendelest, ellenorzi az adatokat, statuszt modosit, a rendszer auditot es keszletvaltozast rogzít."),
        ("UC4", "Admin helyszini vasarlast rogzit", "Az admin mentett vasarlot valaszt vagy uj adatot ad meg, tetelet ad a vasarlashoz, majd PDF bizonylatot general."),
        ("UC5", "Dolgozo termeket kezel", "A dolgozo a szamara engedelyezett admin funkciokkal termeket tolthet fel es keszletet ellenorizhet."),
        ("UC6", "Felhasznalo AI asszisztenst hasznal", "A felhasznalo katalogushoz vagy epulelgepeszethez kapcsolodo kerdest tesz fel, a rendszer korlatos valaszt ad."),
    ]
    add_table(doc, ["Use case", "Nev", "Leiras"], use_cases, "4. tablazat - Use case-ek rovid osszefoglalasa")
    for code, name, desc in use_cases:
        add_heading(doc, f"3.1.{code[-1]}. {name}", 3)
        add_paragraph(doc, desc)
        add_paragraph(doc, "Sikeres lefutas eseten a rendszer egyertelmu visszajelzest ad, az adat pedig a megfelelo Firestore gyujtemenyben vagy kliensoldali allapotban megjelenik. Hibaag eseten a felhasznalo validacios vagy jogosultsagi uzenetet kap, nem pedig csendes hibaval talalkozik.")
    add_figure_placeholder(doc, "Use case diagram", "Ide illesztheto a vasarlo, admin, dolgozo es AI asszisztens kapcsolatát bemutato use case abra.")


def add_gui_ux(doc):
    add_heading(doc, "4. GUI/UX tervezes", 1)
    for p in [
        "A felulet tervezesenel a cel egy olyan modern, sotet alapu, ipari hangulatu arculat kialakitasa volt, amely illeszkedik az epulelgepeszeti temahoz. A dark mode alapertelmezett megjelenes, mert a TDLWebshop logoja es a technikai jellegu termekkor jol mukodik ebben a vizualis kornyezetben. A light mode ugyanakkor fontos kenyelmi es uzleti alternativa.",
        "A design nem pusztan dekorativ szerepet tolt be. A felhasznaloi folyamatokban a keresosav, kategoriak, kosar, profil, termekkartyak es CTA gombok elhelyezese azt szolgalja, hogy a vasarlo gyorsan eljusson a termekhez es a rendeleshez. Az admin felulet mas logikat kovet: ott a surubb, tablazatosabb es funkciokozpontu elrendezes a hatekonyabb.",
        "A GUI/UX fejezetben a vegleges dolgozatba tenyleges kepernyokepeket kell beszurni. Ezek nem csak illusztraciok, hanem bizonyitekok arra, hogy a rendszer nem egy izolalt kodreszlet, hanem hasznalhato felhasznaloi felulettel rendelkezo webalkalmazas.",
    ]:
        add_paragraph(doc, p)
    add_table(
        doc,
        ["Kepernyo", "Cel", "Fontos allapot", "Bizonyitek"],
        [
            ["Kezdolap", "Elso benyomas, kategoria es akcio kiemeles.", "Dark mode, kategoriak dropdown.", "Kepernyokep"],
            ["Termeklista", "Bongeszes, kereses, szures.", "Tobb termek, szurok, akcios termekek.", "Kepernyokep"],
            ["Termekadatlap", "Reszletes termekadatok es kosarba helyezes.", "Kepgaleria, ar, keszlet.", "Kepernyokep"],
            ["Kosar", "Mennyiseg, osszeg, tovabblepes.", "Tobb termek, torles es mennyiseg.", "Kepernyokep"],
            ["Checkout", "Adatbekeres es validacio.", "Hibas email/telefon, sikeres rendeles.", "Kepernyokep"],
            ["Profil", "Rendelestortenet es adatok.", "Korabbi rendelesek, statusz.", "Kepernyokep"],
            ["Admin", "Belso folyamatok kezelese.", "Termekimport, rendelesek, keszlet.", "Kepernyokep"],
            ["AI asszisztens", "Katalogushoz kotott segitseg.", "Domain kerdes es valasz.", "Kepernyokep"],
        ],
        "5. tablazat - GUI/UX kepernyok es bizonyitekok",
    )
    screenshots = [
        ("Kezdolap dark mode, kategoriak legordulo menuvel", "Fooldal, kategoriak menu nyitva."),
        ("Kezdolap AI asszisztenssel", "Fooldal, AI ablak nyitva, egy relevans epulelgepeszeti kerdessel."),
        ("Termeklista szures/kereses kozben", "Termekek oldal, keresoszo vagy kategoria szuro aktiv."),
        ("Termekadatlap", "Egy konkret termek oldala, termekkep, ar, keszlet, kosar gomb lathato."),
        ("Kosar tobb termekkel", "Kosar oldal legalabb ket tetellel."),
        ("Checkout validacio", "Hibas email vagy telefonszam peldaval."),
        ("Checkout sikeres rendeles", "Sikeres leadas utani visszajelzes."),
        ("Profil es rendeleskovetes", "Vasarloi profil oldal korabbi rendelessel."),
        ("Kivansaglista", "Kivansaglista oldal mentett termekkel."),
        ("Admin attekintes", "Admin fo nezet statisztikai kartyakkal."),
        ("Admin termekkezeles es CSV import", "Admin termekek ful CSV import resszel."),
        ("Admin keszletfigyeles", "Keszlet ful alacsony keszlet jelzessel."),
        ("Helyszini vasarlas", "Mentett vasarlo kivalasztva, tetelek hozzaadva."),
        ("PDF szamla vagy bizonylat", "General PDF megnyitva, fejlec es tetelsorok lathatok."),
        ("Admin felhasznalo es jogosultsag kezeles", "Felhasznalok ful admin/dolgozo szerepkorokkal."),
        ("GitHub Actions zold CI", "GitHub Actions lista legfrissebb zold futassal."),
    ]
    for caption, instruction in screenshots:
        add_figure_placeholder(doc, caption, instruction)


def add_technology(doc):
    add_heading(doc, "5. Technologiai hatter", 1)
    for p in [
        "A TDLWebshop Angular alapu egyoldalas webalkalmazaskent keszult. Az Angular valasztasa azert indokolt, mert komponensalapu felépítést, eros TypeScript tamogatast es strukturalt szolgaltatasreteget ad. Egy webshop es admin rendszer eseteben ez kulonosen hasznos, mert a termeklista, kosar, checkout, profil es admin felulet kulon komponensekre bonthato.",
        "Az adattarolasi es hitelesitesi reteget Firebase szolgaltatasok adjak. A Firestore dokumentumalapu modellje jol illeszkedik a termekek, rendelesek, felhasznaloi profilok, kuponok es auditbejegyzesek tarolasahoz. A Firebase Authentication segitsegevel a bejelentkezes es szerepkorhoz kotott feluletkezeles is egyszerubben megvalosithato.",
        "A rendszerben megjeleno AI asszisztens kulso LLM-szolgaltatasra epul, de a kliens nem kozvetlenul hivja az OpenRouter API-t. A kulcs vedelme erdekeben egy Cloudflare Worker proxy kezeli a szerveroldali hivasokat. Ez fontos biztonsagi dontes, mert API kulcsot nem szabad nyilvanos frontend kodba egetni.",
    ]:
        add_paragraph(doc, p)
    add_table(
        doc,
        ["Technologia", "Szerep a rendszerben", "Indoklas"],
        [
            ["Angular", "Frontend keretrendszer", "Komponensalapu, TypeScript alapu, nagyobb felulethez jol strukturalt."],
            ["TypeScript", "Alkalmazaslogika", "Tipusossag, attekinthetobb szolgaltatasok es komponensek."],
            ["Firebase Authentication", "Felhasznaloazonositas", "Emailes belepes, szerepkorhoz kotheto profilok."],
            ["Cloud Firestore", "Adattarolas", "Dokumentumalapu modell, realis webshop entitasokhoz illesztheto."],
            ["Firestore Rules", "Biztonsagi retegek", "Jogosultsagok es payload validacio adatbazis szinten."],
            ["Firebase Hosting", "Publikalas", "Egyszeru deploy es webes eleres."],
            ["Cloudflare Worker", "AI proxy", "API kulcs vedelme es CORS kezeles."],
            ["GitHub Actions", "CI ellenorzes", "Build es teszt futtatas bizonyitasa."],
        ],
        "6. tablazat - A rendszerben hasznalt technologiak",
    )
    add_heading(doc, "5.1. Alternativak es dontesek", 2)
    for p in [
        "A backend megvalositasa tortenhetett volna sajat Node.js/Express szerverrel is. A Firebase valasztasa azert volt kedvezo, mert a szakdolgozati celhoz gyorsabban adott hitelesitest, hostingot, adatbazist es biztonsagi szabalyokat. Ez nem jelenti azt, hogy minden uzleti logika idealisan kliensoldalon maradhat; a dolgozatban ezt MVP-korlatkent is rogziteni kell.",
        "A bankkartya-elfogadas es teljes penzugyi integracio tudatosan kimaradt. Ennek oka, hogy egy valos fizetesi szolgaltato bekotese jogi, biztonsagi es tesztelesi kovetelmenyeket hozna be. A rendszerben a fizetesi modok adminisztrativan jelennek meg, de a valos penzmozgast a projekt nem kezeli.",
        "Az AI asszisztens eseteben szinten fontos volt a hatarhuzas. A cel nem egy teljes szakmai tanacsado rendszer, hanem egy katalogushoz kotott, segito jellegu funkcio. Ha a kerdes tul tavoli vagy nincs pontos termektalalat, a rendszernek nem szabad magabiztosan kitalalt termeket ajanlania.",
    ]:
        add_paragraph(doc, p)


def add_architecture(doc):
    add_heading(doc, "6. Architektura es adatáramlas", 1)
    for p in [
        "A rendszer architekturaja harom fo retegre bonthato. Az elso a kliensoldali Angular alkalmazas, amely a felhasznaloi feluletet, a komponenseket es a szolgaltatasokat tartalmazza. A masodik a Firebase retege, ahol a hitelesites, Firestore adatbazis, biztonsagi szabalyok es hosting jelenik meg. A harmadik a kulso integracios reteg, amelybe az AI proxy es az OpenRouter szolgaltatas tartozik.",
        "Az Angular alkalmazasban a komponensek a megjelenitesert es felhasznaloi interakcioert felelnek, mig a szolgaltatasok kezelik az adatmuveleteket. Ilyen szolgaltatas a kosar, rendeles, szamla, profil, AI asszisztens vagy termekkezeles logikaja. Ez a szetvalasztas segit abban, hogy a felulet ne kozvetlenul tartalmazza az osszes uzleti logikat.",
        "Az adatáramlas peldaja a checkout folyamaton jol bemutathato. A vasarlo a feluleten megadja az adatokat, a komponens validal, majd a rendelesi szolgaltatas letrehozza a Firestore dokumentumot. Az admin felulet kesobb ugyanezt a rendelest olvassa, statuszt modosit, auditbejegyzest keszit es bizonyos esetekben keszletvaltozast hajt vegre.",
    ]:
        add_paragraph(doc, p)
    add_figure_placeholder(doc, "Komponens-architektura abra", "Angular frontend, Firebase/Firestore, Firestore Rules, Cloudflare Worker es OpenRouter kapcsolatainak bemutatasa.")
    add_figure_placeholder(doc, "Checkout szekvencia abra", "Vasarlo -> Angular checkout -> OrderService -> Firestore -> admin rendeleslista folyamat.")
    add_heading(doc, "6.1. Modulok", 2)
    add_table(
        doc,
        ["Modul", "Feladat", "Kapcsolodo adat"],
        [
            ["Home / Product UI", "Kezdolap, kategoriak, termekkartyak.", "products"],
            ["Cart", "Kosar allapot, mennyiseg, osszegzes.", "local state, products"],
            ["Checkout", "Adatbekeres, validacio, rendeles letrehozasa.", "orders"],
            ["Profile", "Felhasznaloi adatok es rendelesek.", "users, orders"],
            ["Admin", "Termek, rendeles, keszlet es vasarlo kezeles.", "products, orders, savedCustomers"],
            ["Invoice", "PDF bizonylat generalasa.", "orders, order items"],
            ["AI assistant", "Katalogushoz kotott valaszadas.", "products, Worker proxy"],
        ],
        "7. tablazat - Fo modulok es kapcsolataik",
    )
    add_code_placeholder(doc, "src/pages/checkout/checkout.ts", "kb. 367. sortol", "A rendelés veglegesitesenek folyamata, validacio es adatosszeallitas.")
    add_code_placeholder(doc, "src/app/services/order.service.ts", "41-127. sor", "Statusz, audit es keszletvaltozas tranzakcios kezelese.")
    add_code_placeholder(doc, "workers/openrouter-proxy/src/index.js", "teljes proxy lenyegi resze", "OpenRouter hivas szerveroldali kulcskezelessel es CORS korlatozassal.")


def add_data_model(doc):
    add_heading(doc, "7. Adatmodell", 1)
    for p in [
        "A Firestore dokumentumalapu adatmodellje nem relacios tablakbol, hanem gyujtemenyekbol es dokumentumokbol epul fel. Ez gyors fejlesztest tesz lehetove, de megkoveteli, hogy az entitasok es mezok tudatosan legyenek meghatarozva. A TDLWebshop eseteben a legfontosabb entitasok a termek, kosartetel, rendeles, felhasznaloi profil, mentett vasarlo, kupon, szamla es auditbejegyzes.",
        "Az adatmodell tervezesenel fontos szempont volt, hogy a vasarloi es admin felulet ugyanarra az adatbazisra epuljon. A termekeket a vasarlo listakent latja, az admin pedig szerkesztheti, importalhatja vagy keszlet szempontbol vizsgalhatja. A rendeleseket a vasarlo sajat profiljabol koveti, az admin pedig statuszt es teljesitest kezel rajtuk.",
        "A dokumentumalapu modell egyik kockazata, hogy a tul laza szerkezet adatminosegi problemakhoz vezethet. Emiatt a Firestore rules es az alkalmazas oldali validacio kulonosen fontos: nem eleg az adatot elmenteni, azt is ellenorizni kell, hogy a mezok megfelelo tipusuak es jogosultsaggal irhatok.",
    ]:
        add_paragraph(doc, p)
    add_table(
        doc,
        ["Entitas / gyujtemeny", "Fontos mezok", "Kapcsolat / szerep"],
        [
            ["products", "name, sku, category, price, stock, images, discount", "Termeklista, admin termekkezeles, AI katalogus."],
            ["orders", "customer, items, total, status, paymentMethod, createdAt", "Checkout es admin rendeleskezeles."],
            ["orderStatusAudit", "orderId, oldStatus, newStatus, actor, createdAt", "Statuszvaltas nyomon kovetese."],
            ["users", "email, role, disabled, profile data", "Jogosultsag es profiladatok."],
            ["savedCustomers", "name, email, phone, company, taxNumber, disabled", "Helyszini vasarlas es mentett vasarlo kezeles."],
            ["coupons", "code, discount, active, validity", "Kedvezmeny logika."],
            ["newsletter", "email, createdAt", "Hirlevel feliratkozas admin oldali lathatosaga."],
        ],
        "8. tablazat - Adatmodell fo entitasai",
    )
    add_figure_placeholder(doc, "Adatmodell diagram", "Mermaid vagy PlantUML diagram a products, orders, users, savedCustomers, coupons es audit kapcsolataival.")
    add_code_placeholder(doc, "src/app/services/order.service.ts", "269-302. sor", "Szamlaszam generalas logikaja es rendeléshez kapcsolasa.")


def add_implementation(doc):
    add_heading(doc, "8. Megvalositas kulcsfolyamatai", 1)
    for p in [
        "A megvalositas fejezet celja, hogy ne csak a hasznalt technologiakat sorolja fel, hanem bemutassa a rendszer legfontosabb mukodesi pontjait. A TDLWebshop eseteben ilyen kulcsfolyamat a checkout, a rendelesek statuszkezelese, a helyszini vasarlas, a PDF bizonylat, a CSV termekimport, a jogosultsagkezeles es az AI asszisztens.",
        "A checkout folyamatban a felhasznalo adatai, a kosar tartalma, a szallitasi es fizetesi mod, valamint az esetleges kupon egy kozos rendelesi objektumma all ossze. A rendszer validalja a kotelezo mezoket, majd a rendeles letrehozasaval a vasarloi ut lezárul. A sikeres rendelés utan a vasarlo visszajelzest kap, az admin pedig latja a rendelest.",
        "A rendeles statuszkezelesnel fontos mernoki szempont, hogy a statuszvaltas ne elszigetelt mezomodositas legyen. A rendszer auditbejegyzest keszit, es bizonyos esetekben keszletvaltozast is vegrehajt. Ez a logika a dolgozatban kulon hangsulyozhato, mert megmutatja, hogy a rendszer az üzleti kovetkezmenyeket is figyelembe veszi.",
    ]:
        add_paragraph(doc, p)
    add_heading(doc, "8.1. Checkout es rendelesmentes", 2)
    for p in [
        "A checkout a vasarloi oldal legkritikusabb folyamatainak egyike. Itt derul ki, hogy a termekbongeszes es kosar logika valos rendelese alakul-e. A feluleten a felhasznalo megadja a szamlazasi, szallitasi es kapcsolattartasi adatokat, a rendszer pedig ellenorzi az email es telefonszam formatumat is.",
        "MVP-korlatkent fontos megemliteni, hogy a webes rendelés eseteben bizonyos ar- es kedvezmenyadatok kliensoldalrol erkeznek. Ez szakdolgozati prototipusnal elfogadhatoan dokumentalhato, de eles webshopban szerveroldali ujraszamolasra lenne szukseg. A dolgozatban ezt nem elrejteni kell, hanem mernoki tudatossaggal leirni.",
    ]:
        add_paragraph(doc, p)
    add_code_placeholder(doc, "src/pages/checkout/checkout.ts", "kb. 367. sortol", "Rendeles veglegesitese, checkout adatok osszeallitasa es hibaagak.")
    add_heading(doc, "8.2. Statusz, audit es keszlet", 2)
    add_paragraph(doc, "A rendelesek teljesitese es torlese keszletoldali kovetkezmenyekkel jarhat. A helyes megoldas celja, hogy a rendeles allapota, az auditbejegyzes es a keszletlogika osszhangban legyen. Ez szakdolgozati szempontbol az egyik legerosebb resz, mert nem pusztan feluleti CRUD muveletrol van szo.")
    add_code_placeholder(doc, "src/app/services/order.service.ts", "41-127. sor", "Statuszvaltas, auditbejegyzes es keszlet tranzakcios kezelese.")
    add_heading(doc, "8.3. Helyszini vasarlas", 2)
    add_paragraph(doc, "A helyszini vasarlas funkcio a rendszer domainhez illeszkedo kulonlegessege. Az admin vagy dolgozo mentett vasarlot valaszthat, uj adatokat adhat meg, termeket kereshet, mennyiseget allithat, majd a vasarlas utan PDF bizonylatot tolthet le. Ez a folyamat a B2B es boltban torteno ertekesitesi helyzeteket modellezi.")
    add_code_placeholder(doc, "src/app/services/order.service.ts", "222-267. sor", "Helyszini rendeles tranzakcios mentese es termektetelek kezelese.")
    add_heading(doc, "8.4. CSV import", 2)
    add_paragraph(doc, "A termekfeltoltesnel a kezi rogzitest kiegesziti a CSV import. Ez kulonosen akkor hasznos, ha tobb tucat termeket kell feltolteni kepekkel, arakkal, SKU-val es kategoriakkal. A validacios lepes azert fontos, mert importnal egy hibas oszlop vagy elvalaszto karakter sok dokumentumot ronthatna el.")
    add_code_placeholder(doc, "src/pages/admin/admin.ts", "1181-1257. sor", "CSV import validacio, hibas sorok jelzese es mentheto termekek feldolgozasa.")
    add_heading(doc, "8.5. PDF bizonylat", 2)
    add_paragraph(doc, "A PDF bizonylat/szamla generalasa az admin folyamatok egyik lathato eredmenye. A dokumentum tartalmazza a kiallito es vevo adatait, a rendeles azonositojat, a tetelek listajat, a netto es brutto osszegeket, valamint a fizetesi es szallitasi adatokat.")
    add_code_placeholder(doc, "src/app/services/invoice.service.ts", "9-154. sor", "PDF szamla felépítése, fejlec, vevo/kiallito blokkok, tetelsorok es osszesites.")
    add_figure_placeholder(doc, "General PDF bizonylat", "A PDF-et nyisd meg es keszits kepernyokepet ugy, hogy a fejlec, vevo es tetelsorok latszodjanak.")
    add_heading(doc, "8.6. AI asszisztens", 2)
    add_paragraph(doc, "Az AI asszisztens a sajat termekkatalogusbol kiindulva probal segitseget adni. A cel az, hogy a valasz ne legyen fuggetlen a webshop adataitol. Ha nincs pontos talalat, a rendszernek inkabb altalanos szakmai iranyt kell adnia, es jeleznie kell, hogy pontos ajanlathoz emailes vagy szemelyes egyeztetes szukseges.")
    add_code_placeholder(doc, "src/app/services/chatbot-llm.service.ts", "26-84. es 217-236. sor", "AI domainkorlatozas, kataloguslogika es nem relevans kerdesek kezelese.")


def add_security(doc):
    add_heading(doc, "9. Biztonsag es adatvedelem", 1)
    for p in [
        "A webes rendszerekben a biztonsag nem kiegeszito elem, hanem alapkovetelmeny. A TDLWebshop eseteben kulonosen fontos a felhasznaloi adatok, rendelési adatok, admin jogosultsagok, kuponok, PDF-ek es AI API kulcsok kezelese. A dolgozatban ezeket nem eleg megemliteni: be kell mutatni, milyen kockazatot jelentenek es milyen megoldas csokkenti a kockazatot.",
        "A Firebase/Firestore hasznalata miatt a biztonsagi szabalyok kulcsszerepet kapnak. A frontend kod onmagaban nem ved elegge, mert egy rosszindulatu kliens kozvetlenul is probalkozhat adatbazis-muveletekkel. Emiatt a Firestore rules feladata, hogy szerepkor, aktiv felhasznalo es payload szerkezet alapjan korlatozza az olvasast es irast.",
        "A titokkezeles szempontjabol kulon fontos, hogy API kulcs vagy jelszo ne keruljon a repositoryba. Az OpenRouter kulcs a Cloudflare Worker secretjekent tarolando, a repoban pedig csak .env.example es dokumentalt konfiguracios valtozok szerepelhetnek.",
    ]:
        add_paragraph(doc, p)
    add_table(
        doc,
        ["Kockazat", "Leiras", "Kezeles a projektben", "Tovabbi javaslat"],
        [
            ["Jogosulatlan admin muvelet", "Normal user admin adatot modositana.", "Szerepkor alapu rules es frontend guard.", "Custom claim es token revocation eles rendszerben."],
            ["Tiltott felhasznalo", "Disabled user tovabb probalkozik.", "Aktiv user ellenorzes szabalyokban.", "Auth szintu tiltás elesben."],
            ["Kupon visszaeles", "Kliensoldali kuponmanipulacio.", "Kuponvalidacio es dokumentalt MVP-korlat.", "Szerveroldali ujraszamolas."],
            ["PDF adatvedelem", "Szemelyes adatok szerepelnek a bizonylaton.", "Csak jogosult feluleten generalhato.", "Letoltesi naplozas es tarolasi szabaly."],
            ["AI API kulcs kiszivargas", "Kulcs frontendbe vagy gitbe kerul.", "Cloudflare Worker secret, .env.example.", "Rate limit es kvota."],
            ["Vendeg rendelés azonositas", "Email-alapu osszekotes kockazata.", "Dokumentalt korlat.", "Regisztralt fiokhoz kotott rendelestortenet erosítese."],
        ],
        "9. tablazat - Biztonsagi minimum es kockazatok",
    )
    add_code_placeholder(doc, "firestore.rules", "25-76. sor", "Aktiv felhasznalo, admin es dolgozo jogosultsag ellenorzese.")
    add_code_placeholder(doc, "firestore.rules", "294-361. sor", "products, orders, users, savedCustomers es audit szabalyok.")
    add_code_placeholder(doc, "src/pages/admin/admin.ts", "607-748. sor", "Admin es dolgozoi jogosultsagok feluleti kezelese.")
    add_heading(doc, "9.1. MVP-korlatok biztonsagi szempontbol", 2)
    for p in [
        "A szakdolgozati rendszer nem vallalja egy teljes erteku penzugyi backend szerepet. A bankkartya elfogadas, szerveroldali arujraszamolas, keszletfoglalas es ERP-szintu integracio tovabbfejlesztesi irany. Fontos, hogy ez nem hiba, ha a dolgozatban vilagosan szerepel mint tudatos MVP-hatar.",
        "Az AI asszisztensnel a CORS beallitas onmagaban nem teljes visszaeles elleni vedelem. Ezt rate limit, kvota vagy felhasznaloi azonositashoz kotott hivaslimit erositheti. A dolgozatban ezt a reszt ugy erdemes bemutatni, mint mukodo prototipust, amelynel elesites elott tovabbi vedelmi lepesek kellenek.",
    ]:
        add_paragraph(doc, p)


def add_testing(doc):
    add_heading(doc, "10. Teszteles es validacio", 1)
    for p in [
        "A teszteles celja annak bizonyitasa, hogy a fo felhasznaloi es adminisztracios folyamatok nem csak elmeletben leteznek, hanem ellenorizhetoen mukodnek. A TDLWebshop eseteben a teszteles ket szinten tortent: automata tesztekkel es kezi tesztjegyzokonyvvel.",
        "Az automata tesztek elonye, hogy a CI folyamatban is futtathatok, igy egy GitHub Actions zold allapot kulso bizonyitek a build es tesztek sikeressegere. A kezi teszteles ezzel szemben a teljes felhasznaloi elmenyt, a kepernyoallapotokat, hibauzeneteket es reszponzivitast vizsgalja.",
        "A dolgozatban kulonosen fontos megmutatni, hogy a kritikus folyamatok ellenorizve voltak: checkout, rendelés letrejotte, keszletvaltozas, admin statuszvaltas, tiltott jogosultsag, kuponvalidacio, CSV import es AI asszisztens.",
    ]:
        add_paragraph(doc, p)
    add_table(
        doc,
        ["Tesztazonosito", "Folyamat", "Elvart eredmeny", "Bizonyitek"],
        [
            ["T1", "Regisztracio es belepes", "Felhasznalo profilja letrejon, tiltott user uzenetet kap.", "Kezi teszt"],
            ["T2", "Termekkereses", "A lista szurodik, nincs talalat eseten ures allapot latszik.", "Kezi teszt"],
            ["T3", "Kosar mennyiseg", "Osszeg frissul, torles mukodik.", "Automata/kezi"],
            ["T4", "Checkout validacio", "Hibas email/telefon nem enged tovabb.", "Kepernyokep"],
            ["T5", "Sikeres rendeles", "Rendeles letrejon Firestore-ban es adminban latszik.", "Kezi teszt"],
            ["T6", "Admin statuszvaltas", "Audit es keszletvaltozas rogzodik.", "Automata/kezi"],
            ["T7", "CSV import", "Hibas sorok jelzodnek, ervenyes sorok menthetoek.", "Kezi teszt"],
            ["T8", "PDF bizonylat", "PDF tartalmazza az adatokat es nem csuszik ossze.", "Kepernyokep"],
            ["T9", "AI asszisztens", "Relevans kerdesre domainvalasz, irrelevansra elhatarolas.", "Kezi teszt"],
            ["T10", "CI", "GitHub Actions build/test zold.", "Kepernyokep"],
        ],
        "10. tablazat - Tesztelesi terv osszefoglalasa",
    )
    add_heading(doc, "10.1. Build es CI", 2)
    add_paragraph(doc, "A projekt CI folyamata GitHub Actions segitsegevel fut. A workflow celja, hogy friss commit utan ellenorizze a telepithetoseget, buildet es teszteket. A zold CI allapot a dolgozatban kulon kepernyokepkent szerepeljen, mert gyorsan bizonyitja, hogy a repo nem csak lokalis gepen mukodik.")
    add_figure_placeholder(doc, "GitHub Actions zold CI futas", "GitHub Actions oldalon a legfrissebb zold workflow run latszodjon, commit nevvvel es idotartammal.")
    add_heading(doc, "10.2. Kezi tesztjegyzokonyv", 2)
    add_paragraph(doc, "A kezi tesztjegyzokonyvet a dolgozat mellekleteben vagy a teszteles fejezetben erdemes bemutatni. A teszteles soran nem csak a sikeres utakat, hanem a hibaagakat is ellenorizni kell. Ilyen pelda a hibas email, tiltott felhasznalo, keszlethiany, jogosulatlan admin muvelet vagy nem relevans AI kerdes.")
    add_figure_placeholder(doc, "Checkout validacios hiba", "Checkout oldal hibas email vagy telefonszam beirasaval.")


def add_repro_ai_summary(doc):
    add_heading(doc, "11. Reprodukalhatosag, CI es deploy", 1)
    for p in [
        "A konzulensi visszajelzes alapjan a reprodukalhatosag kulcsfontossagu. Egy szakdolgozati repo akkor biralhato nyugodtan, ha egy masik kornyezetben is elindithato, es a szukseges lepesek nem csak a fejleszto sajat gepen ismertek. Emiatt a README, .env.example, telepitesi leiras, demo szerepkorok es CI bizonyitek kiemelt szerepet kapnak.",
        "A repo nem tartalmazhat node_modules, build mappakat, valodi jelszavakat, API kulcsokat vagy gepfuggo allomanyokat. A fuggosegek package.json es lock file alapjan telepithetok, a titkok pedig kulon kornyezeti valtozokban vagy szolgaltatoi secretkent tarolandok.",
        "A deploy ket kulon reszbol allhat: a frontend Firebase Hostingra kerul, az AI asszisztens szerveroldali proxyja pedig Cloudflare Workerkent mukodik. A dolgozatban ezt a szetvalasztast erdemes roviden elmagyarazni, mert jol mutatja a kulcsvedelem es a kliens/szerver felelossegi hatarainak megerteset.",
    ]:
        add_paragraph(doc, p)
    add_table(
        doc,
        ["Lepes", "Parancs / ellenorzes", "Cel"],
        [
            ["Fuggosegek telepitese", "npm install", "Angular projekt fuggosegeinek telepitese."],
            ["Build", "npm run build", "Frontend fordithatosaganak ellenorzese."],
            ["Teszt", "npm test -- --watch=false", "Automata tesztek futtatasa."],
            ["Hosting deploy", "firebase deploy --only hosting", "Frontend publikalasa."],
            ["Firestore rules deploy", "firebase deploy --only firestore:rules", "Biztonsagi szabalyok elesitese."],
            ["Worker deploy", "npx wrangler deploy", "OpenRouter proxy publikalasa."],
        ],
        "11. tablazat - Reprodukcios es deploy lepesek",
    )

    add_heading(doc, "12. Mesterseges intelligencia hasznalata a fejlesztes soran", 1)
    for p in [
        "A szakdolgozat keszitese soran mesterseges intelligencia tobb ponton segitette a munkat. Fontos kulonvalasztani ket teruletet: az egyik a fejlesztesi folyamatban hasznalt AI-tamogatas, a masik pedig maga a TDLWebshopban megjeleno vasarloi AI asszisztens. A ketto nem ugyanaz: az elobbi fejlesztoi segedeszkoz, az utobbi a rendszer egyik funkcioja.",
        "A fejlesztes soran AI-tamogatas elsosorban otletelesre, kodreview-ra, hibakeresesre, dokumentacios szerkezet kialakitasara, tesztforgatokonyvek osszeallitasara es nyelvi javitasokra lett hasznalva. A kimenetek nem automatikusan kerultek elfogadasra: a fejleszto feladata maradt a kod megertese, ellenorzese, futtatasa es a dontesek meghozatala.",
        "Az AI hasznalata soran kulon figyelmet kellett forditani a titkok kezelesere. API kulcsot vagy jelszot nem szabad a kodba vagy nyilvanos repoba helyezni. Egy korabbi kulcsot vissza kellett vonni, majd a veglegesebb megoldasban az OpenRouter kulcs Cloudflare Worker secretkent kerult kezelesre.",
        "Az AI korlatai a fejlesztesben is megjelentek. Előfordulhat, hogy egy javaslat elavult, nem illeszkedik a projekt szerkezetebe, vagy tul altalanos megfogalmazast ad. Emiatt a vegleges szakdolgozati szoveget sajat nyelvezetre kell atirni, es minden allitast a mukodo rendszerrel, tesztekkel vagy dokumentacioval kell alatamasztani.",
        "A rendszerben mukodo AI asszisztens szinten korlatos. Nem celja, hogy szakmernoki felelosseggel teljes muszaki ajanlatot adjon. A cel inkabb az, hogy a termekkatalogus alapjan iranyt mutasson, es ha nincs pontos talalat vagy biztos adat, akkor javasolja az emailes vagy szemelyes egyeztetest.",
    ]:
        add_paragraph(doc, p)
    add_table(
        doc,
        ["AI-hasznalat terulete", "Mire szolgalt", "Validacio"],
        [
            ["Otleteles", "Funkciok, MVP-hatar es dokumentacios fejezetek atgondolasa.", "Konzulensi elvarasokkal osszevetve."],
            ["Kodreview", "Hibalehetosegek, jogosultsagok, validacio es tranzakciok vizsgalata.", "Kod olvasasa, build es teszt."],
            ["Dokumentacio", "Fejezetvaz, tesztjegyzokonyv es biztonsagi minimum elokeszitese.", "Sajat atiras es szakmai ellenorzes."],
            ["AI asszisztens funkcio", "Felhasznaloi kerdesek katalogushoz kotott megvalaszolasa.", "Domainkerdes, termekkerdes es irrelevans kerdes tesztje."],
        ],
        "12. tablazat - AI-hasznalat es ellenorzes",
    )

    add_heading(doc, "13. Osszefoglalas, korlatok es tovabbfejlesztes", 1)
    for p in [
        "A TDLWebshop egy epulelgepeszeti temaju webshop es adminisztracios rendszer MVP-je. A projekt bemutatja, hogyan epulhet fel egy vasarloi oldalbol, admin feluletbol, Firestore adatmodellbol, jogosultsagi szabalyokbol, PDF bizonylatbol, CSV importbol, keszletkezelesbol es AI asszisztensbol allo webes rendszer.",
        "A munka legerosebb resze a funkciok osszekapcsolasa: a vasarloi rendeles nem onmagaban all, hanem megjelenik az admin feluleten, statuszt kap, PDF bizonylattal es keszletlogikaval kapcsolodik a belso folyamatokhoz. A helyszini vasarlas es mentett vasarlo kezeles kulon domainerteket ad a rendszernek.",
        "A rendszernek vannak tudatos korlatai. Ilyen a teljes penzugyi integracio hianya, a webes rendelesek szerveroldali arujraszamolasanak tovabbfejlesztesi igenye, a keszletfoglalas ipari szintu kezelese es az AI asszisztens rate limit/kvota oldali erositesenek szuksegessege. Ezeket nem hianyossagkent, hanem tovabbfejlesztesi iranykent kell bemutatni.",
        "[SAJAT ZARO REFLEXIO HELYE] Ide 1-2 bekezdesben sajat hangon ird le, mit tanultal a projektbol, mi volt a legnehezebb, mit csinalnal maskent, es melyik reszere vagy a legbuszkebb. Ez legyen teljesen sajat megfogalmazas, mert a vedesen is ebbol tudsz a legtermeszetesebben beszelni.",
    ]:
        add_paragraph(doc, p)

    add_heading(doc, "Irodalomjegyzek", 1)
    refs = [
        "Angular dokumentacio - https://angular.dev/",
        "Firebase dokumentacio - https://firebase.google.com/docs",
        "Cloud Firestore Security Rules dokumentacio - https://firebase.google.com/docs/firestore/security/get-started",
        "Cloudflare Workers dokumentacio - https://developers.cloudflare.com/workers/",
        "OpenRouter dokumentacio - https://openrouter.ai/docs",
        "GitHub Actions dokumentacio - https://docs.github.com/actions",
        "[Tovabbi piaci osszehasonlitasban hasznalt webshopok forrasai]",
        "[Intézményi szakdolgozati formai és tartalmi útmutató]",
    ]
    for ref in refs:
        add_paragraph(doc, ref)

    add_heading(doc, "Mellekletek", 1)
    add_paragraph(doc, "M1. Kezi tesztjegyzokonyv kitoltott allapotban.")
    add_paragraph(doc, "M2. Kiemelt kódrészletek: checkout, OrderService, Firestore rules, InvoiceService, AI proxy.")
    add_paragraph(doc, "M3. Kepernyokepek gyujtemenye: vasarloi oldal, admin oldal, AI asszisztens, CI.")
    add_paragraph(doc, "M4. Reprodukcios lepesek es kornyezeti valtozok mintaja.")


def add_deepening_pages(doc):
    """Adds thesis-style elaboration so the working draft has real body text, not only an outline."""
    blocks = [
        ("A vasarloi ut reszletes ertelmezese", [
            "A vasarloi ut a kezdolapon indul, ahol a felhasznalo gyorsan kepet kap a webshop temajarol es a fobb kategoriakrol. A keresosav kiemelt szerepet kap, mert egy epulelgepeszeti webshopban a vasarlo sokszor konkret termeknevet, meretet vagy cikkszamot keres. A kategoriak dropdown formaja egyszerre ad gyors navigaciot es modern feluleti elmenyt.",
            "A termeklista celja az, hogy a vasarlo ossze tudja hasonlitani a kinalatot. A termekkartyakon az ar, keszlet es termeknev egyutt jelenik meg. A kedvezmenyes vagy uj termekek jelolese segiti a dontest, de a felulet nem viszi tulzasba a vizualis elemeket, mert a muszaki termekeknel az olvashatosag fontosabb, mint a pusztan reklam jellegu megjelenes.",
            "A kosar oldal a dontes es a rendelés kozotti atmenet. Itt a felhasznalo ellenorizheti a mennyiseget, torolhet tetelet, majd tovabblephet a checkoutra. A kosar logikaja azert fontos, mert az osszegzesnek mindig kovetnie kell a mennyisegvaltozast, es nem szabad instabil termekazonositokra epulnie.",
        ]),
        ("Adminisztracios folyamatok reszletes ertelmezese", [
            "Az admin felulet nem egyszeru masolata a vasarloi oldalnak. Mas celcsoportot szolgal ki, ezert surubb informacios elrendezest, gyors muveleti gombokat es tobb adatot tartalmaz. Egy adminnak nem marketinges termekbemutatot kell latnia, hanem olyan adatokat, amelyek alapjan rendelest, keszletet es vasarlot tud kezelni.",
            "A helyszini vasarlas funkcio kulonosen fontos a projektben, mert egy valos epulelgepeszeti boltban a rendelesek nem mindig weben erkeznek. Az adminnak vagy dolgozonak gyorsan ki kell tudnia valasztani egy mentett vasarlot, termeket kell keresnie, mennyiseget kell megadnia, majd bizonylatot kell generalnia. Ez a funkcio emiatt a domainhez illeszkedo, nem altalanos webshopos kiegeszites.",
            "A dolgozoi jogosultsagok bevezetese azt mutatja, hogy a rendszer nem egyetlen admin felhasznalora epul. A dolgozo tud termeket es keszletet kezelni, vasarlast rogzithet, de nem kap minden jogosultsagot. Ez kozelebb viszi a rendszert egy valos szervezeti mukodeshez.",
        ]),
        ("Adatbiztonsagi es minosegi szempontok reszletes ertelmezese", [
            "A Firestore szabalyokkal kapcsolatban fontos felismeres, hogy a kliensoldali ellenorzes csak felhasznaloi kenyelmi es UX szempontbol eleg. Biztonsagi szempontbol az adatbazis oldali szabalyok jelentik az igazi vedelmi vonalat. Emiatt a szerepkorok, aktiv felhasznaloi allapot es payload mezok ellenorzese szakdolgozati szinten is ertekes resze a projektnek.",
            "A titkok kezelese kulon fejezetet erdemel, mert modern webes fejlesztesnel az API kulcsok es jelszavak veletlen commitolasa gyakori hiba. A projektben a vegleges irany az, hogy a repoban csak pelda konfiguracio szerepel, az eles kulcsok pedig szolgaltatoi secretben tarolodnak. Ez mernoki es biztonsagi szempontbol is helyes irany.",
            "A teszteles es CI nem csak technikai formalitas. A zold CI azt bizonyitja, hogy a projekt egy kulso futtatokornyezetben is le tud fordulni es a tesztek lefutnak. Ez kulonosen fontos a konzulensi visszajelzes alapjan, mert a repo rendezettsége es reprodukalhatosaga a vegleges ertekelesben is szerepet kap.",
        ]),
        ("Tovabbfejlesztesi lehetosegek reszletes ertelmezese", [
            "Eles kornyezetben a legfontosabb tovabbfejlesztes a szerveroldali uzleti logika erosítese lenne. A webes rendelesek arat, kedvezmenyet es keszletet nem idealis teljes mertekben kliensoldali adatokra bizni. Egy Cloud Function vagy kulon backend endpoint ujraszamolhatna a vegosszeget es ellenorizhetne a keszletet.",
            "Az AI asszisztens tovabbfejlesztesehez hasznos lenne keresesi index, termekleirasokbol keszult strukturalt tudastar es rate limit. Ez csokkentene a pontatlan ajanlasok es visszaelesek kockazatat. A rendszer jelenlegi allapotaban prototipusnak tekintheto, amely megmutatja az iranyt, de nem helyettesiti a szakmai ugyfelszolgalatot.",
            "A keszletfigyeles tovabb bovitheto fogyasi trenddel, utanrendelesi javaslattal es admin ertesitesekkel. Ez kulonosen jol illeszkedik az epulelgepeszeti kereskedeshez, ahol bizonyos szerelesi anyagok es gyakran hasznalt alkatreszek keszleten tartasa uzletileg fontos.",
        ]),
    ]
    for title, paragraphs in blocks:
        add_heading(doc, title, 2)
        for p in paragraphs:
            add_paragraph(doc, p)
        add_paragraph(doc, "A vegleges dolgozatban ezt a reszt erdemes sajat tapasztalatokkal kiegesziteni: milyen hibat talaltal, mit javitottal, milyen dontest hoztal, es a vedesen hogyan tudnad sajat szavaiddal elmagyarazni.")


def add_extended_thesis_body(doc):
    """Adds longer thesis-level discussion blocks so the draft is closer to a 40-50 page document with figures."""
    sections = [
        (
            "A projekt valos felhasznaloi problemaja",
            "A TDLWebshop alapgondolata nem pusztan az volt, hogy keszuljon egy termekeket listazo weboldal. Az epulelgepeszeti kereskedelemben a vasarloi es uzemeltetoi folyamatok egyszerre jelennek meg: a vasarlo termeket keres, rendelest ad le, kedvezmenyt hasznalna, a dolgozo pedig keszletet ellenoriz, rendelest rogzit es bizonylatot keszit. Ez a kettos igeny adta a projekt szakdolgozati erteket, mert a rendszernek nem egyetlen oldalt, hanem tobb, egymassal osszefuggo munkafolyamatot kellett kezelnie.",
            [
                "A dolgozatban ezt azert fontos reszletesen bemutatni, mert egy webshop csak akkor tekintheto termekszeru MVP-nek, ha a felhasznalo nem akad el a kulcsfolyamatokban. A kezdolap, termeklista, kosar es checkout onmagaban meg nem eleg, ha nincs mellette adminisztracio, rendelestortenet, jogosultsagkezeles es keszletlogika. A TDLWebshop ezek kozul tobb teruletet is osszekapcsol, ezert a projekt kozelebb all egy valos uzleti alkalmazashoz, mint egy egyszeru tanuloprojekthez.",
                "A valos problema masik oldala az, hogy az epulelgepeszeti termekeknel a vasarlo sokszor nem csak nevet vagy arat nez. Fontos lehet a kategoria, a muszaki jellemzo, a keszlet, a szallitasi lehetoseg es az, hogy helyszini vasarlas soran is rogzithetok legyenek a tetelek. Emiatt a rendszer nem csak latvanyos feluletet kapott, hanem olyan admin oldali funkciokat is, amelyekkel a bolt munkafolyamatai bizonyithatok.",
                "A vegleges dolgozatban ide erdemes beilleszteni egy rovid sajat tortenetet arrol, miert pont epulelgepeszeti webshop lett a tema. Ez segit abban, hogy a dolgozat ne altalanos szoftverleiraskent hasson, hanem latszodjon mogotte a szemelyes motivacio es a domain ismerete.",
            ],
        ),
        (
            "Az MVP hataranak indoklasa",
            "A szakdolgozati MVP hatara tudatosan ugy lett kijelolve, hogy a rendszer a vasarloi utat es az adminisztracios utat egyarant bizonyitsa. Nem celja a projektnek, hogy teljes erteku vallalatiranyitasi rendszer, bankkartya-elfogado rendszer vagy szamlazo program legyen. A cel az, hogy a rendelestol a keszlet- es admin kezelesig bemutathato legyen egy mukodo, bovithetove tervezett webshop alap.",
            [
                "Ez a hatar azert vedheto, mert a szakdolgozat nem egy eles kereskedelmi szolgaltatas teljes jogi es penzugyi megfeleleset vallalja, hanem egy mernokileg indokolt MVP-t. A bankkartya fizetes peldaul nem teljes online fizetesi integracio, hanem a fizetesi modok modellezese es a rendelesi folyamat resze. A PDF bizonylat sem hivatalos e-szamla szolgaltatas, hanem a rendelest osszefoglalo dokumentum, amely bemutatja a generalt kimenet logikajat.",
                "Az MVP-hatar resze az is, hogy bizonyos kockazatokat a dolgozat korlatkent nevez meg. Ide tartozik a webes rendelesek szerveroldali vegosszeg-ujraszamitasa, az AI asszisztens pontos termekajanlasi korlatja, valamint az eles fizetesi szolgaltato hianya. Ezek nem elhallgatando hibak, hanem olyan tovabbfejlesztesi pontok, amelyek egy szakdolgozatban kifejezetten jol mutatjak a realis onertekelest.",
                "A biralo szamara kulonosen hasznos, ha az MVP-hatar tablazatban is megjelenik: mi keszult el, mi maradt prototipus, es mi tudatosan nem lett bevallalva. Ezzel elkerulheto, hogy a projekt felkesz rendszernek hasson, mikozben valojaban meghatarozott celra keszult.",
            ],
        ),
        (
            "Piaci osszehasonlitas szerepe",
            "A piaci osszehasonlitas celja nem az, hogy a TDLWebshop minden nagy webaruhaznal jobb legyen. A cel az, hogy latszodjon: milyen megoldasok jellemzoek a hasonlo webshopokra, es ehhez kepest a sajat rendszer milyen szakdolgozati tobbletet ad. A konzulensi visszajelzes alapjan ez a fejezet kulonosen fontos, mert megmutatja, hogy a projekt nem elszigetelt kodgyakorlat.",
            [
                "Az osszehasonlitasban erdemes 2-4 epulelgepeszeti vagy muszaki webshopot vizsgalni. A szempontok lehetnek: termekkereses, kategoria struktura, kosar, checkout, regisztralt profil, admin jellegu funkciok lathatosaga, keszletinformacio, szakmai tamogatas, akcios termekek es mobilos hasznalhatosag. A sajat rendszer erteke nem feltetlenul abban van, hogy tobb termeket tartalmaz, hanem abban, hogy a vasarloi es admin oldali funkciokat egy szakdolgozati MVP-ben osszekapcsolja.",
                "A TDLWebshop egyik erosebb pontja a helyszini vasarlas rogzitese, a mentett vasarlok kezelese es a dolgozoi szerepkor megjelenese. Ezek olyan funkciok, amelyek sok publikus webshopon nem lathatok, mert belso vallalati feluletekhez tartoznak. Szakdolgozati szempontbol viszont pont ezek bizonyitjak, hogy a rendszer nem csak kirakatoldal.",
                "A piaci fejezetet a vegleges dolgozatban egy tablazattal erdemes zarnod. A tablazatban ne csak pipakat hasznalj, hanem rovid megjegyzeseket is: miert fontos az adott szempont, es a TDLWebshop hogyan kezeli.",
            ],
        ),
        (
            "Kovetelmenyek es traceability",
            "A kovetelmenyek fejezete kot ossze a celkituzes es a megvalositas kozott. Ha egy rendszerben sok funkcio van, konnyen elveszik, hogy melyik mire szolgalt. A traceability tablazat ezt oldja meg: kovetelmeny, use case, megvalositott modul es teszt bizonyitek egy sorban latszik.",
            [
                "A TDLWebshop eseteben kulon kell kezelni a vasarloi kovetelmenyeket, az adminisztracios kovetelmenyeket es a biztonsagi kovetelmenyeket. A vasarloi oldalhoz tartozik a termekkereses, kosar, checkout, profil es rendelestortenet. Az admin oldalhoz tartozik a termekfeltoltes, CSV import, rendeleskezeles, keszletfigyeles, helyszini vasarlas, mentett vasarlok es jogosultsagkezeles.",
                "A biztonsagi kovetelmenyek nem csak technikai reszletek. Ide tartozik, hogy tiltott felhasznalo ne tudjon tovabb rendelni, normal vasarlo ne erjen el admin funkciokat, a dolgozo csak a sajat szerepkorenek megfelelo feluletet lassa, es API kulcs ne szerepeljen a repoban. Ezeket a dolgozatban kulon erdemes kifejteni, mert a webes rendszereknel gyakori biralati pont a jogosultsagkezeles.",
                "A traceability tablazat azert is jo, mert a teszteles fejezethez is hidat ad. Ha leirod, hogy peldaul a checkout validaciohoz melyik kepernyokep es melyik teszt tartozik, akkor a biralo gyorsan latja, hogy nem csak allitod a mukodest, hanem bizonyitod is.",
            ],
        ),
        (
            "GUI/UX dontesek indoklasa",
            "A felulet tervezesenel a cel egy modern, ipari hangulatu webshop volt. A dark mode, a kek-piros accent szinek es a TDLWebshop logo vilaga egyutt adja azt a vizualis iranyt, amely az epulelgepeszeti temahoz illik. A fontos dontes az volt, hogy a latvany ne menjen az olvashatosag rovasara.",
            [
                "A kezdolap szerepe az elso benyomas. Itt a kategoria dropdown, a keresosav, a hero szekcio es a kiemelt termekek vezetik a felhasznalot. A termeklista es termekkartya reszeknel mar a funkcionalitas erosebb: termeknev, ar, keszlet, kosarba helyezes es reszletek gomb. A dolgozatban a kepernyokepeken keresztul erdemes megmutatni, hogy ugyanaz a vizualis rendszer tobb oldalon is visszater.",
                "A light mode kulon kihivas volt, mert nem eleg a hatteret feherre es a szoveget feketere valtani. A kontraszt, gombok, fejlec es kartyak arnyekolasa kulon figyelmet igenyel. Ezt a reszt a dolgozatban UX donteskent lehet leirni: a layout azonos marad, de a szinek es arnyekok alkalmazkodnak a temahoz.",
                "Az admin feluletnel a design mas szempontot kovet. Itt a surubb informacio, a gyors gombok, tablazatok, filterek es statuszok a fontosak. A dolgozatban jo, ha kulon kepernyokep van a vasarloi oldalrol es az admin oldalrol, mert igy latszik, hogy ket celcsoporthoz mas feluleti logika tartozik.",
            ],
        ),
        (
            "Angular komponensalapu felepites",
            "Az Angular hasznalata azert indokolt, mert a rendszer tobb oldalbol, szolgaltatasbol es allapotbol all. A komponensek es service-ek elkulonitese segit abban, hogy a megjelenites es az adatkezeles ne keveredjen teljesen ossze. Egy webshopnal ez kulonosen fontos, mert ugyanazt a termekadatot tobb helyen is hasznalni kell.",
            [
                "A termeklista, kosar, checkout, profil es admin oldalak kulon felelossegeket kapnak. A service retegek, peldaul a product, cart, order, auth vagy chatbot service, az adatbetoltes es uzleti logika kozelebbi reszeit kezelik. Ez a felepites a dolgozat architektura fejezeteben jol bemutathato, mert latszik belole a kliensoldali alkalmazas szerkezete.",
                "A komponensalapu felbontas masik elonye a tesztelhetoseg. Bizonyos reszek onalloan vizsgalhatok, peldaul a kosar osszegszamitas, kuponlogika vagy validacio. A szakdolgozatban nem kell minden tesztet reszletesen bemutatni, de a kritikus folyamatokhoz erdemes egy-egy peldat hozni.",
                "A vegleges szovegben erdemes roviden kiterni arra is, hogy az Angular valasztasa milyen alternativakhoz kepest tortent. Peldaul React vagy Vue is alkalmas lett volna, de az Angular strukturaltsaga es service-orientalt felepitese illeszkedett a nagyobb alkalmazas jellegu projekthez.",
            ],
        ),
        (
            "Firebase es Firestore szerepe",
            "A Firebase a projektben backend-szeru szolgaltatasokat ad: hitelesites, Firestore adatbazis, hosting es biztonsagi szabalyok. Ez lehetove tette, hogy a szakdolgozat fokusza ne teljesen egy sajat backend infrastrukturara menjen el, hanem a webshop uzleti folyamataira es adatmodelljere.",
            [
                "A Firestore dokumentumalapu modellje jol illeszkedik a termekek, rendelesek, felhasznaloi profilok, kuponok es mentett vasarlok tarolasahoz. Ugyanakkor a dokumentumalapu adatbazis nem ugyanazt a gondolkodast igenyli, mint egy relacios adatbazis. A dolgozatban erdemes megmutatni, mely entitasok kulon collectionbe kerultek, es milyen adatok kapcsoljak oket ossze.",
                "A biztonsagi szabalyok kulon fejezetet erdemelnek. A kliensoldali kod onmagaban nem ved elegge, mert egy rosszindulatu felhasznalo megprobalhat kozvetlenul adatbazishoz fordulni. Emiatt a Firestore rules oldalon is kezelni kell, ki mit olvashat es irhat. Ez a projekt egyik erosebb mernoki pontja.",
                "A Firebase config es API kulcs kerdese kulon magyarazatot igenyel. Webes Firebase alkalmazasoknal a kliensoldali config reszben publikus, de ettol meg a valodi vedelem a security rules es a jogosultsagi modell. A dolgozatban erdemes leirni, hogy titkok, OpenRouter kulcsok es jelszavak nem kerulhetnek repoba.",
            ],
        ),
        (
            "Rendeleskezeles es keszletlogika",
            "A rendeleskezeles a webshop egyik legfontosabb folyamata. A vasarloi oldalon a checkout gyujti ossze a kosarat, szallitasi adatokat, fizetesi modot es kedvezmenyt. Az admin oldalon a rendeles statusza kovetheto, modosithato, es a keszletvaltozas is kapcsolodik hozza.",
            [
                "A helyszini vasarlasnal kulonosen fontos volt, hogy a rendeles, keszletvaltozas es bizonylat ne teljesen kulon mozogjon. Ha egy admin rogzit egy eladast, a rendszernek figyelembe kell vennie a kivalasztott termekeket, mennyisegeket es vasarlot. Ez a folyamat szakdolgozati szempontbol ertekesebb, mint egy egyszeru kosarmentes, mert belso bolti munkafolyamatot modellez.",
                "A webes rendelesehez kapcsolodo korlatot oszinten erdemes leirni. Ha bizonyos arak es kedvezmenyek kliensoldalrol jonnek, akkor eles rendszerben szerveroldali ellenorzesre lenne szukseg. Ez nem feltetlenul gyengiti a szakdolgozatot, ha a korlat vilagosan szerepel es tovabbfejlesztesi terv is tartozik hozza.",
                "A statuszvaltas es audit naplo azt mutatja, hogy a rendeles allapotvaltozasai kovethetok. Ez a dolgozatban jo pelda arra, hogyan lehet egy uzleti folyamatot nem csak vegeredmenykent, hanem tortenettel egyutt kezelni.",
            ],
        ),
        (
            "PDF bizonylat es szamlageneralas",
            "A PDF generalas a projektben a rendelesi folyamat kezzelfoghato kimenete. A vasarlo vagy admin nem csak adatbazis-bejegyzest kap, hanem letoltheto dokumentumot is. Ez szakdolgozati szempontbol jol demonstralja, hogy a rendszer kepes strukturalt adatbol dokumentumot eloallitani.",
            [
                "A bizonylat felepitese tartalmazza a rendeles azonositojat, vevoi adatokat, termekeket, fizetesi es szallitasi modot, valamint a vegosszeget. A korabbi visszajelzes alapjan fontos volt az elrendezes javitasa, hogy az osszeg ne csusszon bele mas blokkokba. Ezt a dolgozatban akar fejlesztesi tapasztalatkent is meg lehet emliteni.",
                "Eles kornyezetben egy jogilag ervenyes e-szamla mas kovetelmenyeket jelentene. Ehhez kulso szamlazo szolgaltatas, jogszabalyi megfeleles es adatszolgaltatasi logika kellene. A TDLWebshop jelen allapotaban PDF bizonylatot general, ami MVP szinten bemutatja a dokumentumkeszites elvet.",
                "A vegleges dolgozatban ide mindenkeppen kell egy kep a generalt PDF-rol. A kep mellett roviden magyaraa el, mely adatok a rendelesbol szarmaznak, es melyek fix cegadatok vagy szamlaformazasi elemek.",
            ],
        ),
        (
            "AI asszisztens korlatai es szerepe",
            "Az AI asszisztens nem azert kerult a rendszerbe, hogy szakembert vagy ugyfelszolgalatot teljesen helyettesitsen. A cel az volt, hogy a termekkatalogushoz es epulelgepeszeti temakhoz kapcsolodoan segitse a tajekozodast. Ez fontos kulonbseg, mert egy szakdolgozati prototipusban az AI-hasznalatot korlatokkal egyutt kell bemutatni.",
            [
                "A rendszerben az OpenRouter hivas nem kozvetlenul a kliensbol tortenik, hanem Cloudflare Worker proxyn keresztul. Ez azert fontos, mert az API kulcs nem kerulhet a bongeszobe vagy a repoba. A dolgozatban ez biztonsagi es architekturalis donteskent is szerepelhet.",
                "A valaszoknal kulon szabaly, hogy az asszisztens ne ajanljon random termeket akkor, ha nincs megfelelo katalogus-talalat. Ilyenkor adhat altalanos szakmai iranyt, de jeleznie kell, hogy pontos ajanlatert vagy beszerezhetosegert emailben vagy szemelyesen erdemes egyeztetni. Ez csokkenti a teves termekajanlas kockazatat.",
                "A dolgozatban kulon kell valasztani a fejlesztes soran hasznalt AI-t es a webshopba beepitett AI asszisztenst. Az elobbi munkaeszkoz volt, az utobbi a kesz rendszer egyik funkcioja. Ez a kulonvalasztas megfelel a konzulensi elvarasnak is.",
            ],
        ),
        (
            "Tesztelesi bizonyitas es kezi ellenorzes",
            "A teszteles fejezet celja, hogy a rendszer mukodeset ne csak allitasok tamasszak ala. A build, az automata tesztek es a kezi tesztjegyzokonyv egyutt adnak bizonyitekot. Ez kulonosen fontos, mert a szakdolgozatban tobb felhasznaloi szerepkor es tobb kritikus folyamat szerepel.",
            [
                "Az automata tesztek kozul erdemes kiemelni a kosar, checkout/kuponlogika, admin statuszvaltas, PDF generalas, AI asszisztens es validacio tesztjeit. Nem kell minden assertet bemutatni, de a dolgozatban legyen egy tablazat: tesztelt funkcio, modszertan, vart eredmeny, tenyleges eredmeny.",
                "A kezi tesztek a vizualis es folyamatbeli reszek miatt szuksegesek. Ilyen pelda a mobilos megjelenes, admin tablazatok kezelhetosege, checkout hibauzenetek, mentett vasarlo kiválasztasa vagy PDF letoltes. Ezeket screenshotokkal es pipalhato tesztlistaval lehet igazolni.",
                "A GitHub Actions zold CI kepernyokepe eros bizonyitek, mert megmutatja, hogy a projekt nem csak a sajat gepen futott. A vegleges dolgozatban ezt a reprodukalhatosagi fejezethez erdemes tenni.",
            ],
        ),
        (
            "Repo-higienia es beadashoz tartozo tisztasag",
            "A konzulensi visszajelzes egyik fo pontja a repo rendezettsége volt. Ez nem mellekes adminisztracio, hanem a szakdolgozat biralhatosaganak feltetele. Ha a repoban node_modules, build mappak, jelszavak vagy helyi segedfajlok vannak, akkor a projekt kevesbe professzionalisnak tunik.",
            [
                "A vegleges repoban a futtathato kodnak, konfiguracios mintaknak, dokumentacionak es teszteknek kell szerepelniuk. A szemelyes segedanyagok, hosszu atirasi alapok es munkapeldanyok lokalisan maradhatnak, de nem kell oket GitHubra feltolteni. Ez kulonosen fontos, mert a dolgozat sajat nyelvezetre huzasa a hallgato feladata.",
                "A .env.example szerepe az, hogy megmutassa, milyen kornyezeti valtozokra van szukseg, de ne tartalmazzon valodi kulcsot. Az OpenRouter API kulcsot a Worker secret kezeli, nem a repoban tarolt fajl. Ez a titokkezelesrol szolo reszben jo pelda.",
                "A README-nek olyan szemelynek is segitenie kell, aki nem vett reszt a fejlesztesben. Telepites, inditas, build, teszt, Firebase, Worker es demo szerepkorok: ezek legyenek egyertelmuek. A dolgozat reprodukalhatosagi fejezete erre hivatkozhat.",
            ],
        ),
        (
            "Sajat ertekeles es tanulsagok",
            "A dolgozat vegi reflexio nem egyszeru lezaro bekezdes. Itt kell megmutatni, hogy mit tanultal a projektbol, milyen dontesek voltak nehezek, es hol latod a rendszer korlatait. A jo szakdolgozat nem hibamentesnek mutatja a projektet, hanem realisan ertekeli.",
            [
                "Erdekes tanulsag lehet, hogy a webshop fejlesztese kozben a latvanyos felulet csak az egyik resz. Legalabb ilyen fontos lett a jogosultsagkezeles, Firestore szabalyok, tesztek, adatmodell es dokumentacio. Ez jol mutatja, hogy egy szoftveres projektben a kod mellett a bizonyithatosag is mernoki munka.",
                "A nehezsegek koze bekerulhet a checkout validacio, a PDF elrendezes, a helyszini vasarlas mentese, a CSV import, az AI proxy biztonsagos bekotese vagy a GitHub CI stabilizalasa. Ezek konkret, vedesen is jol elmondhato fejlesztesi tapasztalatok.",
                "A tovabbfejlesztesi iranyok kozott realis marad a szerveroldali rendelesellenorzes, eles fizetesi szolgaltato, fejlettebb keszlet-elorejelzes, admin riportok es AI asszisztens rate limit. Ezeket nem kell most mind megvalositani, de jol zarjak a dolgozatot.",
            ],
        ),
    ]

    add_heading(doc, "Reszletes szakdolgozati kifejtesek az osszedolgozott munkapeldanyhoz", 1)
    add_paragraph(
        doc,
        "Az alabbi alfejezetek azt a celt szolgaljak, hogy a dolgozat ne csak rovid fejezetvaz legyen, hanem a konzulensi elvarasoknak megfeleloen reszletesen is bemutassa a rendszer szakmai hatteret, donteseit, korlatait es bizonyitekait. A vegleges valtozatban ezeket a reszeket sajat megfogalmazasra kell huzni, de tartalmilag a TDLWebshop mernoki tortenetet kovetik.",
    )
    for title, lead, paragraphs in sections:
        add_heading(doc, title, 2)
        add_paragraph(doc, lead)
        for paragraph in paragraphs:
            add_paragraph(doc, paragraph)
        add_paragraph(
            doc,
            "Kapcsolodo bizonyitek a vegleges dolgozatban: ide kerulhet kepernyokep, rovid kodreszlet vagy tablazat, amely az adott allitast alatamasztja. A cel, hogy a szoveg ne onmagaban alljon, hanem lathato rendszerreszhez, teszthez vagy dokumentalt donteshez kapcsolodjon.",
        )


def add_full_length_expansion(doc):
    """Adds chapter-by-chapter text that can be moved into the final thesis chapters."""
    add_heading(doc, "Fejezetenkenti reszletes kidolgozasi alap", 1)
    add_paragraph(
        doc,
        "Ez a resz nem kulon mellekletnek keszult, hanem olyan hosszabb, szakdolgozati stilusu szovegalap, amelyet a vegleges dokumentumban a megfelelo fejezetekbe lehet athelyezni vagy ott megtartani. A cel az, hogy a dolgozat terjedelme es reszletezettsege kozelebb keruljon egy 40-50 oldalas szoftveres szakdolgozathoz, mikozben a tartalom a TDLWebshop valos kodjahoz, tesztjeihez es konzulensi elvarasaihoz kapcsolodik.",
    )

    chapter_material = [
        (
            "Bevezetes es problemafelvetes bovites",
            "A bevezetesben azt kell bizonyitani, hogy a TDLWebshop nem oncelu webes felulet, hanem egy valos kereskedelmi es adminisztracios problema modellezese.",
            [
                "A felhasznaloi oldal akkor tekintheto teljesnek, ha a vasarlo a kezdolaprol eljut a termekekig, a kosarig, majd a sikeres rendelesig.",
                "Az adminisztracios oldal azert fontos, mert egy webshop eletkepesseget nem csak a vasarlok, hanem a belso uzemeltetesi folyamatok is meghatarozzak.",
                "Az epulelgepeszeti termekkor sajatossaga, hogy a vasarlo sokszor szakmai dontesi helyzetben van, ezert a kategoria, keszlet es ajanlasi logika is lenyeges.",
                "A projekt soran a legfontosabb kompromisszum az volt, hogy a rendszer szakdolgozati MVP legyen, ne pedig minden reszleteben eles vallalati termek.",
                "A dolgozatban a problemafelvetesnek ossze kell kotnie a szakmai domaint, a felhasznaloi igenyeket es a megvalositott szoftveres megoldast.",
            ],
        ),
        (
            "MVP es kovetelmenyek bovites",
            "Az MVP meghatarozasa a dolgozat egyik vedelmi pontja: vilagossa teszi, hogy mi keszult el tudatosan, es mi maradt tovabbfejlesztesi irany.",
            [
                "A vasarloi MVP resze a bongeszes, kereses, kosar, checkout, rendelestortenet es profiladatok kezelese.",
                "Az admin MVP resze a termekkezeles, CSV import, rendeleskezeles, keszletfigyeles, mentett vasarlo es helyszini vasarlas.",
                "A dolgozoi szerepkor megjelenese azert ertekes, mert nem csak admin es vasarlo kozotti egyszeru jogosultsagi modellt mutat be.",
                "A kovetelmenyeket celszeru azonosithato kodokkal ellatni, mert igy a megvalositas es teszteles fejezeteiben vissza lehet rajuk hivatkozni.",
                "A kovetelmenyek kozott nem csak funkcionalis, hanem biztonsagi, hasznalhatosagi es reprodukalhatosagi elvarasokat is szerepeltetni kell.",
            ],
        ),
        (
            "Piaci osszehasonlitas bovites",
            "A piaci fejezet azt mutatja meg, hogy a TDLWebshop milyen hasonlo rendszerekhez kepest ertelmezheto, es milyen sajat erteket ad.",
            [
                "A hasonlo webshopok elemzesenel nem az a cel, hogy pontos uzleti titkokat vagy belso admin feluleteket hasonlitsunk ossze, hanem a publikus felhasznaloi elmenyt.",
                "A TDLWebshop sajatossaga, hogy a publikus webshop mellett admin es dolgozoi folyamatokat is bemutat, ami egy szakdolgozatban eros technikai bizonyitek.",
                "A piaci tablazatban erdemes kulon jelolni, hogy mely funkciok lathatok a vasarloi oldalon, es melyek belso uzemeltetesi oldalhoz kapcsolodnak.",
                "A mobilos hasznalhatosag es a keresesi elmeny olyan szempont, amelyet a kepernyokepek is jol alatamasztanak.",
                "A piaci elemzes vegere erdemes rovid kovetkeztetest irni arrol, hogy a sajat rendszer nem termekmennyisegben, hanem folyamat-osszefogasban eros.",
            ],
        ),
        (
            "GUI es UX bovites",
            "A GUI/UX fejezetben a feluletet nem csak kepkent, hanem dontesek sorozatakent kell bemutatni.",
            [
                "A dark mode az alap vizualis irany, mert a logo es az ipari hangulat ehhez illeszkedik legjobban.",
                "A light mode nem kulon layout, hanem ugyanannak a feluletnek vilagosabb, uzletiesebb valtozata.",
                "A termekkartyaknal az ar, keszlet es kosarba helyezes gyors felismerhetosege fontosabb, mint a tulzott dekoracio.",
                "Az admin feluleten a gyors muveletek, tablazatok es allapotjelzesek dominálnak, mert ott a hatekony munkavegzes a fo cel.",
                "A kepernyokepeknek kulonbozo allapotokat is mutatniuk kell: ures lista, hiba, sikeres rendeles, validacios hiba es jogosultsagi elteres.",
            ],
        ),
        (
            "Architektura es adatmodell bovites",
            "Az architektura fejezetben azt kell megmutatni, hogyan mukodnek egyutt az Angular komponensek, a service reteg, a Firebase es az OpenRouter proxy.",
            [
                "A kliensoldali alkalmazas felelos a megjelenitesert, a felhasznaloi interakciokert es az egyes service-ek meghivasaert.",
                "A Firestore az adatok tarolasat es a jogosultsagi szabalyok egy reszet biztosítja, ezert az adatmodell es security rules egyutt ertelmezendo.",
                "A Cloudflare Worker az AI asszisztensnel biztonsagi koztes reteg, mert az OpenRouter API kulcs nem kerulhet a kliensoldalra.",
                "Az adatmodellben kulon figyelmet erdemel az Order es OrderItem kapcsolat, mert a kosar tartalma itt alakul tartos rendelesi adattá.",
                "A savedCustomers es users elkulonitese azert erdekes, mert a helyszini vasarlas es a regisztralt webshopos vasarlo nem teljesen ugyanaz a fogalom.",
            ],
        ),
        (
            "Megvalositas bovites",
            "A megvalositas fejezetnek nem eleg kodreszleteket felsorolnia; azt kell bemutatnia, hogy a kod milyen uzleti folyamatot valosit meg.",
            [
                "A checkout folyamatban a validacio, vegosszeg-szamitas, kuponkezeles es rendelesmentes egymashoz kapcsolodo lepesek.",
                "A statuszvaltas es audit log azt bizonyitja, hogy a rendeles nem csak statikus rekord, hanem kovetheto allapotvaltozasokkal rendelkezo entitas.",
                "A helyszini vasarlas rogzitese az admin felulet egyik legerosebb funkcioja, mert belso bolti munkafolyamatot modellez.",
                "A CSV import a tomeges termekfeltoltest segiti, es jol mutatja, hogy a rendszer nagyobb termekmennyisegre is felkeszitett.",
                "A PDF generalas es AI asszisztens olyan kiegeszito funkciok, amelyek a projektet termekszerubb iranyba viszik, de korlataikat is le kell irni.",
            ],
        ),
        (
            "Biztonsag es adatvedelem bovites",
            "A biztonsagi fejezetben a hangsuly azon van, hogy a rendszer milyen visszaeleseket probal megelozni, es hol maradnak MVP-szintu korlatok.",
            [
                "A jogosultsagkezelesnek kliensoldalon es Firestore rules szinten is meg kell jelennie, mert a kliensoldali elrejtes onmagaban nem eleg.",
                "A tiltott felhasznalo kezelese fontos uzleti eset, mert egy nem fizeto vagy problematikus vasarlo nem kaphat korlatlan hozzaferest.",
                "A kuponoknal a visszaeles lehetoseget ugy kell vizsgalni, hogy ki hozhat letre kupont, ki hasznalhatja, es hogyan ellenorizheto a kedvezmeny.",
                "A PDF dokumentum szemelyes adatokat tartalmazhat, ezert a letoltes es eleres kerdeset adatvedelmi szempontbol is meg kell emliteni.",
                "Az AI proxy CORS beallitasa hasznos, de onmagaban nem teljes vedelmi rendszer, ezert rate limit vagy kvota tovabbfejleszteskent szerepelhet.",
            ],
        ),
        (
            "Teszteles es validacio bovites",
            "A tesztelesi fejezetnek ki kell mutatnia, hogy a kritikus folyamatokat nem csak egyszer kiprobaltad, hanem modszertan szerint ellenorizted.",
            [
                "A build sikeressege azt bizonyitja, hogy az alkalmazas forditasa hibatlanul lefut egy tiszta allapotban.",
                "Az automata tesztek a kod bizonyos logikai reszeit ellenorzik, peldaul kosar, kupon, admin statuszvaltas vagy validacio.",
                "A kezi tesztek a teljes felhasznaloi utat fedik le, ideertve a UI allapotokat es a szerepkorok kozotti kulonbseget.",
                "A GitHub Actions zold futasa kulon bizonyitek arra, hogy a projekt nem csak lokalisan mukodott.",
                "A teszteles vegen erdemes maradek kockazatokat is felsorolni, mert ettol a fejezet mernokileg hitelesebb lesz.",
            ],
        ),
        (
            "Reprodukalhatosag es repo-higienia bovites",
            "A reprodukalhatosag azt jelenti, hogy a biralo vagy temavezeto a repobol meg tudja erteni es el tudja inditani a rendszert.",
            [
                "A README-ben szerepelnie kell a telepitesnek, inditasnak, buildnek, teszteknek es a szukseges kornyezeti valtozoknak.",
                "A .env.example nem titkok tarolasara van, hanem arra, hogy megmutassa, milyen konfiguracios kulcsok szuksegesek.",
                "A node_modules es build mappak kizarasa azert fontos, mert ezek generalt vagy gepfuggo allomanyok.",
                "A dokumentacios fajloknak a repoban a projektet kell alatamasztaniuk, a szemelyes segedanyagokat kulon kell kezelni.",
                "A CI kepernyokep es a teszteredmeny a dolgozatban azt bizonyitja, hogy a kod aktualis allapota ellenorizheto volt.",
            ],
        ),
        (
            "MI-hasznalat bovites",
            "Az MI-hasznalati fejezetnek oszinten kell bemutatnia, hol segitett AI, es hol maradt a dontes, ellenorzes es felelosseg a fejlesztonel.",
            [
                "Az AI a fejlesztes soran otletelesben, hibakeresesben, dokumentacios vazlatokban es ellenorzo listakban segitett.",
                "A kod es dokumentacio nem tekintheto automatikusan helyesnek csak azert, mert AI javasolta; minden kimenetet futtatassal vagy sajat ellenorzessel kellett validalni.",
                "A dolgozatban kulon kell emliteni, hogy a webshopba beepitett AI asszisztens mas kategoria, mint a fejlesztest segito AI eszkozok.",
                "Az AI asszisztensnel fontos korlat, hogy nem adhat korlatlan vagy ellenorizetlen termekajanlast, hanem a katalogushoz es domainhez kotott valaszokat kell adnia.",
                "A vegleges szovegben kerulni kell a marketinges altalanossagokat; konkret peldakat erdemes irni arra, hol gyorsitott es hol kellett javitani az AI javaslatat.",
            ],
        ),
        (
            "Osszegzes es tovabbfejlesztes bovites",
            "Az osszegzes akkor eros, ha nem csak megismetli a fejezeteket, hanem sajat ertekelest is ad a projektrol.",
            [
                "A projekt fo eredmenye egy olyan webshop MVP, amely a vasarloi es admin oldali folyamatokat egy rendszerben mutatja be.",
                "A legfontosabb mernoki eredmenyek kozott szerepel a jogosultsagi modell, Firestore rules, CSV import, PDF bizonylat, AI proxy es tesztelheto workflow.",
                "A korlatok kozott erdemes megemliteni az eles fizetesi integracio hianyat, a webes rendeles szerveroldali ujraszamitasanak tovabbfejleszteset es az AI rate limitet.",
                "A sajat reflexio reszben leirhato, hogy a projekt soran a legtobb munka nem az elso latvanyos felulet, hanem a hibak, edge case-ek es dokumentalas korul jelent meg.",
                "A tovabbfejlesztesi iranyok legyenek realisak: online fizetes, szamlazo integracio, fejlettebb keszletstatisztika, admin riportok, pontosabb AI tudastar.",
            ],
        ),
    ]

    perspective_templates = [
        "A fejezet vegleges valtozataban ezt a gondolatot erdemes egy konkret TDLWebshop peldaval alatamasztani. Peldaul egy olyan kepernyokeppel vagy kodreszlettel, ahol latszik, hogy a tervezesi dontes nem csak elmeleti, hanem a rendszer mukodeseben is megjelenik.",
        "Szakdolgozati szempontbol ennek az a jelentosege, hogy a rendszer nem kulonallo funkciok listaja, hanem osszefuggo folyamatokbol allo alkalmazas. A biralo szamara ez akkor lesz egyertelmu, ha a szovegben a cel, a megvalositas es a teszt bizonyiteka egymasra hivatkozik.",
        "A vegleges szovegben ezt a reszt sajat nyelvezetre kell huzni. Erdemes beleirni, hogy a fejlesztes kozben milyen donteshelyzet alakult ki, milyen alternativat merlegeltel, es miert a jelenlegi megoldas mellett dontottel.",
        "A kapcsolodo abra vagy tablazat ne diszites legyen, hanem bizonyitek. Ha egy funkcio mukodeset allitod, akkor mellette jelenjen meg kepernyokep, teszteredmeny, adatmodell-reszlet vagy rovid kodreszlet is.",
    ]

    for heading, intro, points in chapter_material:
        add_heading(doc, heading, 2)
        add_paragraph(doc, intro)
        for point in points:
            add_paragraph(doc, point)
            for template in perspective_templates:
                add_paragraph(doc, template)

    add_heading(doc, "Kiemelt bizonyitekok osszefoglalo tablazata", 2)
    add_table(
        doc,
        ["Bizonyitek", "Hova keruljon", "Mit igazol"],
        [
            ["GitHub Actions zold CI kepernyokep", "Reprodukalhatosag es teszteles fejezet", "A projekt tiszta kornyezetben is buildelheto es tesztelheto."],
            ["Firestore rules kodreszlet", "Biztonsagi fejezet", "A jogosultsagok nem csak kliensoldalon, hanem adatbazis szinten is kezeltek."],
            ["Checkout validacio kepernyokep", "GUI/UX es teszteles fejezet", "A hibas bemeneteket a rendszer felhasznaloi szinten is kezeli."],
            ["Helyszini vasarlas kepernyokep", "Megvalositas fejezet", "Az admin oldali belso bolti folyamat is mukodik."],
            ["PDF bizonylat kepernyokep", "Megvalositas es eredmenyek fejezet", "A rendelési adatbol strukturalt dokumentum keszul."],
            ["AI asszisztens kepernyokep", "MI-hasznalat es megvalositas fejezet", "A rendszerben mukodo AI funkcio elkulonul a fejlesztest segito AI-hasznalattol."],
        ],
        "A dolgozatba javasolt legfontosabb bizonyitekok",
    )


def main():
    doc = Document()
    configure_document(doc)
    add_front_matter(doc)
    add_intro(doc)
    add_market(doc)
    add_requirements(doc)
    add_gui_ux(doc)
    add_technology(doc)
    add_architecture(doc)
    add_data_model(doc)
    add_implementation(doc)
    add_security(doc)
    add_testing(doc)
    add_deepening_pages(doc)
    add_extended_thesis_body(doc)
    add_full_length_expansion(doc)
    add_repro_ai_summary(doc)
    doc.save(OUT)

    text_chars = sum(len(p.text) for p in doc.paragraphs)
    print(f"Created: {OUT}")
    print(f"Paragraphs: {len(doc.paragraphs)}")
    print(f"Tables: {len(doc.tables)}")
    print(f"Text characters: {text_chars}")


if __name__ == "__main__":
    main()
