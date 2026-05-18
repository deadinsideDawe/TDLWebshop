from pathlib import Path
from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Cm, Pt, RGBColor


OUT_DIR = Path(__file__).resolve().parent
PART1 = OUT_DIR / "TDLWebshop_szakdolgozat_1_resz_elmeleti_tervezesi_alap.docx"
PART2 = OUT_DIR / "TDLWebshop_szakdolgozat_2_resz_megvalositas_teszteles_zaras.docx"


def setup_document() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.5)

    styles = doc.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"].font.size = Pt(12)
    styles["Normal"].paragraph_format.line_spacing = 1.15
    styles["Normal"].paragraph_format.space_after = Pt(6)

    for name, size in [("Title", 22), ("Heading 1", 16), ("Heading 2", 14), ("Heading 3", 13)]:
        styles[name].font.name = "Times New Roman"
        styles[name].font.size = Pt(size)
        styles[name].font.bold = True
        styles[name].font.color.rgb = RGBColor(17, 24, 39)

    return doc


def add_cover(doc: Document, part_title: str, subtitle: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Szegedi Tudományegyetem\nInformatikai Intézet")
    run.bold = True
    run.font.size = Pt(14)

    doc.add_paragraph("\n\n")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("SZAKDOLGOZAT\n")
    run.bold = True
    run.font.size = Pt(24)
    run = p.add_run("TDLWebshop épületgépészeti webáruház és adminisztrációs rendszer")
    run.bold = True
    run.font.size = Pt(16)

    doc.add_paragraph("\n")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(part_title)
    run.bold = True
    run.font.size = Pt(15)
    doc.add_paragraph(subtitle).alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("\n\n\n")
    table = doc.add_table(rows=2, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.cell(0, 0).text = "Készítette:"
    table.cell(0, 1).text = "Tóth Dávid László"
    table.cell(1, 0).text = "Témavezető:"
    table.cell(1, 1).text = "Dr. Bilicki Vilmos, egyetemi docens"
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(12)

    doc.add_paragraph("\n\n\n")
    p = doc.add_paragraph("Szeged\n2026")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()


def add_toc_note(doc: Document) -> None:
    doc.add_heading("Tartalomjegyzék", level=1)
    doc.add_paragraph(
        "A végleges Word dokumentumban ide automatikus tartalomjegyzéket kell beszúrni. "
        "A jelenlegi dokumentum munkapéldány: a fejezetcímek és alcímek már Heading stílusokkal szerepelnek, "
        "ezért a tartalomjegyzék Wordben a Hivatkozások / Tartalomjegyzék menüponttal frissíthető."
    )
    doc.add_page_break()


def para(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def bullet(doc: Document, text: str) -> None:
    doc.add_paragraph(text, style="List Bullet")


def add_figure_placeholder(doc: Document, caption: str, instruction: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"[ÁBRA HELYE] {caption}")
    run.bold = True
    run.font.color.rgb = RGBColor(37, 99, 235)
    para(doc, f"Beillesztendő képernyőkép/ábra: {instruction}")


def add_code_placeholder(doc: Document, title: str, file_ref: str, explanation: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(f"[KÓDRÉSZLET HELYE] {title}")
    run.bold = True
    run.font.color.rgb = RGBColor(220, 38, 38)
    para(doc, f"Beillesztendő kódképernyő: {file_ref}. {explanation}")


def add_table(doc: Document, headers, rows) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = str(value)
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(10)


def expand_topic(doc: Document, title: str, focus: str, details: list[str], evidence: str) -> None:
    doc.add_heading(title, level=2)
    para(
        doc,
        f"A {title.lower()} fejezet célja, hogy a TDLWebshop rendszer egy konkrét, szakdolgozati szempontból is "
        f"értelmezhető részét ne csak felsorolásszerűen, hanem mérnöki indoklással mutassa be. {focus} "
        "A bemutatás során különösen fontos, hogy a rendszer ne elszigetelt képernyők gyűjteményeként jelenjen meg, "
        "hanem olyan alkalmazásként, amelyben a vásárlói, adminisztrátori és dolgozói folyamatok ugyanarra az adatmodellre "
        "és jogosultsági logikára épülnek."
    )
    para(
        doc,
        "A fejlesztés során a fő szempont az volt, hogy a webáruház használható MVP-ként működjön. Ez azt jelenti, "
        "hogy a rendszerben nem minden éles vállalati webshop funkciója készült el teljes mélységben, ugyanakkor a "
        "legfontosabb végigvihető folyamatok működnek: termékek böngészése, kosárkezelés, rendelésleadás, adminisztrátori "
        "termék- és rendeléskezelés, jogosultsági elkülönítés, PDF bizonylat, készletkezelés és dokumentált tesztelés. "
        "A szakdolgozatban ezért minden fejezetnél érdemes külön jelezni, hogy az adott megoldás miért tartozik az MVP-be, "
        "és mi az, ami tudatos továbbfejlesztési irányként maradt meg."
    )
    for item in details:
        para(
            doc,
            f"A konkrét megvalósítás egyik lényeges pontja: {item} Ez a döntés azért fontos, mert az épületgépészeti "
            "termékkörben a vásárló gyakran nem csak egy terméket keres, hanem műszaki szempontból összefüggő megoldást. "
            "A webáruház felületének ezért egyszerre kell egyszerűnek maradnia és elegendő szakmai információt adnia. "
            "A választott megoldás a dolgozat szempontjából azért is hasznos, mert jól bemutatható rajta a frontend, "
            "az adatbázis, a jogosultságkezelés és a validáció együttműködése."
        )
    para(
        doc,
        f"A fejezet bizonyítékaként a dolgozatban {evidence} szerepeljen. Ezek nem csak illusztrációk: a bíráló számára "
        "azt mutatják meg, hogy a leírt funkció ténylegesen futó rendszerben is megjelenik, illetve a kódban is "
        "visszakövethető. A képernyőképek és kódrészletek mellett rövid magyarázatot kell adni arról is, hogy az adott "
        "rész milyen felhasználói igényt vagy mérnöki problémát old meg."
    )
    para(
        doc,
        "A végleges szövegben ezt a részt érdemes saját példákkal kiegészíteni. Ilyen lehet egy konkrét rendelési "
        "folyamat végigvezetése, egy adminisztrátori művelet leírása, vagy annak bemutatása, hogy hibás adat megadása "
        "esetén hogyan reagál a felület. Ettől a dolgozat nem általános technológiai ismertető lesz, hanem a saját "
        "fejlesztés működését és döntéseit bemutató mérnöki dokumentum."
    )


def add_part1() -> tuple[Document, int]:
    doc = setup_document()
    add_cover(doc, "I. rész: elméleti, tervezési és követelményi alap", "Munkapéldány a végleges szakdolgozati szöveghez")

    doc.add_heading("Feladatkiírás jellegű összefoglaló", level=1)
    para(
        doc,
        "A szakdolgozat célja egy épületgépészeti termékeket értékesítő webáruház és a hozzá tartozó adminisztrációs "
        "felület megtervezése, megvalósítása és ellenőrzése. A rendszer munkacíme TDLWebshop. Az alkalmazás célja, "
        "hogy a vásárlók termékeket kereshessenek, kosárba helyezhessenek, rendelést adhassanak le, valamint "
        "bejelentkezés után megtekinthessék saját adataikat és rendeléseiket. A rendszer adminisztrációs oldala "
        "lehetővé teszi a termékek, készletek, rendelési állapotok, felhasználói szerepkörök, mentett vásárlók és "
        "helyszíni értékesítések kezelését."
    )
    para(
        doc,
        "A feladat külön hangsúlyt helyez a reprodukálható fejlesztői környezetre, a jogosultságkezelésre, az adatmodell "
        "átláthatóságára, a biztonsági minimumok dokumentálására, valamint a tesztelési bizonyítékokra. A projekt nem "
        "teljes vállalatirányítási rendszerként készült, hanem olyan MVP-ként, amely egy valós webáruházi és adminisztrációs "
        "folyamat fő részeit demonstrálja."
    )
    doc.add_page_break()

    doc.add_heading("Tartalmi összefoglaló", level=1)
    para(
        doc,
        "A dolgozat egy Angular és Firebase alapú épületgépészeti webáruház fejlesztését mutatja be. A rendszerben a "
        "vásárlói oldal mellett külön adminisztrációs felület is készült, amely támogatja a rendeléskezelést, "
        "termékfeltöltést, CSV importot, készletfigyelést, mentett vásárlókat, PDF bizonylatot, jogosultsági szerepköröket "
        "és egy termékkatalógushoz kötött AI asszisztenst. A dolgozat bemutatja a felhasználói folyamatokat, a választott "
        "technológiákat, az adatmodellt, a biztonsági megfontolásokat és a tesztelési eredményeket."
    )
    para(
        doc,
        "A megvalósítás fő értéke, hogy nem csak egy bemutató jellegű terméklista készült, hanem egy több szerepkörös, "
        "rendelési és adminisztrációs logikával rendelkező rendszer. A dolgozat ennek megfelelően nem kizárólag a "
        "felületet ismerteti, hanem a döntések indoklását, az MVP-határ kijelölését, a biztonsági korlátokat és a "
        "tesztelési bizonyítékokat is tartalmazza."
    )
    add_toc_note(doc)

    doc.add_heading("1. Bevezetés, problémafelvetés és célkitűzés", level=1)
    topics = [
        (
            "1.1 A webáruházak szerepe a kis- és középvállalkozások működésében",
            "A kereskedelmi folyamatok jelentős része ma már digitális felületeken zajlik, ezért egy vállalkozás számára nem elegendő pusztán termékkatalógust megjeleníteni.",
            [
                "a webáruház a termékbemutatás mellett rendeléskezelési és ügyfélkommunikációs feladatokat is ellát",
                "a vásárlói bizalom szempontjából fontos a készlet, ár, szállítási mód és rendelési állapot átlátható megjelenítése",
                "a dolgozatban a webshop nem marketingfelületként, hanem működő üzleti folyamatként kerül bemutatásra",
            ],
            "kezdőlapi és terméklistás képernyőképek, valamint a vásárlói út rövid leírása",
        ),
        (
            "1.2 Az épületgépészeti termékkör sajátosságai",
            "Az épületgépészeti termékek esetében a vásárlói döntés sokszor szakmai szempontoktól, kompatibilitástól, készlettől és beszerzési lehetőségtől függ.",
            [
                "a termékekhez gyakran műszaki leírás, kategória, készlet és ár tartozik",
                "a vásárló sok esetben nem csak egy terméket, hanem egy rendszerhez illeszkedő megoldást keres",
                "a rendszernek ezért támogatnia kell a keresést, kategóriákat, szűrést és az adminisztrátori karbantartást",
            ],
            "termékkategória és termékadatlap képernyőkép",
        ),
        (
            "1.3 A TDLWebshop fejlesztési célja",
            "A TDLWebshop célja egy olyan szakdolgozati MVP elkészítése volt, amely a vásárlói és adminisztrátori folyamatokat egy rendszerben kapcsolja össze.",
            [
                "a vásárlói oldalon a böngészés, keresés, kosár és rendelésleadás kerül előtérbe",
                "az admin oldalon a termékek, rendelési állapotok, készlet és felhasználók kezelése a fő cél",
                "a dolgozói szerepkör külön kezeli azokat a műveleteket, amelyekhez nem szükséges teljes admin jogosultság",
            ],
            "MVP összefoglaló táblázat és admin áttekintő képernyőkép",
        ),
        (
            "1.4 Az MVP határai és tudatos korlátai",
            "A projekt tudatosan nem teljes körű éles webshop infrastruktúraként készült, hanem szakdolgozati MVP-ként, amely bizonyítja a fő üzleti folyamatok működését.",
            [
                "a bankkártyás fizetés a dolgozatban nem teljes fizetési szolgáltatói integrációként szerepel",
                "a webes rendelés egyes összegzési adatai kliensoldalról indulnak, ezt a dolgozat MVP-korlátként is értékeli",
                "az AI asszisztens tájékoztató jellegű szakmai és katalógus alapú segítséget ad, nem helyettesít szakembert",
            ],
            "MVP-határ és továbbfejlesztési irányok táblázat",
        ),
        (
            "1.5 A dolgozat felépítése",
            "A dolgozat felépítése a konzulensi elvárásokhoz igazodik: probléma, cél, piac, követelmények, use case-ek, UX, architektúra, megvalósítás, biztonság, tesztelés és összefoglalás.",
            [
                "az első nagy rész a tervezési és követelményi alapot mutatja be",
                "a második nagy rész a megvalósítást, ellenőrzést és értékelést részletezi",
                "a mellékletekbe kerülhetnek a hosszabb tesztjegyzőkönyvek, képernyőképek és reprodukciós leírások",
            ],
            "a dolgozat végleges tartalomjegyzéke",
        ),
    ]
    for t in topics:
        expand_topic(doc, *t)

    doc.add_heading("2. Piaci és területi összehasonlítás", level=1)
    expand_topic(
        doc,
        "2.1 Hasonló webshopok és rendszerek vizsgálata",
        "A piaci összehasonlítás célja annak bemutatása, hogy a TDLWebshop milyen ismert megoldásokhoz képest pozicionálható.",
        [
            "a nagyobb épületgépészeti webshopok erősek termékválasztékban, de a dolgozatban bemutatott admin és szakdolgozati dokumentáltság más jellegű érték",
            "az általános webshopmotorok gyors indulást adnak, de kevésbé mutatják meg a saját mérnöki döntéseket",
            "a TDLWebshop előnye, hogy a vásárlói és admin folyamatok saját fejlesztésként, átlátható kódbázisban jelennek meg",
        ],
        "piaci összehasonlító táblázat, amely 2-4 hasonló rendszert hasonlít össze",
    )
    add_table(
        doc,
        ["Szempont", "Általános webshopmotor", "Nagy épületgépészeti webshop", "TDLWebshop MVP"],
        [
            ["Termékkatalógus", "Kész megoldás, kevés saját logika", "Nagy választék", "Saját adatmodell, CSV import, kategóriák"],
            ["Admin felület", "Motorfüggő", "Belső rendszer, nem látható", "Saját admin, dolgozói és admin szerepkör"],
            ["Készletkezelés", "Bővítményfüggő", "Éles ERP integráció", "MVP szintű készlet, helyszíni eladásnál tranzakciós kezelés"],
            ["Dokumentálhatóság", "Kevés saját döntés", "Nem hozzáférhető", "Architektúra, adatmodell, teszt és biztonság bemutatható"],
            ["AI asszisztens", "Nem alapfunkció", "Ritkán látható", "Katalógushoz kötött, korlátozott szakmai segítő"],
        ],
    )
    expand_topic(
        doc,
        "2.2 A saját megoldás értéke",
        "A saját fejlesztés értéke nem abban áll, hogy minden piaci szereplőnél nagyobb funkcionalitást nyújt, hanem abban, hogy a működés mérnöki szempontból végigkövethető.",
        [
            "a rendszerben látható a frontend és backend szolgáltatások közötti felelősségmegosztás",
            "a jogosultságkezelés és Firestore szabályok külön dokumentálhatóak",
            "a tesztelés és CI bizonyítja, hogy a fő folyamatok ellenőrzött módon működnek",
        ],
        "rövid táblázat arról, hogy mely funkciók adják a szakdolgozati értéket",
    )
    expand_topic(
        doc,
        "2.3 Versenyelőny szakdolgozati értelemben",
        "A konzulensi visszajelzés alapján a rendszer erőssége a termékszerű MVP, ezért a dolgozatban ezt kell hangsúlyozni.",
        [
            "a webáruház nem csak CRUD-felület, mert a rendelési és adminisztrációs folyamatok összekapcsolódnak",
            "a helyszíni értékesítés és mentett vásárlók kezelése domainközeli kiegészítés",
            "az AI asszisztens modern, de korlátozott és dokumentált funkcióként jelenik meg",
        ],
        "kezdőlap AI asszisztenssel és admin helyszíni vásárlás képernyőkép",
    )

    doc.add_heading("3. Követelmények és traceability", level=1)
    req_topics = [
        (
            "3.1 Funkcionális követelmények",
            "A funkcionális követelmények azt rögzítik, hogy a rendszernek milyen konkrét műveleteket kell támogatnia.",
            [
                "a vásárló termékeket kereshet, szűrhet, kosárba tehet és rendelést adhat le",
                "az admin termékeket tölthet fel, CSV importot indíthat és rendelési állapotot módosíthat",
                "a dolgozó korlátozott jogosultsággal kezelhet helyszíni vásárlást és készletinformációt",
            ],
            "követelmény-use case-modul-teszt traceability táblázat",
        ),
        (
            "3.2 Nem funkcionális követelmények",
            "A nem funkcionális követelmények a használhatóságra, biztonságra, reprodukálhatóságra és karbantarthatóságra vonatkoznak.",
            [
                "a felület legyen reszponzív és sötét/világos módban is használható",
                "a repo legyen tiszta, node_modules és valódi titkok nélkül",
                "a build és tesztelés futtatható legyen helyi gépen és GitHub Actions környezetben",
            ],
            "GitHub Actions zöld CI képernyőkép és README részlet",
        ),
        (
            "3.3 Jogosultsági követelmények",
            "A rendszer három fő szerepkört különböztet meg: vásárló, dolgozó és adminisztrátor.",
            [
                "a vásárló a saját rendeléseit és profiladatait láthatja",
                "a dolgozó csak az operatív adminisztráció egy részéhez fér hozzá",
                "az admin a teljes admin felületet, felhasználókezelést és jóváhagyásokat is kezelheti",
            ],
            "admin felhasználókezelés képernyőkép és Firestore rules jogosultsági kódrészlet",
        ),
        (
            "3.4 Adatvédelmi és biztonsági követelmények",
            "A rendszer személyes adatokat is kezel, ezért külön figyelmet kap a jogosultság és a titkok kezelése.",
            [
                "a valódi API kulcsok nem kerülhetnek verziókezelésbe",
                "a Firestore szabályok korlátozzák, hogy ki milyen collectiont olvashat vagy írhat",
                "a tiltott felhasználó nem folytathat aktív műveletet a rendszerben",
            ],
            "security_minimum dokumentáció és firestore.rules releváns részlete",
        ),
    ]
    for t in req_topics:
        expand_topic(doc, *t)
    add_table(
        doc,
        ["Követelmény", "Use case", "Megvalósított modul", "Bizonyíték"],
        [
            ["Termékkeresés és böngészés", "Vásárló terméket keres", "Products page, product service", "Terméklista képernyőkép, kézi teszt"],
            ["Kosárkezelés", "Vásárló kosarat módosít", "Cart service, cart page", "Kosár teszt, kosár képernyőkép"],
            ["Checkout validáció", "Rendelés leadása", "Checkout page", "Hibás email/telefon képernyőkép, teszt"],
            ["Admin státuszváltás", "Rendelés teljesítése", "Order service, admin page", "Kódrészlet és admin képernyőkép"],
            ["CSV import", "Tömeges termékfeltöltés", "Admin products import", "CSV import képernyőkép"],
            ["AI asszisztens", "Termékkatalógushoz kötött tanácsadás", "Chatbot service, Worker proxy", "AI ablak képernyőkép, kódrészlet"],
        ],
    )

    doc.add_heading("4. Use case-ek és fő folyamatok", level=1)
    use_case_topics = [
        (
            "4.1 Vásárlói rendelési folyamat",
            "A vásárlói rendelési folyamat a termékkereséstől a sikeres rendelésig vezeti végig a felhasználót.",
            [
                "a folyamat több állapotot kezel: üres kosár, több termék, hibás adatok, kupon, sikeres rendelés",
                "a checkout oldalon az összegzés és a szállítási adatok egy felhasználói döntési pontot alkotnak",
                "a rendelés létrejötte után a profilban követhető a rendelés státusza",
            ],
            "kosár, checkout validáció és sikeres rendelés képernyőképek",
        ),
        (
            "4.2 Adminisztrátori rendeléskezelés",
            "Az adminisztrátori folyamat célja a beérkező rendelések áttekintése, állapotváltása és bizonylat kezelése.",
            [
                "az admin rendelési listából dolgozik, ahol státusz, összeg és vevői adatok is megjelennek",
                "a státuszváltások auditálhatóak, mert a rendszer rögzíti a változtatásokat",
                "a PDF bizonylat letöltése a rendelési folyamat dokumentálásának része",
            ],
            "admin rendelések és PDF számla képernyőkép",
        ),
        (
            "4.3 Helyszíni vásárlás és mentett vásárlók",
            "A helyszíni vásárlás domainközeli funkció, mert épületgépészeti kereskedésnél gyakori a személyes vagy telefonos rendelés.",
            [
                "az admin vagy dolgozó mentett vásárlót választhat, amely automatikusan kitölti az adatokat",
                "céges vásárlónál adószám és cégnév is szerepelhet a bizonylaton",
                "a rendszer kezeli a fizetési módot és a fizetési határidőt is",
            ],
            "helyszíni vásárlás mentett vásárlóval képernyőkép",
        ),
        (
            "4.4 Dolgozói szerepkör",
            "A dolgozói szerepkör azt mutatja be, hogy az adminisztráció nem feltétlenül egyszintű jogosultsági modell.",
            [
                "a dolgozó láthat készletet és rögzíthet vásárlást, de nem rendelkezik teljes admin jóváhagyási joggal",
                "a szerepkör csökkenti annak kockázatát, hogy minden belső felhasználó teljes hozzáférést kapjon",
                "a dolgozói működés Firestore szabályokkal és frontend guard logikával is dokumentálható",
            ],
            "dolgozói nézet és jogosultsági kódrészlet",
        ),
        (
            "4.5 AI asszisztens use case",
            "Az AI asszisztens célja nem az, hogy helyettesítse a szakembert, hanem hogy a katalógus és általános épületgépészeti kérdések alapján segítséget adjon.",
            [
                "a felület nem engedi a modell szabad kiválasztását a felhasználónak",
                "a válaszadásnál a rendszer előnyben részesíti a saját termékkatalógus adatait",
                "ha nincs pontos terméktalálat, a válaszban jelezni kell, hogy pontos ajánlatért egyeztetés szükséges",
            ],
            "kezdőlap AI ablak nyitva és chatbot service kódrészlet",
        ),
    ]
    for t in use_case_topics:
        expand_topic(doc, *t)

    doc.add_heading("5. GUI/UX tervezés és képernyők", level=1)
    ux_topics = [
        (
            "5.1 Arculat és dark/light mód",
            "A TDLWebshop vizuális iránya a logóhoz és az épületgépészeti témához igazodik, ezért a dark mód ipari, modern hangulatot kapott.",
            [
                "a sötét mód erősebb kontrasztot és kék-piros brand színeket használ",
                "a világos mód ugyanazt a layoutot tartja meg, csak üzletiesebb színvilággal",
                "a theme switcher nem változtatja meg a funkcionális elrendezést",
            ],
            "kezdőlap dark és light mód képernyőkép",
        ),
        (
            "5.2 Navbar, kategóriák és keresés",
            "A navigációban a keresés és a kategóriák kiemelt szerepet kapnak, mert a termékkatalógusban ezek vezetik a vásárlót.",
            [
                "a kategóriák legördülő menüben jelennek meg",
                "a kereső a terméknév, cikkszám és kategória alapján segíti a találatokat",
                "mobil nézetben a navigáció tömörebb, de a fő funkciók elérhetőek maradnak",
            ],
            "kezdőlap kategória lenyílóval és mobil kezdőlap képernyőkép",
        ),
        (
            "5.3 Terméklista és termékkártyák",
            "A terméklista a webshop egyik legfontosabb képernyője, mert itt dől el, hogy a vásárló gyorsan megtalálja-e a megfelelő terméket.",
            [
                "a termékkártyán kép, név, ár, kategória, készletinformáció és kosár funkció jelenik meg",
                "az akciós és új termékek vizuálisan megkülönböztethetőek",
                "a részletek gomb és a kívánságlista ikon a használhatósági visszajelzések alapján módosult",
            ],
            "terméklista szűréssel és akciós termékekkel",
        ),
        (
            "5.4 Checkout és hibakezelés",
            "A checkout felületnél a legfontosabb UX cél az, hogy a felhasználó egyértelmű visszajelzést kapjon hibás vagy hiányzó adatok esetén.",
            [
                "az email és telefonszám mezők validációja megakadályozza a nyilvánvalóan hibás rendeléseket",
                "a rendelésösszegzés külön blokkban segíti az ellenőrzést",
                "sikeres rendelés után a rendszer visszajelzést ad a felhasználónak",
            ],
            "checkout hibás email/telefon példával és sikeres rendelés képernyőkép",
        ),
        (
            "5.5 Admin felület használhatósága",
            "Az admin felület célja a belső munkafolyamatok gyors támogatása, ezért a sűrűbb adatmegjelenítés és a státuszok áttekinthetősége fontosabb, mint a marketing jellegű látvány.",
            [
                "a rendelési lista, termékkezelés és készletfigyelés külön funkcionális egységekként jelennek meg",
                "a helyszíni vásárlásnál mentett vásárló keresése és automatikus kitöltése gyorsítja a munkát",
                "az admin és dolgozó szerepkör eltérő műveleteket lát",
            ],
            "admin áttekintés, CSV import, készletfigyelés és helyszíni vásárlás képernyőkép",
        ),
    ]
    for t in ux_topics:
        expand_topic(doc, *t)

    add_figure_placeholder(doc, "Kezdőlap dark mode kategória menüvel", "Nyisd meg a főoldalt dark módban, a Kategóriák menü legyen lenyitva.")
    add_figure_placeholder(doc, "AI asszisztens a főoldalon", "Nyisd meg a főoldalt, kattints az AI segítő gombra, és legyen látható egy termékkérdésre adott válasz.")
    add_figure_placeholder(doc, "Checkout validáció", "Checkout oldalon írj be hibás emailt vagy telefonszámot, majd készíts képernyőképet a hibaüzenetről.")
    add_figure_placeholder(doc, "Admin CSV import", "Admin felület Termékek/CSV import részén készíts képernyőképet valid importtal.")

    doc.add_heading("6. Technológiai háttér és döntések", level=1)
    tech_topics = [
        (
            "6.1 Angular frontend",
            "Az Angular használata azért indokolt, mert komponensalapú felépítést, routingot, service réteget és jól tesztelhető TypeScript kódot biztosít.",
            [
                "a komponensek külön kezelik a vásárlói és adminisztrátori képernyőket",
                "a service-ek felelnek az adatelérésért és üzleti logikáért",
                "a TypeScript típusossága segíti a nagyobb kódbázis karbantarthatóságát",
            ],
            "komponens és service struktúrát bemutató ábra",
        ),
        (
            "6.2 Firebase és Firestore",
            "A Firebase választása a szakdolgozati MVP szempontjából azért kedvező, mert gyorsan biztosít hitelesítést, adatbázist és hostingot.",
            [
                "a Firestore dokumentum alapú modellje illeszkedik a product, order és user profile entitásokhoz",
                "a Firestore rules segítségével szabályszinten is dokumentálható a jogosultságkezelés",
                "a Firebase Hosting egyszerű publikálási lehetőséget ad a frontendhez",
            ],
            "adatmodell ábra és Firestore rules kódrészlet",
        ),
        (
            "6.3 Cloudflare Worker és OpenRouter proxy",
            "Az AI asszisztensnél a kliensoldali API kulcs használata nem lenne biztonságos, ezért a hívás egy Worker proxyn keresztül történik.",
            [
                "az OpenRouter kulcs nem kerül a frontend kódba",
                "a Worker CORS és payload ellenőrzést is végezhet",
                "a megoldás MVP-ben megfelelő, de éles környezetben rate limit és kvóta kezelés is szükséges",
            ],
            "workers/openrouter-proxy/src/index.js kódrészlet és AI működési ábra",
        ),
        (
            "6.4 GitHub Actions és reprodukálhatóság",
            "A CI célja, hogy a projekt ne csak lokálisan, hanem tiszta környezetben is buildelhető és tesztelhető legyen.",
            [
                "a CI futtatja a telepítést, buildet és teszteket",
                "a zöld GitHub Actions futás bizonyítékként szerepelhet a dolgozatban",
                "a README és .env.example segíti a reprodukálható indítást",
            ],
            "GitHub Actions zöld CI képernyőkép",
        ),
    ]
    for t in tech_topics:
        expand_topic(doc, *t)

    doc.add_page_break()
    doc.add_heading("Az első rész összegzése", level=1)
    para(
        doc,
        "Az első rész a TDLWebshop szakdolgozat problémafelvetési, piaci, követelményi, use case, UX és technológiai "
        "alapját foglalja össze. A végleges dolgozatban ezt a részt célszerű képernyőképekkel, traceability táblával, "
        "piaci összehasonlítással és saját megfogalmazású indoklással kiegészíteni. A második részben a konkrét "
        "architektúra, adatmodell, megvalósítás, biztonság, tesztelés, reprodukálhatóság és AI-használat kerül részletesen bemutatásra."
    )
    return doc, sum(len(p.text) for p in doc.paragraphs)


def add_part2() -> tuple[Document, int]:
    doc = setup_document()
    add_cover(doc, "II. rész: megvalósítás, tesztelés, biztonság és lezárás", "Munkapéldány a végleges szakdolgozati szöveghez")
    add_toc_note(doc)

    doc.add_heading("7. Architektúra és adatáramlás", level=1)
    arch_topics = [
        (
            "7.1 Rendszerarchitektúra áttekintése",
            "A TDLWebshop architektúrája frontend, Firebase szolgáltatások és külső AI proxy együttműködésére épül.",
            [
                "az Angular alkalmazás kezeli a felhasználói felületet és a kliensoldali állapotokat",
                "a Firebase Auth a bejelentkezést és felhasználói azonosítást biztosítja",
                "a Firestore tárolja a termékeket, rendeléseket, profilokat, kuponokat és audit bejegyzéseket",
                "a Cloudflare Worker az AI asszisztens szerveroldali kulcskezelését végzi",
            ],
            "komponensarchitektúra ábra és adatáramlási ábra",
        ),
        (
            "7.2 Frontend komponensek és service réteg",
            "A frontend felépítésében különválik a megjelenítés és az adatelérés, ami karbantarthatóbbá teszi a kódot.",
            [
                "a pages mappában a fő képernyők találhatóak, például terméklista, checkout, profil és admin",
                "az app/services mappában a rendelés, kosár, autentikáció, számla és AI asszisztens logika szerepel",
                "a komponensek a service-ekből kapott adatokat jelenítik meg, és eseményeken keresztül indítanak műveleteket",
            ],
            "projektstruktúra képernyőkép vagy komponenslista",
        ),
        (
            "7.3 Firebase adatáramlás",
            "A Firestore adatáramlás a felhasználói műveletekhez kötődik: rendelésleadás, állapotváltás, CSV import és profilkezelés.",
            [
                "a rendelésleadás új order dokumentumot hoz létre",
                "az admin státuszváltás audit bejegyzéssel és készletmozgással kapcsolódik össze",
                "a termékfeltöltés és CSV import validáció után módosítja a products collectiont",
            ],
            "szekvenciaábra checkout és admin státuszváltás folyamatokra",
        ),
    ]
    for t in arch_topics:
        expand_topic(doc, *t)
    add_figure_placeholder(doc, "Komponensarchitektúra", "Mermaid vagy PlantUML ábra: Angular frontend, Firebase Auth, Firestore, Hosting, Cloudflare Worker, OpenRouter.")
    add_figure_placeholder(doc, "Checkout szekvenciaábra", "Vásárló -> Checkout komponens -> OrderService -> Firestore -> Profil/rendeléskövetés.")

    doc.add_heading("8. Adatmodell", level=1)
    data_topics = [
        (
            "8.1 Fő entitások",
            "Az adatmodell fő célja, hogy a webshop és admin funkciók ugyanarra az adatszerkezetre épüljenek.",
            [
                "a Product entitás tartalmazza a termék nevét, árát, készletét, kategóriáját, képeit és megjelenítési állapotát",
                "az Order és OrderItem entitások a rendelés fejléc és tételadatokat választják szét",
                "a UserProfile és SavedCustomer entitások a regisztrált és mentett vásárlói adatokat kezelik",
                "a Coupon, Invoice és audit adatok külön felelősséget kapnak",
            ],
            "adatmodell diagram és entitás táblázat",
        ),
        (
            "8.2 Termék és készlet modell",
            "A termékmodellben a készlet nem csak megjelenítési adat, hanem adminisztrációs és rendelési folyamatokhoz is kapcsolódik.",
            [
                "a készletfigyelés segíti az adminisztrátort az alacsony készlet azonosításában",
                "a helyszíni vásárlásnál a készletmozgás tranzakciósabb logikával jelenik meg",
                "a CSV import lehetővé teszi több termék gyors feltöltését és frissítését",
            ],
            "admin készletfigyelés és CSV import képernyőkép",
        ),
        (
            "8.3 Rendelés és audit modell",
            "A rendelésmodell a vásárlói adatok, tételek, árak, státusz és audit információk összekapcsolására szolgál.",
            [
                "a rendelés tételei külön listában szerepelnek, így a többtermékes rendelés is kezelhető",
                "az állapotváltás auditálása visszakövethetővé teszi az admin műveleteket",
                "a PDF számla vagy bizonylat a rendelési adatokból generálódik",
            ],
            "rendelés részletei és PDF bizonylat képernyőkép",
        ),
        (
            "8.4 AI asszisztens kontextus",
            "Az AI asszisztens adatmodell szempontból nem önálló üzleti entitás, hanem a termékkatalógus és a felhasználói kérdés alapján dolgozó kiegészítő szolgáltatás.",
            [
                "a kérdéshez csak szükséges termékadatok kerülnek a promptba",
                "a válasz nem ír közvetlenül adatbázist",
                "a szakdolgozatban fontos rögzíteni, hogy az AI válasz tájékoztató jellegű",
            ],
            "AI működési kódrészlet és felületi képernyőkép",
        ),
    ]
    for t in data_topics:
        expand_topic(doc, *t)
    add_table(
        doc,
        ["Entitás", "Fő mezők", "Kapcsolat", "Szakdolgozati szerep"],
        [
            ["Product", "name, sku, category, price, stock, images", "OrderItem hivatkozik rá", "Termékkatalógus és készlet alapja"],
            ["Order", "customer, items, total, status, payment", "UserProfile/SavedCustomer adatokkal kapcsolódik", "Rendelési folyamat központja"],
            ["OrderItem", "productId, name, quantity, price", "Order része", "Tételes rendelés és PDF alapja"],
            ["UserProfile", "uid, email, role, disabled", "Auth felhasználóhoz kapcsolódik", "Jogosultság és profilkezelés"],
            ["SavedCustomer", "name, email, phone, company, taxNumber", "Helyszíni vásárlásnál használható", "B2B és visszatérő vevők támogatása"],
            ["Coupon", "code, type, value, active", "Checkout használja", "Kedvezmény és visszaélési pont"],
            ["Invoice", "invoiceNumber, dates, totals", "Orderből készül", "PDF bizonylat/számla dokumentálása"],
            ["Audit", "orderId, oldStatus, newStatus, actor", "Order státuszváltáshoz kötött", "Visszakövethetőség"],
        ],
    )

    doc.add_heading("9. Megvalósítás", level=1)
    impl_topics = [
        (
            "9.1 Checkout és rendelésmentés",
            "A checkout a vásárlói oldal legkritikusabb folyamata, mert itt válik a böngészés üzleti eseménnyé.",
            [
                "a checkout validálja a megadott adatokat és összegzi a rendelés tételeit",
                "a rendelés mentése után a felhasználó visszajelzést kap",
                "a folyamatnál külön kell dokumentálni, hogy az MVP milyen szerveroldali és kliensoldali ellenőrzéseket tartalmaz",
            ],
            "checkout.ts rendelés véglegesítési kódrészlet és checkout képernyőkép",
        ),
        (
            "9.2 Státuszváltás, audit és készlet",
            "Az admin rendeléskezelésnél a státuszváltás nem önmagában álló mezőmódosítás, hanem üzleti jelentésű művelet.",
            [
                "a státusz módosítása audit bejegyzéssel együtt értelmezhető",
                "a készletváltozásnál fontos, hogy ne lehessen észrevétlenül inkonzisztens állapotot létrehozni",
                "a szakdolgozatban a helyszíni rendelés tranzakciósabb kezelése külön bemutatható",
            ],
            "OrderService státusz/audit/készlet kódrészlet",
        ),
        (
            "9.3 PDF számla és bizonylat generálása",
            "A PDF generálás azért lényeges, mert a rendelés adataiból letölthető dokumentum készül.",
            [
                "a PDF tartalmazza a kiállító, vevő, rendelés, tételek és végösszeg adatait",
                "a korábbi elcsúszások javítása után a végösszeg külön, olvasható blokkban jelenik meg",
                "a végleges dolgozatban érdemes egy generált PDF-et ábraként szerepeltetni",
            ],
            "invoice.service.ts kódrészlet és generált PDF képernyőkép",
        ),
        (
            "9.4 CSV import és termékkezelés",
            "A CSV import a tömeges termékfeltöltés problémájára ad MVP szintű megoldást.",
            [
                "az import validálja a sorokat, és jelzi a hibás adatokat",
                "a termékek SKU vagy egyedi azonosító alapján frissíthetőek",
                "a több képes termékeknél a képek kezelését külön szabály szerint kell kitölteni",
            ],
            "admin.ts CSV import validáció és mentés kódrészlet",
        ),
        (
            "9.5 Jogosultságkezelés",
            "A jogosultságkezelés két szinten jelenik meg: a felületen és az adatbázis szabályaiban.",
            [
                "a frontend guard megakadályozza, hogy a felhasználó számára nem releváns oldalra jusson",
                "a Firestore rules a kliens megkerülése esetén is korlátozza a hozzáférést",
                "a disabled állapot megakadályozza, hogy tiltott felhasználó aktív műveletet végezzen",
            ],
            "firestore.rules és admin.ts jogosultsági kódrészletek",
        ),
        (
            "9.6 AI asszisztens megvalósítása",
            "Az AI asszisztens kiegészítő funkcióként működik, amely a saját termékkatalógus és épületgépészeti kérdések alapján válaszol.",
            [
                "a frontend nem enged modellválasztást a felhasználónak",
                "a Worker proxy szerveroldalon kezeli az OpenRouter API kulcsot",
                "ha nincs pontos terméktalálat, a válaszban nem szabad véletlenszerű terméket ajánlani",
            ],
            "chatbot-llm.service.ts és workers/openrouter-proxy/src/index.js kódrészlet",
        ),
    ]
    for t in impl_topics:
        expand_topic(doc, *t)

    code_refs = [
        ("Rendelés véglegesítése", "src/pages/checkout/checkout.ts 367. sortól", "A checkout folyamat lezárása és rendelés mentése."),
        ("Státusz, audit és készlet", "src/app/services/order.service.ts 41-127. sor", "Admin státuszváltás és készletmozgás összekapcsolása."),
        ("Helyszíni rendelés tranzakció", "src/app/services/order.service.ts 222-267. sor", "Helyszíni vásárlás mentése és készletkezelése."),
        ("Számlaszám generálás", "src/app/services/order.service.ts 269-302. sor", "PDF bizonylathoz/számlához kapcsolódó sorszámozás."),
        ("PDF számla felépítése", "src/app/services/invoice.service.ts 9-154. sor", "A letölthető dokumentum szerkezete."),
        ("Aktív felhasználó és szerepkörök", "firestore.rules 25-76. sor", "Admin, dolgozó és aktív felhasználói ellenőrzések."),
        ("Fő collection szabályok", "firestore.rules 294-361. sor", "Products, orders, users, savedCustomers és audit szabályok."),
        ("Admin/dolgozói jogosultságok", "src/pages/admin/admin.ts 607-748. sor", "Felületi jogosultsági döntések."),
        ("CSV import validáció", "src/pages/admin/admin.ts 1181-1257. sor", "Tömeges termékfeltöltés validációja."),
        ("AI asszisztens logika", "src/app/services/chatbot-llm.service.ts 26-84. és 217-236. sor", "Domain- és katalóguslogika."),
        ("OpenRouter proxy", "workers/openrouter-proxy/src/index.js", "Szerveroldali kulcskezelés és AI API hívás."),
    ]
    for title, ref, exp in code_refs:
        add_code_placeholder(doc, title, ref, exp)

    doc.add_heading("10. Biztonság és adatvédelem", level=1)
    security_topics = [
        (
            "10.1 Hitelesítés és szerepkörök",
            "A rendszer biztonságának alapja, hogy a felhasználók szerepköre szerint eltérő műveletek érhetőek el.",
            [
                "az admin teljes körű rendszerkezelést végezhet",
                "a dolgozó szűkebb jogosultságot kap",
                "a vásárló csak a saját profiljához és rendeléseihez fér hozzá",
            ],
            "jogosultsági táblázat és Firestore rules részlet",
        ),
        (
            "10.2 Titkok és API kulcsok kezelése",
            "A konzulensi és intézményi elvárások alapján valódi jelszó, token vagy API kulcs nem kerülhet nyilvános repóba.",
            [
                "az OpenRouter kulcs a Worker secretként kezelendő",
                "a .env.example csak a szükséges változók nevét tartalmazza",
                "a korábban használt kulcsot vissza kell vonni, ha valaha fájlba került",
            ],
            ".env.example és Worker konfiguráció rövid bemutatása",
        ),
        (
            "10.3 Firestore szabályok",
            "A Firestore rules nem csak technikai részlet, hanem a szakdolgozat biztonsági bizonyításának egyik fontos eleme.",
            [
                "nincs általános allow read, write: if true szabály",
                "a szabályok ellenőrzik az aktív felhasználót és szerepkört",
                "az audit log módosítása és törlése korlátozott",
            ],
            "firestore.rules kódrészlet és magyarázat",
        ),
        (
            "10.4 MVP korlátok és kockázatok",
            "A dolgozatban őszintén szerepeltetni kell azokat a pontokat, amelyek éles rendszerben további erősítést igényelnének.",
            [
                "a checkout összegzés egyes adatai kliensoldalról indulnak, ezért éles rendszerben szerveroldali újraszámolás szükséges",
                "az AI proxyhoz rate limit vagy kvóta javasolt",
                "a vendég rendelés email alapú azonosítása adatvédelmi és integritási szempontból körültekintést igényel",
            ],
            "biztonsági minimum táblázat",
        ),
    ]
    for t in security_topics:
        expand_topic(doc, *t)

    doc.add_heading("11. Tesztelés és validáció", level=1)
    test_topics = [
        (
            "11.1 Automata tesztek és build",
            "Az automata tesztek célja a kritikus üzleti logikák ellenőrzése és a regressziók csökkentése.",
            [
                "a build futtatása ellenőrzi, hogy a TypeScript és Angular projekt fordítható",
                "a tesztek lefedik a kosár, checkout, kupon, admin státusz, PDF és AI asszisztens fontos részeit",
                "a GitHub Actions zöld futása a reprodukálhatóság egyik bizonyítéka",
            ],
            "npm run build, npm test és GitHub Actions képernyőkép",
        ),
        (
            "11.2 Kézi tesztjegyzőkönyv",
            "A kézi tesztelés azért szükséges, mert a felhasználói felület teljes működését nem minden esetben fedi automata unit teszt.",
            [
                "a regisztráció, bejelentkezés és tiltott felhasználó ellenőrzendő",
                "a termékkeresés, kosár, checkout és profil folyamatokat végig kell próbálni",
                "az admin és dolgozói jogosultságokat külön fiókokkal kell ellenőrizni",
            ],
            "kézi teszt checklist kitöltött állapotban",
        ),
        (
            "11.3 Hibás bemenetek és negatív tesztek",
            "A hibás bemenetek kezelése különösen fontos webshopnál, mert a hibás rendelési adatok üzleti problémát okozhatnak.",
            [
                "hibás email és telefonszám esetén a felület ne engedjen tovább",
                "érvénytelen kupon esetén egyértelmű hibaüzenet jelenjen meg",
                "jogosulatlan felhasználó ne férjen admin művelethez",
            ],
            "checkout validáció és admin jogosultsági negatív teszt képernyőkép",
        ),
    ]
    for t in test_topics:
        expand_topic(doc, *t)
    add_table(
        doc,
        ["Tesztelt folyamat", "Teszt típusa", "Elvárt eredmény", "Bizonyíték"],
        [
            ["Regisztráció/bejelentkezés", "Kézi", "Felhasználó belép, tiltott felhasználó hibaüzenetet kap", "Profil képernyőkép"],
            ["Termékkeresés", "Kézi", "Találatok megjelennek", "Terméklista képernyőkép"],
            ["Kosár módosítás", "Automata + kézi", "Összeg frissül", "Teszt és kosár képernyőkép"],
            ["Checkout validáció", "Automata + kézi", "Hibás adatnál nincs rendelés", "Checkout hiba képernyőkép"],
            ["Admin státuszváltás", "Automata + kézi", "Audit és státusz frissül", "Admin képernyőkép"],
            ["CSV import", "Kézi", "Valid sorok menthetőek, hibák listázódnak", "CSV import képernyőkép"],
            ["AI asszisztens", "Kézi + service teszt", "Domain kérdésre válaszol, nem releváns kérdést elutasít", "AI ablak képernyőkép"],
        ],
    )

    doc.add_heading("12. Reprodukálhatóság és üzembe helyezés", level=1)
    repro_topics = [
        (
            "12.1 Repo-higiénia",
            "A repo-higiénia célja, hogy a bíráló tiszta, átlátható kódbázist lásson.",
            [
                "a node_modules és build mappák nem részei a verziózott forrásnak",
                "a segédanyagok elkülönülnek a beadandó dokumentációtól",
                "a README tartalmazza a futtatási és tesztelési lépéseket",
            ],
            "GitHub repo fájlstruktúra és README képernyőkép",
        ),
        (
            "12.2 Lokális futtatás",
            "A lokális futtatási útmutató a reprodukálhatóság alapja.",
            [
                "a függőségek telepítése npm install paranccsal történik",
                "a fejlesztői szerver npm run start paranccsal indítható",
                "a tesztek és build külön parancsokkal ellenőrizhetőek",
            ],
            "README futtatási részlet",
        ),
        (
            "12.3 Firebase és Worker deploy",
            "Az éles vagy demonstrációs publikálás Firebase Hosting és Cloudflare Worker segítségével történhet.",
            [
                "a Firebase hosting a frontend publikálását biztosítja",
                "a Firestore rules deploy külön ellenőrizhető lépés",
                "a Worker deploy az AI asszisztens proxyját frissíti",
            ],
            "deploy parancsok és GitHub/Firebase képernyőkép",
        ),
    ]
    for t in repro_topics:
        expand_topic(doc, *t)

    doc.add_heading("13. Mesterséges intelligencia használata a fejlesztés során", level=1)
    ai_paragraphs = [
        "A szakdolgozat készítése során mesterséges intelligencia alapú eszközöket fejlesztéstámogató módon használtam. "
        "A használat célja nem az volt, hogy a rendszer tervezési és fejlesztési felelősségét átadjam egy külső eszköznek, "
        "hanem az, hogy gyorsabban tudjak hibákat keresni, alternatív megoldásokat összehasonlítani, dokumentációs szempontokat "
        "összegyűjteni és a már elkészült kódot ellenőrizni. A döntéseket, a végleges működést és a leadott anyagért való felelősséget "
        "minden esetben én vállalom.",
        "Az AI-eszközöket elsősorban ötletelésre, kódreview jellegű ellenőrzésre, hibakeresésre, szövegezési vázlatok készítésére "
        "és dokumentációs ellenőrzőlisták összeállítására használtam. Például a Firestore szabályoknál, a jogosultsági modelleknél, "
        "az admin felület működésénél és a tesztelési bizonyítékoknál segítséget adtak abban, hogy milyen kockázatokat érdemes "
        "külön megvizsgálni. Ezeket a javaslatokat nem automatikusan fogadtam el, hanem a projekt kódja, a futtatott tesztek és a "
        "konzulensi elvárások alapján ellenőriztem.",
        "A fejlesztés közben előfordult, hogy egy AI által javasolt megoldás nem volt közvetlenül megfelelő. Ilyen esetben a javaslatot "
        "kiindulópontként kezeltem, majd a projekt meglévő szerkezetéhez igazítottam. Különösen fontos volt az API kulcsok kezelése, "
        "mert a kliensoldali beégetés biztonsági kockázatot jelentett volna. Emiatt az AI asszisztens OpenRouter hívása végül nem "
        "közvetlenül a böngészőből, hanem Cloudflare Worker proxyn keresztül valósult meg.",
        "A dolgozatban külön kell választani a fejlesztés közbeni AI-használatot és a rendszerben működő AI asszisztenst. Az előbbi "
        "fejlesztői segédeszköz volt, az utóbbi pedig a TDLWebshop egyik funkciója. A beépített AI asszisztens a saját termékkatalógus "
        "és épületgépészeti kérdések alapján ad tájékoztató jellegű választ, de nem helyettesít szakembert, és nem tekinthető hivatalos "
        "műszaki ajánlatnak.",
        "Szándékosan nem használtam AI-t olyan módon, hogy ellenőrzés nélkül kerüljön be végleges kód vagy szakmai állítás. A kritikus "
        "részeknél, például jogosultságkezelésnél, Firestore szabályoknál, rendelésmentésnél és titokkezelésnél különösen fontos volt "
        "a saját ellenőrzés. A szakdolgozat végleges szövegét saját megfogalmazásra kell igazítani, hogy a dolgozat a saját munkámat, "
        "gondolkodásomat és döntéseimet tükrözze.",
        "A folyamatból azt tanultam, hogy az AI hatékonyan gyorsíthatja a fejlesztést és a dokumentáció előkészítését, de csak akkor, "
        "ha a fejlesztő képes ellenőrizni a javaslatokat. Az AI nem mentesít a tesztelés, a biztonsági gondolkodás és a saját döntések "
        "megindoklása alól. A következő projektben még tudatosabban vezetnék naplót arról, hogy melyik döntésnél milyen segítséget "
        "használtam, és hogyan validáltam a kapott eredményt.",
    ]
    for p in ai_paragraphs:
        para(doc, p)

    doc.add_heading("14. Értékelés, korlátok és továbbfejlesztési lehetőségek", level=1)
    eval_topics = [
        (
            "14.1 Az elkészült rendszer értékelése",
            "A TDLWebshop szakdolgozati szempontból azért tekinthető erős MVP-nek, mert több összefüggő felhasználói és admin folyamatot tartalmaz.",
            [
                "a vásárlói út és admin út egyaránt működőképes",
                "a rendszerben láthatóak a domainhez illeszkedő funkciók, például helyszíni vásárlás és mentett vásárlók",
                "a dokumentáció, tesztek és CI bizonyítják, hogy a projekt nem csak lokális próbálkozás",
            ],
            "összefoglaló képernyőkép-válogatás és CI eredmény",
        ),
        (
            "14.2 Korlátok",
            "A dolgozatban a korlátokat nem gyengeségként, hanem mérnöki őszinteségként kell kezelni.",
            [
                "éles fizetési szolgáltatói integráció nem része az MVP-nek",
                "a checkout szerveroldali újraszámítása továbbfejlesztési irány",
                "az AI asszisztenshez éles környezetben rate limit, naplózás és költségkontroll szükséges",
            ],
            "MVP-korlátok táblázat",
        ),
        (
            "14.3 Továbbfejlesztési lehetőségek",
            "A továbbfejlesztési irányok azt mutatják meg, hogy a projekt hogyan válhatna élesebb üzleti rendszerré.",
            [
                "szerveroldali rendelés- és árvalidáció bevezetése",
                "komolyabb raktárkezelési és beszerzési modul",
                "fizetési szolgáltatói integráció",
                "admin riportok és üzleti statisztikák",
                "AI asszisztens naplózása, rate limitje és pontosabb katalóguskeresése",
            ],
            "továbbfejlesztési roadmap táblázat",
        ),
    ]
    for t in eval_topics:
        expand_topic(doc, *t)

    doc.add_heading("15. Összefoglalás és személyes reflexió", level=1)
    summary_texts = [
        "A szakdolgozatban bemutatott TDLWebshop egy olyan épületgépészeti webáruház és adminisztrációs rendszer, amely "
        "a vásárlói és belső adminisztrációs folyamatokat egy rendszerben kezeli. A fejlesztés során a legfontosabb cél "
        "az volt, hogy a rendszer ne csak látványos felületként, hanem működő MVP-ként legyen bemutatható.",
        "A munka során külön figyelmet kaptak a rendelési folyamatok, a kosár és checkout működése, a termék- és készletkezelés, "
        "a jogosultsági szerepkörök, a PDF bizonylat, a CSV import, a mentett vásárlók és az AI asszisztens. Ezek együtt olyan "
        "funkcionalitási volument adnak, amely szakdolgozati szinten már nem puszta gyakorlófeladat, hanem termékszerű rendszer.",
        "A projekt egyik legfontosabb tanulsága az volt, hogy egy webáruház fejlesztése nem csak képernyők elkészítéséből áll. "
        "A használható rendszerhez adatmodell, jogosultságkezelés, hibakezelés, tesztelés, reprodukálható futtatás és biztonsági "
        "gondolkodás is szükséges. A dolgozat véglegesítésénél ezért a kód mellett a bizonyítékok: képernyőképek, ábrák, "
        "teszteredmények és saját magyarázatok is ugyanolyan fontosak.",
        "Személyes reflexióként ide a végleges dolgozatban saját szöveget kell írni arról, hogy a fejlesztés során melyik rész "
        "volt a legnehezebb, melyik döntés bizonyult hasznosnak, mit tanultál az Angular, Firebase, jogosultságkezelés, tesztelés "
        "és AI integráció kapcsán, valamint mit csinálnál másként egy következő projektben.",
    ]
    for p in summary_texts:
        para(doc, p)

    doc.add_heading("Irodalomjegyzék és hivatkozások helye", level=1)
    para(
        doc,
        "A végleges dolgozatba ide kerüljenek a hivatkozások: Angular dokumentáció, Firebase dokumentáció, Cloudflare Workers "
        "dokumentáció, OpenRouter dokumentáció, GitHub Actions dokumentáció, valamint a piaci összehasonlításban szereplő "
        "webshopok vagy rendszerek weboldalai. A hivatkozásokat az intézményi formai elvárások szerint kell egységesíteni."
    )

    doc.add_heading("Mellékletek helye", level=1)
    bullet(doc, "Kézi tesztjegyzőkönyv.")
    bullet(doc, "GitHub Actions zöld CI képernyőkép.")
    bullet(doc, "Fontosabb Firestore rules részletek.")
    bullet(doc, "Adatmodell és architektúra diagramok.")
    bullet(doc, "Generált PDF bizonylat mintája.")
    bullet(doc, "AI-használati nyilatkozat és rövid napló.")

    return doc, sum(len(p.text) for p in doc.paragraphs)


def save_doc(doc: Document, path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)


def main() -> None:
    part1_doc, part1_chars = add_part1()
    part2_doc, part2_chars = add_part2()
    save_doc(part1_doc, PART1)
    save_doc(part2_doc, PART2)

    guide = OUT_DIR / "00_olvasd_el_hasznalat_elott.md"
    guide.write_text(
        "# TDLWebshop szakdolgozati munkapéldányok\n\n"
        "Ez a mappa helyi segédanyag. Ne commitold GitHubra, mert munkapéldány és átírási alap.\n\n"
        "Fájlok:\n"
        f"- `{PART1.name}`: elméleti, piaci, követelményi, use case, GUI/UX és technológiai rész.\n"
        f"- `{PART2.name}`: architektúra, adatmodell, megvalósítás, biztonság, tesztelés, AI-használat és összefoglalás.\n\n"
        "Használat:\n"
        "1. A két dokumentum szövegét egy végleges Word fájlba kell összefésülni.\n"
        "2. A [ÁBRA HELYE] és [KÓDRÉSZLET HELYE] jelöléseket tényleges képernyőképre, diagramra vagy rövid kódrészletre kell cserélni.\n"
        "3. A végleges beadás előtt a szöveget saját nyelvezetre kell húzni.\n"
        "4. A címlapot, tartalomjegyzéket, feladatkiírást, irodalomjegyzéket és oldalszámokat Wordben kell véglegesíteni.\n",
        encoding="utf-8",
    )
    print(f"PART1={PART1}")
    print(f"PART2={PART2}")
    print(f"PART1_CHARS={part1_chars}")
    print(f"PART2_CHARS={part2_chars}")
    print(f"TOTAL_CHARS={part1_chars + part2_chars}")


if __name__ == "__main__":
    main()
