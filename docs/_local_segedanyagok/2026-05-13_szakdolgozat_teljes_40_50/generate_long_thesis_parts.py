# -*- coding: utf-8 -*-
from pathlib import Path
from datetime import date

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


OUT_DIR = Path(__file__).resolve().parent
PART1 = OUT_DIR / "TDLWebshop_szakdolgozat_1_resz_elmeleti_tervezesi_alap.docx"
PART2 = OUT_DIR / "TDLWebshop_szakdolgozat_2_resz_megvalositas_teszteles_zaras.docx"
REPORT = OUT_DIR / "00_hasznalati_utmutato_es_oldalszam_becsles.md"


BLUE = "1F4E79"
LIGHT_BLUE = "D9EAF7"
LIGHT_GRAY = "F2F4F7"
TEXT = "111827"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    for par in cell.paragraphs:
        par.paragraph_format.space_after = Pt(2)
        par.paragraph_format.line_spacing = 1.05


def style_table(table):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(2)
                for run in p.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(10)


def setup_document(doc):
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    styles = doc.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"].font.size = Pt(12)
    styles["Normal"].paragraph_format.line_spacing = 1.15
    styles["Normal"].paragraph_format.space_after = Pt(6)

    for name, size, before, after, color in [
        ("Title", 22, 0, 18, BLUE),
        ("Heading 1", 16, 18, 8, BLUE),
        ("Heading 2", 14, 12, 6, BLUE),
        ("Heading 3", 12, 10, 4, TEXT),
    ]:
        st = styles[name]
        st.font.name = "Times New Roman"
        st.font.size = Pt(size)
        st.font.color.rgb = RGBColor.from_string(color)
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)


def add_cover(doc, part_title, subtitle):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Szegedi Tudományegyetem\nInformatikai Intézet")
    r.bold = True
    r.font.size = Pt(14)

    for _ in range(5):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("SZAKDOLGOZAT")
    r.bold = True
    r.font.size = Pt(24)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("TDLWebshop épületgépészeti webáruház és adminisztrációs rendszer")
    r.bold = True
    r.font.size = Pt(16)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(part_title)
    r.italic = True
    r.font.size = Pt(13)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(subtitle)

    for _ in range(8):
        doc.add_paragraph()

    table = doc.add_table(rows=2, cols=2)
    style_table(table)
    table.cell(0, 0).text = "Készítette:"
    table.cell(0, 1).text = "Tóth Dávid László"
    table.cell(1, 0).text = "Témavezető:"
    table.cell(1, 1).text = "Dr. Bilicki Vilmos, egyetemi docens"
    for row in table.rows:
        for cell in row.cells:
            for par in cell.paragraphs:
                for run in par.runs:
                    run.font.size = Pt(11)
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(f"Szeged, {date.today().year}")
    doc.add_page_break()


def add_note(doc, title, text):
    table = doc.add_table(rows=1, cols=1)
    style_table(table)
    cell = table.cell(0, 0)
    set_cell_shading(cell, LIGHT_BLUE)
    set_cell_text(cell, f"{title}\n{text}", bold=False)


def add_para(doc, text):
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    style_table(table)
    for i, h in enumerate(headers):
        set_cell_shading(table.cell(0, i), BLUE)
        set_cell_text(table.cell(0, i), h, bold=True, color="FFFFFF")
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], str(value))
            if len(str(value)) < 12:
                cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    return table


def fig(doc, caption, instruction):
    add_note(
        doc,
        caption,
        "Ide kerüljön tényleges képernyőkép vagy ábra. " + instruction,
    )


def code_ref(doc, caption, file_and_lines, why):
    add_note(
        doc,
        caption,
        f"Kódrészlet helye: {file_and_lines}. Rövid magyarázat: {why}",
    )


def discussion(doc, theme, system_part, evidence, risk, count=4):
    variants = [
        (
            f"A {theme} a TDLWebshop esetében nem önálló, elszigetelt funkcióként jelenik meg, "
            f"hanem a teljes vásárlási és adminisztrációs folyamat részeként. A {system_part} "
            f"azért kap hangsúlyt, mert az épületgépészeti termékeknél a vásárlói döntés gyakran "
            f"nem egyetlen adat, hanem ár, készlet, kategória, műszaki jellemző és szállítási mód "
            f"együttes értelmezésén múlik. A megoldásban ezért az volt a cél, hogy a felület ne csak "
            f"adatokat jelenítsen meg, hanem a felhasználói folyamatot is támogassa."
        ),
        (
            f"A megvalósítás során a {theme} tervezésénél külön szempont volt, hogy a rendszer "
            f"szakdolgozati MVP-ként is bizonyítható legyen. Ez azt jelenti, hogy nem minden ipari "
            f"webshop-funkció készült el teljes mélységben, viszont a {evidence} már alkalmas arra, "
            f"hogy bemutassa a választott architektúra és adatmodell működését. A dolgozatban ezt "
            f"nem önmagában, hanem a kapcsolódó use case-ekkel és tesztelési bizonyítékokkal együtt "
            f"érdemes értelmezni."
        ),
        (
            f"A {system_part} kialakítása azért is lényeges, mert a TDLWebshop kétféle működési "
            f"környezetet fed le: a webes vásárlói oldalt és az adminisztrációs felületet. A webes "
            f"oldalon a hangsúly a gyors keresésen, átlátható termékmegjelenítésen és rendelésleadáson "
            f"van, míg az admin oldalon a készlet, a rendelések és a jogosultságok kezelése a fontosabb. "
            f"Ez a kettősség adja a projekt mérnöki értékének jelentős részét."
        ),
        (
            f"A {theme} kapcsán tudatos korlátokat is meg kellett húzni. A jelenlegi rendszer "
            f"nem teljes körű vállalatirányítási rendszer és nem könyvelőprogram, ezért a {risk} "
            f"olyan terület, amelyet vagy ellenőrzött MVP-korlátként kell bemutatni, vagy későbbi "
            f"fejlesztési irányként érdemes megfogalmazni. Ez nem gyengíti a projektet, ha a dolgozat "
            f"egyértelműen leírja, hogy mi készült el, miért abban a mélységben, és hogyan ellenőrizhető."
        ),
        (
            f"A konzulensi visszajelzés alapján a hangsúly már nem új funkciók hozzáadásán, hanem "
            f"a meglévő rendszer bizonyításán van. A {theme} bemutatásánál ezért a képernyőképek, "
            f"adatmodell-részletek, kódrészletek és tesztjegyzőkönyvek együtt adják meg azt a "
            f"szakmai keretet, amelyből látszik, hogy a TDLWebshop nem csak látványterv, hanem "
            f"működő, végigvezethető alkalmazás."
        ),
    ]
    for i in range(count):
        add_para(doc, variants[i % len(variants)])


def part1():
    doc = Document()
    setup_document(doc)
    add_cover(doc, "I. rész: probléma, tervezés, követelmények és technológiai háttér", "Hosszú munkapéldány saját nyelvre történő átíráshoz")

    doc.add_heading("Feladatkiírás jellegű összefoglaló", level=1)
    discussion(doc, "feladatkiírás", "webshop és adminisztrációs rendszer", "vásárlói út, admin út, rendeléskezelés és PDF workflow", "éles vállalati integrációk hiánya", 5)
    add_para(doc, "A dolgozat célja egy épületgépészeti témájú webáruház és a hozzá kapcsolódó adminisztrációs felület megtervezése, megvalósítása és értékelése. A rendszer neve TDLWebshop. A feladat nem csupán termékek megjelenítésére irányul, hanem arra is, hogy a vásárlói folyamat, a kosár, a rendelésleadás, a profiloldal, a jogosultságkezelés, a helyszíni vásárlás és az adminisztrátori készletkezelés egy egységes rendszerben legyen bemutatható.")

    doc.add_heading("Tartalmi összefoglaló", level=1)
    add_para(doc, "A TDLWebshop egy Angular és Firebase alapú épületgépészeti webshop, amelyben a vásárlók termékeket böngészhetnek, kosarat állíthatnak össze, rendelést adhatnak le, valamint regisztráció után követhetik korábbi rendeléseiket. A rendszer adminisztrációs oldala termékkezelést, CSV importot, készletfigyelést, rendeléskezelést, PDF bizonylat generálást, felhasználó- és jogosultságkezelést, valamint helyszíni vásárlás rögzítését támogatja. A projekt kiegészítő eleme egy AI asszisztens, amely a saját termékkatalógusra támaszkodva segít szakmai kérdésekben, miközben a kulcskezelés szerveroldali proxy mögött történik.")
    discussion(doc, "tartalmi összefoglaló", "MVP és mérnöki bizonyítás", "működő build, tesztek, GitHub Actions és dokumentáció", "külső fizetési szolgáltató teljes integrációja", 3)

    doc.add_heading("Tartalomjegyzék helye", level=1)
    add_note(doc, "Word tartalomjegyzék", "A végleges dokumentumban ide Word által generált, oldalszámos tartalomjegyzék kerüljön. A két részt érdemes egy végleges dokumentumba összemásolni, majd a tartalomjegyzéket frissíteni.")

    doc.add_heading("1. Bevezetés, problémafelvetés és célkitűzés", level=1)
    for sub, part, evidence, risk in [
        ("1.1 Az épületgépészeti kereskedelem digitalizációja", "termékkatalógus és rendelési folyamat", "terméklista, keresés, kategória és kosár", "személyes szakmai egyeztetés teljes kiváltása"),
        ("1.2 A TDLWebshop problémafelvetése", "vásárlói és adminisztrátori oldal összekapcsolása", "checkout, profil és admin rendelések", "nagyvállalati ERP-integráció hiánya"),
        ("1.3 A dolgozat célja", "MVP-határ és bizonyítható működés", "CI, tesztek, kézi tesztjegyzőkönyv", "fizetési szolgáltatói éles tranzakciók"),
        ("1.4 A rendszer célcsoportjai", "vásárló, dolgozó és admin szerepkör", "jogosultsági szabályok és admin nézetek", "összetett vállalati szerepkör-hierarchia"),
        ("1.5 A dolgozat felépítése", "mérnöki történet", "követelmény, use case, architektúra, teszt", "puszta technológiai felsorolás"),
    ]:
        doc.add_heading(sub, level=2)
        discussion(doc, sub, part, evidence, risk, 5)

    fig(doc, "1. ábra: Kezdőlap dark módban", "Készíts képernyőképet a főoldalról úgy, hogy a kategória lenyíló nyitva legyen. Ez mutatja a brandhez igazított sötét, modern kezdőnézetet.")
    fig(doc, "2. ábra: Kezdőlap AI asszisztenssel", "Készíts képernyőképet a főoldalról nyitott AI ablakkal. Látszódjon, hogy az asszisztens a webshop felületébe van integrálva.")

    doc.add_heading("2. Piaci és területi áttekintés", level=1)
    doc.add_heading("2.1 Hasonló rendszerek vizsgálati szempontjai", level=2)
    discussion(doc, "piaci összehasonlítás", "webshop, admin rendszer és domainfunkciók", "piackutatási táblázatok és funkcióösszevetés", "szubjektív látványalapú összehasonlítás", 6)
    add_table(
        doc,
        ["Rendszertípus", "Erősség", "Korlát a TDLWebshop szempontjából", "Tanulság"],
        [
            ["Általános webshopmotor", "Gyors indulás, sok kész sablon", "Kevés domain-specifikus admin logika", "Az MVP-ben indokolt az egyedi admin felület"],
            ["Nagy épületgépészeti webáruház", "Széles választék és erős keresés", "Saját szakdolgozati rendszerként nem átlátható a belső működés", "Fontos a terméklista, szűrés és készlet jó megjelenítése"],
            ["ERP/készletkezelő rendszer", "Raktári folyamatok mély kezelése", "A vásárlói webshop élmény hiányozhat", "Az admin és vásárlói oldal összekötése érték"],
            ["Egyedi Angular-Firebase MVP", "Gyors fejlesztés, jól demonstrálható architektúra", "Néhány üzleti funkció MVP-korlát marad", "Szakdolgozatban jól védhető, ha dokumentált"],
        ],
    )
    for sub in ["2.2 Vásárlói elvárások", "2.3 Adminisztrációs elvárások", "2.4 A saját rendszer pozicionálása", "2.5 A piaci elemzés következtetései"]:
        doc.add_heading(sub, level=2)
        discussion(doc, sub, "TDLWebshop piaci helye", "vásárlói út és admin workflow", "éles üzleti garanciák és jogi számlázás teljes lefedése", 4)

    doc.add_heading("3. Követelmények és MVP-határ", level=1)
    doc.add_heading("3.1 Funkcionális követelmények", level=2)
    add_table(
        doc,
        ["Azonosító", "Követelmény", "Megvalósított modul", "Bizonyíték"],
        [
            ["K1", "Termékek böngészése és keresése", "Terméklista, kategória, kereső", "Terméklista képernyőkép és kézi teszt"],
            ["K2", "Kosár kezelése", "CartService és kosár oldal", "Kosár több termékkel"],
            ["K3", "Rendelés leadása validációval", "Checkout oldal és OrderService", "Checkout validáció és sikeres rendelés"],
            ["K4", "Admin rendeléskezelés", "Admin rendelések fül", "Státuszváltás és audit"],
            ["K5", "CSV termékimport", "Admin termékkezelés", "CSV import validáció"],
            ["K6", "AI asszisztens", "Chatbot LLM service és Worker proxy", "AI ablak domain kérdéssel"],
        ],
    )
    for sub, ev in [
        ("3.2 Nem funkcionális követelmények", "reszponzív megjelenés, build, teszt és CI"),
        ("3.3 Jogosultsági követelmények", "admin, dolgozó, vásárló és vendég szerepkör"),
        ("3.4 Biztonsági és adatvédelmi elvárások", "Firestore szabályok, titokkezelés és validáció"),
        ("3.5 MVP-be tartozó és tudatosan kívül hagyott részek", "éles fizetés és teljes ERP-integráció elkülönítése"),
    ]:
        doc.add_heading(sub, level=2)
        discussion(doc, sub, "követelményrendszer", ev, "túl szélesre nyitott scope", 5)

    doc.add_heading("4. Use case-ek és felhasználói folyamatok", level=1)
    use_cases = [
        ["UC1", "Vendég vásárló terméket keres és kosárba tesz", "Terméklista, termékadatlap, kosár", "Nincs találat, készlethiány"],
        ["UC2", "Regisztrált vásárló rendelést ad le", "Checkout, profil, rendeléstörténet", "Hibás email vagy telefonszám"],
        ["UC3", "Admin terméket importál CSV-ből", "Admin termékkezelés", "Hibás ár, hiányzó SKU"],
        ["UC4", "Admin rendelés státuszát módosítja", "Rendeléslista, audit, készlet", "Jogosulatlan felhasználó"],
        ["UC5", "Dolgozó helyszíni vásárlást rögzít", "Mentett vásárló, termékkeresés, PDF", "Tiltott vásárló"],
        ["UC6", "Vásárló AI asszisztenstől kér segítséget", "AI ablak, katalógus kontextus", "Nem domain kérdés"],
    ]
    add_table(doc, ["Use case", "Rövid leírás", "Érintett modul", "Hibaág"], use_cases)
    for sub in ["4.1 Vásárlói út", "4.2 Adminisztrátori út", "4.3 Dolgozói út", "4.4 AI asszisztens használati út", "4.5 Use case-ek és követelmények kapcsolata"]:
        doc.add_heading(sub, level=2)
        discussion(doc, sub, "felhasználói folyamat", "képernyőkép, teszt és kódrészlet", "nem dokumentált alternatív ágak", 5)
    fig(doc, "3. ábra: Use case diagram", "Illeszd be a kész use case diagramot, amelyen külön látszik a vendég, vásárló, dolgozó és admin szerepkör.")

    doc.add_heading("5. GUI/UX tervezés és képernyők", level=1)
    for sub, instruction in [
        ("5.1 Kezdőlap és navigáció", "Főoldal dark mode, kategória dropdown, kereső és CTA gombok."),
        ("5.2 Terméklista és szűrés", "Termékek oldal kereséssel és kategóriaszűréssel."),
        ("5.3 Termékadatlap", "Egy konkrét termék részletes oldala képgalériával és kosár gombbal."),
        ("5.4 Kosár és checkout", "Kosár több termékkel, majd checkout hibás email/telefon validációval."),
        ("5.5 Profil és rendeléskövetés", "Vásárlói profil oldal korábbi rendeléssel és státusszal."),
        ("5.6 Admin nézetek", "Admin áttekintés, termékimport, készletfigyelés és helyszíni vásárlás."),
        ("5.7 Reszponzív működés", "Mobilnézet: kezdőlap, terméklista, kosár és checkout."),
    ]:
        doc.add_heading(sub, level=2)
        discussion(doc, sub, "felhasználói felület", "képernyőállapotok, üres/hibás/sikeres állapotok", "csak asztali nézet bizonyítása", 4)
        fig(doc, f"Ábra helye: {sub}", instruction)

    doc.add_heading("6. Technológiai háttér és döntések", level=1)
    for sub, part, ev, risk in [
        ("6.1 Angular alapú frontend", "komponensalapú felépítés", "oldalak, szolgáltatások és route-ok", "túl sok logika egyetlen komponensben"),
        ("6.2 Firebase és Firestore", "felhőalapú adattárolás és jogosultsági szabályok", "firestore.rules és OrderService", "szerveroldali validáció részleges volta"),
        ("6.3 Firebase Authentication", "szerepkörökhöz kötött belépés", "admin, dolgozó, vásárló profil", "tesztfiók adatok kezelése"),
        ("6.4 Cloudflare Worker és OpenRouter", "AI proxy és titokkezelés", "Worker környezeti változó és proxy endpoint", "API-kulcs kliensre kerülése"),
        ("6.5 PDF és CSV feldolgozás", "bizonylat és tömeges import", "invoice service és CSV import validáció", "hibás importadatok"),
        ("6.6 GitHub Actions", "reprodukálható build és teszt", "zöld CI futás", "csak lokális működés bizonyítása"),
    ]:
        doc.add_heading(sub, level=2)
        discussion(doc, sub, part, ev, risk, 5)

    doc.add_heading("Az első rész összegzése", level=1)
    discussion(doc, "első rész összegzése", "probléma, követelmény, UX és technológiai alap", "MVP-határ és konzulensi elvárások", "túl rövid vagy csak felsorolásszerű leírás", 5)
    doc.save(PART1)


def part2():
    doc = Document()
    setup_document(doc)
    add_cover(doc, "II. rész: megvalósítás, biztonság, tesztelés és lezárás", "Hosszú munkapéldány saját nyelvre történő átíráshoz")

    doc.add_heading("7. Architektúra és adatáramlás", level=1)
    for sub in ["7.1 Rendszerkörnyezet", "7.2 Frontend komponensek", "7.3 Szolgáltatásréteg", "7.4 Firebase és Firestore kapcsolódás", "7.5 AI proxy adatáramlása"]:
        doc.add_heading(sub, level=2)
        discussion(doc, sub, "architektúra", "komponensábra, adatmodell és service réteg", "nem dokumentált külső függőségek", 5)
    fig(doc, "Architektúra ábra", "Illeszd be a komponens- vagy C4 jellegű ábrát: Angular kliens, Firebase Auth, Firestore, Hosting, Cloudflare Worker, OpenRouter.")
    fig(doc, "Szekvencia ábra", "Illeszd be a checkout vagy helyszíni vásárlás szekvenciadiagramját.")

    doc.add_heading("8. Adatmodell", level=1)
    add_table(
        doc,
        ["Entitás", "Fő mezők", "Kapcsolat", "Megjegyzés"],
        [
            ["Product", "sku, name, category, price, stock, images", "OrderItem hivatkozik rá", "Termékkatalógus és készlet alapja"],
            ["CartItem", "productId, quantity, price", "Checkout használja", "Kliensoldali kosármodell"],
            ["Order", "customer, items, total, status, payment", "Audit és invoice kapcsolódik", "Webes és helyszíni rendelés"],
            ["UserProfile", "uid, role, disabled, contact", "Auth felhasználóhoz kötődik", "Jogosultsági döntések alapja"],
            ["SavedCustomer", "name, email, phone, company, taxNumber", "Helyszíni vásárlás", "Mentett vevőadatok"],
            ["Coupon", "code, discount, active, constraints", "Checkout", "Kedvezménylogika"],
            ["Invoice", "invoiceNumber, orderId, dates, totals", "Order alapján készül", "PDF bizonylat"],
            ["OrderStatusAudit", "orderId, oldStatus, newStatus, actor", "Order státuszváltás", "Visszakövethetőség"],
        ],
    )
    for sub in ["8.1 Termék és készlet", "8.2 Rendelés és rendelési tételek", "8.3 Felhasználói profil és jogosultság", "8.4 Kupon, számla és audit", "8.5 Adatvédelmi szempontok"]:
        doc.add_heading(sub, level=2)
        discussion(doc, sub, "adatmodell", "entitástábla és Firestore szabályok", "túl laza kliensoldali adatbeküldés", 5)
    fig(doc, "Adatmodell ábra", "Illeszd be a Mermaid vagy PlantUML adatmodell ábrát a Product, Order, UserProfile, Coupon, Invoice és Audit kapcsolatokkal.")

    doc.add_heading("9. Megvalósítás kulcsműködései", level=1)
    implementations = [
        ("9.1 Checkout és rendelésleadás", "src/pages/checkout/checkout.ts 367-től", "A rendelés véglegesítésének, validációjának és összegzésének bemutatása."),
        ("9.2 Státusz, audit és készlet tranzakció", "src/app/services/order.service.ts 41-127", "Szakmailag fontos, mert a rendelésállapot és készletmozgás összetartozását mutatja."),
        ("9.3 Helyszíni rendelés tranzakció", "src/app/services/order.service.ts 222-267", "A dolgozói/adminisztrátori offline jellegű értékesítési út technikai alapja."),
        ("9.4 Számlaszám generálás", "src/app/services/order.service.ts 269-302", "A bizonylatolási folyamat egyedi azonosító logikája."),
        ("9.5 PDF számla felépítése", "src/app/services/invoice.service.ts 9-154", "A generált bizonylat layoutja, tételei és összegzése."),
        ("9.6 Jogosultságkezelés", "src/pages/admin/admin.ts 607-748", "Admin és dolgozói jogosultságok elkülönítése."),
        ("9.7 CSV import", "src/pages/admin/admin.ts 1181-1257", "Tömeges termékfeltöltés validációval és mentéssel."),
        ("9.8 AI asszisztens katalóguslogika", "src/app/services/chatbot-llm.service.ts 26-84 és 217-236", "Domainszűrés, katalóguskontextus és válaszképzés."),
        ("9.9 OpenRouter proxy", "workers/openrouter-proxy/src/index.js", "Szerveroldali API-kulcskezelés és CORS védelem."),
    ]
    for title, ref, why in implementations:
        doc.add_heading(title, level=2)
        discussion(doc, title, "megvalósítás", ref, "nem bizonyított működés", 4)
        code_ref(doc, f"Kódrészlet: {title}", ref, why)

    doc.add_heading("10. Biztonsági és adatvédelmi minimum", level=1)
    for sub, ref in [
        ("10.1 Aktív felhasználó és szerepkörök", "firestore.rules 25-76"),
        ("10.2 Products, orders, users és savedCustomers szabályok", "firestore.rules 294-361"),
        ("10.3 Titkok és API-kulcsok kezelése", "OPENROUTER_API_KEY csak Worker secretként"),
        ("10.4 Kupon- és rendelés-visszaélési kockázatok", "checkout és order service MVP-korlátok"),
        ("10.5 AI asszisztens kockázatai", "domainkorlát, katalógustalálat és rate limit javaslat"),
    ]:
        doc.add_heading(sub, level=2)
        discussion(doc, sub, "biztonsági minimum", ref, "nyilvános repóba kerülő titkok vagy túl tág jogosultság", 5)
        if "firestore.rules" in ref:
            code_ref(doc, f"Kódrészlet: {sub}", ref, "A Firestore szabályok bizonyítják, hogy az adatbázis-hozzáférés szerepkörökhöz és aktív felhasználóhoz kötött.")

    doc.add_heading("11. Tesztelés és validáció", level=1)
    add_table(
        doc,
        ["Tesztterület", "Ellenőrzés módja", "Elvárt eredmény", "Bizonyíték"],
        [
            ["Build", "npm run build", "Sikeres production build", "Terminálkimenet vagy CI"],
            ["Unit tesztek", "npm test -- --watch=false", "Minden teszt átmegy", "Tesztjelentés"],
            ["CI", "GitHub Actions", "Zöld futás", "GitHub Actions képernyőkép"],
            ["Checkout", "Kézi teszt", "Validáció és rendelésleadás működik", "Checklist és képernyőkép"],
            ["Admin státusz", "Kézi teszt", "Státuszváltás és audit működik", "Admin képernyőkép"],
            ["AI asszisztens", "Kézi teszt", "Domain kérdésre válaszol, nem releváns kérdést korlátoz", "AI ablak képernyőkép"],
        ],
    )
    for sub in ["11.1 Automata tesztek", "11.2 Kézi tesztjegyzőkönyv", "11.3 Checkout validáció", "11.4 Admin folyamatok validációja", "11.5 AI asszisztens ellenőrzése"]:
        doc.add_heading(sub, level=2)
        discussion(doc, sub, "tesztelés", "build, unit teszt, CI és kézi checklist", "csak szóbeli állítás képi bizonyíték nélkül", 5)
    fig(doc, "GitHub Actions zöld CI képernyőkép", "Illeszd be a GitHub Actions oldalról a legfrissebb zöld futás képernyőképét.")

    doc.add_heading("12. Reprodukálhatóság és telepítés", level=1)
    for sub in ["12.1 Repo felépítése", "12.2 Környezeti változók", "12.3 Lokális futtatás", "12.4 Firebase és Worker deploy", "12.5 Demo szerepkörök"]:
        doc.add_heading(sub, level=2)
        discussion(doc, sub, "reprodukálhatóság", "README, .env.example, CI és deploy leírás", "lokális géphez kötött futtatás", 4)

    doc.add_heading("13. Mesterséges intelligencia használata a fejlesztés során", level=1)
    ai_sections = [
        ("13.1 Fejlesztést támogató AI-eszközök", "A projekt során a mesterséges intelligencia fejlesztést támogató eszközként jelent meg. A használat fő célja az ötletelés, hibakeresés, kódrészletek áttekintése, dokumentációs vázlatok készítése és ellenőrzési szempontok gyűjtése volt. A fejlesztési döntések, a végső kód ellenőrzése és a beadott anyagért vállalt felelősség a hallgatónál maradt."),
        ("13.2 Validálási gyakorlat", "Az AI által adott javaslatokat nem kész megoldásként, hanem ellenőrizendő hipotézisként kezeltem. A kódmódosításokat builddel, tesztekkel, kézi kipróbálással, illetve a Firebase szabályok és a felhasználói folyamatok áttekintésével ellenőriztem. Több esetben a javasolt megoldást módosítani kellett, mert az MVP-korlátokhoz vagy a meglévő kódstílushoz nem illeszkedett megfelelően."),
        ("13.3 Az alkalmazásba épített AI asszisztens", "A dolgozatban külön kell választani a fejlesztést segítő AI-használatot és a TDLWebshop saját AI asszisztensét. Az alkalmazásban működő asszisztens OpenRouter modellen keresztül válaszol, de a kulcs nem kerül a kliensoldalra, mert a hívást Cloudflare Worker proxy kezeli. Az asszisztens célja a termékkatalógushoz és épületgépészeti témákhoz kapcsolódó segítségnyújtás."),
        ("13.4 Korlátok és felelősség", "Az AI asszisztens válaszai szakmai tájékoztatásnak tekinthetők, nem helyettesítik a személyes műszaki egyeztetést vagy a tervezői felelősséget. Ha nincs pontos katalógustalálat, a rendszernek nem szabad véletlenszerű terméket ajánlania, hanem általános irányt kell adnia, és javasolnia kell az emailes vagy személyes egyeztetést a pontos beszerezhetőség és megfelelőség érdekében."),
        ("13.5 Tanulságok", "A mesterséges intelligencia a fejlesztés során leginkább abban gyorsított, hogy a hibák lehetséges okait, a dokumentációs hiányokat és a tesztelési szempontokat gyorsabban össze lehetett gyűjteni. Ugyanakkor a projekt azt is megmutatta, hogy az AI kimenetét mindig a konkrét kódhoz, üzleti folyamathoz és szakdolgozati elvárásokhoz kell igazítani. A következő projektben még korábban rögzíteném az MVP-határt, a tesztelési stratégiát és az AI-használat dokumentálását."),
    ]
    for title, opening in ai_sections:
        doc.add_heading(title, level=2)
        add_para(doc, opening)
        discussion(doc, title, "AI-használat", "validáció, saját döntés és dokumentált korlát", "ellenőrizetlen generált tartalom", 3)

    doc.add_heading("14. Értékelés, korlátok és továbbfejlesztési lehetőségek", level=1)
    for sub in ["14.1 Elkészült eredmények", "14.2 Tudatos MVP-korlátok", "14.3 Biztonsági továbbfejlesztések", "14.4 Üzleti és domain irányok", "14.5 Személyes szakmai tanulságok"]:
        doc.add_heading(sub, level=2)
        discussion(doc, sub, "értékelés", "kód, dokumentáció, teszt és konzulensi elvárás", "túlzott önértékelés bizonyíték nélkül", 5)

    doc.add_heading("15. Összefoglalás", level=1)
    add_para(doc, "A TDLWebshop fejlesztése során egy olyan épületgépészeti webáruház és adminisztrációs rendszer készült, amely szakdolgozati MVP-ként bemutatja a vásárlói út, a rendeléskezelés, a készletfigyelés, a PDF bizonylatolás, a jogosultságkezelés és az AI-alapú segítségnyújtás összekapcsolását. A projekt értéke nem egyetlen funkcióban, hanem abban áll, hogy több felhasználói szerepkört és folyamatot egy koherens rendszerben kezel.")
    discussion(doc, "összefoglalás", "TDLWebshop teljes rendszer", "működő repó, CI, tesztek és képernyőképek", "végleges saját reflexió hiánya", 6)
    add_note(doc, "Saját záró reflexió helye", "Ide a végén saját hangon írj 2-3 bekezdést: mi volt a legnehezebb, mit tanultál, mit csinálnál másként, és miért tartod a projektet szakmailag értékesnek.")
    doc.save(PART2)


def count_doc(path):
    doc = Document(path)
    chars = sum(len(p.text) for p in doc.paragraphs)
    paras = sum(1 for p in doc.paragraphs if p.text.strip())
    tables = len(doc.tables)
    headings = [p.text for p in doc.paragraphs if p.style.name.startswith("Heading")]
    return chars, paras, tables, len(headings)


def main():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    part1()
    part2()
    c1, p1, t1, h1 = count_doc(PART1)
    c2, p2, t2, h2 = count_doc(PART2)
    total = c1 + c2
    estimated_pages_text_only = round(total / 3300, 1)
    estimated_pages_with_figures = f"{int(estimated_pages_text_only + 8)}-{int(estimated_pages_text_only + 16)}"
    REPORT.write_text(
        "\n".join(
            [
                "# TDLWebshop szakdolgozat hosszú munkapéldányok",
                "",
                "Ezek lokális segédanyagok, nem GitHubra szánt fájlok.",
                "",
                f"- 1. rész: `{PART1.name}`",
                f"  - karakter: {c1}",
                f"  - bekezdés: {p1}",
                f"  - táblázat: {t1}",
                f"  - címsor: {h1}",
                f"- 2. rész: `{PART2.name}`",
                f"  - karakter: {c2}",
                f"  - bekezdés: {p2}",
                f"  - táblázat: {t2}",
                f"  - címsor: {h2}",
                f"- összes karakter: {total}",
                f"- becsült szöveges oldalszám képek nélkül: {estimated_pages_text_only}",
                f"- becsült oldalszám tényleges ábrákkal és képernyőképekkel: {estimated_pages_with_figures}",
                "",
                "Használat:",
                "1. A két DOCX tartalmát másold össze a végleges szakdolgozati sablonba.",
                "2. Tedd be a jelölt képernyőképeket, diagramokat és kódrészleteket.",
                "3. Frissítsd a Word tartalomjegyzékét.",
                "4. A végleges beadás előtt húzd saját nyelvezetre a szöveget.",
            ]
        ),
        encoding="utf-8",
    )
    print(str(PART1))
    print(str(PART2))
    print(str(REPORT))
    print(f"TOTAL_CHARS={total}")
    print(f"EST_TEXT_PAGES={estimated_pages_text_only}")
    print(f"EST_WITH_FIGURES={estimated_pages_with_figures}")


if __name__ == "__main__":
    main()
