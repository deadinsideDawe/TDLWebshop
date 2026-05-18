from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


OUT = Path(__file__).with_name("TDLWebshop_szakdolgozat_konzulensi_review_munkaverzio.docx")


BLUE = "1f4e79"
LIGHT_BLUE = "d9eaf7"
LIGHT_GRAY = "f2f5f9"
TEXT = "1f2937"


def set_font(run, size=12, bold=False, italic=False, color=TEXT):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def borders(cell, color="c8d3e0", size="6"):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = "w:" + edge
        element = tc_borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tc_borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def configure(doc):
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.4)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2.3)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.space_after = Pt(6)

    for name, size, color in [
        ("Title", 22, BLUE),
        ("Heading 1", 16, BLUE),
        ("Heading 2", 14, BLUE),
        ("Heading 3", 12, TEXT),
    ]:
        style = styles[name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(12 if name != "Heading 3" else 8)
        style.paragraph_format.space_after = Pt(6)


def para(doc, text="", style=None, align=None):
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(6)
    if text:
        r = p.add_run(text)
        set_font(r)
    return p


def heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for r in p.runs:
        set_font(r, size={1: 16, 2: 14, 3: 12}.get(level, 12), bold=True, color=BLUE if level < 3 else TEXT)
    return p


def page_break(doc):
    doc.add_page_break()


def table(doc, headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        shade(hdr[i], BLUE)
        borders(hdr[i])
        for p in hdr[i].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                set_font(r, size=10, bold=True, color="ffffff")
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            borders(cells[i])
            if len(t.rows) % 2 == 0:
                shade(cells[i], LIGHT_GRAY)
            for p in cells[i].paragraphs:
                p.paragraph_format.space_after = Pt(2)
                for r in p.runs:
                    set_font(r, size=9)
    if widths:
        for row in t.rows:
            for i, width in enumerate(widths):
                row.cells[i].width = Cm(width)
    para(doc)
    return t


def figure(doc, caption, note):
    p = para(doc)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("[IDE KERUL A KEPERNYOKEP / ABRA]")
    set_font(r, size=10, bold=True, color=BLUE)
    p = para(doc, caption)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in p.runs:
        set_font(r, size=10, italic=True)
    if note:
        p = para(doc, note)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs:
            set_font(r, size=9, italic=True, color="6b7280")


def code_place(doc, title, path, lines, reason):
    p = para(doc)
    r = p.add_run(f"{title} ({path}, {lines})")
    set_font(r, size=10, bold=True, color=BLUE)
    p = para(doc, f"[IDE KERUL A ROVID KODRESZLET: {path}, {lines}]")
    shade_paragraph(p, "eef6ff")
    p = para(doc, reason)
    for r in p.runs:
        set_font(r, size=10, italic=True, color="4b5563")


def shade_paragraph(paragraph, fill):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)


def add_cover(doc):
    for _ in range(3):
        para(doc)
    p = para(doc, "Szegedi Tudomanyegyetem", align=WD_ALIGN_PARAGRAPH.CENTER)
    for r in p.runs:
        set_font(r, size=14, bold=True)
    para(doc, "[Kar / Intezet / Tanszek neve]", align=WD_ALIGN_PARAGRAPH.CENTER)
    for _ in range(5):
        para(doc)
    p = para(doc, "TDLWebshop", align=WD_ALIGN_PARAGRAPH.CENTER)
    for r in p.runs:
        set_font(r, size=24, bold=True, color=BLUE)
    p = para(doc, "Epületgepeszeti webshop es adminisztracios rendszer fejlesztese", align=WD_ALIGN_PARAGRAPH.CENTER)
    for r in p.runs:
        set_font(r, size=16, bold=True)
    para(doc)
    p = para(doc, "Szakdolgozat", align=WD_ALIGN_PARAGRAPH.CENTER)
    for r in p.runs:
        set_font(r, size=15, bold=True)
    for _ in range(6):
        para(doc)
    table(doc, ["Szerzo", "Temavezeto"], [["Toth David Laszlo", "Dr. Bilicki Vilmos, egyetemi docens"]], [7, 8])
    for _ in range(4):
        para(doc)
    para(doc, "Szeged, 2026", align=WD_ALIGN_PARAGRAPH.CENTER)
    page_break(doc)


def add_front_matter(doc):
    heading(doc, "Feladatkiiras helye", 1)
    para(doc, "Ide kerul az intezmenyi feladatkiiras vagy annak mellekletre valo hivatkozasa. A vegleges valtozatban ezt a reszt az aktualis kari/tanszeki formai kovetelmenyek szerint kell rendezni.")
    page_break(doc)

    heading(doc, "Tartalmi kivonat", 1)
    for text in [
        "A szakdolgozat celja egy epületgepeszeti termekekre specializalt webshop es adminisztracios rendszer megtervezese, megvalositasa es ertekelese. A fejlesztes soran a hangsuly nem csupan egy termeklista vagy kosar elkeszitesen volt, hanem egy olyan MVP kialakitasan, amely a vasarloi es az adminisztratori oldal legfontosabb folyamatait is lefedi.",
        "A TDLWebshop Angular alapu kliensalkalmazasra, Firebase szolgaltatasokra, Firestore adatbazisra es Cloudflare Workerrel vedett AI proxyra epul. A rendszer tartalmaz termekbongeszeshez, kosarhoz, checkout folyamatokhoz, rendelestortenethez, admin termek- es keszletkezeleshez, CSV importhoz, kuponlogikahoz, PDF bizonylathoz es AI asszisztenshez kapcsolodo funkciokat.",
        "A dolgozat bemutatja a problema hatteret, az MVP hatarait, a piaci osszehasonlitast, a kovetelmenyeket, a felhasznaloi eseteket, az architekturat, az adatmodellt, a megvalositasi reszleteket, a biztonsagi megfontolasokat es a teszteles eredmenyeit. A fejlesztes soran kulon figyelmet kaptak a jogosultsagok, az inputvalidacio, a titokkezeles, a reprodukalhato futtatas es a dokumentalt AI-hasznalat.",
    ]:
        para(doc, text)
    heading(doc, "Abstract", 1)
    para(doc, "This thesis presents the design and implementation of TDLWebshop, a domain-focused e-commerce and administration system for building services products. The project demonstrates a product-like MVP with customer workflows, administrative order and product management, inventory-related features, CSV import, coupon validation, PDF invoice generation and an AI assistant connected through a protected proxy layer.")
    para(doc, "The work discusses requirements, use cases, system architecture, data model, implementation details, security considerations, testing and reproducibility. The goal is to provide not only a working application, but also engineering evidence that explains the design decisions, the limitations of the MVP and the validation process.")
    page_break(doc)

    heading(doc, "Tartalomjegyzek", 1)
    para(doc, "[A vegleges Word fajlban ide automatikus tartalomjegyzek kerul oldalszamokkal.]")
    page_break(doc)


def add_intro(doc):
    heading(doc, "1. Bevezetes", 1)
    for text in [
        "Az online kereskedelem ma mar nem csak a klasszikus fogyasztasi termekeknel meghatarozo. Egyre tobb olyan szakmai terulet jelenik meg a weben, ahol a vasarlo nem egyszeruen arat keres, hanem muszaki szempontbol is ertheto, rendszerezett es gyorsan ellenorizheto informaciot var. Az epületgepeszeti termekek ilyen teruletet kepviselnek, mert a futeshez, vizszereleshez, szellozeshez vagy hutestechnikahoz kapcsolodo termekek kivalasztasa gyakran szakmai dontes is.",
        "A projekt otlete abbol indult ki, hogy egy epületgepeszeti webshopnak egyszerre kell kiszolgalnia a lakossagi vasarlot es a szakmai felhasznalot. A lakossagi felhasznalo altalaban egyszeru, attekintheto termekkeresest, kosarat, biztonsagos checkoutot es rendeleseinek koveteset igenyli. A szerelo, dolgozo vagy adminisztrator oldalon viszont fontos a keszletatlatas, a gyors helyszini vasarlas rogzitese, a CSV-s termekfeltoltes, a jogosultsagkezeles es a bizonylatkeszites.",
        "A TDLWebshop ennek a ket oldalnak az osszekotesere keszult. A rendszer celja nem egy minden reszleteben kereskedelmi szintu vallalatiranyitasi rendszer megvalositasa volt, hanem egy olyan szakdolgozati MVP, amely valos webshop-logikat mutat be, es kozben elegendoen tag ahhoz, hogy architekturalis, adatmodellezesi, biztonsagi es tesztelesi szempontbol is ertekelheto legyen.",
        "A fejlesztes soran fontos szempont volt, hogy az alkalmazas ne csak mukodjon, hanem bemutathato es reprodukalhato is legyen. Ezert a kod mellett dokumentacio, tesztelesi jegyzokonyv, adatmodell, architekturaleiras, biztonsagi minimum es AI-hasznalati fejezet is keszult. Ezek egyutt bizonyitjak, hogy a rendszer nem elszigetelt kepernyok halmaza, hanem osszefuggo szoftvertermek.",
    ]:
        para(doc, text)
    figure(doc, "1. abra: A TDLWebshop kezdolapja dark modban, kategoria lenyiloval.", "Kepernyokep: kezdolap, kategoria menu nyitva.")

    heading(doc, "1.1. Problemafelvetes", 2)
    for text in [
        "Egy hagyomanyos webshopnal gyakran elegendo, ha a felhasznalo gyorsan megtalalja a termeket es le tudja adni a rendelest. Az epületgepeszeti teruleten azonban a termekekhez kapcsolodo muszaki jellemzok, raktarallapot, ar, kategoria, felhasznalasi terulet es a kesobbi adminisztracios folyamat is lenyeges.",
        "A problema egyik resze felhasznaloi: a vasarlo szeretne attekintheto termeklistat, keresest, kosarat, rendelestortenetet es visszajelzest kapni a rendeles allapotarol. A masik resze adminisztracios: a webshop uzemeltetoje szeretne termekeket feltolteni, keszletet figyelni, rendelest kezelni, helyszini vasarlast rogzitani, szamlat vagy bizonylatot generalni, es jogosultsag alapjan elkuloniteni az admin, dolgozo es vasarlo szerepkoroket.",
        "A szakdolgozatban bemutatott megoldas ezekre a problemakra ad egy mukodo, de tudatosan korlatozott MVP valaszt. A rendszer nem vallalja peldaul egy teljes NAV-kompatibilis szamlazo, egy bankkartya-elszamolasi backend vagy egy teljes vallalati ERP kivaltasat. Ehelyett a webshop es admin felulet azon funkcioit helyezi eloterbe, amelyek egy szoftveres szakdolgozatban jol bemutathatok es ellenorizhetok.",
    ]:
        para(doc, text)

    heading(doc, "1.2. Célkitűzés", 2)
    table(doc, ["Cel", "Megvalositas a rendszerben", "Bizonyitek"], [
        ["Vasarloi webshop folyamat", "Termeklista, termekadatlap, kosar, checkout, profil", "Kepernyokepek es checkout tesztek"],
        ["Adminisztracios folyamat", "Termekkezeles, CSV import, rendeleskezeles, keszletfigyeles", "Admin kepernyokepek es unit tesztek"],
        ["Jogosultsagi modell", "Admin, dolgozo, vasarlo es tiltott felhasznalo szerepkorok", "Firestore rules es admin felulet"],
        ["PDF bizonylat", "Rendeleshez kapcsolt PDF szamla/bizonylat generalas", "Invoice service kodreszlet es PDF kep"],
        ["AI asszisztens", "Katalogushoz kotott epületgepeszeti valaszadas proxy-n keresztul", "AI ablak kepernyokep es proxy kodreszlet"],
    ], [4, 7, 4])

    heading(doc, "1.3. MVP-határ", 2)
    para(doc, "Az MVP hataranak kijelolese azert fontos, mert igy a rendszer nem tunik befejezetlen funkciok halmazanak. A TDLWebshop MVP-je a vasarloi es adminisztratori alapu webshopmukodest bizonyitja. Ide tartozik a termekbongeszes, kosar, checkout, rendelesmentes, profil, admin termekkezeles, keszletfigyeles, CSV import, felhasznaloi szerepkorok, PDF bizonylat es AI asszisztens. Tudatosan kivul maradt a teljes fizetesi szolgaltatoi integracio, a NAV-kompatibilis eles szamlazas, a teljes vallalati keszletgazdalkodasi rendszer es a valos logisztikai partnerintegracio.")
    figure(doc, "2. abra: MVP-hatar es tovabbfejlesztesi iranyok attekintese.", "Abraszeru osszefoglalo: belul az MVP, kivul a kesobbi lehetosegek.")


def add_market_requirements(doc):
    heading(doc, "2. Piaci es teruleti osszehasonlitas", 1)
    for text in [
        "A piaci osszehasonlitas celja nem az volt, hogy a TDLWebshopot egy teljes kereskedelmi rendszerrel azonos szintre helyezze, hanem hogy lathato legyen, milyen funkciok jellemzoek egy epületgepeszeti webshopra, es ezekbol a szakdolgozati MVP melyeket valosit meg.",
        "A vizsgalt rendszerek altalaban eros termekkatalogussal, kategoriakkal es keresessel rendelkeznek, de a nyilvanos vasarloi feluleten kevesbe lathato az adminisztracios oldal. A TDLWebshop erteke abban jelenik meg, hogy a dolgozatban a vasarloi oldal mellett az adminisztracios es jogosultsagi folyamatok is bemutathatok.",
    ]:
        para(doc, text)
    table(doc, ["Szempont", "Altalanos webshop", "Epületgepeszeti webshop", "TDLWebshop MVP"], [
        ["Termekkatalogus", "Alap termeklista es kategoriak", "Muszaki kategoriak, keszletinformacio", "Kategoriak, termekkartyak, akcios termekek"],
        ["Admin folyamat", "Nem mindig lathato", "Keszlet es rendeleskezeles fontos", "Admin dashboard, CSV import, keszletfigyeles"],
        ["Rendeleskovetes", "Gyakori funkcio", "Ugyfelszolgalati szempontbol fontos", "Profilban rendelestortenet es statusz"],
        ["Bizonylat", "Szamlazo rendszerrel integralt", "Helyszini eladasnal is lenyeges", "PDF bizonylat szakdolgozati MVP-ben"],
        ["AI tamogatas", "Ritkabban domainhez kotott", "Hasznos lehet termekvalasztasnal", "Katalogushoz kotott AI asszisztens"],
    ], [4, 4, 4, 5])
    para(doc, "A TDLWebshop nem a piaci szereplok teljes funkcionalitasat probalja masolni. A cel inkabb az volt, hogy a webshop domainhez kotott legyen, es tartalmazzon olyan elemeket is, amelyek a szakdolgozatban mernoki szempontbol bemutathatok: adatmodell, jogosultsag, validacio, CI, teszteles, titokkezeles es reprodukalhatosag.")

    heading(doc, "3. Kovetelmenyek es felhasznaloi esetek", 1)
    para(doc, "A kovetelmenyek meghatarozasakor a rendszer fo szerepkorei szerint bontottam fel a feladatokat. A vasarlo a termekeket bongeszi, kosarat kezel, rendelest ad le es koveti a rendeleseit. A dolgozo helyszini vasarlasokat rogzithet, termekeket kezelhet es keszletet lathat. Az admin teljesebb jogosultsaggal rendelkezik: kezelheti a felhasznalokat, rendeleseket, kuponokat, keszletet es adminisztracios adatokat.")
    table(doc, ["Azonosito", "Kovetelmeny", "Use case", "Modul", "Teszt / bizonyitek"], [
        ["F-01", "Termekek listazasa es keresese", "Termek bongeszese", "Products page", "Kepernyokep, unit teszt"],
        ["F-02", "Kosar kezelese", "Kosar modositasa", "Cart service", "Kosar teszt"],
        ["F-03", "Rendeles leadasa validacioval", "Checkout", "Checkout page", "Checkout validacios teszt"],
        ["F-04", "Rendeles statusz kovetese", "Profil megtekintese", "Profile page, Order service", "Profil kepernyokep"],
        ["F-05", "Admin statuszvaltas audittal", "Rendeles kezelese", "Admin page, Order service", "Order service teszt"],
        ["F-06", "CSV termekimport", "Tomeges termekfeltoltes", "Admin CSV import", "Admin kepernyokep"],
        ["F-07", "PDF bizonylat", "Szamla letoltese", "Invoice service", "PDF kep es teszt"],
        ["F-08", "AI asszisztens", "Termekajanlas kerese", "Chatbot service, Worker", "AI kepernyokep"],
    ], [2, 5, 4, 4, 4])

    heading(doc, "3.1. Fo use case-ek", 2)
    table(doc, ["Use case", "Sikeres lefutas", "Hibaag", "Jogosultsag"], [
        ["Vasarloi checkout", "Kosarbol rendeles jon letre", "Hibas email/telefon vagy ures kosar", "Vendeg vagy bejelentkezett vasarlo"],
        ["Admin statuszvaltas", "Statusz es audit bejegyzes frissul", "Keszlethiany vagy jogosultsagi hiba", "Admin"],
        ["Helyszini vasarlas", "Mentett vasarlo es termekek alapjan rendeles/PDF keszul", "Tiltott vasarlo vagy keszlethiany", "Admin vagy dolgozo"],
        ["CSV import", "Valid sorok termekkent mentodnek", "Hibas ar, kategoria vagy SKU", "Admin/dolgozo jogosultsag"],
        ["AI asszisztens", "Domainhez kotott valasz es katalogus-talalat", "Nem relevans kerdesre korlatozo valasz", "Nyilvanos felulet"],
    ], [4, 5, 5, 4])
    figure(doc, "3. abra: Use case diagram a vasarlo, dolgozo es admin szerepkorokkal.", "A diagramot a kesz PlantUML/Mermaid segedanyag alapjan erdemes beszurni.")


def add_gui_tech_arch(doc):
    heading(doc, "4. GUI/UX tervezes", 1)
    for text in [
        "A felulet kialakitasanal a cel egy modern, de szakmai hangulatu webshop volt. A dark mode alapertelmezett megjelenese a TDLWebshop logo ipari, technologiai karakterehez igazodik. A light mode ugyanazt az elrendezest koveti, de vilagosabb, uzletiesebb megjelenessel.",
        "A navigacio a webshop alapfolyamataira epul: kezdolap, kategoriak, termekek, akciok, ujdonsagok, kapcsolat, kivansaglista, profil es kosar. A kategoriak lenyilo menut kaptak, mert ez az a pont, ahol sok alrendszer es termekcsoport jelenhet meg. A tobbi elem gombkent vagy egyszeru navigacios hivatkozaskent maradt, hogy a fejlec ne legyen tulzsufolt.",
        "A fo felhasznaloi utat a termeklista, termekadatlap, kosar es checkout kepernyok adjak. Ezeknel a prioritas az olvashatosag, a keszlet es ar kiemelese, valamint a hibas bemenetek egyertelmu jelzese volt. Az admin feluletnel ezzel szemben a surubb informacio es a gyors muveletvegzes kerult eloterbe.",
    ]:
        para(doc, text)
    for cap, note in [
        ("4. abra: Kezdolap dark mode-ban, nyitott kategoria menüvel.", "Webshop kezdolap, desktop nezeti kepernyokep."),
        ("5. abra: Kezdolap AI asszisztenssel.", "AI Segito gomb megnyitva, egy domain kerdesre adott valasszal."),
        ("6. abra: Termeklista kereses es szures kozben.", "Termekek oldal, kategoria/kereso allapot."),
        ("7. abra: Termekadatlap kepgaleriaval es kosar gombbal.", "Egy konkret termek oldala."),
        ("8. abra: Kosar oldal tobb termekkel.", "Mennyisegmodositas es vegosszeg lathato."),
        ("9. abra: Checkout validacios hiba.", "Hibas email vagy telefonszam megadasa."),
        ("10. abra: Sikeres rendeles utan megjeleno visszajelzes.", "Checkout sikeres allapot."),
        ("11. abra: Profil oldal rendelestortenettel.", "Bejelentkezett vasarlo profilja."),
        ("12. abra: Admin attekintes es rendeleskezeles.", "Admin dashboard desktop nezet."),
        ("13. abra: CSV termekimport admin feluleten.", "Admin termekkezeles, CSV import blokk."),
        ("14. abra: Helyszini vasarlas mentett vasarlo kivalasztasaval.", "Admin rendelesrogzites kepernyo."),
        ("15. abra: Generalt PDF bizonylat.", "PDF szamla/bizonylat letoltve."),
        ("16. abra: GitHub Actions zold CI futas.", "GitHub Actions kepernyokep a zold runrol."),
    ]:
        figure(doc, cap, note)

    heading(doc, "5. Technologiai hatter", 1)
    for text in [
        "Az alkalmazas frontend oldala Angularra epul. Az Angular valasztasat az indokolta, hogy komponensalapu, strukturalt es jol tesztelheto keretrendszer, amely alkalmas osszetettebb webshop es admin feluletek kialakitasara. A komponensek, szolgaltatasok es routing egyertelmu modulhatarokat adnak.",
        "Az adatbazis es hitelesites Firebase alapokon keszult. A Firestore dokumentumalapu adatmodellje jol illeszkedik a termekek, rendelesek, felhasznaloi profilok, kuponok es adminisztracios naplok tarolasahoz. A Firebase Authentication a bejelentkezes es szerepkor alapu hozzaferes egyik alapjat adja.",
        "A projektben a PDF generalas kliensoldali szolgaltatasbol tortenik, amely szakdolgozati MVP szinten elegendoen bemutatja a bizonylat kesziteset. Az AI asszisztens OpenRouter modellen keresztul mukodik, de az API kulcs nem kerul a kliensoldali kodba: a hivas Cloudflare Worker proxy-n at tortenik.",
    ]:
        para(doc, text)
    table(doc, ["Technologia", "Szerep", "Indoklas"], [
        ["Angular", "Frontend es komponenslogika", "Strukturalt, tesztelheto kliensalkalmazas"],
        ["Firebase Auth", "Hitelesites", "Gyors MVP es szerepkorok kezelese"],
        ["Firestore", "Adattarolas", "Dokumentumalapu webshop adatokhoz megfelelo"],
        ["Firebase Hosting", "Publikalas", "Egyszeru deploy es statikus hosting"],
        ["Cloudflare Worker", "OpenRouter proxy", "API kulcs vedelme Firebase Blaze nelkul"],
        ["GitHub Actions", "CI", "Build es teszt reprodukalhato ellenorzese"],
    ], [4, 5, 7])

    heading(doc, "6. Architektura", 1)
    para(doc, "Az architektura kliensoldali Angular alkalmazasra, Firebase szolgaltatasokra es egy kulso Worker proxy-ra bonthato. A frontend kezeli a felhasznaloi interakciokat, a szolgaltatasretegek pedig elkulonitik a kosar, rendeles, termek, kupon, profil, PDF es AI asszisztens logikajat. A Firestore biztonsagi szabalyok kulcsszerepet kapnak, mert ezek hatarozzak meg, hogy mely szerepkor milyen adathoz ferhet hozza.")
    figure(doc, "17. abra: Komponens architektura.", "Angular kliens, Firebase Auth, Firestore, Hosting, Cloudflare Worker es OpenRouter kapcsolata.")
    figure(doc, "18. abra: Checkout szekvencia diagram.", "Kosar ellenorzes, validacio, rendelesmentes, email/PDF es visszajelzes lepesei.")
    para(doc, "A rendszer egyik fontos tervezesi dontese, hogy a szerepkorok nem csak a feluleten jelennek meg, hanem a Firestore rules szinten is ervenyesulnek. Ez csokkenti annak kockazatat, hogy egy kozvetlen klienshivas megkerulje a feluleti korlatozasokat.")

    heading(doc, "7. Adatmodell", 1)
    para(doc, "Az adatmodell a webshop domain entitasaira epul. A Product a katalogus alapja, az Order es OrderItem a rendelesi adatokat fogja ossze, a UserProfile a felhasznaloi szerepkort es profiladatokat tartalmazza, a Coupon a kedvezmeny logikat, az Invoice pedig a bizonylathoz kapcsolodo adatokat kepviseli. Az admin oldalon SavedCustomer, audit naplok es keszletinformaciok is megjelennek.")
    table(doc, ["Entitas", "Fontos mezok", "Kapcsolat / szerep"], [
        ["Product", "name, sku, category, price, stock, images", "Termeklista, kosar, rendeles tetelek"],
        ["Order", "customer, items, totals, status, paymentMethod", "Checkout es admin rendeleskezeles"],
        ["OrderItem", "productId, name, quantity, unitPrice", "Rendeleshez tartozo tetelek"],
        ["UserProfile", "uid, email, role, disabled", "Jogosultsag es profil"],
        ["Coupon", "code, type, value, active", "Kedvezmeny validacio"],
        ["Invoice", "invoiceNumber, orderId, issuedAt", "PDF bizonylat"],
        ["SavedCustomer", "name, email, phone, company, disabled", "Helyszini vasarlas gyorsitasa"],
        ["Audit", "orderId, oldStatus, newStatus, changedBy", "Statuszvaltas nyomkovetese"],
    ], [4, 6, 6])
    figure(doc, "19. abra: Adatmodell diagram.", "Product, Order, OrderItem, UserProfile, Coupon, Invoice, SavedCustomer es Audit kapcsolatok.")


def add_implementation(doc):
    heading(doc, "8. Megvalositas", 1)
    para(doc, "A megvalositasi fejezetben azokat a reszeket mutatom be, amelyek szakdolgozati szempontbol a legtobb mernoki tartalmat hordozzak. Ilyen a checkout validacio es rendelesmentes, a rendelesstatusz audit es keszletkezeles, a PDF bizonylat, a CSV termekimport, a jogosultsagkezeles es az AI asszisztens proxy megoldasa.")

    heading(doc, "8.1. Checkout es rendelesmentes", 2)
    para(doc, "A checkout folyamat soran a rendszer ellenorzi a vasarloi adatokat, a kosar tartalmat, a fizetesi es szallitasi adatokat, majd a rendelesi objektumot Firestore-ba menti. A hibas email vagy telefonszam mezoszintu validacios visszajelzest ad, mert ez a felhasznaloi elmeny es az adatminoseg miatt is fontos.")
    code_place(doc, "Checkout rendeles veglegesitese", "src/pages/checkout/checkout.ts", "kb. 367. sortol", "Ez a resz mutatja be, hogyan all ossze a rendelesei adatcsomag, es hogyan kezeli a felulet a sikeres vagy hibas leadast.")

    heading(doc, "8.2. Rendelesstatusz, audit es keszlet", 2)
    para(doc, "Az admin oldali statuszvaltas nem egyszeru mezofrissiteskent jelenik meg. A statusz mellett audit informacio is keszul, ami visszakereshetove teszi, hogy ki es mikor modositta a rendelest. A keszletkezeles a helyszini ertekesitesnel szigorubb, mert ott a termek azonnal kikerul a keszletbol.")
    code_place(doc, "Statusz, audit es keszlet tranzakcio", "src/app/services/order.service.ts", "41-127", "A kodreszlet azt bizonyitja, hogy a rendelesallapot es kapcsolodo adminisztracios adatok egyutt kezelodnek.")
    code_place(doc, "Helyszini rendeles mentese", "src/app/services/order.service.ts", "222-267", "Ez a resz a mentett vasarlo es a helyszini tetelek alapjan letrehozott rendelest mutatja be.")
    code_place(doc, "Szamlaszam generalas", "src/app/services/order.service.ts", "269-302", "A szamlaszam kepzese a PDF bizonylat es rendeles adminisztracio miatt fontos.")

    heading(doc, "8.3. PDF bizonylat", 2)
    para(doc, "A PDF bizonylat celja, hogy a rendeleshez ember altal olvashato, letoltheto dokumentum keszuljon. A dolgozatban ezt nem NAV-kompatibilis eles szamlazokent kezelem, hanem olyan MVP bizonylatkent, amely demonstralja a rendelesei adatok dokumentumma alakitását.")
    code_place(doc, "PDF szamla felepitese", "src/app/services/invoice.service.ts", "9-154", "A kodreszletben latszik a fejlec, vevo, elado, tetelek es osszesites elrendezese.")

    heading(doc, "8.4. CSV import es admin termekkezeles", 2)
    para(doc, "A CSV import a termekfeltoltes gyakorlati problemajara ad valaszt. Sok termek kezi felvitele lassu es hibalehetosegekkel teli lenne, ezert a rendszer kepes validalt CSV sorok alapjan termekeket felvenni vagy frissiteni. Az importnal a hibas sorok visszajelzest kapnak, a valid sorok pedig menthetok.")
    code_place(doc, "CSV import validacio es mentes", "src/pages/admin/admin.ts", "1181-1257", "Ez a resz mutatja be, hogyan kulonulnek el a hibas es ervenyes import sorok.")

    heading(doc, "8.5. Jogosultsagkezeles", 2)
    para(doc, "A jogosultsagkezelesben harom fo szerepkor jelenik meg: admin, dolgozo es vasarlo. Az admin teljesebb hozzaferessel rendelkezik, a dolgozo operativ muveleteket vegezhet, a vasarlo pedig sajat profiljat es rendeleseit kezeli. A tiltott felhasznalo kulon allapot, amely bejelentkezes vagy helyszini vasarlas eseteben is korlatozast jelent.")
    code_place(doc, "Admin es dolgozoi jogosultsagok", "src/pages/admin/admin.ts", "607-748", "A kodreszlet a feluleti jogosultsagi logika es szerepkorok megjelenitesenek bemutatasara hasznalhato.")
    code_place(doc, "Firestore aktiv felhasznalo es szerepkor szabalyok", "firestore.rules", "25-76", "A szabalyok bizonyitjak, hogy a jogosultsag nem csak kliensoldali elem.")
    code_place(doc, "Firestore products/orders/users szabalyok", "firestore.rules", "294-361", "A legfontosabb collection-szintu jogosultsagi dontesek itt jelennek meg.")

    heading(doc, "8.6. AI asszisztens", 2)
    para(doc, "Az AI asszisztens celja, hogy a vasarlo szakmai vagy katalogushoz kapcsolodo kerdesekben gyors iranymutatast kapjon. A rendszer nem helyettesiti a szakembert, es nem ad kotelezo ervenyu muszaki dontest. Ha nincs pontos katalogus-talalat, altalanos iranyt ad, es jelzi, hogy pontos ajanlatert vagy beszerezhetosegert emailben vagy szemelyesen erdemes egyeztetni.")
    para(doc, "A megoldas lenyege, hogy az OpenRouter API kulcs nem kerul a bongeszobe. A kliens a Cloudflare Worker proxy-t hivja, amely szerveroldali titokkent kezeli a kulcsot. Ez a Firebase Spark csomag korlatai mellett is biztonsagosabb irany, mint a kulcs kliensoldali tarolasa.")
    code_place(doc, "AI asszisztens domain- es kataloguslogika", "src/app/services/chatbot-llm.service.ts", "26-84 es 217-236", "A kodreszlet azt mutatja, hogyan korlatozodik az asszisztens a webshop es epületgepeszeti temak kore.")
    code_place(doc, "OpenRouter proxy", "workers/openrouter-proxy/src/index.js", "teljes fajl roviditett reszlete", "Ez a resz a szerveroldali kulcskezeles es CORS ellenorzes miatt fontos.")


def add_security_testing(doc):
    heading(doc, "9. Biztonsagi megfontolasok", 1)
    for text in [
        "A rendszer biztonsagi modellje tobb retegbol all. A felhasznaloi feluleten csak a szerepkornek megfelelo muveletek jelennek meg, de ez onmagaban nem elegendo. Ezert a Firestore rules szinten is szerepkor es aktiv felhasznalo szerinti hozzaferes kerult kialakitasra.",
        "A titokkezelesnel kulon szempont volt, hogy API kulcs, jelszo, token vagy .env fajl ne keruljon a publikus repoba. A projektben .env.example szerepel, amely csak a szukseges valtozok nevet mutatja. Az OpenRouter kulcs Cloudflare Worker secretkent kerul beallitasra.",
        "A projekt MVP jellege miatt vannak tudatos korlatok is. A webes checkoutnal bizonyos ar- es kosaradatok kliensoldalrol indulnak, ezert eles kereskedelmi rendszerben erosebb szerveroldali ar- es keszletellenorzesre lenne szukseg. A dolgozatban ezt korlatkent es tovabbfejlesztesi iranykent kezelem.",
    ]:
        para(doc, text)
    table(doc, ["Kockazat", "Kezeles az MVP-ben", "Tovabbfejlesztes"], [
        ["Jogosulatlan admin muvelet", "Firestore rules es admin guard", "Custom claim es token revocation"],
        ["API kulcs kiszivarog", "Worker secret, nincs kliensoldali kulcs", "Rate limit es kulcsrotacio"],
        ["Kupon visszaeles", "Validacio es aktiv allapot ellenorzes", "Szerveroldali teljes ujraszamolas"],
        ["Keszleteltérés", "Admin/helyszini tranzakcios logika", "Checkout szerveroldali lockolas"],
        ["Szemelyes adatok vedelme", "Szerepkor szerinti olvasas/iras", "Adatmegorzesi szabalyok formalizalasa"],
    ], [4, 6, 6])

    heading(doc, "10. Teszteles es validacio", 1)
    para(doc, "A teszteles celja az volt, hogy a kritikus folyamatok mukodese bizonyithato legyen. A projektben automata tesztek es kezi tesztjegyzokonyv egyarant szerepelnek. Az automata tesztek a kod bizonyos logikai reszeit ellenorzik, mig a kezi tesztek a felhasznaloi folyamatok vegigjarasat dokumentaljak.")
    table(doc, ["Tesztelt terulet", "Ellenorzes modja", "Elvart eredmeny"], [
        ["Build", "npm run build", "Sikeres production build"],
        ["Unit tesztek", "npm test -- --watch=false", "Minden relevans teszt zold"],
        ["Checkout validacio", "Kezi + unit", "Hibas adatnal nem enged tovabb"],
        ["Kuponlogika", "Unit teszt", "Ervenyes/ervenytelen kupon kezelese"],
        ["Admin statuszvaltas", "Unit + kezi teszt", "Audit es statusz frissul"],
        ["PDF generalas", "Unit + kezi teszt", "Letoltheto, olvashato PDF"],
        ["AI asszisztens", "Kezi teszt", "Domainhez kotott valasz, irrelevans kerdes szurese"],
        ["CI", "GitHub Actions", "Zold workflow run"],
    ], [4, 5, 7])
    figure(doc, "20. abra: GitHub Actions sikeres CI futas.", "Kepernyokep a legfrissebb zold workflow runrol.")
    figure(doc, "21. abra: Kezi teszt checklist reszlete.", "Kitoltott tesztjegyzokonyv reszlet a vegleges mellekletbol.")

    heading(doc, "10.1. Kezi tesztforgatokonyvek", 2)
    table(doc, ["Azonosito", "Forgatokonyv", "Lepesek roviden", "Eredmeny"], [
        ["T-01", "Regisztracio es bejelentkezes", "Uj felhasznalo letrehozasa, bejelentkezes, profil megnyitasa", "Sikeres"],
        ["T-02", "Termekkereses", "Kategoria es keresoszo hasznalata", "Talalatok szurodnek"],
        ["T-03", "Kosar muveletek", "Hozzaadas, mennyisegvaltas, torles", "Osszeg frissul"],
        ["T-04", "Checkout hibas adattal", "Hibas email/telefon megadasa", "Hibauzenet jelenik meg"],
        ["T-05", "Admin CSV import", "CSV feltoltes, validacio, mentes", "Valid sorok mentodnek"],
        ["T-06", "Helyszini vasarlas", "Mentett vasarlo es termek kivalasztasa", "Rendeles es PDF keszul"],
        ["T-07", "Dolgozoi jogosultsag", "Dolgozoi fiokkal admin oldal megnyitasa", "Csak engedelyezett funkciok lathatok"],
        ["T-08", "AI asszisztens", "Domain es nem domain kerdes feltevese", "Korlatozott, megfelelo valasz"],
    ], [2, 4, 7, 3])


def add_repro_ai_summary(doc):
    heading(doc, "11. Reprodukalhatosag es repo-higienia", 1)
    for text in [
        "A konzulensi visszajelzes alapjan kulon figyelmet kapott a repo tisztasaga. A vegleges repoban a forraskod, konfiguracios mintafajlok, tesztek, dokumentacio es CI workflow maradnak. A generalt vagy gepfuggo allomanyok, peldaul node_modules, build mappak, lokalis cache-ek es szemelyes segedanyagok nem a beadando repo reszei.",
        "A README feladata, hogy a biralo vagy temavezeto tiszta kornyezetben is el tudja inditani az alkalmazast. Ezert tartalmaznia kell a telepitesi lepeseket, az npm parancsokat, a Firebase es Worker konfiguracio lenyeget, a demo szerepkorok leirasat, a tesztek futtatasat es a deploy alapjait.",
    ]:
        para(doc, text)
    table(doc, ["Parancs / fajl", "Szerep"], [
        ["npm install", "Fuggosegek telepitese"],
        ["npm run build", "Production build ellenorzese"],
        ["npm test -- --watch=false", "Automata tesztek futtatasa"],
        [".env.example", "Szukseges kornyezeti valtozok bemutatasa titkok nelkul"],
        [".github/workflows/ci.yml", "GitHub Actions build es teszt pipeline"],
        ["workers/openrouter-proxy", "AI proxy kulon deployolhato resze"],
    ], [5, 10])

    heading(doc, "12. Mesterseges intelligencia hasznalata a fejlesztes soran", 1)
    for text in [
        "A fejlesztes soran mesterseges intelligenciat ket kulon szerepben hasznaltam. Az egyik a fejlesztest tamogato eszkozhasznalat volt, ahol az AI otletelesben, hibakeresesben, kodreview jellegu ellenorzesben es dokumentacios vazlatok kesziteseben segitett. A masik maga a rendszerbe beepitett AI asszisztens, amely a TDLWebshop felhasznaloi feluleten jelenik meg.",
        "A fejlesztest tamogato AI-hasznalatnal az eredmenyeket nem vegleges igazsagkent kezeltem. A javasolt megoldasokat futtatassal, builddel, tesztekkel, sajat ellenorzessel es a konzulensi visszajelzesek alapjan validaltam. Voltak olyan esetek, amikor egy javasolt megoldast modositani kellett, peldaul jogosultsagi, titokkezelesi vagy UI elrendezesi okok miatt.",
        "Az AI-t nem arra hasznaltam, hogy a szakmai felelosseget atadjam. A projekt donteseiert, a kod mukodeseert, a tesztelesert es a vegleges dolgozati allitasokert en felelek. A vegleges szoveget sajat megfogalmazasra kell atirni, mert a dolgozatnak nem csak tartalmilag, hanem szemelyes szakmai reflexiokent is vallalhatonak kell lennie.",
        "A beepitett AI asszisztensnel kulon korlat, hogy nem adhat kotelezo ervenyu muszaki vagy jogi tanacsot. Ha nincs pontos termek vagy elegendo meretezesi adat, akkor altalanos iranyt ad, es felhivja a figyelmet arra, hogy pontos ajanlatert vagy beszerezhetosegert emailben vagy szemelyesen erdemes egyeztetni.",
    ]:
        para(doc, text)
    table(doc, ["AI-hasznalat terulete", "Mire hasznaltam", "Validacio"], [
        ["Otleteles", "Funkcio- es dokumentacios javaslatok", "Konzulensi elvarasokhoz illesztes"],
        ["Kodreview", "Hibalehetosegek keresese", "Futtatas, teszt, kezi ellenorzes"],
        ["Dokumentacio", "Vazlatok es strukturak", "Sajat nyelvre atiras es ellenorzes"],
        ["Beepitett AI", "Felhasznaloi asszisztens", "Domainkorlatozas es proxy kulcsvedelem"],
    ], [4, 6, 6])

    heading(doc, "13. Osszefoglalas", 1)
    for text in [
        "A szakdolgozat eredmenyekent elkeszult a TDLWebshop, amely egy epületgepeszeti termekekre specializalt webshop es adminisztracios rendszer MVP-je. A projekt fo erteke, hogy nem csak a vasarloi oldalt valositja meg, hanem az adminisztracios folyamatokat, jogosultsagokat, CSV importot, PDF bizonylatot, keszletfigyelest es AI asszisztenst is bemutatja.",
        "A fejlesztes soran megtapasztaltam, hogy egy webshopnal a felhasznaloi felulet mellett legalabb olyan fontos az adatmodell, a jogosultsagkezeles, a validacio, a teszteles es a reprodukalhato mukodes. A legnehezebb reszek koze a rendelesi logika, a keszletkezeles, a jogosultsagok es a dolgozati bizonyitas osszeallitasa tartoztak.",
        "A rendszer jelenlegi allapotaban szakdolgozati MVP-kent vedheto. Eles kereskedelmi hasznalat elott tovabbi fejlesztesre lenne szukseg, peldaul szerveroldali ar- es keszletellenorzesre, fizetesi szolgaltatoi integraciora, erosebb rate limitre, adatmegorzesi szabalyokra es teljesebb admin riportokra. Ezeket a dolgozat tovabbfejlesztesi iranykent kezeli.",
        "[IDE KERUL A SAJAT ZARO REFLEXIOD: mit tanultal, mi volt nehez, mit csinalnal maskent, mire vagy buszke a projektben.]",
    ]:
        para(doc, text)

    heading(doc, "Irodalomjegyzek", 1)
    for item in [
        "Angular dokumentacio: https://angular.dev/",
        "Firebase dokumentacio: https://firebase.google.com/docs",
        "Cloudflare Workers dokumentacio: https://developers.cloudflare.com/workers/",
        "OpenRouter dokumentacio: https://openrouter.ai/docs",
        "GitHub Actions dokumentacio: https://docs.github.com/actions",
        "Konzulensi mintacsomag es szakdolgozati tartalmi segedanyagok.",
        "[IDE KERULNEK A PIACI OSSZEHASONLITASBAN HASZNALT WEBSHOPOK ES EGYEB FORRASOK.]",
    ]:
        para(doc, item)

    heading(doc, "Mellekletek", 1)
    para(doc, "A mellekletbe kerulhet a kezi tesztjegyzokonyv, a hosszabb kodreszletek, az adatmodell teljes valtozata, valamint a fontosabb konfiguracios reszletek titkok nelkul.")


def main():
    doc = Document()
    configure(doc)
    add_cover(doc)
    add_front_matter(doc)
    add_intro(doc)
    add_market_requirements(doc)
    add_gui_tech_arch(doc)
    add_implementation(doc)
    add_security_testing(doc)
    add_repro_ai_summary(doc)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
