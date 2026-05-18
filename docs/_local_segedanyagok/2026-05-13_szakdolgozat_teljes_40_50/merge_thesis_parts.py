from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK


BASE_DIR = Path(__file__).resolve().parent
PART_1 = BASE_DIR / "TDLWebshop_szakdolgozat_1_resz_elmeleti_tervezesi_alap.docx"
PART_2 = BASE_DIR / "TDLWebshop_szakdolgozat_2_resz_megvalositas_teszteles_zaras.docx"
OUT = BASE_DIR / "TDLWebshop_szakdolgozat_egyesitett_munkapeldany.docx"


def append_document(target: Document, source: Document) -> None:
    target_body = target.element.body
    source_body = source.element.body

    # The final sectPr element stores page settings; keep the target document's
    # section settings and append only real content nodes from the second file.
    for child in source_body:
        if child.tag.endswith("}sectPr"):
            continue
        target_body.append(deepcopy(child))


def main() -> None:
    target = Document(PART_1)
    source = Document(PART_2)

    page_break_paragraph = target.add_paragraph()
    page_break_paragraph.add_run().add_break(WD_BREAK.PAGE)
    target.add_paragraph("II. rész - Megvalósítás, tesztelés és zárás")

    append_document(target, source)
    target.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
