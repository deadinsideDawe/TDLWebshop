from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt


OUT_DIR = Path(__file__).resolve().parent
DOCX_PATH = OUT_DIR / "TDLWebshop_word_beillesztesi_utmutato_kepek_kodok.docx"
MD_PATH = OUT_DIR / "TDLWebshop_word_beillesztesi_utmutato_kepek_kodok.md"

PROJECT_ROOT = Path(r"C:\Users\Dell\webshop")
FIG_DIR = PROJECT_ROOT / "docs" / "_local_segedanyagok" / "word_abrak"


def p(path: Path) -> str:
    return str(path)


SCREENSHOTS = [
    ("Kezdőlap sötét módban, kategória lenyílóval", "GUI/UX fejezet, kezdőlap bemutatása", "Nyisd meg a főoldalt dark módban, a Kategóriák menü legyen lenyitva."),
    ("Kezdőlap AI asszisztenssel", "GUI/UX vagy AI-asszisztens fejezet", "A főoldalon legyen nyitva az AI ablak, egy szakmai kérdésre adott válasszal."),
    ("Terméklista kereséssel és szűréssel", "GUI/UX fejezet, termékböngészés", "A termékek oldalon legyen látható keresés vagy kategóriaszűrés."),
    ("Termékadatlap", "GUI/UX fejezet, termékrészletek", "Nyiss meg egy konkrét terméket, látszódjon kép, ár, készlet és kosár gomb."),
    ("Kosár több termékkel", "GUI/UX fejezet, kosárkezelés", "A kosárban legalább két termék legyen, látszódjon mennyiség és összegzés."),
    ("Checkout validáció hibás e-mail vagy telefonszám példával", "GUI/UX / validáció fejezet", "Csak akkor tedd be, ha a validáció tényleg működik. Ha még hibás, inkább MVP-korlátként írd le."),
    ("Sikeres rendelés / rendelés visszaigazolása", "Megvalósítás vagy tesztelés fejezet", "Rendelésleadás után látszódjon a sikeres állapot."),
    ("Profil és rendeléstörténet", "GUI/UX vagy megvalósítás fejezet", "Bejelentkezett vásárló profiloldala korábbi rendelésekkel."),
    ("Kívánságlista", "GUI/UX fejezet", "Egy-két kedvencként megjelölt termék látszódjon."),
    ("Admin áttekintés", "Admin funkciók fejezet", "Admin főnézet statisztikákkal vagy rendelésáttekintéssel."),
    ("Admin termékkezelés és CSV import", "Megvalósítás fejezet, admin termékkezelés", "Látszódjon a CSV import blokk és terméklista."),
    ("Admin rendeléskezelés státuszváltással", "Megvalósítás / tesztelés fejezet", "Egy rendelésnél látszódjon státusz módosítási lehetőség."),
    ("Helyszíni vásárlás mentett vásárlóval", "Admin / helyszíni értékesítés fejezet", "Mentett vásárló kiválasztása és tételek hozzáadása legyen látható."),
    ("PDF számla vagy bizonylat", "Megvalósítás fejezet, PDF generálás", "A generált PDF első oldala legyen képként beszúrva."),
    ("GitHub Actions zöld CI", "Tesztelés és reprodukálhatóság fejezet", "GitHub Actions oldalon a legfrissebb zöld futás legyen látható."),
]


FIGURES = [
    ("Komponens-architektúra ábra", FIG_DIR / "01_komponens_architektura.svg", "Architektúra fejezet", "A TDLWebshop komponens-architektúrája."),
    ("Checkout szekvencia ábra", FIG_DIR / "02_checkout_szekvencia.svg", "Architektúra vagy megvalósítás fejezet", "A checkout folyamat szekvenciája."),
    ("Adatmodell diagram", FIG_DIR / "03_adatmodell_diagram.svg", "Adatmodell fejezet", "A TDLWebshop fő adatmodellje."),
    ("Jogosultsági modell", FIG_DIR / "04_jogosultsagi_modell.svg", "Biztonság fejezet", "Szerepkörök és jogosultsági szintek."),
    ("Tesztelési bizonyítási folyamat", FIG_DIR / "05_teszteles_bizonyitasi_folyamat.svg", "Tesztelés fejezet", "A tesztelési és validációs folyamat áttekintése."),
    ("AI asszisztens adatfolyam", FIG_DIR / "06_ai_asszisztens_adatfolyam.svg", "AI-asszisztens / megvalósítás fejezet", "Az AI asszisztens adatfolyama."),
    ("Piaci összehasonlító táblázat", FIG_DIR / "07_piaci_osszehasonlito_tablazat.md", "Piaci összehasonlítás fejezet", "Épületgépészeti webshopok összehasonlítása saját piackutatás alapján."),
    ("Piaci összehasonlító ábra", FIG_DIR / "07_piaci_osszehasonlito_abra.svg", "Piaci összehasonlítás fejezet", "A TDLWebshop pozicionálása a vizsgált rendszerekhez képest."),
    ("Use case diagram", FIG_DIR / "08_use_case_diagram.svg", "Use case fejezet", "A TDLWebshop fő use case-ei és szereplői."),
    ("Követelmény-traceability táblázat", FIG_DIR / "09_kovetelmeny_traceability_tablazat.md", "Követelmények fejezet", "Követelmény-traceability összefoglaló."),
]


CODE_SNIPPETS = [
    ("Checkout validáció és rendelésmentés", PROJECT_ROOT / "src" / "pages" / "checkout" / "checkout.ts", "393-422 és 575-651", "A checkout űrlap ellenőrzése, rendelés előkészítése és mentése."),
    ("Rendelés státusz, audit és készlet tranzakció", PROJECT_ROOT / "src" / "app" / "services" / "order.service.ts", "41-126", "A státuszváltás, auditnapló és készletmódosítás összetartozó kezelése."),
    ("Helyszíni rendelés tranzakció", PROJECT_ROOT / "src" / "app" / "services" / "order.service.ts", "238-284", "Helyszíni vásárlás mentése készletellenőrzéssel."),
    ("Számlaszám generálása", PROJECT_ROOT / "src" / "app" / "services" / "order.service.ts", "313-346", "Éves számlasorszám előállítása Firestore tranzakcióval."),
    ("PDF számla felépítése", PROJECT_ROOT / "src" / "app" / "services" / "invoice.service.ts", "81-154", "PDF bizonylat szerkezete: fejlécek, vevő, tételek, összegek."),
    ("Firestore jogosultsági segédfüggvények", PROJECT_ROOT / "firestore.rules", "12-67", "Aktív felhasználó, admin és dolgozó szerepkör ellenőrzése."),
    ("Firestore fő collection szabályok", PROJECT_ROOT / "firestore.rules", "288-356", "Products, orders, users, customerProfiles és audit szabályok."),
    ("Admin és dolgozói jogosultságok", PROJECT_ROOT / "src" / "pages" / "admin" / "admin.ts", "614-700", "Admin oldali jogosultsági logika."),
    ("CSV import validáció és mentés", PROJECT_ROOT / "src" / "pages" / "admin" / "admin.ts", "1184-1258", "Tömeges termékfeltöltés validációja és mentése."),
    ("AI asszisztens kliensoldali logika", PROJECT_ROOT / "src" / "app" / "services" / "chatbot-llm.service.ts", "31-84", "Katalógushoz kötött AI kérés előkészítése, modellválasztás nélkül a felületen."),
    ("OpenRouter proxy", PROJECT_ROOT / "workers" / "openrouter-proxy" / "src" / "index.js", "153-213", "Szerveroldali API kulcskezelés, CORS és OpenRouter hívás."),
]


def markdown_content() -> str:
    lines = []
    lines.append("# TDLWebshop - Word beillesztési útmutató képekhez, ábrákhoz és kódrészletekhez")
    lines.append("")
    lines.append("Ez a fájl külön segédanyag. Nem beadandó szöveg, hanem munkalista ahhoz, hogy a `TDLWebshop_szakdolgozat_javitott_v2.docx` dokumentumban a placeholdereket tényleges tartalomra tudd cserélni.")
    lines.append("")
    lines.append("## Fontos arány")
    lines.append("")
    lines.append("- Nem kell minden kódrészletet és minden képernyőt betenni.")
    lines.append("- Review-ready állapothoz elég kb. 12-15 képernyőkép, 4-6 ábra/táblázat és 5-7 rövid kódrészlet.")
    lines.append("- A kódrészlet ne legyen túl hosszú: fél oldal, legfeljebb egy oldal. Mindig legyen előtte vagy utána 3-5 mondatos magyarázat.")
    lines.append("- A végleges PDF-ben ne maradjon ilyen szöveg: `[KÉP / ÁBRA HELYE]`, `[KÓDRÉSZLET HELYE]`, `ide tedd`, `placeholder`, `ÁTÍRVA`, emoji vagy munkautasítás.")
    lines.append("")
    lines.append("## Ábrák és táblázatok")
    lines.append("")
    for name, path, where, caption in FIGURES:
        lines.append(f"### {name}")
        lines.append(f"- Hova: {where}")
        lines.append(f"- Fájl: `{p(path)}`")
        lines.append(f"- Javasolt ábrafelirat: {caption}")
        lines.append("")
    lines.append("## Képernyőképek")
    lines.append("")
    for title, where, how in SCREENSHOTS:
        lines.append(f"### {title}")
        lines.append(f"- Hova: {where}")
        lines.append(f"- Hogyan: {how}")
        lines.append(f"- Javasolt felirat: {title}.")
        lines.append("")
    lines.append("## Kódrészletek")
    lines.append("")
    lines.append("A dolgozatba 5-7 darabot válassz ki. A legerősebb kombináció: checkout, OrderService tranzakció, Firestore rules, PDF generálás, CSV import, OpenRouter proxy.")
    lines.append("")
    for title, path, rows, why in CODE_SNIPPETS:
        lines.append(f"### {title}")
        lines.append(f"- Fájl: `{p(path)}`")
        lines.append(f"- Sorok: {rows}")
        lines.append(f"- Miért érdekes: {why}")
        lines.append("")
    lines.append("## Javasolt végső sorrend a Wordben")
    lines.append("")
    lines.append("1. Címlap, témavezető, A4, tartalomjegyzék rendbetétele.")
    lines.append("2. Piaci összehasonlító táblázat és követelmény-traceability táblázat beillesztése.")
    lines.append("3. Use case, komponens-architektúra, adatmodell és checkout szekvencia ábra beillesztése.")
    lines.append("4. GUI/UX képernyőképek beszúrása.")
    lines.append("5. Megvalósítás fejezetben 5-7 rövid kódrészlet beszúrása.")
    lines.append("6. Tesztelési fejezethez GitHub Actions zöld CI és kézi tesztelési eredmények hozzáadása.")
    lines.append("7. Összefoglalás végére saját reflexió megírása: mit tanultál, mi volt nehéz, mit csinálnál másként.")
    lines.append("")
    lines.append("## Kritikus megjegyzés")
    lines.append("")
    lines.append("Ha egy funkcióról tudod, hogy még hibás, ne mutasd be kész funkcióként. Vagy javítsd ki előtte, vagy írd le őszintén MVP-korlátként. Ez szakmailag sokkal védhetőbb, mint egy hibás működés eltakarása.")
    return "\n".join(lines)


def set_document_style(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    styles = doc.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"].font.size = Pt(11)
    for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
        styles[style_name].font.name = "Times New Roman"


def add_bullets(doc: Document, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_table(doc: Document, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for idx, header in enumerate(headers):
        hdr[idx].text = header
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value
    return table


def create_docx():
    doc = Document()
    set_document_style(doc)

    title = doc.add_heading("TDLWebshop - Word beillesztési útmutató", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_paragraph("Képek, ábrák, táblázatok és kódrészletek helye a szakdolgozatban")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(
        "Ez a dokumentum külön segédanyag. Nem beadandó szöveg, hanem munkalista ahhoz, "
        "hogy a szakdolgozat Word fájljában a placeholdereket tényleges bizonyítékokra "
        "tudd cserélni. A fő cél az, hogy a konzulensi visszajelzésben kért ábrák, "
        "képernyőképek, kódrészletek és bizonyítékok rendezett formában bekerüljenek."
    )

    doc.add_heading("1. Fontos arány", level=1)
    add_bullets(
        doc,
        [
            "Nem kell minden kódrészletet és minden képernyőt betenni.",
            "Review-ready állapothoz kb. 12-15 képernyőkép, 4-6 ábra/táblázat és 5-7 rövid kódrészlet elég.",
            "A kódrészlet fél oldal, legfeljebb egy oldal legyen, és mindig tartozzon hozzá magyarázat.",
            "A végleges PDF-ben ne maradjon placeholder, munkautasítás vagy AI-jellegű szerkesztési megjegyzés.",
        ],
    )

    doc.add_heading("2. Ábrák és táblázatok", level=1)
    add_table(
        doc,
        ["Elem", "Hova kerüljön", "Fájl", "Javasolt felirat"],
        [(name, where, p(path), caption) for name, path, where, caption in FIGURES],
    )

    doc.add_page_break()
    doc.add_heading("3. Képernyőképek", level=1)
    doc.add_paragraph(
        "A képernyőképek célja nem az, hogy minden oldalt túlzsúfoljanak, hanem hogy "
        "bizonyítsák a fő felhasználói és admin folyamatok működését. Az alábbi lista "
        "prioritási sorrendként is használható."
    )
    add_table(
        doc,
        ["Képernyőkép", "Hova kerüljön", "Hogyan készítsd el"],
        [(title, where, how) for title, where, how in SCREENSHOTS],
    )

    doc.add_page_break()
    doc.add_heading("4. Kódrészletek", level=1)
    doc.add_paragraph(
        "A dolgozatba 5-7 kódrészletet érdemes betenni. A legerősebb kombináció: "
        "checkout validáció, OrderService tranzakció, Firestore rules, PDF generálás, "
        "CSV import és OpenRouter proxy. A többi csak akkor kell, ha marad hely."
    )
    add_table(
        doc,
        ["Kódrészlet", "Fájl", "Sorok", "Miért érdekes"],
        [(title, p(path), rows, why) for title, path, rows, why in CODE_SNIPPETS],
    )

    doc.add_heading("5. Javasolt végső sorrend a Wordben", level=1)
    add_bullets(
        doc,
        [
            "Címlap, témavezető, A4-es lapméret és tartalomjegyzék rendbetétele.",
            "Piaci összehasonlító táblázat és követelmény-traceability táblázat beillesztése.",
            "Use case, komponens-architektúra, adatmodell és checkout szekvencia ábra beillesztése.",
            "GUI/UX képernyőképek beszúrása a megfelelő fejezetekhez.",
            "Megvalósítás fejezetben 5-7 rövid kódrészlet beszúrása.",
            "Tesztelési fejezethez GitHub Actions zöld CI és kézi tesztelési eredmények hozzáadása.",
            "Összefoglalás végére saját reflexió: mit tanultál, mi volt nehéz, mit csinálnál másként.",
        ],
    )

    doc.add_heading("6. Kritikus megjegyzés", level=1)
    doc.add_paragraph(
        "Ha egy funkcióról tudod, hogy még hibás, ne mutasd be kész funkcióként. "
        "Vagy javítsd ki előtte, vagy írd le őszintén MVP-korlátként. Ez szakmailag "
        "sokkal védhetőbb, mint egy hibás működés eltakarása."
    )

    doc.save(DOCX_PATH)


def main():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    MD_PATH.write_text(markdown_content(), encoding="utf-8")
    create_docx()
    print(f"Created: {DOCX_PATH}")
    print(f"Created: {MD_PATH}")


if __name__ == "__main__":
    main()
