from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.shared import Cm, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


ROOT = Path(r"C:\Users\Dell\webshop")
OUT_DIR = ROOT / "docs" / "_local_segedanyagok" / "word_beillesztesi_utmutato"
OUT_DOCX = OUT_DIR / "TDLWebshop_word_beillesztesi_utmutato_kepek_kodok.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(9)


def add_note(document, text):
    p = document.add_paragraph()
    p.style = "Intense Quote"
    run = p.add_run(text)
    run.font.size = Pt(10)


def add_bullets(document, items):
    for item in items:
        p = document.add_paragraph(style="List Bullet")
        p.add_run(item)


def add_table(document, headers, rows):
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, header in enumerate(headers):
        set_cell_text(hdr[i], header, bold=True)
        set_cell_shading(hdr[i], "D9EAF7")
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], str(value))
    document.add_paragraph()
    return table


def add_code_style(document):
    styles = document.styles
    if "Code Block" not in styles:
        style = styles.add_style("Code Block", 1)
        style.font.name = "Consolas"
        style.font.size = Pt(9)
        style.font.color.rgb = RGBColor(30, 41, 59)
        style.paragraph_format.left_indent = Cm(0.5)
        style.paragraph_format.space_after = Pt(6)


def main():
    OUT_DIR.mkdir(parents=True, exist_ok=True)

    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)
    add_code_style(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("TDLWebshop - Word beillesztési útmutató")
    run.bold = True
    run.font.size = Pt(18)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("Képernyőképek, ábrák, kódrészletek és fejezeti beillesztési javaslatok").italic = True

    add_note(
        doc,
        "Ez a dokumentum külön segédanyag. A szakdolgozat fő Word fájlját nem módosítja, "
        "csak azt mutatja meg, hogy a TDLWebshop_szakdolgozat_javitott_v2.docx dokumentumban "
        "hova milyen bizonyítékot érdemes beilleszteni."
    )

    doc.add_heading("1. Általános beillesztési szabályok", level=1)
    add_bullets(doc, [
        "A végleges szakdolgozatból töröld az összes placeholdert: [KÉP / ÁBRA HELYE], [KÓDRÉSZLET HELYE], ide illeszthető, beillesztendő tartalom.",
        "Review-ready állapothoz általában elég 12-15 képernyőkép, 4-6 ábra vagy táblázat és 5-7 rövid kódrészlet.",
        "A kódrészletek ne legyenek forráskód-dumpok: egy részlet fél oldal, legfeljebb egy oldal legyen.",
        "Minden ábra, kép és táblázat kapjon számozott feliratot, és a szövegben legyen rá hivatkozás.",
        "Ha egy funkció még nem működik tökéletesen, azt ne bizonyítékként mutasd be, hanem vagy javítsuk, vagy írd le MVP-korlátként."
    ])

    doc.add_heading("2. Fejezetenkénti beillesztési terv", level=1)
    add_table(doc, ["Fejezet", "Mit tegyél be?", "Forrás / hely", "Rövid magyarázat"], [
        ["Bevezetés, célkitűzés", "Nem szükséges kép", "-", "A probléma, cél és MVP-határ legyen világos."],
        ["Piaci összehasonlítás", "Piaci összehasonlító táblázat vagy ábra", r"docs\_local_segedanyagok\word_abrak\07_piaci_osszehasonlito_tablazat.svg", "Mutasd meg, miben tér el a TDLWebshop a hasonló rendszerektől."],
        ["Követelmények", "Traceability táblázat", r"docs\_local_segedanyagok\word_abrak\09_kovetelmeny_traceability_tablazat.md", "Követelmény, use case, modul és teszt kapcsolatát bizonyítja."],
        ["Use case-ek", "Use case diagram", r"docs\_local_segedanyagok\word_abrak\08_use_case_diagram.svg", "Vásárló, admin, dolgozó és AI-asszisztens szerepkörök bemutatása."],
        ["GUI/UX", "Webshop képernyőképek", "Saját böngészős képernyőképek", "A működő felületet és fő állapotokat bizonyítja."],
        ["Architektúra", "Komponens-architektúra ábra", r"docs\_local_segedanyagok\word_abrak\01_komponens_architektura.svg", "Angular, Firebase, Firestore Rules, Worker és OpenRouter kapcsolata."],
        ["Adatmodell", "Adatmodell diagram", r"docs\_local_segedanyagok\word_abrak\03_adatmodell_diagram.svg", "Products, orders, users, coupons, customerProfiles és audit kapcsolatok."],
        ["Megvalósítás", "5-7 rövid kódrészlet", "Lásd a 4. fejezetet ebben a segédanyagban", "A kulcsműködések technikai bizonyítása."],
        ["Biztonság", "Jogosultsági ábra + Firestore rules részlet", r"docs\_local_segedanyagok\word_abrak\04_jogosultsagi_modell.svg", "Admin, dolgozó és vásárló jogosultságok elkülönítése."],
        ["Tesztelés", "Tesztelési folyamat ábra + zöld CI kép", r"docs\_local_segedanyagok\word_abrak\05_teszteles_bizonyitasi_folyamat.svg", "Build, teszt, kézi ellenőrzés és GitHub Actions bizonyítása."],
        ["AI-asszisztens", "AI adatfolyam ábra + Worker proxy kódrészlet", r"docs\_local_segedanyagok\word_abrak\06_ai_asszisztens_adatfolyam.svg", "API-kulcs kliensoldalról való leválasztása és katalógushoz kötött válaszadás."],
        ["Összefoglalás", "Saját záró reflexió", "-", "Mit tanultál, mi volt nehéz, mit fejlesztenél tovább."]
    ])

    doc.add_heading("3. Pontos képernyőkép-lista", level=1)
    add_table(doc, ["Sorszám", "Képernyőkép", "Hol készüljön?", "Mit mutasson?", "Javasolt felirat"], [
        [1, "Kezdőlap dark mode", "Főoldal", "Logo, kereső, kategóriák, hero szekció", "A kezdőlap sötét módban, kategória lenyílóval."],
        [2, "AI asszisztens", "Főoldal", "Nyitott AI ablak egy épületgépészeti kérdéssel", "Vásárlói AI asszisztens a kezdőlapon."],
        [3, "Terméklista", "Termékek oldal", "Keresés vagy kategóriaszűrés", "Terméklista keresési és szűrési lehetőséggel."],
        [4, "Termékadatlap", "Egy konkrét termék oldala", "Kép, ár, készlet, kosár gomb", "Termékadatlap részletes termékinformációkkal."],
        [5, "Kosár", "Kosár oldal", "Több termék, mennyiség és végösszeg", "Kosár oldal mennyiségkezeléssel."],
        [6, "Checkout validáció", "Checkout oldal", "Hibás e-mail/telefon vagy kötelező mező hibája", "Checkout validáció hibás bemeneti adattal."],
        [7, "Sikeres rendelés", "Checkout után", "Sikeres leadás visszajelzése", "Sikeres rendelésleadás visszajelzése."],
        [8, "Profil", "Profil/rendelések oldal", "Rendeléstörténet és állapot", "Vásárlói profil és rendeléskövetés."],
        [9, "Kívánságlista", "Kívánságlista oldal", "Felhasználóhoz kötött mentett termékek", "Kívánságlista mentett termékekkel."],
        [10, "Admin áttekintés", "Admin főnézet", "Statisztikák és fő admin blokkok", "Admin áttekintő felület."],
        [11, "CSV import", "Admin termékek fül", "CSV import előnézet vagy mentés", "Admin termékkezelés és CSV import."],
        [12, "Rendeléskezelés", "Admin rendeléslista", "Státusz módosítása", "Admin rendeléskezelés státuszváltással."],
        [13, "Helyszíni vásárlás", "Admin helyszíni vásárlás", "Mentett vásárló + tételek", "Helyszíni vásárlás rögzítése mentett vásárlóval."],
        [14, "PDF bizonylat", "Letöltött PDF", "Vevő, tételek, végösszeg", "Generált PDF bizonylat."],
        [15, "GitHub Actions", "GitHub Actions oldal", "Legfrissebb zöld CI futás", "Sikeres GitHub Actions CI futás."]
    ])

    doc.add_heading("4. Ajánlott kódrészletek", level=1)
    add_table(doc, ["Fájl", "Sorok", "Miért érdekes?"], [
        [r"C:\Users\Dell\webshop\src\pages\checkout\checkout.ts", "393-422", "Rendelés véglegesítése, kupon, profiladatok és hibakezelés bekötése."],
        [r"C:\Users\Dell\webshop\src\pages\checkout\checkout.ts", "575-651", "Checkout validáció és tisztított űrlapadatok előállítása."],
        [r"C:\Users\Dell\webshop\src\app\services\order.service.ts", "41-126", "Rendelés státuszváltás audit- és készlettranzakcióval."],
        [r"C:\Users\Dell\webshop\src\app\services\order.service.ts", "238-284", "Helyszíni vásárlás tranzakciós mentése és készletcsökkentése."],
        [r"C:\Users\Dell\webshop\src\app\services\order.service.ts", "313-345", "Számlaszám generálása éves futószámmal."],
        [r"C:\Users\Dell\webshop\src\app\services\invoice.service.ts", "81-154", "PDF számla/bizonylat felépítése."],
        [r"C:\Users\Dell\webshop\src\pages\admin\admin.ts", "614-700", "Admin/dolgozói jogosultságok kliensoldali segédlogikája."],
        [r"C:\Users\Dell\webshop\src\pages\admin\admin.ts", "1184-1258", "CSV import validáció és mentés."],
        [r"C:\Users\Dell\webshop\firestore.rules", "12-67", "Aktív felhasználó, admin és dolgozó jogosultsági helper szabályok."],
        [r"C:\Users\Dell\webshop\firestore.rules", "288-356", "Products, orders, users, customerProfiles és audit szabályok."],
        [r"C:\Users\Dell\webshop\workers\openrouter-proxy\src\index.js", "153-213", "OpenRouter hívás szerveroldali kulcskezeléssel és CORS korlátozással."],
        [r"C:\Users\Dell\webshop\src\app\services\chatbot-llm.service.ts", "31-84", "AI asszisztens kliensoldali híváslogikája, modellválasztás felületről elrejtve."]
    ])

    add_note(
        doc,
        "A dolgozatba nem kell mind a 12 kódrészlet. A legerősebb 6-7: checkout, OrderService tranzakció, Firestore rules, PDF generálás, CSV import, Worker proxy és opcionálisan az AI klienslogika."
    )

    doc.add_heading("5. Beadás előtti ellenőrzőlista", level=1)
    add_bullets(doc, [
        "Címlap kitöltve: szak, intézmény, kar/tanszék, témavezető: Dr. Bilicki Vilmos, egyetemi docens.",
        "A dokumentum A4-es lapméretű.",
        "A tartalomjegyzék Wordben frissítve van oldalszámokkal.",
        "Minden placeholder ki van törölve.",
        "Minden képnek, ábrának és táblázatnak van felirata.",
        "Van feladatkiírás jellegű rész vagy melléklet.",
        "Van irodalomjegyzék.",
        "A tesztelési fejezetben szerepel build, teszt és kézi ellenőrzés.",
        "Az AI-használati fejezet saját megfogalmazású, és nem munkautasítás jellegű.",
        "Az összefoglalás végén szerepel saját reflexió."
    ])

    doc.add_heading("6. Rövid záró javaslat", level=1)
    p = doc.add_paragraph()
    p.add_run(
        "A dolgozat akkor lesz konzulensi review-ra vállalható, ha a jelenlegi vázlatból eltűnnek a placeholder részek, "
        "és helyükre valódi képek, ábrák, rövid kódrészletek és saját magyarázat kerül. "
        "A kódoldali bizonyítékok már elég erősek, ezért most a bemutatás minősége a döntő."
    )

    doc.save(OUT_DOCX)
    print(OUT_DOCX)


if __name__ == "__main__":
    main()
