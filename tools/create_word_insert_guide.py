from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(r"C:\Users\Dell\webshop")
OUT_DIR = ROOT / "docs" / "_local_segedanyagok" / "word_beillesztesi_utmutato"
DOCX_OUT = OUT_DIR / "TDLWebshop_word_beillesztesi_utmutato_kepek_kodok.docx"
MD_OUT = OUT_DIR / "TDLWebshop_word_beillesztesi_utmutato_kepek_kodok.md"


def set_cell(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(9)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_note(doc, title, body):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    cell._tc.get_or_add_tcPr()
    p = cell.paragraphs[0]
    r = p.add_run(title)
    r.bold = True
    r.font.color.rgb = RGBColor(37, 99, 235)
    r.font.size = Pt(10)
    p.add_run("\n" + body).font.size = Pt(9)
    doc.add_paragraph()


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for idx, h in enumerate(headers):
        set_cell(hdr[idx], h, bold=True)
        if widths:
            hdr[idx].width = Cm(widths[idx])
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell(cells[idx], value)
            if widths:
                cells[idx].width = Cm(widths[idx])
    doc.add_paragraph()
    return table


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = "Calibri"
        if level == 1:
            run.font.color.rgb = RGBColor(15, 23, 42)
        elif level == 2:
            run.font.color.rgb = RGBColor(37, 99, 235)
    return p


def p(doc, text="", style=None):
    par = doc.add_paragraph(text, style=style)
    for run in par.runs:
        run.font.size = Pt(10.5)
    return par


def bullet(doc, text):
    par = doc.add_paragraph(style="List Bullet")
    run = par.add_run(text)
    run.font.size = Pt(10)
    return par


def build_doc():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.3)
    section.right_margin = Cm(2.3)

    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(10.5)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("TDLWebshop - Word beillesztesi utmutato")
    r.bold = True
    r.font.size = Pt(20)
    r.font.color.rgb = RGBColor(15, 23, 42)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = subtitle.add_run("Kepek, abrak, kodreszletek es feliratok a szakdolgozat review-ready verziojahoz")
    sr.font.size = Pt(11)
    sr.font.color.rgb = RGBColor(75, 85, 99)
    doc.add_paragraph()

    add_note(
        doc,
        "Fontos",
        "Ez kulon segedanyag, nem beadando dokumentum. A lenyege az, hogy a "
        "TDLWebshop_szakdolgozat_javitott_v2.docx fajlban gyorsan tudd cserelni "
        "a placeholder szovegeket valodi kepekre, abrakra es rovid kodreszletekre.",
    )

    add_heading(doc, "1. Gyors munkamenet", 1)
    for item in [
        "Nyisd meg a TDLWebshop_szakdolgozat_javitott_v2.docx fajlt, es kapcsold be a bekezdesjelek megjeleniteset.",
        "A [KEP / ABRA HELYE] es [KODRESZLET HELYE] jeloleseket egyesevel csereld valodi tartalomra.",
        "Eloszor az abrakat es tablazatokat rakd be, utana a webshop screenshotokat, vegul a kodreszleteket.",
        "Minden kep ala keruljon rovid abrafelirat, peldaul: '5.1. abra: A TDLWebshop komponens-architekturaja'.",
        "A vegen frissitsd a Word tartalomjegyzeket, ellenorizd az A4 lapmeretet, es torolj minden megmaradt placeholdert.",
    ]:
        bullet(doc, item)

    add_note(
        doc,
        "Mennyi kep es kod eleg?",
        "Nem kell tulzasba vinni. Review-ready dolgozathoz altalaban 12-15 webshop kepernyokep, "
        "4-6 diagram/tabilazat es 5-7 rovid kodreszlet eleg. A kodreszletek legyenek celzottak, "
        "ne tobb oldalas blokkok.",
    )

    add_heading(doc, "2. Fejezetenkenti beillesztesi terv", 1)
    chapter_rows = [
        ("Piaci osszehasonlitas", "07_piaci_osszehasonlito_tablazat.md vagy 07_piaci_osszehasonlito_tablazat.svg", "2.1. tablazat: Epületgepeszeti webshopok osszehasonlitasa sajat piackutatas alapjan."),
        ("Kovetelmenyek", "09_kovetelmeny_traceability_tablazat.md", "3. tablazat: Kovetelmeny-traceability osszefoglalo."),
        ("Use case-ek", "08_use_case_diagram.svg es 08_use_case_reszletes.md", "3.1. abra: A TDLWebshop fo use case-ei es szereploi."),
        ("GUI/UX bemutatas", "Valodi webshop kepernyokepek", "4.x. abra: A bemutatott oldal vagy allapot neve."),
        ("Architektura", "01_komponens_architektura.svg", "5.1. abra: A TDLWebshop komponens-architekturaja."),
        ("Adatmodell", "03_adatmodell_diagram.svg", "5.2. abra: A TDLWebshop fo adatmodellje."),
        ("Checkout folyamat", "02_checkout_szekvencia.svg", "5.3. abra: A checkout folyamat szekvenciaja."),
        ("Biztonsag", "04_jogosultsagi_modell.svg + firestore.rules kodreszlet", "7.1. abra: Jogosultsagi modell es Firestore vedelmi reteg."),
        ("Teszteles", "05_teszteles_bizonyitasi_folyamat.svg + GitHub Actions kepernyokep", "8.1. abra: Build, teszteles es CI bizonyitasi folyamat."),
        ("AI asszisztens", "06_ai_asszisztens_adatfolyam.svg + Worker proxy kodreszlet", "9.1. abra: Az AI asszisztens adatfolyama es kulcskezelesi hatara."),
    ]
    add_table(doc, ["Fejezet", "Mit rakj be", "Javasolt felirat"], chapter_rows, [4.0, 6.3, 6.2])

    add_heading(doc, "3. Kepernyokepek pontos listaja", 1)
    screenshot_rows = [
        ("Kezdolap", "Dark mode, navbar, kategoria lenyilo nyitva", "A kezdolap es a liquid glass kategoria menu."),
        ("Kezdolap AI-val", "AI asszisztens ablaka nyitva, egy szakmai kerdes valasszal", "A beepitett AI asszisztens mukodes kozben."),
        ("Termeklista", "Kereses vagy kategoria szerinti szures latszodjon", "Termeklista keresessel es kategoriaval."),
        ("Termekadatlap", "Egy konkret termek kepe, ara, keszlete, kosar gomb", "Termekadatlap termekinformaciokkal."),
        ("Kosar", "Legalabb ket termek, mennyisegmodositas latszodjon", "Kosar oldal tobb termekkel."),
        ("Checkout validacio", "Hibas email vagy telefonszam es validacios uzenet", "Checkout mezoszintu validacio."),
        ("Sikeres rendeles", "Rendeles leadasa utani sikeres allapot", "Sikeresen leadott rendeles visszajelzese."),
        ("Profil", "Korabbi rendelesek es statuszkovetes", "Vasarloi profil es rendeleskovetes."),
        ("Kivansaglista", "Bejelentkezett felhasznalo kedvenc termekei", "Kivansaglista felhasznaloi nezetben."),
        ("Admin attekintes", "Statisztikai kartyak, rendelesek osszefoglaloja", "Admin attekinto nezet."),
        ("Admin termekkezeles", "CSV import vagy termeklista admin oldalon", "Admin termekkezeles es CSV import."),
        ("Admin rendeleskezeles", "Statuszvaltas vagy rendeles reszletei", "Admin rendeleskezeles statuszmodositassal."),
        ("Helyszini vasarlas", "Mentett vasarlo kivalasztva, tetelek hozzaadva", "Helyszini vasarlas rogzitese mentett vasarloval."),
        ("PDF bizonylat", "Generalt szamla/bizonylat elrendezese", "PDF bizonylat a rendeles adataival."),
        ("GitHub Actions", "Legfrissebb zold CI run", "GitHub Actions sikeres CI futas."),
    ]
    add_table(doc, ["Oldal / allapot", "Mit mutasson", "Abrafelirat otlet"], screenshot_rows, [3.3, 6.5, 6.7])

    add_heading(doc, "4. Kodreszletek - mit es honnan erdemes betenni", 1)
    p(doc, "A dolgozatba nem kell az osszes alabbit betenni. A legjobb valasztas: 5-7 rovid, jol magyarazhato kodreszlet.")
    code_rows = [
        ("Checkout folyamat", r"C:\Users\Dell\webshop\src\pages\checkout\checkout.ts", "376-422", "Rendeles veglegesitese: validacio, kosar, kupon es rendelespayload osszeallitasa."),
        ("Checkout validacio", r"C:\Users\Dell\webshop\src\pages\checkout\checkout.ts", "590-651", "Email, telefon es kupon ellenorzese. Ha meg javitas alatt van, csak javitas utan fotozd."),
        ("Statusz + audit + keszlet", r"C:\Users\Dell\webshop\src\app\services\order.service.ts", "41-126", "Tranzakcios rendelesstatusz-valtas, audit naplo es keszletmodositas."),
        ("Helyszini vasarlas", r"C:\Users\Dell\webshop\src\app\services\order.service.ts", "244-276", "Helyszini rendeles mentese keszletellenorzessel."),
        ("Szamlaszam generalas", r"C:\Users\Dell\webshop\src\app\services\order.service.ts", "313-345", "Invoice counter es szamlaszam kepzese tranzakcioban."),
        ("PDF bizonylat", r"C:\Users\Dell\webshop\src\app\services\invoice.service.ts", "81-154", "PDF tartalmi felepites, disclaimer es osszegzo reszek."),
        ("CSV import", r"C:\Users\Dell\webshop\src\pages\admin\admin.ts", "1183-1258", "CSV validacio es termekimport admin feluleten."),
        ("Jogosultsagi szabalyok", r"C:\Users\Dell\webshop\firestore.rules", "12-67", "Admin, dolgozo, aktiv felhasznalo es jogosultsagi segedfuggvenyek."),
        ("Adat-hozzaferesi szabalyok", r"C:\Users\Dell\webshop\firestore.rules", "288-356", "Products, orders, users, savedCustomers es audit collection szabalyai."),
        ("OpenRouter proxy", r"C:\Users\Dell\webshop\workers\openrouter-proxy\src\index.js", "153-213", "Szerveroldali API kulcskezeles, CORS es OpenRouter hivas lenyegi resze."),
        ("AI kliens logika", r"C:\Users\Dell\webshop\src\app\services\chatbot-llm.service.ts", "31-84", "Frontend endpoint-hivas ugy, hogy a modell nem allithato a feluleten."),
    ]
    add_table(doc, ["Tema", "Fajl", "Sorok", "Miert erdekes"], code_rows, [3.0, 5.2, 1.8, 6.8])

    add_heading(doc, "5. Minimum csomag, ha gyorsan kell review-ready verzio", 1)
    for item in [
        "Abrak: komponens-architektura, adatmodell, use case diagram, checkout szekvencia.",
        "Kepernyokepek: kezdolap, termeklista, termekadatlap, kosar, checkout, profil, admin, PDF, CI.",
        "Kodreszletek: checkout veglegesites, OrderService tranzakcio, Firestore rules, PDF generalas, OpenRouter proxy.",
        "Szoveges bizonyitas: minden abra utan 1-2 bekezdes magyarazat arrol, mit bizonyit a kep.",
    ]:
        bullet(doc, item)

    add_heading(doc, "6. Mit ne rakj be a vegleges dolgozatba", 1)
    for item in [
        "Ne rakd be a segedanyag mappak neveit vagy azokat a fajlokat, amelyek csak munkatamogatasra keszultek.",
        "Ne maradjon benne [KEP / ABRA HELYE] vagy [KODRESZLET HELYE] placeholder.",
        "Ne legyenek tul hosszu kodblokkok. Egy kodreszlet legyen inkabb fel oldal, mint ket oldal.",
        "Ne allitsd azt, hogy az AI teljesen onalloan keszitette a kodot. A helyes megfogalmazas: fejlesztesi es ellenorzesi tamogatasra hasznaltad.",
        "Ne commitold a kulon segedanyagokat a beadando repoba, ha azok csak neked szolnak.",
    ]:
        bullet(doc, item)

    add_heading(doc, "7. Utolso Word-ellenorzes beadashoz", 1)
    checklist = [
        "A4 lapmeret beallitva.",
        "Cimlap kitoltve: szak, intezmeny, kar/tanszek, temavezeto: Dr. Bilicki Vilmos, egyetemi docens.",
        "Word tartalomjegyzek frissitve oldalszamokkal.",
        "Minden kepnek es tablazatnak van sorszamozott felirata.",
        "Van feladatkiiras vagy melleklet.",
        "Van irodalomjegyzek.",
        "Van sajat zaro reflexio: mit tanultal, mi volt nehez, mit csinalnal maskepp.",
        "AI-hasznalati fejezet benne van, de sajat nyelvezetre at van huzva.",
        "Minden placeholder torolve.",
    ]
    for item in checklist:
        bullet(doc, "[ ] " + item)

    doc.add_section(WD_SECTION.NEW_PAGE)
    add_heading(doc, "8. Beillesztesi forrasfajlok helye", 1)
    for item in [
        r"Abrak: C:\Users\Dell\webshop\docs\_local_segedanyagok\word_abrak",
        r"Ez a segedanyag: C:\Users\Dell\webshop\docs\_local_segedanyagok\word_beillesztesi_utmutato",
        r"Fo szakdolgozati Word fajl: C:\Users\Dell\Downloads\TDLWebshop_szakdolgozat_javitott_v2.docx",
    ]:
        bullet(doc, item)

    doc.save(DOCX_OUT)
    MD_OUT.write_text(build_markdown(), encoding="utf-8")


def build_markdown():
    return """# TDLWebshop - Word beillesztesi utmutato

Ez kulon segedanyag, nem beadando dokumentum. A `TDLWebshop_szakdolgozat_javitott_v2.docx` fajlban segit lecserelni a placeholder helyeket valodi kepekre, abrakra es kodreszletekre.

## Minimum mennyiseg

- 12-15 webshop kepernyokep
- 4-6 diagram vagy tablazat
- 5-7 rovid kodreszlet

## Legfontosabb abrak

- `docs/_local_segedanyagok/word_abrak/01_komponens_architektura.svg`
- `docs/_local_segedanyagok/word_abrak/02_checkout_szekvencia.svg`
- `docs/_local_segedanyagok/word_abrak/03_adatmodell_diagram.svg`
- `docs/_local_segedanyagok/word_abrak/08_use_case_diagram.svg`
- `docs/_local_segedanyagok/word_abrak/09_kovetelmeny_traceability_tablazat.md`
- `docs/_local_segedanyagok/word_abrak/07_piaci_osszehasonlito_tablazat.md`

## Javasolt kodreszletek

- `src/pages/checkout/checkout.ts` 376-422
- `src/pages/checkout/checkout.ts` 590-651
- `src/app/services/order.service.ts` 41-126
- `src/app/services/invoice.service.ts` 81-154
- `src/pages/admin/admin.ts` 1183-1258
- `firestore.rules` 12-67 es 288-356
- `workers/openrouter-proxy/src/index.js` 153-213

## Utolso ellenorzes

- Minden placeholder torolve.
- A4 lapmeret.
- Tartalomjegyzek frissitve.
- Minden kephez abrafelirat.
- Irodalomjegyzek es sajat zaro reflexio megvan.
"""


if __name__ == "__main__":
    build_doc()
    print(DOCX_OUT)
