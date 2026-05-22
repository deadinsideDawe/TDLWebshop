from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "kodreszlet_kepalairasok_utvonalakkal_frissitett.docx"


ITEMS = [
    {
        "id": "M2.1",
        "title": "Útvonalkezelés és lazy loading",
        "image": "docs/code-snippet-images/M2_1_utvonalkezeles_es_lazy_loading.png",
        "sources": ["src/app/app.routes.ts:5-21"],
        "caption": (
            "Útvonalkezelés és lazy loading megvalósítása az Angular alkalmazásban. "
            "A kódrészlet azt mutatja be, hogyan vannak elkülönítve a publikus, "
            "vásárlói és adminisztrációs felületek, valamint hogyan töltődnek be "
            "az egyes oldalak csak szükség esetén."
        ),
    },
    {
        "id": "M2.2",
        "title": "Admin útvonalvédő",
        "image": "docs/code-snippet-images/M2_2_admin_route_guard.png",
        "sources": ["src/app/guards/admin.guard.ts:7-23"],
        "caption": (
            "Admin útvonalvédő működése. A kódrészlet szemlélteti, hogy az admin "
            "felület csak megfelelő jogosultságú felhasználó számára érhető el."
        ),
    },
    {
        "id": "M2.3",
        "title": "Szerepkör-ellenőrzés",
        "image": "docs/code-snippet-images/M2_3_szerepkor_ellenorzes.png",
        "sources": ["src/app/services/auth.service.ts:104-122"],
        "caption": (
            "Felhasználói szerepkör ellenőrzése. A részlet azt mutatja, hogyan "
            "különbözteti meg a rendszer a vásárlói, dolgozói és adminisztrátori "
            "jogosultságokat."
        ),
    },
    {
        "id": "M2.4a",
        "title": "Firestore admin és aktív user helper",
        "image": "docs/code-snippet-images/M2_4a_firestore_admin_es_aktiv_user_helper.png",
        "sources": ["firestore.rules:1-32"],
        "caption": (
            "Firestore admin és aktív felhasználó ellenőrző függvények. A részlet "
            "bemutatja, hogyan azonosítja a szabályrendszer az admin e-mail címeket, "
            "a bejelentkezett felhasználót, valamint azt, hogy a felhasználói "
            "dokumentum alapján aktívnak tekinthető-e a user."
        ),
    },
    {
        "id": "M2.4b",
        "title": "Firestore admin jogosultság helper",
        "image": "docs/code-snippet-images/M2_4b_firestore_admin_jogosultsag_helper.png",
        "sources": ["firestore.rules:34-45"],
        "caption": (
            "Firestore admin jogosultság ellenőrzése. A kódrészlet azt szemlélteti, "
            "hogy a rendszer tokenben szereplő admin e-mail vagy felhasználói "
            "dokumentumban tárolt admin szerepkör alapján engedélyezi az admin "
            "műveleteket."
        ),
    },
    {
        "id": "M2.4c",
        "title": "Firestore termék és tartalom szabályok",
        "image": "docs/code-snippet-images/M2_4c_firestore_termek_es_tartalom_szabalyok.png",
        "sources": ["firestore.rules:288-309"],
        "caption": (
            "Firestore termék- és tartalomkezelési szabályok. A részlet bemutatja, "
            "hogy a termékek nyilvánosan olvashatók, de létrehozásuk, módosításuk "
            "és törlésük admin vagy megfelelő dolgozói jogosultsághoz kötött."
        ),
    },
    {
        "id": "M2.4d",
        "title": "Firestore rendelés szabályok",
        "image": "docs/code-snippet-images/M2_4d_firestore_rendeles_szabalyok.png",
        "sources": ["firestore.rules:311-316"],
        "caption": (
            "Firestore rendeléskezelési szabályok. A kódrészlet azt mutatja, hogyan "
            "korlátozza a rendszer a rendelések létrehozását, olvasását, módosítását "
            "és törlését admin, dolgozói vagy rendeléstulajdonosi jogosultság alapján."
        ),
    },
    {
        "id": "M2.5",
        "title": "Rendelés adatmodell",
        "image": "docs/code-snippet-images/M2_5_rendeles_adatmodell.png",
        "sources": ["src/app/models/order.model.ts:3-63"],
        "caption": (
            "Rendelés adatmodellje. A részletben látható, hogy a rendelés milyen "
            "mezőkből épül fel, beleértve a vásárlói adatokat, tételeket, "
            "összegeket és státuszokat."
        ),
    },
    {
        "id": "M2.6a",
        "title": "Checkout validáció és összegszámítás",
        "image": "docs/code-snippet-images/M2_6a_checkout_validacio_es_osszegszamitas.png",
        "sources": ["src/pages/checkout/checkout.ts:457-486"],
        "caption": (
            "Checkout validáció és összegszámítás. A kódrészlet azt szemlélteti, "
            "hogyan ellenőrzi a rendszer a rendeléshez szükséges adatokat, majd "
            "hogyan számolja ki a szállítási díjat, kedvezményt és végösszeget."
        ),
    },
    {
        "id": "M2.6b",
        "title": "Rendelés objektum összeállítása",
        "image": "docs/code-snippet-images/M2_6b_rendeles_objektum_osszeallitasa.png",
        "sources": ["src/pages/checkout/checkout.ts:488-545"],
        "caption": (
            "Rendelés objektum összeállítása. A részlet bemutatja, hogy a checkout "
            "folyamat során a rendszer milyen vásárlói, szállítási, számlázási, "
            "fizetési és kuponadatokat rendez egy menthető rendelési objektumba."
        ),
    },
    {
        "id": "M2.6c",
        "title": "Rendelés összegzés és profilmentés",
        "image": "docs/code-snippet-images/M2_6c_rendeles_osszegzes_es_profilmentes.png",
        "sources": ["src/pages/checkout/checkout.ts:546-618"],
        "caption": (
            "Rendelés összegzés és profilmentés. A kódrészlet azt mutatja, hogyan "
            "készül el a sikeres rendeléshez szükséges összegző adatcsomag, illetve "
            "hogyan kapcsolódik a rendelés mentése a bejelentkezett felhasználó "
            "profiladatainak frissítéséhez."
        ),
    },
    {
        "id": "M2.7a",
        "title": "Checkout űrlaphibák gyűjtése",
        "image": "docs/code-snippet-images/M2_7a_checkout_urlaphibak_gyujtese.png",
        "sources": ["src/pages/checkout/checkout.ts:659-693"],
        "caption": (
            "Checkout űrlaphibák gyűjtése. A részlet bemutatja, hogyan ellenőrzi "
            "a checkout oldal a kötelező mezőket, az e-mail címet, a telefonszámot, "
            "az adószámot és a szállítási adatokat."
        ),
    },
    {
        "id": "M2.7b",
        "title": "E-mail validáció helper",
        "image": "docs/code-snippet-images/M2_7b_email_validacio_helper.png",
        "sources": ["src/app/utils/form-validators.ts:1-38"],
        "caption": (
            "E-mail validációs segédfüggvény. A kódrészlet azt mutatja, hogyan "
            "szűri ki a rendszer az üres, túl hosszú, hibásan tagolt vagy "
            "érvénytelen domainnel rendelkező e-mail címeket."
        ),
    },
    {
        "id": "M2.7c",
        "title": "Telefonszám validáció helper",
        "image": "docs/code-snippet-images/M2_7c_telefonszam_validacio_helper.png",
        "sources": ["src/app/utils/form-validators.ts:40-89"],
        "caption": (
            "Telefonszám validációs segédfüggvény. A részlet bemutatja, hogyan "
            "kezeli a rendszer a magyar telefonszámformátumokat, a körzetszámokat "
            "és a nem megengedett karaktereket."
        ),
    },
    {
        "id": "M2.8a",
        "title": "Rendelésstátusz tranzakció indítása",
        "image": "docs/code-snippet-images/M2_8a_rendelesstatusz_tranzakcio_inditasa.png",
        "sources": ["src/app/services/order.service.ts:41-70"],
        "caption": (
            "Rendelésstátusz tranzakció indítása. A kódrészlet bemutatja, hogyan "
            "keresi meg a rendszer a módosítandó rendelést, hogyan ellenőrzi az "
            "aktuális státuszt, és mikor szükséges készletkorrekciót végezni."
        ),
    },
    {
        "id": "M2.8b",
        "title": "Készletkorrekció státuszváltáskor",
        "image": "docs/code-snippet-images/M2_8b_keszletkorrekcio_statuszvaltaskor.png",
        "sources": ["src/app/services/order.service.ts:72-116"],
        "caption": (
            "Készletkorrekció státuszváltáskor. A részlet azt szemlélteti, hogyan "
            "olvassa ki a rendszer a rendelés tételeihez tartozó termékeket, hogyan "
            "számolja ki a készletváltozást, és hogyan akadályozza meg a negatív "
            "készlet kialakulását."
        ),
    },
    {
        "id": "M2.8c",
        "title": "Rendelésfrissítés és audit napló",
        "image": "docs/code-snippet-images/M2_8c_rendelesfrissites_es_audit_naplo.png",
        "sources": ["src/app/services/order.service.ts:118-134"],
        "caption": (
            "Rendelésfrissítés és audit napló. A kódrészlet bemutatja, hogyan menti "
            "a rendszer az új rendelésstátuszt, valamint hogyan rögzíti az admin "
            "műveletet az audit naplóban visszakövethetőség céljából."
        ),
    },
    {
        "id": "M2.9",
        "title": "Admin státuszváltás meghívása",
        "image": "docs/code-snippet-images/M2_9_admin_statuszvaltas_meghivasa.png",
        "sources": ["src/pages/admin/admin.ts:1577-1614"],
        "caption": (
            "Admin rendelésstátusz-váltás meghívása a felületről. A részlet azt "
            "mutatja, hogyan indítja el az admin felület a rendelés állapotának "
            "módosítását."
        ),
    },
    {
        "id": "M2.10a",
        "title": "Helyszíni vásárlás összeg és címke",
        "image": "docs/code-snippet-images/M2_10a_helyszini_vasarlas_osszeg_es_cimke.png",
        "sources": ["src/pages/admin/admin.ts:2335-2361"],
        "caption": (
            "Helyszíni vásárlás összeg- és kedvezménycímkéi. A részlet bemutatja, "
            "hogyan számolja ki az admin felület a helyszíni vásárlás részösszegét, "
            "kedvezményét és végösszegét."
        ),
    },
    {
        "id": "M2.10b",
        "title": "Mentett vásárló előző rendelései",
        "image": "docs/code-snippet-images/M2_10b_mentett_vasarlo_elozo_rendelesek.png",
        "sources": ["src/pages/admin/admin.ts:2363-2398"],
        "caption": (
            "Mentett vásárló előző rendeléseinek kezelése. A kódrészlet azt mutatja, "
            "hogyan keresi vissza a rendszer a kiválasztott vásárló korábbi rendeléseit, "
            "és hogyan számolja ki az előzmények alapján járó kedvezményt."
        ),
    },
    {
        "id": "M2.10c",
        "title": "Mentett vásárló kiválasztás és kedvezmény",
        "image": "docs/code-snippet-images/M2_10c_mentett_vasarlo_kivalasztas_es_kedvezmeny.png",
        "sources": ["src/pages/admin/admin.ts:2400-2431"],
        "caption": (
            "Mentett vásárló kiválasztása és kedvezménylogika. A részlet bemutatja, "
            "hogyan nyílik meg a vásárlói előzményablak, illetve hogyan dönt a rendszer "
            "a nagykereskedelmi vagy törzsvásárlói kedvezményről."
        ),
    },
    {
        "id": "M2.10d",
        "title": "Helyszíni vásárlás elővalidáció",
        "image": "docs/code-snippet-images/M2_10d_helyszini_vasarlas_elovalidacio.png",
        "sources": ["src/pages/admin/admin.ts:2433-2478"],
        "caption": (
            "Helyszíni vásárlás elővalidációja. A kódrészlet azt mutatja, hogyan "
            "ellenőrzi a rendszer a belső jogosultságot, a kosár tételeit, a vásárlói "
            "adatokat, az e-mail címet, a telefonszámot és az adószámot."
        ),
    },
    {
        "id": "M2.10e",
        "title": "Helyszíni tételek és összegek",
        "image": "docs/code-snippet-images/M2_10e_helyszini_tetelek_es_osszegek.png",
        "sources": ["src/pages/admin/admin.ts:2480-2498"],
        "caption": (
            "Helyszíni vásárlás tételeinek és összegeinek összeállítása. A részlet "
            "bemutatja, hogyan alakítja át a felület a kiválasztott termékeket "
            "rendelési tételekké, és hogyan számolja ki a fizetendő végösszeget."
        ),
    },
    {
        "id": "M2.10f",
        "title": "Helyszíni rendelés objektum",
        "image": "docs/code-snippet-images/M2_10f_helyszini_rendeles_objektum.png",
        "sources": ["src/pages/admin/admin.ts:2500-2549"],
        "caption": (
            "Helyszíni rendelés objektumának létrehozása. A kódrészlet bemutatja, "
            "milyen vásárlói, szállítási, számlázási, fizetési és kedvezményadatokkal "
            "jön létre a helyszíni vásárlás rendelési rekordja."
        ),
    },
    {
        "id": "M2.10g",
        "title": "Helyszíni számla és PDF indítás",
        "image": "docs/code-snippet-images/M2_10g_helyszini_szamla_es_pdf_inditas.png",
        "sources": ["src/pages/admin/admin.ts:2551-2557"],
        "caption": (
            "Helyszíni számla és PDF generálás indítása. A részlet azt mutatja, "
            "hogyan kéri le a rendszer a számlaazonosítót, majd hogyan indítja el "
            "a helyszíni vásárláshoz tartozó PDF bizonylat letöltését."
        ),
    },
    {
        "id": "M2.11",
        "title": "Helyszíni vásárlás tranzakció",
        "image": "docs/code-snippet-images/M2_11_helyszini_vasarlas_tranzakcio.png",
        "sources": ["src/app/services/order.service.ts:238-304"],
        "caption": (
            "Helyszíni vásárlás tranzakciós mentése. A részlet azt szemlélteti, "
            "hogyan menti a rendszer a helyszíni vásárlást, miközben kezeli a "
            "készletváltozást és az adatkonzisztenciát."
        ),
    },
    {
        "id": "M2.12a",
        "title": "Termék törlés és import előkészítés",
        "image": "docs/code-snippet-images/M2_12a_termek_torles_es_import_elokeszites.png",
        "sources": ["src/app/services/product.service.ts:59-77"],
        "caption": (
            "Termék törlés és import előkészítés. A kódrészlet bemutatja az egyedi "
            "terméktörlést, valamint az adminisztrációs újratöltéshez használt teljes "
            "termékkollekció törlésének előkészítését."
        ),
    },
    {
        "id": "M2.12b",
        "title": "CSV termékek normalizálása",
        "image": "docs/code-snippet-images/M2_12b_csv_termekek_normalizalasa.png",
        "sources": ["src/app/services/product.service.ts:78-94"],
        "caption": (
            "CSV termékek normalizálása és szűrése. A részlet azt mutatja, hogyan "
            "alakítja át a rendszer az importált termékadatokat egységes formára, "
            "majd hogyan hagyja ki a hiányos rekordokat."
        ),
    },
    {
        "id": "M2.12c",
        "title": "CSV insert mód termékmentés",
        "image": "docs/code-snippet-images/M2_12c_csv_insert_mod_termekmentes.png",
        "sources": ["src/app/services/product.service.ts:96-111"],
        "caption": (
            "CSV insert módú termékmentés. A kódrészlet bemutatja, hogyan hoz létre "
            "a rendszer új termékdokumentumokat, amikor az importálás egyszerű "
            "beszúrási módban fut."
        ),
    },
    {
        "id": "M2.12d",
        "title": "CSV upsert SKU alapú frissítés",
        "image": "docs/code-snippet-images/M2_12d_csv_upsert_sku_alapu_frissites.png",
        "sources": ["src/app/services/product.service.ts:113-166"],
        "caption": (
            "CSV upsert SKU alapú frissítés. A részlet bemutatja, hogyan párosítja "
            "a rendszer a CSV-ben érkező termékeket a meglévő Firestore "
            "dokumentumokkal cikkszám alapján, majd hogyan frissít vagy hoz létre "
            "terméket batch műveletekkel."
        ),
    },
    {
        "id": "M2.13a",
        "title": "Számlaszám generálás tranzakcióban",
        "image": "docs/code-snippet-images/M2_13a_szamlaszam_generalas_tranzakcioban.png",
        "sources": ["src/app/services/order.service.ts:327-362"],
        "caption": (
            "Számlaszám generálása tranzakcióban. A kódrészlet bemutatja, hogyan "
            "használ a rendszer éves számlaszámlálót, hogyan kerüli el a duplikált "
            "számlaszámot, és hogyan menti vissza a számlaadatokat a rendeléshez."
        ),
    },
    {
        "id": "M2.13b",
        "title": "PDF bizonylat adatainak előkészítése",
        "image": "docs/code-snippet-images/M2_13b_pdf_bizonylat_adatainak_elokeszitese.png",
        "sources": ["src/app/services/invoice.service.ts:8-40"],
        "caption": (
            "PDF bizonylat adatainak előkészítése. A részlet azt mutatja, hogyan "
            "állítja össze a rendszer a számlán megjelenő vevői, szállítási, "
            "fizetési, tétel- és összegadatokat."
        ),
    },
    {
        "id": "M2.13c",
        "title": "PDF letöltés böngészőben",
        "image": "docs/code-snippet-images/M2_13c_pdf_letoltes_bongeszoben.png",
        "sources": ["src/app/services/invoice.service.ts:41-56"],
        "caption": (
            "PDF letöltés indítása a böngészőben. A kódrészlet bemutatja, hogyan "
            "készül a PDF byte tömbből Blob, majd hogyan jön létre az ideiglenes "
            "letöltési hivatkozás a számla fájlnévvel történő mentéséhez."
        ),
    },
    {
        "id": "M2.14a",
        "title": "AI kérdés domain szűrése",
        "image": "docs/code-snippet-images/M2_14a_ai_kerdes_domain_szurese.png",
        "sources": ["src/app/services/chatbot-llm.service.ts:39-59"],
        "caption": (
            "AI kérdés domain szűrése. A kódrészlet bemutatja, hogyan engedi tovább "
            "a rendszer csak a webshophoz és épületgépészeti témákhoz kapcsolódó "
            "kérdéseket, illetve hogyan ad visszajelzést irreleváns kérdés esetén."
        ),
    },
    {
        "id": "M2.14b",
        "title": "AI proxy hívás és válaszfeldolgozás",
        "image": "docs/code-snippet-images/M2_14b_ai_proxy_hivas_es_valaszfeldolgozas.png",
        "sources": ["src/app/services/chatbot-llm.service.ts:61-88"],
        "caption": (
            "AI proxy hívás és válaszfeldolgozás. A részlet azt mutatja, hogyan "
            "küldi el a kliensoldal a felhasználói üzenetet és a releváns "
            "termékkatalógust a proxy felé, majd hogyan dolgozza fel a válaszban "
            "kapott szöveget és termékajánlásokat."
        ),
    },
    {
        "id": "M2.14c",
        "title": "AI releváns termékkatalógus",
        "image": "docs/code-snippet-images/M2_14c_ai_relevans_termekkatalogus.png",
        "sources": ["src/app/services/chatbot-llm.service.ts:92-114"],
        "caption": (
            "AI releváns termékkatalógus összeállítása. A kódrészlet bemutatja, "
            "hogyan normalizálja a rendszer a felhasználói üzenetet, hogyan szűri "
            "a készleten lévő és kategóriában illeszkedő termékeket, majd hogyan "
            "pontozza és rendezi az ajánlható termékeket."
        ),
    },
    {
        "id": "M2.14d",
        "title": "AI katalógus termék DTO",
        "image": "docs/code-snippet-images/M2_14d_ai_katalogus_termek_dto.png",
        "sources": ["src/app/services/chatbot-llm.service.ts:116-128"],
        "caption": (
            "AI katalógus termék DTO létrehozása. A részlet bemutatja, hogyan "
            "alakítja át a rendszer a belső termékmodellt az AI-asszisztens számára "
            "küldhető, csak szükséges mezőket tartalmazó katalógusobjektummá."
        ),
    },
    {
        "id": "M2.15a",
        "title": "Worker CORS és metódusvédelem",
        "image": "docs/code-snippet-images/M2_15a_worker_cors_es_metodus_vedelem.png",
        "sources": ["workers/openrouter-proxy/src/index.js:153-160"],
        "caption": (
            "Worker CORS és metódusvédelem. A kódrészlet bemutatja, hogyan kezeli "
            "a proxy az OPTIONS preflight kéréseket, és hogyan engedélyezi csak a "
            "POST metódust az AI-asszisztens használatához."
        ),
    },
    {
        "id": "M2.15b",
        "title": "Worker rate limit és kulcsellenőrzés",
        "image": "docs/code-snippet-images/M2_15b_worker_rate_limit_es_kulcsellenorzes.png",
        "sources": ["workers/openrouter-proxy/src/index.js:162-180"],
        "caption": (
            "Worker rate limit és API-kulcs ellenőrzés. A részlet azt mutatja, "
            "hogyan korlátozza a rendszer a túl gyakori AI-kéréseket, és hogyan "
            "akadályozza meg a működést hiányzó OpenRouter kulcs esetén."
        ),
    },
    {
        "id": "M2.15c",
        "title": "Worker üzenet és domain validáció",
        "image": "docs/code-snippet-images/M2_15c_worker_uzenet_es_domain_validacio.png",
        "sources": ["workers/openrouter-proxy/src/index.js:182-197"],
        "caption": (
            "Worker üzenet- és domainvalidáció. A kódrészlet bemutatja, hogyan "
            "olvassa ki a proxy a felhasználói üzenetet és a termékkatalógust, "
            "valamint hogyan szűri ki a webshop témakörén kívüli kérdéseket."
        ),
    },
    {
        "id": "M2.15d",
        "title": "OpenRouter API hívás",
        "image": "docs/code-snippet-images/M2_15d_openrouter_api_hivas.png",
        "sources": ["workers/openrouter-proxy/src/index.js:198-216"],
        "caption": (
            "OpenRouter API hívás. A részlet azt szemlélteti, hogyan hívja meg "
            "a Cloudflare Worker az OpenRouter chat completions végpontját, miközben "
            "a titkos API-kulcs a szerveroldali környezetben marad."
        ),
    },
    {
        "id": "M2.15e",
        "title": "OpenRouter válasz parse",
        "image": "docs/code-snippet-images/M2_15e_openrouter_valasz_parse.png",
        "sources": ["workers/openrouter-proxy/src/index.js:217-224"],
        "caption": (
            "OpenRouter válasz feldolgozása. A kódrészlet bemutatja, hogyan ellenőrzi "
            "a proxy a külső AI-szolgáltatás válaszát, majd hogyan alakítja át a "
            "modell által adott szöveget a frontend számára értelmezhető eredménnyé."
        ),
    },
]

GROUP_NOTES = [
    (
        "M2.4a-M2.4d",
        "Az M2.4 melléklet a Firestore biztonsági szabályokat több részletben mutatja be. "
        "Az M2.4a-M2.4d ábrák sorrendben a jogosultsági segédfüggvényeket, az admin "
        "ellenőrzést, a termék- és tartalomszabályokat, valamint a rendelési szabályokat szemléltetik."
    ),
    (
        "M2.6a-M2.6c",
        "Az M2.6 melléklet a checkout és rendelésmentés folyamatát több részletben mutatja be "
        "az olvashatóság érdekében. Az M2.6a-M2.6c ábrák ugyanannak a folyamatnak egymás "
        "utáni kódrészletei."
    ),
    (
        "M2.7a-M2.7c",
        "Az M2.7 melléklet az űrlapvalidációt bontja részekre: a checkout oldali hibagyűjtésre, "
        "az e-mail ellenőrzésre és a telefonszám-validációra."
    ),
    (
        "M2.8a-M2.8c",
        "Az M2.8 melléklet a rendelésstátusz-váltást, a készletkorrekciót és az auditnaplózást "
        "három egymást követő kódrészletként mutatja be."
    ),
    (
        "M2.10a-M2.10g",
        "Az M2.10 melléklet a helyszíni vásárlás admin felületi folyamatát több részletben "
        "mutatja be. Az M2.10a-M2.10g ábrák sorrendben a kedvezményszámítást, az elővalidációt, "
        "a tételek összeállítását, a rendelésobjektum létrehozását és a PDF-generálás indítását szemléltetik."
    ),
    (
        "M2.12a-M2.12d",
        "Az M2.12 melléklet a CSV-importot több részletben mutatja be: előkészítés, normalizálás, "
        "insert mód és SKU alapú upsert frissítés."
    ),
    (
        "M2.13a-M2.13c",
        "Az M2.13 melléklet a számlaszám-generálást és a PDF-bizonylat készítését három részre "
        "bontja: számlaszám tranzakció, PDF-adatok előkészítése és böngészős letöltés."
    ),
    (
        "M2.14a-M2.14d",
        "Az M2.14 melléklet az AI-asszisztens kliensoldali működését négy egymást követő "
        "részletben mutatja be: domain-szűrés, proxy hívás, releváns katalógus összeállítása "
        "és AI-nak küldhető termékobjektum."
    ),
    (
        "M2.15a-M2.15e",
        "Az M2.15 melléklet az OpenRouter proxy Worker működését több részletben mutatja be: "
        "CORS/metódusvédelem, rate limit és API-kulcs ellenőrzés, üzenetvalidáció, külső API-hívás "
        "és válaszfeldolgozás."
    ),
]


def set_run_font(run, size=10.5, bold=False, color=None):
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_heading(document, text, level=1):
    paragraph = document.add_heading("", level=level)
    run = paragraph.add_run(text)
    set_run_font(run, 16 if level == 1 else 13, True, (31, 78, 121))
    return paragraph


def add_source_paragraph(document, item):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(2)
    run = paragraph.add_run("Forrás a projektben: ")
    set_run_font(run, 10, True)
    run = paragraph.add_run("; ".join(item["sources"]))
    set_run_font(run, 10)

    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(8)
    run = paragraph.add_run("Képfájl: ")
    set_run_font(run, 10, True)
    run = paragraph.add_run(item["image"])
    set_run_font(run, 10)


def add_group_notes(document):
    add_heading(document, "Bontott ábrák jelölése", 1)
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(8)
    run = paragraph.add_run(
        "A több részre bontott kódrészleteket alábraként érdemes kezelni. "
        "A főszövegben például így hivatkozz rájuk: "
        "\"a checkout folyamat főbb lépéseit az M2.6a-M2.6c ábrák mutatják be\". "
        "A mellékletben minden alábra külön képaláírást és forráshelyet kap."
    )
    set_run_font(run, 10.5)

    for label, note in GROUP_NOTES:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.left_indent = Cm(0.4)
        paragraph.paragraph_format.space_after = Pt(4)
        run = paragraph.add_run(f"{label}: ")
        set_run_font(run, 10, True, (31, 78, 121))
        run = paragraph.add_run(note)
        set_run_font(run, 10)


def build_doc():
    document = Document()
    section = document.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Cm(29.7)
    section.page_height = Cm(21.0)
    section.top_margin = Cm(1.7)
    section.bottom_margin = Cm(1.7)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Kódrészlet-képek képaláírásai és forráshelyei")
    set_run_font(run, 18, True, (31, 78, 121))

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Mellékletbe illeszthető szövegek a szakdolgozathoz")
    set_run_font(run, 11, False, (90, 90, 90))

    intro = document.add_paragraph()
    intro.paragraph_format.space_after = Pt(12)
    run = intro.add_run(
        "Az alábbi lista az M2 mellékletben szereplő kódrészlet-képekhez ad "
        "bemásolható képaláírást, valamint megadja, hogy a részlet a projekt "
        "forráskódjában melyik fájlban és melyik sorszám-tartományban található."
    )
    set_run_font(run, 10.5)

    add_group_notes(document)

    table = document.add_table(rows=1, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    headers = ["Azonosító", "Téma", "Képaláírás", "Forráshely", "Képfájl"]
    for index, text in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(text)
        set_run_font(run, 9.5, True, (255, 255, 255))
        cell._tc.get_or_add_tcPr().append(parse_shading("1F4E79"))

    widths = [1.4, 3.0, 7.1, 4.3, 4.0]
    for item in ITEMS:
        cells = table.add_row().cells
        values = [
            item["id"],
            item["title"],
            f"{item['id']}. ábra: {item['caption']}",
            "\n".join(item["sources"]),
            item["image"],
        ]
        for index, value in enumerate(values):
            cell = cells[index]
            cell.width = Cm(widths[index])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            run = paragraph.add_run(value)
            set_run_font(run, 8.5 if index in (3, 4) else 9)

    detail_section = document.add_section(WD_SECTION.NEW_PAGE)
    detail_section.orientation = WD_ORIENT.PORTRAIT
    detail_section.page_width = Cm(21.0)
    detail_section.page_height = Cm(29.7)
    detail_section.top_margin = Cm(1.8)
    detail_section.bottom_margin = Cm(1.8)
    detail_section.left_margin = Cm(2.0)
    detail_section.right_margin = Cm(2.0)
    add_heading(document, "Bemásolható képaláírások", 1)

    add_group_notes(document)

    for item in ITEMS:
        add_heading(document, f"{item['id']} - {item['title']}", 2)
        caption = document.add_paragraph()
        caption.paragraph_format.space_after = Pt(4)
        run = caption.add_run(f"{item['id']}. ábra: {item['caption']}")
        set_run_font(run, 10.5)
        add_source_paragraph(document, item)

    document.save(OUTPUT)
    return OUTPUT


def parse_shading(fill):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    element = OxmlElement("w:shd")
    element.set(qn("w:fill"), fill)
    return element


if __name__ == "__main__":
    print(build_doc())
