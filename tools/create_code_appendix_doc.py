from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


OUT = Path(
    r"C:\Users\Dell\OneDrive\Asztali gép\szakdoga\kodreszlet-melleklet-javaslat_szakdolgozathoz.docx"
)


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(8.3)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP


doc = Document()
section = doc.sections[0]
section.top_margin = Inches(0.7)
section.bottom_margin = Inches(0.7)
section.left_margin = Inches(0.65)
section.right_margin = Inches(0.65)

styles = doc.styles
styles["Normal"].font.name = "Calibri"
styles["Normal"].font.size = Pt(10)
styles["Title"].font.name = "Calibri"
styles["Title"].font.size = Pt(20)
styles["Heading 1"].font.name = "Calibri"
styles["Heading 1"].font.size = Pt(15)

title = doc.add_paragraph()
title.style = "Title"
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.add_run("Kódrészletek mellékletbe rendezve")

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.add_run(
    "TDLWebshop szakdolgozat - forráskód-melléklet javasolt tartalma és hivatkozásai"
).italic = True

doc.add_heading("Használati javaslat", level=1)
doc.add_paragraph(
    "A kódrészleteket ne a fő fejezetekbe tedd, hanem külön mellékletként a "
    "szakdolgozat végére. A főszövegben elég röviden hivatkozni rájuk, például: "
    '"A checkout folyamat releváns forráskódrészlete az M2.6 mellékletben látható."'
)
doc.add_paragraph(
    "A mellékletben minden kódrészlethez szerepeljen: mellékletazonosító, "
    "fájlútvonal, sortartomány, rövid cím és 2-4 mondatos magyarázat arról, "
    "hogy mit bizonyít a részlet."
)

items = [
    (
        "M2.1",
        "Útvonalkezelés és lazy loading",
        "src/app/app.routes.ts",
        "5-21",
        "Az alkalmazás fő route-jai és az admin route guard bekötése.",
        "A projekt képernyőszerkezetét és az admin felület védett útvonalát mutatja.",
    ),
    (
        "M2.2",
        "Admin route guard",
        "src/app/guards/admin.guard.ts",
        "7-23",
        "Az admin/dolgozói felületre lépés előtt megvárja az auth állapotot, majd jogosultság alapján enged vagy loginra irányít.",
        "Rövid, jól érthető példa a frontend route-védelemre.",
    ),
    (
        "M2.3",
        "Szerepkör-ellenőrzés",
        "src/app/services/auth.service.ts",
        "104-122",
        "Admin e-mail, admin szerepkör, dolgozói szerepkör és belső staff jogosultság eldöntése.",
        "A jogosultságkezelés kliensoldali döntési pontját mutatja.",
    ),
    (
        "M2.4",
        "Firestore biztonsági szabályok",
        "firestore.rules",
        "1-45 és 288-316",
        "Admin helper függvények, products és orders jogosultsági szabályok.",
        "Ez legyen az egyik kiemelt melléklet, mert bizonyítja a szerveroldali jogosultságvédelmet.",
    ),
    (
        "M2.5",
        "Rendelés adatmodell",
        "src/app/models/order.model.ts",
        "3-63",
        "A rendelés adatszerkezete: vásárló, szállítás, számlázás, fizetés, kupon, tételek, státusz.",
        "Az adatmodell-magyarázat forráskódos bizonyítéka.",
    ),
    (
        "M2.6",
        "Checkout és rendelésmentés",
        "src/pages/checkout/checkout.ts",
        "381-481",
        "Kosár ellenőrzése, validáció, összegszámítás, rendelés objektum összeállítása és Firestore mentés.",
        "A fő vásárlói folyamat forráskódja, ezért mellékletben nagyon hasznos.",
    ),
    (
        "M2.7",
        "Űrlapvalidáció",
        "src/pages/checkout/checkout.ts + src/app/utils/form-validators.ts",
        "596-654 és 1-89",
        "Checkout validáció, e-mail és telefonszám ellenőrzés.",
        "A kézi tesztek negatív eseteihez kapcsolható.",
    ),
    (
        "M2.8",
        "Rendelésstátusz, audit és készlet",
        "src/app/services/order.service.ts",
        "41-134",
        "Firestore tranzakcióban történő státuszváltás, audit bejegyzés és készletkorrekció.",
        "A legerősebb üzleti logikai melléklet: konzisztencia, audit és készlet egyszerre.",
    ),
    (
        "M2.9",
        "Admin státuszváltás meghívása",
        "src/pages/admin/admin.ts",
        "1600-1610",
        "Az admin felület meghívja az auditált státuszváltó service metódust, admin UID/e-mail átadással.",
        "Rövid kapcsolódási pont a felület és a service réteg között.",
    ),
    (
        "M2.10",
        "Helyszíni vásárlás UI-folyamat",
        "src/pages/admin/admin.ts",
        "2335-2455",
        "Helyszíni vásárlás validációja, tételek összeállítása, local-admin salesChannel, rendelés létrehozása, számlagenerálás.",
        "Bemutatja a dolgozói/admin extra funkciót.",
    ),
    (
        "M2.11",
        "Helyszíni vásárlás tranzakció",
        "src/app/services/order.service.ts",
        "238-284",
        "Készletellenőrzés, készletcsökkentés és rendelésmentés Firestore tranzakcióban.",
        "Ha rövidíteni kell, ezt tartsd meg az M2.10 helyett, mert erősebb backend-logikát mutat.",
    ),
    (
        "M2.12",
        "CSV-import termékmentés",
        "src/app/services/product.service.ts",
        "59-151",
        "Valid termékek szűrése, SKU szerinti upsert, batch commit.",
        "Az admin tömeges termékfeltöltés forráskódos bizonyítéka.",
    ),
    (
        "M2.13",
        "Számlaszám és PDF-bizonylat",
        "src/app/services/order.service.ts + src/app/services/invoice.service.ts",
        "313-347 és 8-56",
        "Éves futószámos számlaszám tranzakcióban, majd PDF-letöltési adatok összeállítása.",
        "A bizonylatolási funkció mellékleti bizonyítéka.",
    ),
    (
        "M2.14",
        "AI-asszisztens kliensoldal",
        "src/app/services/chatbot-llm.service.ts",
        "39-88 és 92-128",
        "Domain-korlát, releváns katalógus építése, AI proxy meghívása, termékek visszakötése.",
        "Megmutatja, hogy az AI nem általános chatbot, hanem termékkatalógushoz kötött asszisztens.",
    ),
    (
        "M2.15",
        "OpenRouter proxy Worker",
        "workers/openrouter-proxy/src/index.js",
        "73-100 és 182-224",
        "Domainkérdés-szűrés, katalógus szanitizálás, OpenRouter API hívás, upstream hibakezelés.",
        "Fontos architektúra/biztonság melléklet: az API kulcs nem kerül frontendbe.",
    ),
]

doc.add_heading("Javasolt mellékleti sorrend", level=1)
table = doc.add_table(rows=1, cols=6)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.style = "Table Grid"
headers = ["Melléklet", "Cím", "Fájl", "Sorok", "Mi látható?", "Miért hasznos?"]
for cell, header in zip(table.rows[0].cells, headers):
    set_cell_text(cell, header, True)
    shade(cell, "D9EAF7")

for item in items:
    cells = table.add_row().cells
    for cell, text in zip(cells, item):
        set_cell_text(cell, text)

doc.add_heading("Ha nem fér be mind a 15 kódrészlet", level=1)
recommendations = [
    (
        "Minimum mellékletcsomag",
        "M2.4, M2.6, M2.8, M2.11 vagy M2.10, M2.12, M2.13, M2.15. "
        "Ez lefedi a biztonságot, checkoutot, tranzakciót, admin funkciót, "
        "CSV-importot, PDF-et és AI proxyt.",
    ),
    (
        "Bővített, de még kezelhető csomag",
        "A minimum mellé tedd be M2.1, M2.5, M2.7 és M2.14 részeket. "
        "Így az architektúra, adatmodell, validáció és AI kliensoldal is bizonyított.",
    ),
    (
        "Kerülendő",
        "Ne másold be egyben a teljes admin.ts fájlt. Mellékletben is csak célzott, "
        "értelmezhető részletek legyenek, különben a bíráló nem fogja végigolvasni.",
    ),
]
for label, body in recommendations:
    paragraph = doc.add_paragraph()
    paragraph.add_run(label + ": ").bold = True
    paragraph.add_run(body)

doc.add_heading("Főszövegben használható hivatkozási mondatok", level=1)
references = [
    "A route-struktúra és az admin útvonal védelme az M2.1-M2.2 mellékletben látható.",
    "A Firestore biztonsági szabályok releváns részlete az M2.4 mellékletben szerepel.",
    "A checkoutból létrejövő rendelés objektum összeállítását az M2.6 melléklet mutatja.",
    "A rendelésstátusz, auditnapló és készletkorrekció tranzakciós megoldását az M2.8 melléklet tartalmazza.",
    "A helyszíni vásárlás és bizonylatolás forráskódos részletei az M2.10-M2.13 mellékletekben találhatók.",
    "Az AI-asszisztens katalógushoz kötött működését és a proxy alapú API-hívást az M2.14-M2.15 melléklet mutatja be.",
]
for reference in references:
    doc.add_paragraph(reference)

doc.save(OUT)
print(OUT)
print(OUT.stat().st_size)
