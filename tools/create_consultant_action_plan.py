from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


OUT = Path(
    r"C:\Users\Dell\OneDrive\Asztali gép\szakdoga\konzulensi_javitasok_akcioterv.docx"
)


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def cell_text(cell, text, bold=False, size=8.6):
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP


doc = Document()
section = doc.sections[0]
section.top_margin = Inches(0.7)
section.bottom_margin = Inches(0.7)
section.left_margin = Inches(0.65)
section.right_margin = Inches(0.65)

styles = doc.styles
styles["Normal"].font.name = "Calibri"
styles["Normal"].font.size = Pt(10)
styles["Title"].font.name = "Calibri"
styles["Title"].font.size = Pt(20)
styles["Heading 1"].font.name = "Calibri"
styles["Heading 1"].font.size = Pt(15)

title = doc.add_paragraph()
title.style = "Title"
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.add_run("Konzulensi javítások - akcióterv")

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.add_run(
    "TDLWebshop szakdolgozat véglegesítési lista a 2026.05.22-i konzulensi visszajelzés alapján"
).italic = True

doc.add_heading("Gyors helyzetkép", level=1)
doc.add_paragraph(
    "A dolgozat beadásközeli állapotban van. Új nagy funkciót nem kell fejleszteni. "
    "A fő feladat a főszöveg tehermentesítése, a mérnöki magyarázatok erősítése, "
    "a mellékletek tényleges elérhetőségének rendezése és a végleges export előtti "
    "számozás/frissítés."
)
doc.add_paragraph(
    "A jelenlegi DOCX-ben a 'Hova kerüljön' megfogalmazás már nem található; a bizonyítékos "
    "táblázatban a megfelelőbb 'Hol szerepel / hol ellenőrizhető' forma szerepel. "
    "A fő kockázatok most a placeholder szövegek, az M1-M4 mellékletek pontos útvonala és "
    "a sok kép/kódrészlet megfelelő mellékletbe szervezése."
)

rows = [
    (
        "1",
        "Képernyőképek ritkítása",
        "4. GUI/UX és Mellékletek",
        "A főszövegben csak 4-6 elemzett kép maradjon: pageflow, kezdőlap/terméklista, checkout validáció, admin státuszváltás, helyszíni vásárlás, PDF-bizonylat.",
        "A többi kép menjen M3 mellékletbe vagy docs/ux/screenshots útvonalra.",
        "Magas",
    ),
    (
        "2",
        "Kódrészletek mellékletbe",
        "9. Biztonság, 8. Megvalósítás, Mellékletek",
        "A főszövegben legfeljebb 1-2 rövid kódrészlet maradjon, például AI proxy CORS vagy Firestore rules. A hosszabb checkout/order/admin kódok M2 mellékletbe kerüljenek.",
        "Használd a kodreszlet-melleklet-javaslat_szakdolgozathoz.docx listát.",
        "Magas",
    ),
    (
        "3",
        "Pageflow ábra",
        "4. GUI/UX",
        "A pageflow már szerepel és jó irány. Véglegesítsd úgy, hogy látszódjon a vásárlói fő út és az admin út is.",
        "Ha a saját pageflow képedet használod, legyen rajta S01-Sxx azonosító és rövid élcímke.",
        "Magas",
    ),
    (
        "4",
        "Mérnöki ábrák magyarázata",
        "6. Architektúra, 7. Adatmodell, 8. Megvalósítás",
        "Minden ábra után legyen 1 bekezdés: komponensek, adatirány, validáció, jogosultság, kapcsolódó use case.",
        "A jelenlegi architektúra-ábra magyarázata jó, ezt a mintát vidd végig minden mérnöki ábrán.",
        "Közepes",
    ),
    (
        "5",
        "Adatmodell kapcsolatainak erősítése",
        "7. Adatmodell",
        "A products, orders, users, savedCustomers, coupons, orderStatusAudit/invoiceCounters kapcsolatát konkrét folyamatokhoz kösd.",
        "A jelenlegi 117. bekezdés jó alap, egészítsd ki coupons + audit + invoiceCounters említéssel, ha még nincs a táblázatban.",
        "Magas",
    ),
    (
        "6",
        "Biztonsági kockázat -> megoldás párosítás",
        "9. Biztonság és adatvédelem",
        "Kockázatokat konkrét megoldásokkal párosíts: admin jogosultság -> Firestore rules; AI kulcs -> Worker secret; kupon -> validáció/MVP-korlát; PDF -> személyes adat kezelése.",
        "Javítsd az env.OPENROUTER_KEY hivatkozást env.OPENROUTER_API_KEY-re, mert a worker kódban ez szerepel.",
        "Magas",
    ),
    (
        "7",
        "Tesztelés pozitív/negatív esetekkel",
        "10. Tesztelés és M1",
        "A pozitív és negatív esetek már megjelennek. A főszövegben maradjon rövid összefoglaló, a részletes táblázat M1-ben legyen.",
        "A M1 melléklet útvonala most ne docs/testing/manual-test-log.md legyen, hanem a tényleges M1_Kezi_tesztjegyzokonyv_1.docx vagy repóútvonal.",
        "Magas",
    ),
    (
        "8",
        "Mellékletek tényleges elérhetősége",
        "Mellékletek",
        "M1-M4 hivatkozások legyenek valósak: M1 kézi tesztjegyzőkönyv, M2 kódrészlet melléklet, M3 képernyőképek/UX, M4 repó+reprodukció.",
        "Töltsd ki a GitHub URL-t, commit hash-t, deploy URL-t és Coospace/modulóra azonosítót.",
        "Nagyon magas",
    ),
    (
        "9",
        "Placeholder törlés",
        "4. GUI/UX és Mellékletek",
        "Töröld vagy cseréld a '[IDE KERÜL A KÉP]' és '[TÖLTSD KI]' jelöléseket.",
        "A PDF-bizonylat képhelye jelenleg konkrét placeholder, ezt mindenképp cseréld képre vagy vedd ki.",
        "Nagyon magas",
    ),
    (
        "10",
        "Végleges Word/PDF frissítés",
        "Teljes dokumentum",
        "Frissítsd a tartalomjegyzéket, oldalszámokat, ábra- és táblázatszámozást, majd exportáld PDF-be.",
        "Wordben Ctrl+A, F9; utána ellenőrizd a képaláírásokat és a tartalomjegyzéket.",
        "Nagyon magas",
    ),
]

doc.add_heading("Konkrét teendők", level=1)
table = doc.add_table(rows=1, cols=6)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.style = "Table Grid"
headers = ["#", "Teendő", "Hely", "Mit csinálj?", "Megjegyzés", "Prioritás"]
for cell, header in zip(table.rows[0].cells, headers):
    cell_text(cell, header, bold=True)
    shade(cell, "D9EAF7")

for row in rows:
    cells = table.add_row().cells
    for cell, text in zip(cells, row):
        cell_text(cell, text)

doc.add_heading("Kész szövegek, amelyeket be tudsz illeszteni", level=1)

snippets = [
    (
        "Mellékletek M1-M4 javasolt szövege",
        "M1. Kézi tesztjegyzőkönyv kitöltött állapotban: M1_Kezi_tesztjegyzokonyv_1.docx. "
        "M2. Kiemelt forráskódrészletek: a repóban és a kódrészlet-mellékletben megadott fájlok és sortartományok szerint. "
        "M3. GUI/UX képernyőképek és pageflow: docs/ux/ mappa, különösen docs/ux/pageflow.png, docs/ux/screens.csv és docs/ux/screenshots/. "
        "M4. Elektronikus melléklet és reprodukció: GitHub-repó, main branch, végleges commit hash, deploy URL és README.md futtatási lépések.",
    ),
    (
        "Biztonsági fejezetbe javított kulcsmondat",
        "Az OpenRouter API-kulcs nem kerül a frontend kódba: a Cloudflare Worker környezeti változóként, "
        "env.OPENROUTER_API_KEY néven éri el. Így a böngészőbe letöltött JavaScriptből és a GitHub repóból sem olvasható ki a titkos kulcs.",
    ),
    (
        "Melléklet-hivatkozás kódokra",
        "A hosszabb forráskódrészleteket nem a főszöveg tartalmazza, hanem az M2 melléklet. "
        "A főszövegben csak azok a rövid részletek szerepelnek, amelyekhez közvetlen mérnöki magyarázat kapcsolódik.",
    ),
    (
        "Tesztelési fejezet összekötő mondat",
        "A tesztelésnél minden kritikus folyamatnál szerepelt pozitív és negatív eset is: sikeres rendelés mellett hibás e-mail/telefon, "
        "admin státuszváltás mellett jogosulatlan módosítási kísérlet, érvényes CSV mellett hibás CSV, valamint releváns AI-kérdés mellett irreleváns kérdés.",
    ),
]

for heading, text in snippets:
    p = doc.add_paragraph()
    p.add_run(heading + ":").bold = True
    doc.add_paragraph(text)

doc.add_heading("Aktuálisan talált konkrét hibák a DOCX-ben", level=1)
for issue in [
    "4. fejezetben van egy '[IDE KERÜL A KÉP]' placeholder a PDF-bizonylat képnél.",
    "Mellékletek részben van '[TÖLTSD KI]', '[felhasználónév]' és '[deploy-URL]' placeholder.",
    "M1 jelenlegi hivatkozása docs/testing/manual-test-log.md, miközben a tényleges fájl a szakdoga mappában M1_Kezi_tesztjegyzokonyv_1.docx.",
    "A biztonsági szövegben env.OPENROUTER_KEY szerepel, de a worker kódban env.OPENROUTER_API_KEY a tényleges név.",
    "A fő DOCX-ben 18 inline kép van; ebből a főszövegben csak a magyarázott, legfontosabb képeket érdemes megtartani.",
]:
    doc.add_paragraph(issue)

doc.save(OUT)
print(OUT)
print(OUT.stat().st_size)
